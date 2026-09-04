# -*- coding: utf-8 -*-
"""
文档转换页（T03 / 从 app.py 迁移）
上传源文档 → 上传模板文档 → 转换配置 → 开始转换 → 下载结果 → 使用说明。
"""
import logging
import os
import time
from datetime import datetime
from pathlib import Path

import streamlit as st

from state import app_state

logger = logging.getLogger('WordStyle')


@st.cache_resource
def _get_logo_base64():
    """读取 logo.png 并 base64 编码（进程级缓存，仅读取一次）。"""
    import base64 as _b64
    _logo_path = Path(__file__).parent.parent / "resource" / "logo.png"
    if _logo_path.exists():
        return _b64.b64encode(_logo_path.read_bytes()).decode()
    return None


def _get_user_data():
    """获取当前用户数据（会话缓存，复用 app.py 入口统一初始化结果）。"""
    user_data = st.session_state.get('user_data')
    if user_data is None:
        user_data = {
            'user_id': app_state.get_user_id(),
            'balance': 0.0,
            'paragraphs_remaining': 0,
            'total_paragraphs_used': 0,
            'total_converted': 0,
            'is_active': False,
            'created_at': '',
            'last_login': '',
            'conversion_history': [],
        }
        st.session_state.user_data = user_data
    return user_data


def _render_header():
    """渲染全屏提示和说明信息。"""
    st.markdown("""
    <div style='background-color: #e3f2fd; padding: 10px; border-radius: 5px; margin-bottom: 10px;'>
    💡 <strong>提示：</strong>按 <kbd>F11</kbd> 键可以让浏览器全屏显示，获得更好的体验
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div style='background-color: #fff3cd; padding: 10px; border-radius: 5px; margin-bottom: 10px;'>
     <strong>说明：</strong>文档转换完成后，请及时下载，源文档和结果文档会被自动清理。自行做好标书检查，Good luck！
    </div>
    """, unsafe_allow_html=True)


def _load_user_defaults():
    """加载用户持久化的默认配置到 session_state（幂等）。"""
    _sm = {}
    try:
        uid = app_state.get_user_id()
        if uid and 'file_style_mappings' not in st.session_state:
            from data_manager import load_user_data
            ud = load_user_data(uid)
            if ud and 'style_mappings' in ud:
                _sm = ud.get('style_mappings', {}) or {}
    except Exception:
        _sm = {}

    if 'file_style_mappings' not in st.session_state:
        st.session_state.file_style_mappings = _sm if isinstance(_sm, dict) else {}

    _h = _sm.get('_default_hint_settings', {}) or {}
    _a = _sm.get('_default_answer_config', {}) or {}
    _l = _sm.get('_default_list_config', {}) or {}
    _t = _sm.get('_default_tbl_img_config', {}) or {}

    def _init(key, setter, value):
        if key not in st.session_state:
            setter(value)

    _init('do_answer_config', lambda v: app_state.set_do_answer_config(v), _a.get('do_answer', False))
    _init('answer_text_config', lambda v: app_state.set_answer_text_config(v), _a.get('answer_text', '应答：本投标人理解并满足要求。'))
    _init('answer_style_config', lambda v: app_state.set_answer_style_config(v), _a.get('answer_style', 'Normal'))
    _init('answer_mode_config', lambda v: app_state.set_answer_mode_config(v), _a.get('answer_mode', 'copy_chapter'))
    _init('answer_source_style_config', lambda v: app_state.set_answer_source_style_config(v), _a.get('answer_source_style', ''))
    _init('answer_copy_style_config', lambda v: app_state.set_answer_copy_style_config(v), _a.get('answer_copy_style', ''))

    _init('enable_list_style_config', lambda v: app_state.set_enable_list_style_config(v), _l.get('enable_list', True))
    _init('list_method_config', lambda v: app_state.set_list_method_config(v), _l.get('method', 'bullet'))
    _init('list_bullet_config', lambda v: app_state.set_list_bullet_config(v), _l.get('bullet', '•'))
    _init('list_style_config', lambda v: app_state.set_list_style_config(v), _l.get('style', 'Body Text'))
    _init('list_answer_method_config', lambda v: app_state.set_list_answer_method_config(v), _l.get('answer_method', 'bullet'))
    _init('list_answer_bullet_config', lambda v: app_state.set_list_answer_bullet_config(v), _l.get('answer_bullet', '•'))
    _init('list_answer_style_config', lambda v: app_state.set_list_answer_style_config(v), _l.get('answer_style', 'Body Text'))

    _init('enable_table_style_config', lambda v: app_state.set_enable_table_style_config(v), _t.get('enable_table_style', False))
    _init('table_style_config', lambda v: app_state.set_table_style_config(v), _t.get('table_style', 'Body Text'))
    _init('table_answer_style_config', lambda v: app_state.set_table_answer_style_config(v), _t.get('table_answer_style', ''))
    _init('enable_image_style_config', lambda v: app_state.set_enable_image_style_config(v), _t.get('enable_image_style', False))
    _init('image_style_config', lambda v: app_state.set_image_style_config(v), _t.get('image_style', 'Body Text'))
    _init('image_answer_style_config', lambda v: app_state.set_image_answer_style_config(v), _t.get('image_answer_style', ''))

    _rm = _sm.get('_default_remove_chapter_label', None)
    _init('remove_chapter_label_config', lambda v: app_state.set_remove_chapter_label_config(v), False if _rm is None else bool(_rm))

    _init('do_hint_config', lambda v: app_state.set_do_hint_config(v), _h.get('do_hint', False))
    _init('hint_type_config', lambda v: app_state.set_hint_type_config(v), _h.get('hint_type', 'text'))
    _init('hint_text_config', lambda v: app_state.set_hint_text_config(v), _h.get('hint_text', '招标文件原文'))
    _init('hint_style_config', lambda v: app_state.set_hint_style_config(v), _h.get('hint_style', 'Normal'))
    _init('hint_image_config', lambda v: app_state.set_hint_image_config(v), None)


def _render_usage_instructions():
    """使用说明（可折叠 expander，默认折叠）。"""
    from data_manager import get_config

    with st.expander("📖 使用说明", expanded=False):
        if 'free_paragraphs_display' not in st.session_state:
            try:
                v = get_config('free_paragraphs_daily')
                st.session_state.free_paragraphs_display = f"{int(v):,}" if v else "10,000"
            except Exception:
                st.session_state.free_paragraphs_display = "10,000"
        free_paragraphs_display = st.session_state.free_paragraphs_display

        st.markdown(f"""
### 🎯 本工具能帮你解决什么

如果你也是一名苦逼的售前，是否也被抄写标书这种低级的牛马工作折磨过？要把样式乱七八糟的需求文档或厂家方案，按照公司要求的标书样式重新复制粘贴一遍。为了不引入新格式、确保合稿顺利，还必须粘贴为纯文本，再一点一点调整格式……

**现在一键搞定**——过去数天的工作，现在只需几分钟：配置好源文档与模板的样式映射，工具自动完成全部转换。同时支持：
- 将招标文件的祈使语气自动转为投标人口吻（"应""须""必须"→统统消失）
- 五种应答句插入模式，自动批量生成应答内容
- 图片/文本提示语插入，标注原文位置
剩下的就是舒心检查、把标书打磨得更完美。

---

### ℹ️ 操作步骤

1. **上传源文档**：选择需要转换样式的 Word 文档（支持同时上传多个 `.docx` 文件）
2. **上传模板**：选择定义了目标样式的 Word 模板文档
3. **配置样式映射**（重要‼️）：点击「📊 配置样式映射」按钮，完成四步配置
4. **配置转换选项**：勾选祈使语气转换、插入提示语等
5. **点击「🚀 开始转换」**：系统自动处理并生成结果
6. **下载结果**：转换完成后下载文档

---

### 🎁 段落额度说明

**每日免费额度**：每位用户每天可获得 **{free_paragraphs_display}** 段落的免费转换额度。

**额度规则**：
- 免费额度按日计算，每天自动重置
- 只统计非标题的正文段落
- 转换失败的文件不计入额度消耗

---

### 📊 段落定义

**什么是段落？**
- 段落指 Word 文档中的**正文内容段落**
- **不包括**标题（Heading 1-9、标题 1-9 等样式）
- **包括**普通文本段落、列表项、表格外文字等

---

### 👤 账号绑定（可选）

系统默认使用设备指纹识别用户身份，**无需注册登录即可使用**。

**绑定账号的优势：**
- 更换设备后可登录账号恢复额度
- 用户名易记易识别，方便区分多设备使用场景
- 绑定后每日额度与账号关联，不受设备限制

> ⚠️ 一个设备指纹只能绑定一个账号，用户名不区分大小写且不可重复。
        """)


def render_conversion_page():
    """文档转换页入口（供 st.navigation 调用）。"""
    from components.config_panel import render_conversion_config
    from components.upload import (
        get_template_styles_list,
        detect_missing_heading_styles,
        count_template_styles,
    )
    from config import TEMPLATE_STYLE_THRESHOLD

    from components.sidebar import render_sidebar
    render_sidebar("conversion")

    _render_header()

    user_data = _get_user_data()
    user_id = app_state.get_user_id()

    # ==================== 上传源文档 ====================
    st.subheader("📄 上传源文档")
    source_files = st.file_uploader(
        "选择要转换的 Word 文档（支持 .docx，可多选）",
        type=['docx'],
        help="支持 .docx 格式，可同时选择多个文件",
        accept_multiple_files=True,
        key="source_uploader",
    )
    if source_files:
        app_state.set_current_source_files(list(source_files))
    current_source_files = app_state.get_current_source_files()

    if current_source_files:
        need_analyze = False
        current_file_names = [sf.name for sf in current_source_files]
        analyzed_file_names = list(st.session_state.get('file_styles_map', {}).keys())
        if not analyzed_file_names or set(current_file_names) != set(analyzed_file_names):
            need_analyze = True

        progress_bar = st.progress(0)
        status_text = st.empty()
        _last_progress_ts = [0.0]

        if need_analyze:
            progress_bar.progress(0)
            status_text.text(" 正在分析源文档...")
            start_time = time.time()

            file_styles_map = {}
            file_paragraph_counts = {}
            total_paragraphs = 0
            total_files = len(current_source_files)

            from docx import Document
            from doc_converter import DocumentConverter

            for idx, source_file in enumerate(current_source_files, 1):
                temp_source = f"temp_source_{user_id}_{source_file.name}"
                with open(temp_source, 'wb') as f:
                    f.write(source_file.getbuffer())

                doc = Document(temp_source)
                current_file_total = len(doc.paragraphs)
                file_paragraph_counts[source_file.name] = current_file_total
                total_paragraphs += current_file_total

                styles = set()
                converter_temp = DocumentConverter()
                list_virtual_styles = converter_temp.get_list_virtual_styles(doc)
                styles.update(list_virtual_styles)

                status_text.text(f"🔍 正在分析文件 {idx}/{total_files}: {source_file.name}...")

                for para_idx, para in enumerate(doc.paragraphs):
                    if para.style and para.style.name:
                        styles.add(para.style.name)
                        para_style_lower = para.style.name.lower()
                        if not (para_style_lower.startswith('heading') or para_style_lower.startswith('head')):
                            elem = para._element
                            pPr = elem.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
                            if pPr is not None:
                                outline = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl')
                                if outline is not None:
                                    val = outline.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                                    if val is not None:
                                        try:
                                            level = int(val) + 1
                                            if 1 <= level <= 9:
                                                styles.add(f'[大纲级别 {level}]')
                                        except ValueError:
                                            pass

                    if (para_idx + 1) % 10 == 0 or para_idx == len(doc.paragraphs) - 1:
                        completed_files_progress = (idx - 1) * (100 / total_files)
                        current_file_progress = ((para_idx + 1) / current_file_total) * (100 / total_files)
                        total_progress = completed_files_progress + current_file_progress
                        _now = time.time()
                        if total_progress >= 99.9 or _now - _last_progress_ts[0] >= 0.2:
                            _last_progress_ts[0] = _now
                            progress_bar.progress(min(total_progress / 100, 1.0))

                file_styles_map[source_file.name] = sorted(list(styles))
                if current_file_total == 0:
                    progress_bar.progress(min(idx * (100 / total_files) / 100, 1.0))

            elapsed = time.time() - start_time
            progress_bar.progress(1.0)
            status_text.text(f"[OK] 分析完成！耗时: {elapsed:.1f}秒")

            app_state.set_file_styles_map(file_styles_map)
            app_state.set_file_paragraph_counts(file_paragraph_counts)

            all_styles = set()
            for styles in file_styles_map.values():
                all_styles.update(styles)
            app_state.set_source_styles(sorted(list(all_styles)))
        else:
            file_styles_map = app_state.get_file_styles_map()
            file_paragraph_counts = st.session_state.get('file_paragraph_counts', {})
            all_styles = app_state.get_source_styles()
            progress_bar.progress(1.0)
            status_text.text("[OK] 已分析完成（使用缓存）")

        file_info = [(sf.name, file_paragraph_counts[sf.name]) for sf in current_source_files]
        total_paragraphs = sum(file_paragraph_counts.values())

        with st.expander(f"📄 源文档信息：{len(current_source_files)}个文件 | {len(all_styles)}种样式 | {total_paragraphs:,}段落", expanded=True):
            st.markdown(f"**✅ 已上传:** {len(current_source_files)} 个文件")
            st.markdown(f"**📋 检测到样式:** {len(all_styles)} 种 - {', '.join(all_styles[:10])}{'...' if len(all_styles) > 10 else ''}")
            st.markdown("**ℹ️ 文件详情：**")
            for fname, fpara in file_info:
                st.markdown(f"  • {fname}: {fpara:,} 个段落")
            st.markdown(f"**📊 总段落数:** {total_paragraphs:,}")
            if not st.session_state.get('show_download_buttons', False):
                if total_paragraphs > user_data.get('paragraphs_remaining', 0):
                    st.error(f"❌ 余额不足！需要 {total_paragraphs:,}，剩余 {user_data.get('paragraphs_remaining', 0):,}")

        # P1-1：源文档无标题样式引导提示
        try:
            missing_map = detect_missing_heading_styles(current_source_files, user_id)
            missing_files = [name for name, missing in missing_map.items() if missing]
            if missing_files:
                st.info(
                    "💡 检测到部分源文档未使用标题样式（如 1. / 1.1 / 1.1.1 编号标题以正文格式出现），"
                    "前往「🛠️ 工具箱 → 源文档标题预处理」可自动提取并修复标题级别。"
                )
        except Exception as e:
            logger.warning(f"标题样式检测失败: {e}")

    # ==================== 上传模板文档 ====================
    st.subheader("📋 上传模板文档")
    template_file = st.file_uploader(
        "选择模板文档（支持 .docx）",
        type=['docx'],
        help="用于定义目标样式的 Word 文档",
        key="template_uploader",
    )

    if template_file:
        if 'template_styles' in st.session_state:
            app_state.delete_key('template_styles')
            logger.info("[REFRESH] 清除旧模板样式缓存，准备重新解析")
        temp_template = f"temp_template_{user_id}.docx"
        with open(temp_template, 'wb') as f:
            f.write(template_file.getbuffer())
        app_state.set_current_temp_template(temp_template)
        app_state.set_last_template_name(template_file.name)

    current_temp_template = app_state.get_current_temp_template()
    last_template_name = app_state.get_last_template_name()

    if current_temp_template:
        need_analyze_template = (
            'template_styles' not in st.session_state
            or st.session_state.get('last_template_name') != last_template_name
        )
        template_progress_bar = st.progress(0)
        template_status_text = st.empty()

        if need_analyze_template:
            template_progress_bar.progress(0)
            template_status_text.text("[SEARCH] 正在分析模板样式...")
            template_progress_bar.progress(0.5)
            template_status_text.text("正在提取所有段落样式...")
            template_styles_list = get_template_styles_list(current_temp_template)
            template_progress_bar.progress(1.0)
            template_status_text.text(f"[OK] 已提取 {len(template_styles_list)} 种样式！")
            app_state.set_template_styles(template_styles_list)
            app_state.set_last_template_name(last_template_name)
        else:
            template_progress_bar.progress(1.0)
            template_status_text.text("[OK] 已分析完成（使用缓存）")

        template_styles = app_state.get_template_styles() or []
        with st.expander(f"📋 模板文档信息：{os.path.basename(current_temp_template)} | {len(template_styles)}种样式", expanded=True):
            st.markdown(f"**✅ 已上传:** {os.path.basename(current_temp_template)}")
            st.markdown(f"**📋 检测到样式:** {len(template_styles)} 种 - {', '.join(template_styles[:10])}{'...' if len(template_styles) > 10 else ''}")

        # P1-2：模板样式过多引导提示
        try:
            style_count = count_template_styles(current_temp_template)
            if style_count > TEMPLATE_STYLE_THRESHOLD:
                st.info(
                    f"💡 模板包含 {style_count} 种样式（超过 {TEMPLATE_STYLE_THRESHOLD} 种），"
                    "前往「🛠️ 工具箱 → 模板样式精简」可删除未使用样式，降低映射复杂度。"
                )
        except Exception as e:
            logger.warning(f"模板样式统计失败: {e}")

    # ==================== 转换配置 ====================
    st.markdown("---")
    st.subheader("⚙️ 转换配置")

    _load_user_defaults()

    if 'do_mood_config' not in st.session_state:
        app_state.set_do_mood_config(True)

    st.markdown("---")
    map_col1, map_col2 = st.columns([2, 8])
    with map_col1:
        if st.button("📊 配置样式映射", key="open_style_mapping_btn", use_container_width=True,
                     help="完整的四步样式配置（标题映射、应答句、正文映射、表格/图片/列表兜底）"):
            from components.dialogs.style_mapping import show_style_mapping_dialog
            show_style_mapping_dialog()
    with map_col2:
        if current_source_files and st.session_state.get('template_styles'):
            st.caption("点击按钮配置当前文件的样式映射")
        else:
            st.caption("上传源文档和模板文档后即可配置样式映射")

    result = render_conversion_config()
    do_mood, do_answer, list_bullet, answer_text, answer_style, answer_mode = result[0:6]
    do_hint, hint_type, hint_text, hint_image_path, hint_style = result[6:11]
    answer_source_style, answer_copy_style = result[11:13]
    list_method, list_style, list_answer_method, list_answer_style, list_answer_bullet = result[13:18]
    remove_chapter_label = result[18]
    enable_list_style = result[19] if len(result) > 19 else True

    if not do_answer:
        answer_text = app_state.get_answer_text_config()
        answer_style = app_state.get_answer_style_config()
        answer_mode = app_state.get_answer_mode_config()
        answer_source_style = app_state.get_answer_source_style_config()
        answer_copy_style = app_state.get_answer_copy_style_config()

    # ==================== 开始转换 ====================
    st.markdown("---")
    is_converting = st.session_state.get('is_converting', False)

    if is_converting:
        st.warning("⏳ **正在进行前台转换，请稍后...**\n\n转换期间无法进行其他操作，请耐心等待转换完成。")
        st.info("💡 转换完成后将自动恢复操作权限")
    else:
        if st.button("🚀 开始转换", type="primary", use_container_width=True):
            st.session_state.conversion_file_results = []
            st.session_state.recent_results = []
            st.session_state.show_download_buttons = False

            progress_placeholder = st.empty()
            status_placeholder = st.empty()
            progress_bar = progress_placeholder.progress(0)

            current_source_files = st.session_state.get('current_source_files', None)
            current_temp_template = st.session_state.get('current_temp_template', None)

            if not current_source_files or not current_temp_template:
                st.error("❌ 请上传源文档和模板文档")
                status_placeholder.text("[ERROR] 验证失败：缺少文件")
                progress_bar.progress(0)
                st.stop()
            elif not os.path.exists(current_temp_template):
                st.error("❌ 文件上传失败，请重试")
                status_placeholder.text("[ERROR] 验证失败：文件上传错误")
                progress_bar.progress(0)
                st.stop()
            else:
                st.session_state.is_converting = True
                status_placeholder.text("⏳ 正在验证输入...")
                progress_bar.progress(5)

            if 'file_paragraph_counts' in st.session_state and st.session_state.file_paragraph_counts:
                file_paragraph_counts = st.session_state.file_paragraph_counts
                file_info = [(sf.name, file_paragraph_counts[sf.name]) for sf in current_source_files]
                total_paragraphs = sum(file_paragraph_counts.values())
            else:
                logger.warning("file_paragraph_counts 不存在，使用兜底逻辑重新计算")
                from components.upload import count_paragraphs
                total_paragraphs = 0
                file_info = []
                for sf in current_source_files:
                    temp_source = f"temp_source_{user_id}_{sf.name}"
                    paragraphs = count_paragraphs(temp_source)
                    total_paragraphs += paragraphs
                    file_info.append((sf.name, paragraphs))

            progress_bar.progress(10)
            status_placeholder.text("⏳ 准备转换...")

            source_files_info = []
            for fname, fpara in file_info:
                temp_source = f"temp_source_{user_id}_{fname}"
                source_files_info.append((fname, temp_source, fpara))

            config = {
                'do_mood': do_mood,
                'answer_text': answer_text,
                'answer_style': answer_style,
                'list_bullet': list_bullet if list_bullet else "—",
                'do_answer_insertion': do_answer,
                'answer_mode': answer_mode,
                'answer_source_style': answer_source_style,
                'answer_copy_style': answer_copy_style,
                'list_method': list_method,
                'list_style': list_style,
                'list_answer_method': list_answer_method,
                'list_answer_style': list_answer_style,
                'list_answer_bullet': list_answer_bullet,
                'remove_chapter_label': remove_chapter_label,
                'custom_style_map': st.session_state.get('style_mapping', None),
            }

            try:
                status_placeholder.text("⏳ 正在初始化转换器...")
                progress_bar.progress(10)
                _conv_progress_ts = [0.0]

                # 注入用户自定义语气规则（无配置时用默认规则）
                from data_manager import get_tone_rules
                from doc_converter import DocumentConverter
                user_tone_rules = get_tone_rules(user_id)
                converter = DocumentConverter(tone_rules=user_tone_rules)
                progress_bar.progress(10)

                output_files = []
                success_count = 0
                fail_count = 0
                total_success_paragraphs = 0

                if 'conversion_file_results' not in st.session_state:
                    st.session_state.conversion_file_results = []

                for idx, source_file_obj in enumerate(current_source_files):
                    base_name = os.path.splitext(source_file_obj.name)[0]
                    output_filename = f"result_{base_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                    output_file = os.path.join("conversion_results", output_filename)
                    temp_source = f"temp_source_{user_id}_{source_file_obj.name}"

                    if not os.path.exists(temp_source):
                        logger.warning(f"临时文件 {temp_source} 不存在，从 UploadedFile 重新创建")
                        for sf in st.session_state.current_source_files:
                            if sf.name == source_file_obj.name:
                                with open(temp_source, 'wb') as f:
                                    f.write(sf.getbuffer())
                                break

                    file_paragraphs = 0
                    for fname, fpara in file_info:
                        if fname == source_file_obj.name:
                            file_paragraphs = fpara
                            break

                    status_placeholder.text(f" 正在转换第 {idx + 1}/{len(current_source_files)} 个文件 {source_file_obj.name} ({file_paragraphs:,} 段落)")

                    file_mapping = None
                    file_tbl_img_config = {}
                    file_list_config = {}
                    if 'file_style_mappings' in st.session_state and source_file_obj.name in st.session_state.file_style_mappings:
                        file_mapping_data = st.session_state.file_style_mappings[source_file_obj.name]
                        file_mapping = {k: v for k, v in file_mapping_data.items() if not k.startswith('_')}
                        file_tbl_img_config = file_mapping_data.get('_table_image_style', {})
                        file_list_config = file_mapping_data.get('_list_config', {})

                    if not file_mapping:
                        default_style_map = st.session_state.file_style_mappings.get('_default_style_map', {})
                        if default_style_map:
                            file_mapping = {k: v for k, v in default_style_map.items() if not k.startswith('_')}
                            if file_mapping:
                                st.info(f"📋 {source_file_obj.name}: 使用默认样式映射 ({len(file_mapping)} 个样式)")

                    if not file_tbl_img_config:
                        file_tbl_img_config = st.session_state.file_style_mappings.get('_default_tbl_img_config', {})
                    if not file_list_config:
                        file_list_config = st.session_state.file_style_mappings.get('_default_list_config', {})

                    warnings_list = []
                    def warning_callback(msg):
                        warnings_list.append(msg)

                    def make_progress_callback(file_idx, total_files):
                        def callback(step, message):
                            base_progress = 10 + int((file_idx / total_files) * 70)
                            step_progress = int((step / 7) * (70 / total_files))
                            current_progress = min(base_progress + step_progress, 80)
                            _now = time.time()
                            if _now - _conv_progress_ts[0] >= 0.2:
                                _conv_progress_ts[0] = _now
                                progress_bar.progress(current_progress)
                                status_placeholder.text(f"⏀{message}")
                        return callback

                    source_styles_for_file = st.session_state.file_styles_map.get(source_file_obj.name, None)

                    file_table_answer_style = file_tbl_img_config.get('table_answer_style') or st.session_state.get('table_answer_style_config', '')
                    file_image_answer_style = file_tbl_img_config.get('image_answer_style') or st.session_state.get('image_answer_style_config', '')

                    _list_bullet = file_list_config.get('bullet', list_bullet if list_bullet else "—")
                    _list_method = file_list_config.get('method', list_method)
                    _list_style = file_list_config.get('style', list_style)
                    _list_answer_method = file_list_config.get('answer_method', list_answer_method)
                    _list_answer_style = file_list_config.get('answer_style', list_answer_style)
                    _list_answer_bullet = file_list_config.get('answer_bullet', list_answer_bullet)
                    _enable_list_style = file_list_config.get('enable_list', enable_list_style)

                    success, actual_file, msg = converter.full_convert(
                        source_file=temp_source,
                        template_file=current_temp_template,
                        output_file=output_file,
                        custom_style_map=file_mapping,
                        do_mood=do_mood,
                        answer_text=answer_text,
                        answer_style=answer_style,
                        answer_source_style=answer_source_style,
                        answer_copy_style=answer_copy_style,
                        table_answer_style=file_table_answer_style,
                        list_bullet=_list_bullet,
                        list_method=_list_method,
                        list_style=_list_style,
                        list_answer_method=_list_answer_method,
                        list_answer_style=_list_answer_style,
                        list_answer_bullet=_list_answer_bullet,
                        do_answer_insertion=do_answer,
                        answer_mode=answer_mode,
                        do_hint_insertion=do_hint,
                        hint_type=hint_type,
                        hint_text=hint_text,
                        hint_image_path=hint_image_path,
                        hint_style=hint_style,
                        progress_callback=make_progress_callback(idx, len(current_source_files)),
                        warning_callback=warning_callback,
                        source_styles_cache=source_styles_for_file,
                        table_style_override=file_tbl_img_config.get('table_style') or st.session_state.get('table_style_config', 'Body Text'),
                        enable_table_style=file_tbl_img_config.get('enable_table_style', st.session_state.get('enable_table_style_config', False)),
                        image_style_override=file_tbl_img_config.get('image_style') or st.session_state.get('image_style_config', 'Body Text'),
                        enable_image_style=file_tbl_img_config.get('enable_image_style', st.session_state.get('enable_image_style_config', False)),
                        remove_chapter_label=remove_chapter_label,
                        enable_list_style=_enable_list_style,
                    )

                    if success:
                        output_files.append(actual_file)
                        success_count += 1
                        total_success_paragraphs += file_paragraphs
                        st.session_state.conversion_file_results.append({
                            'name': source_file_obj.name,
                            'status': 'success',
                            'paragraphs': file_paragraphs,
                            'warnings': warnings_list.copy(),
                        })
                    else:
                        fail_count += 1
                        st.session_state.conversion_file_results.append({
                            'name': source_file_obj.name,
                            'status': 'fail',
                            'msg': msg,
                        })

                progress_bar.progress(90)

                if success_count > 0:
                    progress_bar.progress(100)
                    if user_data.get('paragraphs_remaining', 0) >= total_success_paragraphs:
                        user_data['paragraphs_remaining'] -= total_success_paragraphs
                    else:
                        user_data['paragraphs_remaining'] = 0
                    user_data['total_converted'] = user_data.get('total_converted', 0) + success_count
                    user_data['total_paragraphs_used'] = user_data.get('total_paragraphs_used', 0) + total_success_paragraphs

                    conversion_record = {
                        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                        'files': len(current_source_files),
                        'success': success_count,
                        'failed': fail_count,
                        'paragraphs_charged': total_success_paragraphs,
                        'mode': 'foreground',
                    }
                    if 'conversion_history' not in user_data:
                        user_data['conversion_history'] = []
                    user_data['conversion_history'].append(conversion_record)

                    from data_manager import add_conversion_record, save_user_data
                    add_conversion_record(
                        files_count=len(current_source_files),
                        success_count=success_count,
                        failed_count=fail_count,
                        user_id=user_id,
                        paragraphs=total_success_paragraphs,
                    )
                    save_user_data(user_data, user_id)

                    if 'recent_results' not in st.session_state:
                        st.session_state.recent_results = []
                    for output_file in output_files:
                        if os.path.exists(output_file):
                            st.session_state.recent_results.append({
                                'path': output_file,
                                'name': os.path.basename(output_file),
                                'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                            })

                    st.session_state.is_converting = False
                    st.session_state.conversion_summary = {
                        'success_count': success_count,
                        'fail_count': fail_count,
                        'total_paragraphs': total_success_paragraphs,
                    }

                    try:
                        from file_manager import get_file_manager
                        fm = get_file_manager()
                        fm.cleanup_temp_files(user_id)
                    except Exception as cleanup_error:
                        logger.warning(f"临时文件清理失败（不影响转换结果）{cleanup_error}")

                    st.session_state.show_download_buttons = True
                    st.rerun()
                else:
                    status_placeholder.text("[ERROR] 转换失败")
                    progress_bar.progress(100)
                    st.session_state.is_converting = False
                    st.session_state.conversion_summary = {
                        'success_count': success_count,
                        'fail_count': fail_count,
                        'total_paragraphs': total_success_paragraphs,
                    }
                    st.session_state.show_download_buttons = True
                    st.error("❌ 所有文件转换失败，请检查错误信息")
                    st.info("💡 请查看下方的具体错误提示，修正后重试")
                    st.rerun()

            except Exception as e:
                st.session_state.is_converting = False
                error_msg = f"发生错误: {str(e)}"
                st.error(error_msg)
                if 'conversion_file_results' not in st.session_state or not st.session_state.conversion_file_results:
                    st.session_state.conversion_file_results = []
                    for source_file_obj in st.session_state.get('current_source_files', []):
                        st.session_state.conversion_file_results.append({
                            'name': source_file_obj.name,
                            'status': 'fail',
                            'msg': f"系统异常: {error_msg}",
                        })
                st.session_state.conversion_summary = {
                    'success_count': 0,
                    'fail_count': len(st.session_state.get('current_source_files', [])),
                    'total_paragraphs': 0,
                }
                st.session_state.show_download_buttons = True
                import traceback
                with st.expander("📋 查看详细错误堆栈"):
                    st.code(traceback.format_exc())
                st.rerun()

    # ==================== 下载结果 ====================
    if st.session_state.get('show_download_buttons', False):
        if st.session_state.get('conversion_summary'):
            summary = st.session_state.conversion_summary
            st.success(f"🎉 转换完成！成功 {summary['success_count']} 个，失败: {summary['fail_count']} 个")
            if summary['fail_count'] > 0:
                st.warning(f"⚠️ 有 {summary['fail_count']} 个文件转换失败")
            st.info(f"处理 {summary['total_paragraphs']:,} 个段落")

        for result in st.session_state.get('conversion_file_results', []):
            if result['status'] == 'success':
                st.success(f"✅ {result['name']} 转换成功")
                if result.get('warnings'):
                    ole_fail_count = sum(1 for w in result['warnings'] if '无法自动提取预览图' in w)
                    if ole_fail_count > 0:
                        st.warning(
                            f"⚠️ **{result['name']}** 文件中有 {ole_fail_count} 个 OLE 对象无法自动转换，"
                            f"已在文档相应位置标注「[OLE对象，请手动复制]」，请手动处理。"
                        )
            else:
                st.error(f"❌ {result['name']} 转换失败: {result.get('msg', '')}")

        st.subheader("📥 下载转换结果")
        for idx, file_info in enumerate(st.session_state.get('recent_results', [])):
            if os.path.exists(file_info['path']):
                with open(file_info['path'], 'rb') as f:
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        st.download_button(
                            label=f"[DOWNLOAD] 下载: {file_info['name']}",
                            data=f.read(),
                            file_name=file_info['name'],
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                            key=f"download_recent_{file_info['name']}",
                        )
                    with col2:
                        st.caption(f"转换时间: {file_info['time']}")
            else:
                st.warning(f"⚠️ 文件已过期或不存在: {file_info['name']}")

        st.markdown("---")

    # ==================== 使用说明 ====================
    _render_usage_instructions()
