# -*- coding: utf-8 -*-
"""
文档格式转换程序 - 核心处理模块
整合样式转换、祈使语气转换、标题后插入应答句功能
"""
import os
import sys
import re
import io
import logging
from datetime import datetime
from copy import deepcopy
from collections import defaultdict

try:
    from docx import Document
    from docx.table import _Cell
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement, parse_xml
    from lxml import etree
    from docx.shared import Emu
    from docx.enum.style import WD_STYLE_TYPE
    from docx.image.exceptions import UnrecognizedImageError
except ImportError:
    print("错误：未安装 python-docx 库。请运行: pip install python-docx")
    sys.exit(1)

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False


# ==================== 配置常量 ====================
DEFAULT_TARGET = "Normal"
DEFAULT_TABLE_STYLE = "Body Text"
DEFAULT_IMAGE_STYLE = "Body Text"
IMAGE_SCALE_RATIO = 2 / 3
LIST_BULLET_SYMBOL = "● "
TABLE_BORDER_SIZE = '4'
TABLE_BORDER_COLOR = '000000'
ANSWER_TEXT = "应答：本投标人理解并满足要求。"
ANSWER_STYLE = "应答句"
HEADING_STYLE_IDS = {'1', '2', '3', '4', '5', '6', '7', '8', '9'}
HEADING_STYLES = {f"Heading {i}" for i in range(1, 10)}

# 样式映射表
STYLE_MAP = {
    "Heading 1": "Heading 1",
    "Heading 2": "Heading 2",
    "Heading 3": "Heading 3",
    "Heading 4": "Heading 4",
    "Heading 5": "Heading 5",
    "Heading 6": "Heading 6",
    "List Paragraph": "List Paragraph",
}

OUTLINE_STYLE_MAP = {
    1: "Heading 1", 2: "Heading 2", 3: "Heading 3",
    4: "Heading 4", 5: "Heading 5", 6: "Heading 6",
    7: "Heading 7", 8: "Heading 8", 9: "Heading 9",
}

# 祈使语气替换规则
MULTI_IMPERATIVE_TO_STATEMENT = {
    "必须": "将", "不得": "不会", "不应": "不会", "不可": "不会",
    "不能": "无法", "切勿": "不要", "严禁": "禁止", "请勿": "请避免",
    "不许": "不允许",
}

MULTI_EXCEPTIONS = [
    "不可抗力", "不得已", "不由得", "不可通行", "不可开交", "不可理喻", "不可或缺",
    "不得少于", "不得大于", "不得超过", "不得低于", "不得高于", "不得小于", "不得用于",
    "不可否认", "不可避免", "不可逆", "不可分割",
]

SINGLE_REPLACE = {"应": "将", "须": "将"}

EXCEPTION_WORDS_YING = [
    "响应", "应用", "适应", "相应", "供应", "反应", "效应", "对应", "有应", "报应",
    "呼应", "感应", "应邀", "应酬", "应允", "应声", "应景", "应试", "应变", "应付",
    "应急", "应验", "应战", "应征", "应运", "应答", "应对", "应接", "应诺", "应求",
    "应时", "应需",
]
EXCEPTION_WORDS_XU = ["必须", "无须", "无需","须知"]

# "应+对"分离结构标志动词：用于区分复合词"应对"(yìng duì) 和 情态动词"应"+介词"对"(yīng+duì)
# 模式：应 + 对 + <名词短语> + [动词] → 类型B，不应视为例外
_YING_DUI_VERBS = re.compile(
    r'(负责|进行|予以|加以|负有|承担|提供|作出|做出|提交|出示'
    r'|说明|描述|解释|保证|给予|出具|支付|赔偿|归还|返还|退回'
    r'|履行|执行|实施|开展|组织|落实|协调|处理|管理|审核|审批'
    r')'
)
REPLACE_MAP = {
    "投标人需要": "本投标人",
    "投标人需": "本投标人",
    "投标人": "本投标人",
}


def build_word_pattern(word):
    return r'(?<![a-zA-Z0-9])' + re.escape(word) + r'(?![a-zA-Z0-9])'


def clean_list_numbering(text):
    """清理开头数字编号：1、 1） (1) 等"""
    pattern = r'^\s*(?:\d+[、\)）]|[（(]\d+[）\)])\s*'
    cleaned = re.sub(pattern, '', text, count=1)
    return cleaned


MULTI_IMPERATIVE_PATTERNS = [build_word_pattern(w) for w in MULTI_IMPERATIVE_TO_STATEMENT.keys()]
MULTI_IMPERATIVE_REGEX = re.compile('|'.join(MULTI_IMPERATIVE_PATTERNS))

SINGLE_IMPERATIVE_PATTERNS = [build_word_pattern(w) for w in SINGLE_REPLACE.keys()]
SINGLE_IMPERATIVE_REGEX = re.compile('|'.join(SINGLE_IMPERATIVE_PATTERNS))

REPLACE_REGEX = None
if REPLACE_MAP:
    patterns = []
    for word, repl in REPLACE_MAP.items():
        if word.startswith("投标人"):
            pat = r'(?<![本])' + re.escape(word) + r'(?![a-zA-Z0-9])'
        else:
            pat = build_word_pattern(word)
        patterns.append(pat)
    REPLACE_REGEX = re.compile('|'.join(patterns))


class DocumentConverter:
    """文档转换器主类"""
    
    def __init__(self):
        self.logger = None
        self.stats = {"para": 0, "table": 0, "heading": 0}
        self.source_styles = set()  # 源文档中使用的样式
        self.template_styles = set()  # 模板文档中的样式
        self.list_bullet = LIST_BULLET_SYMBOL  # 列表段落符号，默认为配置常量
        
    def setup_logger(self, source_file):
        log_filename = os.path.splitext(source_file)[0] + "_err.log"
        logger = logging.getLogger(f"converter_{os.path.basename(source_file)}")
        logger.setLevel(logging.WARNING)
        if not logger.handlers:
            handler = logging.FileHandler(log_filename, encoding='utf-8')
            formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
            handler.setFormatter(formatter)
            logger.addHandler(handler)
        self.logger = logger
        return logger
    
    def get_all_styles_from_doc(self, doc):
        """获取文档中使用的所有样式（包括虚拟大纲级别样式和列表段落虚拟样式）
        
        Args:
            doc: Document对象
        Returns:
            set: 样式名集合
        """
        styles = set()
        for para in doc.paragraphs:
            if para.style and para.style.name:
                styles.add(para.style.name)
        # 额外收集具有 outlineLvl 但样式为 Normal 的段落，生成虚拟大纲样式名
        outline_styles = self.get_outline_virtual_styles(doc)
        styles.update(outline_styles)
        # 额外收集具有自动编号的列表段落，生成虚拟列表样式名（如 '1 列表段落'、'● 列表段落'）
        list_styles = self.get_list_virtual_styles(doc)
        styles.update(list_styles)
        return styles
    
    def get_outline_virtual_styles(self, doc):
        """检测文档中通过大纲级别（outlineLvl）标记但无独立样式的段落，
        返回虚拟样式名称集合（如 '[大纲级别 1]'、'[大纲级别 2]'）。
        仅统计那些段落应用的样式名称为 'Normal' 或其他无大纲级别的普通样式，
        且段落自身有 outlineLvl 属性（直接设置）的段落。"""
        virtual_styles = set()
        # 常见标题样式名，这些不需要虚拟化
        actual_heading_style_names = {'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5',
                                       'Heading 6', 'Heading 7', 'Heading 8', 'Heading 9',
                                       'heading 1', 'heading 2', 'heading 3', 'heading 4', 'heading 5',
                                       'heading 6', 'heading 7', 'heading 8', 'heading 9',
                                       'Heading1', 'Heading2', 'Heading3', 'Heading4', 'Heading5',
                                       'Heading6', 'Heading7', 'Heading8', 'Heading9',
                                       'head 1', 'head 2', 'head 3', 'head 4', 'head 5',
                                       'head 6', 'head 7', 'head 8', 'head 9',
                                       '标题 1', '标题 2', '标题 3', '标题 4', '标题 5',
                                       '标题 6', '标题 7', '标题 8', '标题 9'}
        
        for para in doc.paragraphs:
            para_style_name = para.style.name if para.style and para.style.name else 'Normal'
            # 如果段落已经有已知的标题样式名，跳过
            if para_style_name in actual_heading_style_names:
                continue
            
            elem = para._element
            pPr = elem.find(qn('w:pPr'))
            if pPr is not None:
                outline = pPr.find(qn('w:outlineLvl'))
                if outline is not None:
                    val = outline.get(qn('w:val'))
                    if val is not None:
                        try:
                            level = int(val) + 1  # 转为 1-9 级别
                            if 1 <= level <= 9:
                                virtual_styles.add(f'[大纲级别 {level}]')
                        except ValueError:
                            pass
        return virtual_styles
    
    def get_template_styles(self, template_doc):
        """获取模板文档中的所有可用样式"""
        styles = set()
        for style in template_doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                styles.add(style.name)
        return styles
    
    def clear_document_content(self, doc):
        """清空文档内容但保留样式"""
        for table in doc.tables:
            table._element.getparent().remove(table._element)
        for para in doc.paragraphs:
            p = para._element
            p.getparent().remove(p)
        doc.add_paragraph()
    
    def get_outline_level(self, paragraph_or_elem, doc=None):
        """获取段落的大纲级别
        
        参数可以是：
        - python-docx 的 Paragraph 对象
        - lxml 的 XML 元素
        """
        # 判断输入类型
        if hasattr(paragraph_or_elem, '_element'):
            # 这是 Paragraph 对象
            elem = paragraph_or_elem._element
        else:
            # 这是 XML 元素
            elem = paragraph_or_elem
        
        # 安全检查：确保elem是XML元素
        if not hasattr(elem, 'tag'):
            return 0
        if elem.tag != qn('w:p'):
            return 0
        
        # 1. 从段落自身 w:pPr/w:outlineLvl 获取
        pPr = elem.find(qn('w:pPr'))
        if pPr is not None:
            outline = pPr.find(qn('w:outlineLvl'))
            if outline is not None:
                val = outline.get(qn('w:val'))
                if val is not None:
                    try:
                        level = int(val) + 1
                        return level
                    except ValueError:
                        pass
        
        # 2. 从段落应用的样式中获取（仅当提供doc时）
        if doc is not None and pPr is not None:
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                style_id = pStyle.get(qn('w:val'))
                if style_id:
                    try:
                        style = doc.styles[style_id]
                        style_elem = style._element
                        style_pPr = style_elem.find(qn('w:pPr'))
                        if style_pPr is not None:
                            outline = style_pPr.find(qn('w:outlineLvl'))
                            if outline is not None:
                                val = outline.get(qn('w:val'))
                                if val is not None:
                                    try:
                                        return int(val) + 1
                                    except ValueError:
                                        pass
                    except KeyError:
                        pass
        
        return 0
    
    def is_toc_paragraph(self, paragraph):
        """检查段落是否为目录（TOC）"""
        # 只通过域代码来判断，不依赖样式名称，避免误判
        elem = paragraph._element
        
        # 检查是否包含 TOC 域指令
        has_toc_instr = False
        for instr_text in elem.findall('.//' + qn('w:instrText')):
            if instr_text.text and instr_text.text.strip().upper().startswith('TOC'):
                has_toc_instr = True
                break
        
        if has_toc_instr:
            return True
        
        # 检查是否有 PAGEREF 域（TOC 中的页码引用）且有超链接
        has_pageref = False
        for instr_text in elem.findall('.//' + qn('w:instrText')):
            if instr_text.text and 'PAGEREF' in instr_text.text.upper():
                has_pageref = True
                break
        
        has_hyperlink = len(elem.findall('.//' + qn('w:hyperlink'))) > 0
        
        # 只有同时有 PAGEREF 和超链接时，才判定为目录
        if has_pageref and has_hyperlink:
            return True
        
        return False
    
    def has_numbering(self, paragraph):
        """检查段落是否有编号
        numId=0 表示无编号（Word 中的特殊值），应视为没有编号。
        """
        pPr = paragraph._element.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                # 检查 numId 是否为 0（0 表示无编号）
                numId = numPr.find(qn('w:numId'))
                if numId is not None:
                    val = numId.get(qn('w:val'))
                    if val == '0':
                        return False
                return True
        return False
    
    def remove_auto_numbering(self, paragraph):
        """移除自动编号"""
        pPr = paragraph._element.get_or_add_pPr()
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            pPr.remove(numPr)
    
    def remove_manual_numbering(self, text):
        """移除手动编号（智能判断中文数字是否为编号）
        
        仅匹配独立的编号格式，不会误伤文本中的数字（如"17号线"中的"17"）。
        改进：支持"1 .总则"（数字与分隔符之间有空格）等变体格式。
        """
        fragment_patterns = [
            # 字母+数字多级编号：如"C5.1.1.1"、"A1.2"、"B3.4.5"等附录/章节编号
            # 允许点之间有空格：如"C5. 1.1.1"
            r'[A-Za-z]\d+(?:\s*\.\s*\d+)*(?:\s*[\.、，,）\)\s]|(?=\D|$))?',
            # 多级阿拉伯数字编号："1.1"、"1.1.1"、"14.2.1"等，点分隔，后面可选分隔符或直接跟文字
            # 允许数字与点之间有空格，如"1 .总则"
            r'\d+\s*\.\s*\d+(?:\s*\.\s*\d+)*(?:\s*[\.、，,）\)\s]|(?=\D|$))',
            # 单级阿拉伯数字编号："1."、"1、"、"1）"等，必须有显式分隔符，避免误伤"17号线"
            # 允许数字与分隔符之间有空格：如"1 .总则"、"1 、总则"
            r'\d+\s*[\.、，,）\)\s]',
            # 中文数字后带分隔符才视为编号
            r'[一二三四五六七八九十]+\s*[、.．)）]',
            # 括号内的中/阿拉伯数字编号（如"（二）"、"（1）"）
            r'（[一二三四五六七八九十0-9]+）',
            r'\([0-9]+\)',
            r'[①-⑩]',
            r'[A-Za-z]\.[\s、，]',
        ]
        pattern = r'^\s*(' + '|'.join(fragment_patterns) + r')[\s、，]*'
        compiled = re.compile(pattern)
        cleaned = text
        while True:
            m = compiled.match(cleaned)
            if m:
                cleaned = cleaned[m.end():]
            else:
                break
        return cleaned
    
    # ==================== 桌面版迁移新增方法 ====================
    
    def get_list_virtual_styles(self, doc):
        """检测文档中具有编号/符号的列表段落，返回分类后的虚拟样式名称集合。
        
        对于有自动编号（numPr）的段落：
          - 如果能获取编号格式信息，根据格式类型生成：
            '1 列表段落'（数字编号）、'● 列表段落'（符号编号）
          - 如果无法获取具体格式信息，统一归为 '● 列表段落'
        
        同一种格式只出现一个虚拟样式名（不按序号区分）。
        """
        virtual_styles = set()
        for para in doc.paragraphs:
            if self.has_numbering(para):
                fmt = self._detect_numbering_format(para)
                virtual_styles.add(fmt)
        return virtual_styles

    def _detect_numbering_format(self, paragraph):
        """检测段落的编号格式类型，返回分类标识字符串。
        
        返回格式：
          - '1 列表段落'：数字编号（1. / 1) / (1) / ①）
          - '● 列表段落'：符号编号（bullet）
        """
        try:
            pPr = paragraph._element.find(qn('w:pPr'))
            if pPr is not None:
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    numId_elem = numPr.find(qn('w:numId'))
                    ilvl_elem = numPr.find(qn('w:ilvl'))
                    if numId_elem is not None:
                        numId = numId_elem.get(qn('w:val'))
                        ilvl = ilvl_elem.get(qn('w:val')) if ilvl_elem is not None else '0'
                        if numId:
                            # ★ 修复：通过 paragraph.part.numbering_part.element 访问 numbering XML，
                            # 而不是 getroottree().getroot()（后者返回的是 document body 根，不包含 numbering）
                            numbering_root = paragraph.part.numbering_part.element
                            if numbering_root is not None:
                                # 查找 num 元素
                                num_elem = None
                                for n in numbering_root.findall(qn('w:num')):
                                    if n.get(qn('w:numId')) == numId:
                                        num_elem = n
                                        break
                                if num_elem is not None:
                                    abstractNumId_elem = num_elem.find(qn('w:abstractNumId'))
                                    if abstractNumId_elem is not None:
                                        abstractNumId = abstractNumId_elem.get(qn('w:val'))
                                        if abstractNumId:
                                            # 查找 abstractNum 定义
                                            abs_num = None
                                            for an in numbering_root.findall(qn('w:abstractNum')):
                                                if an.get(qn('w:abstractNumId')) == abstractNumId:
                                                    abs_num = an
                                                    break
                                            if abs_num is not None:
                                                # 查找对应 ilvl 的级别定义
                                                target_lvl = None
                                                for l in abs_num.findall(qn('w:lvl')):
                                                    if l.get(qn('w:ilvl')) == ilvl:
                                                        target_lvl = l
                                                        break
                                                if target_lvl is None:
                                                    target_lvl = abs_num.find(qn('w:lvl'))
                                                if target_lvl is not None:
                                                    numFmt = target_lvl.find(qn('w:numFmt'))
                                                    if numFmt is not None:
                                                        fmt_val = numFmt.get(qn('w:val'))
                                                        if fmt_val == 'bullet':
                                                            return '● 列表段落'
                                                        elif fmt_val == 'decimal':
                                                            return '1 列表段落'
                                            # 查找替代格式：通过 numStyleLink
                                            styleLink = abs_num.find(qn('w:numStyleLink'))
                                            if styleLink is not None:
                                                val = styleLink.get(qn('w:val'))
                                                if val:
                                                    return '1 列表段落'
        except Exception:
            pass

        text = paragraph.text.strip() if paragraph.text else ''
        if text:
            if text[0].isdigit():
                return '1 列表段落'
            if text[0] in ('●', '◆', '▪', '▸', '➢', '○', '·', '', '-', '–', '*', '+', '·'):
                return '● 列表段落'
            if text.startswith('(') or text.startswith('（'):
                return '1 列表段落'
            if text.startswith('①') or text.startswith('②') or text.startswith('③'):
                return '1 列表段落'
            if text[0] in '一二三四五六七八九十':
                if len(text) > 1 and text[1] in ('、', '，', '　', '.'):
                    return '1 列表段落'
        return '● 列表段落'

    def remove_chapter_section_marking(self, text):
        """移除"第X章/第X节/第X篇/第X部分"等章节标记
        
        匹配如：第一章、第一节、第一篇、第二部分、第二章、第二节等。
        仅对文本中实际存在的章节标记进行清理，不影响自动编号。
        """
        if not text:
            return text
        chapter_pattern = r'^\s*第[一二三四五六七八九十]+(?:部分|[章节篇])[\s、，]*'
        cleaned = re.sub(chapter_pattern, '', text).strip()
        return cleaned

    def _clean_residual_numbering_artifacts(self, text):
        """清理自动编号段落中残留的手动编号痕迹
        
        当段落具有自动编号（numPr）时，原文中可能残留：
        - 前导点号：如".总则"（原始文档中自动编号提供"1"，文本残留"."）
        - 空格+点号：如" .总则"
        - 前导空格：如"  概述"
        - 其他分隔符残留
        
        此方法仅在段落有自动编号时调用，用于清理这些残留字符。
        """
        if not text:
            return text
        cleaned = re.sub(r'^[\s\.．、，,）\)]+', '', text)
        return cleaned

    def _resolve_auto_numbering_text(self, para):
        """解析段落自动编号的文本表示（如numId=4, ilvl=0 → "第二节"）
        
        通过访问源文档的 numbering part（内存中的 XML），查找 numId 对应的
        abstractNumId，再找到对应级别的 lvlText 和 numFmt，结合编号实例的
        当前值，生成完整的编号文本。如果无法解析，返回空字符串。
        
        改进：支持多级编号占位符（%1、%2、%3...）的完整解析。
        例如 lvlText='%1.%2.%3' 时，会分别解析级别0、1、2的编号值，
        正确替换所有占位符，避免出现未替换的"%2"等残留字符。
        """
        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            return ''
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            return ''
        
        numId_elem = numPr.find(qn('w:numId'))
        ilvl_elem = numPr.find(qn('w:ilvl'))
        if numId_elem is None or ilvl_elem is None:
            return ''
        
        numId = numId_elem.get(qn('w:val'))
        ilvl = ilvl_elem.get(qn('w:val'))
        if numId is None or ilvl is None:
            return ''
        
        try:
            doc = para.part.document
            numbering_part = doc.part.numbering_part
            root = numbering_part._element
            
            nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            abstractNumId = None
            for num in root.findall('.//w:num', nsmap):
                nid = num.get(qn('w:numId'))
                if nid == numId:
                    abs_ref = num.find('w:abstractNumId', nsmap)
                    if abs_ref is not None:
                        abstractNumId = abs_ref.get(qn('w:val'))
                    break
            
            if abstractNumId is None:
                return ''
            
            abstractNum = None
            for an in root.findall('.//w:abstractNum', nsmap):
                aid = an.get(qn('w:abstractNumId'))
                if aid == abstractNumId:
                    abstractNum = an
                    break
            
            if abstractNum is None:
                return ''
            
            current_lvl = None
            for l in abstractNum.findall('w:lvl', nsmap):
                li = l.get(qn('w:ilvl'))
                if li == ilvl:
                    current_lvl = l
                    break
            
            if current_lvl is None:
                return ''
            
            lvlText_elem = current_lvl.find('w:lvlText', nsmap)
            if lvlText_elem is None:
                return ''
            lvlText = lvlText_elem.get(qn('w:val'), '')
            
            if not lvlText:
                return ''
            
            placeholders = set(re.findall(r'%(\d+)', lvlText))
            
            result = lvlText
            for ph in placeholders:
                level = int(ph) - 1
                level_str = str(level)
                
                lvl_def = None
                for l in abstractNum.findall('w:lvl', nsmap):
                    li = l.get(qn('w:ilvl'))
                    if li == level_str:
                        lvl_def = l
                        break
                
                if lvl_def is None:
                    lvl_def = current_lvl
                
                start_elem = lvl_def.find('w:start', nsmap)
                start_val = 1
                if start_elem is not None:
                    start_val = int(start_elem.get(qn('w:val'), '1'))
                
                for num in root.findall('.//w:num', nsmap):
                    nid = num.get(qn('w:numId'))
                    if nid == numId:
                        for lo in num.findall('w:lvlOverride', nsmap):
                            loi = lo.get(qn('w:ilvl'))
                            if loi == level_str:
                                so = lo.find('w:startOverride', nsmap)
                                if so is not None:
                                    start_val = int(so.get(qn('w:val'), str(start_val)))
                        break
                
                count = 0
                for p in doc.paragraphs:
                    if p._element is para._element:
                        break
                    ppPr = p._element.find(qn('w:pPr'))
                    if ppPr is not None:
                        pnumPr = ppPr.find(qn('w:numPr'))
                        if pnumPr is not None:
                            pnid = pnumPr.find(qn('w:numId'))
                            if pnid is not None and pnid.get(qn('w:val')) == numId:
                                pilvl = pnumPr.find(qn('w:ilvl'))
                                if pilvl is not None and pilvl.get(qn('w:val')) == level_str:
                                    count += 1
                
                current_num = start_val + count
                
                numFmt_elem = lvl_def.find('w:numFmt', nsmap)
                if numFmt_elem is None:
                    continue
                numFmt = numFmt_elem.get(qn('w:val'), '')
                
                num_text = self._format_numbering_value(current_num, numFmt)
                if num_text is None:
                    continue
                
                result = result.replace('%' + ph, num_text)
            
            if result:
                result = result + ' '
            
            return result
        except Exception:
            return ''

    def _is_chapter_style_numbering(self, para):
        """判断段落的自动编号是否是章节样式（如"第%1节"、"第%1章"等）
        
        返回 True 如果编号模板中包含"第"字，表示是章节标记类编号。
        用于在不勾选"清除章/节/篇"时，仅对章节类编号进行解析和保留。
        """
        result = self._get_numbering_lvl_text(para)
        if result and '第' in result:
            return True
        return False

    def get_style_id_by_name(self, doc, style_name):
        """通过样式名称获取样式ID（处理 name 与 style_id 不一致的情况）"""
        if not style_name:
            return None
        try:
            style = doc.styles[style_name]
            return style.style_id
        except KeyError:
            return None

    def _get_numbering_lvl_text(self, para):
        """获取段落自动编号的 lvlText 模板
        
        返回 lvlText 的 val 属性值（如"第%1节"、"%1"、"%1.%2"等），
        如果无法获取则返回空字符串。
        """
        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            return ''
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            return ''
        numId_elem = numPr.find(qn('w:numId'))
        ilvl_elem = numPr.find(qn('w:ilvl'))
        if numId_elem is None or ilvl_elem is None:
            return ''
        numId = numId_elem.get(qn('w:val'))
        ilvl = ilvl_elem.get(qn('w:val'))
        if numId is None or ilvl is None:
            return ''
        
        try:
            doc = para.part.document
            numbering_part = doc.part.numbering_part
            root = numbering_part._element
            nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            abstractNumId = None
            for num in root.findall('.//w:num', nsmap):
                nid = num.get(qn('w:numId'))
                if nid == numId:
                    abs_ref = num.find('w:abstractNumId', nsmap)
                    if abs_ref is not None:
                        abstractNumId = abs_ref.get(qn('w:val'))
                    break
            if abstractNumId is None:
                return ''
            
            abstractNum = None
            for an in root.findall('.//w:abstractNum', nsmap):
                aid = an.get(qn('w:abstractNumId'))
                if aid == abstractNumId:
                    abstractNum = an
                    break
            if abstractNum is None:
                return ''
            
            current_lvl = None
            for l in abstractNum.findall('w:lvl', nsmap):
                li = l.get(qn('w:ilvl'))
                if li == ilvl:
                    current_lvl = l
                    break
            if current_lvl is None:
                return ''
            
            lvlText_elem = current_lvl.find('w:lvlText', nsmap)
            if lvlText_elem is None:
                return ''
            return lvlText_elem.get(qn('w:val'), '')
        except Exception:
            return ''

    def _format_numbering_value(self, num, fmt):
        """将数字编号值转换为指定格式的文本"""
        if fmt == 'decimal':
            return str(num)
        elif fmt in ('upperRoman', 'upperLetter'):
            return str(num)
        elif fmt in ('lowerRoman', 'lowerLetter'):
            return str(num)
        elif fmt in ('chineseCounting', 'chineseCountingThousand', 'japaneseCounting'):
            chinese_nums = '一二三四五六七八九十'
            if num <= 10:
                return chinese_nums[num - 1]
            elif num < 100:
                tens = num // 10
                ones = num % 10
                if tens == 1:
                    result = '十'
                else:
                    result = chinese_nums[tens - 1] + '十'
                if ones > 0:
                    result += chinese_nums[ones - 1]
                return result
            return str(num)
        elif fmt == 'bullet':
            return ''
        elif fmt == 'none':
            return ''
        else:
            return str(num)

    def _copy_numPr_to_paragraph(self, target_para, src_numPr):
        """复制自动编号定义(numPr)到目标段落"""
        if src_numPr is None:
            return
        new_pPr = target_para._element.find(qn('w:pPr'))
        if new_pPr is None:
            new_pPr = etree.SubElement(target_para._element, qn('w:pPr'))
            target_para._element.insert(0, new_pPr)
        new_numPr = etree.SubElement(new_pPr, qn('w:numPr'))
        for child_tag in ['w:numId', 'w:ilvl']:
            src_child = src_numPr.find(qn(child_tag))
            if src_child is not None:
                val = src_child.get(qn('w:val'))
                if val is not None:
                    new_child = etree.SubElement(new_numPr, qn(child_tag))
                    new_child.set(qn('w:val'), val)

    @staticmethod
    def _elem_has_numbering(elem, doc=None):
        """检查XML元素（w:p）是否有编号
        检查两种方式：
        1. 段落元素自身是否有 w:numPr（直接定义的编号）
        2. 段落样式（pStyle）是否在样式定义中包含 w:numPr（样式自带的编号）
        注意：numId=0 表示无编号，应视为没有编号。
        """
        from docx.oxml.ns import qn as _qn
        pPr = elem.find(_qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(_qn('w:numPr'))
            if numPr is not None:
                numId = numPr.find(_qn('w:numId'))
                if numId is not None:
                    val = numId.get(_qn('w:val'))
                    if val == '0':
                        return False
                return True
            pStyle = pPr.find(_qn('w:pStyle'))
            if pStyle is not None:
                pStyle_val = pStyle.get(_qn('w:val'))
                if pStyle_val and doc is not None:
                    try:
                        style_xml = doc.styles[pStyle_val]._element.xml
                        if '<w:numPr>' in style_xml:
                            return True
                    except Exception:
                        pass
        return False

    def get_target_style(self, original_style_name, template_doc, source_file=""):
        """获取目标样式名称"""
        # 使用实例变量中的样式映射，避免使用全局变量
        style_map = getattr(self, 'current_style_map', STYLE_MAP)
        target = style_map.get(original_style_name)
        if target is not None:
            try:
                template_doc.styles[target]
                return target
            except KeyError:
                try:
                    template_doc.styles[original_style_name]
                    return original_style_name
                except KeyError:
                    return DEFAULT_TARGET
        else:
            try:
                template_doc.styles[original_style_name]
                return original_style_name
            except KeyError:
                return DEFAULT_TARGET
    
    def get_image_size(self, image_bytes):
        """获取图片尺寸"""
        if not PIL_AVAILABLE:
            return None, None
        try:
            img = Image.open(io.BytesIO(image_bytes))
            return img.width, img.height
        except:
            return None, None
    
    def get_image_extent(self, blip_element):
        """从 blip 元素向上查找 wp:inline，并返回图片的原始显示尺寸 (cx, cy) 单位 EMU，若失败则返回 (None, None)"""
        parent = blip_element.getparent()
        while parent is not None:
            if parent.tag == qn('wp:inline'):
                extent = parent.find(qn('wp:extent'))
                if extent is not None:
                    cx = int(extent.get('cx', '0'))
                    cy = int(extent.get('cy', '0'))
                    return (cx, cy)
                break
            parent = parent.getparent()
        return (None, None)
    
    def resize_image_to_fixed_width(self, image_bytes, target_width_emu, dpi=96):
        """调整图片大小"""
        w_px, h_px = self.get_image_size(image_bytes)
        if w_px is None or h_px is None:
            return None, None
        w_emu = int(w_px / dpi * 914400)
        if w_emu <= target_width_emu:
            return w_emu, int(h_px / dpi * 914400)
        scale = target_width_emu / w_emu
        return int(w_emu * scale), int(h_px / dpi * 914400 * scale)
    
    def add_picture(self, run, img_bytes, page_width_emu, available_width_emu, emu_width=None, emu_height=None):
        """
        添加图片。
        若提供了源文档的显示尺寸 (emu_width, emu_height)，则优先使用；
        当图片宽度超出可用宽度时，缩放到页面宽度的 IMAGE_SCALE_RATIO，高度等比缩放。
        如果未提供尺寸，则使用图片的像素尺寸（96 DPI 计算）并做相同处理。
        """
        if not PIL_AVAILABLE:
            run.add_picture(io.BytesIO(img_bytes))
            return

        # 如果有源尺寸，直接使用；否则从像素计算
        if emu_width is not None and emu_height is not None:
            w_emu = emu_width
            h_emu = emu_height
        else:
            try:
                img = Image.open(io.BytesIO(img_bytes))
                w_px, h_px = img.size
                w_emu = int(w_px / 96 * 914400)
                h_emu = int(h_px / 96 * 914400)
            except:
                run.add_picture(io.BytesIO(img_bytes))
                return

        # 如果宽度超出可用宽度，则按页面宽度的 IMAGE_SCALE_RATIO 缩放
        if w_emu > available_width_emu:
            target_w = int(page_width_emu * IMAGE_SCALE_RATIO)
            scale = target_w / w_emu
            new_w = int(w_emu * scale)
            new_h = int(h_emu * scale)
            run.add_picture(io.BytesIO(img_bytes), width=Emu(new_w), height=Emu(new_h))
        else:
            run.add_picture(io.BytesIO(img_bytes), width=Emu(w_emu), height=Emu(h_emu))
    
    def _ole_preview_to_png(self, blob):
        """将 OLE 预览图转为 python-docx 支持的 PNG（WMF/EMF 等需转换）。
        返回 PNG bytes，失败返回 None。"""
        if not PIL_AVAILABLE:
            return None
        try:
            img = Image.open(io.BytesIO(blob))
            fmt = (img.format or '').upper()
            if fmt in ('PNG', 'JPEG', 'GIF', 'BMP', 'TIFF'):
                return blob
            # WMF/EMF 等矢量或非常见格式 → 转 PNG
            img = img.convert('RGB')
            buf = io.BytesIO()
            img.save(buf, 'PNG')
            return buf.getvalue()
        except Exception:
            return None
    
    def _extract_ole_preview(self, part, obj_elem):
        """从 OLE 对象中提取预览图并转为 PNG，返回 (png_bytes, emu_w, emu_h) 或 None。
        显示尺寸从 v:shape 的 style 属性解析（pt → EMU），避免用 PNG 像素反推导致放大失真。"""
        try:
            shape = obj_elem.find('{urn:schemas-microsoft-com:vml}shape')
            if shape is None:
                return None
            imagedata = shape.find('{urn:schemas-microsoft-com:vml}imagedata')
            if imagedata is None:
                return None
            rId = imagedata.get(qn('r:id')) or imagedata.get(qn('r:embed'))
            if not rId:
                return None
            style = shape.get('style') or ''
            m_w = re.search(r'width:\s*([\d.]+)pt', style)
            m_h = re.search(r'height:\s*([\d.]+)pt', style)
            if not m_w or not m_h:
                return None
            emu_w = int(round(float(m_w.group(1)) * 12700))
            emu_h = int(round(float(m_h.group(1)) * 12700))
            blob = part.related_parts[rId].blob
            png_bytes = self._ole_preview_to_png(blob)
            if png_bytes is None:
                return None
            return (png_bytes, emu_w, emu_h)
        except Exception:
            return None
    
    def set_table_width(self, table, width_emu):
        """设置表格宽度"""
        width_dxa = int(width_emu / 635)
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = parse_xml('<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            tbl.insert(0, tblPr)
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = parse_xml('<w:tblW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            tblPr.append(tblW)
        tblW.set(qn('w:w'), str(width_dxa))
        tblW.set(qn('w:type'), 'dxa')
    
    def set_table_borders(self, table):
        """为表格添加边框"""
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                for border_name in ['top', 'left', 'bottom', 'right']:
                    existing = tcPr.find(qn(f'w:{border_name}'))
                    if existing is not None:
                        tcPr.remove(existing)
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), TABLE_BORDER_SIZE)
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), TABLE_BORDER_COLOR)
                    tcPr.append(border)
    
    def copy_element_with_objects(self, source_elem, target_doc, target_style_name,
                                  page_width_emu, available_width_emu, para_idx=None, source_file="",
                                  warning_callback=None):
        """复制元素（包含图片、Visio图、OLE对象等）
        :param warning_callback: 警告回调函数 callback(message)
        """
        # 检查是否为段落
        if hasattr(source_elem, 'tag') and source_elem.tag == qn('w:p'):
            return self.copy_paragraph_with_images(
                source_elem, target_doc, target_style_name,
                page_width_emu, available_width_emu, 
                para_idx if para_idx is not None else 0, source_file,
                warning_callback
            )
        
        # 检查是否为表格
        elif hasattr(source_elem, 'tag') and source_elem.tag == qn('w:tbl'):
            # 这里需要找到对应的表格索引
            return None
        
        # 处理其他类型的元素（如OLE对象、Visio图等）
        else:
            return self.copy_special_element(source_elem, target_doc, target_style_name)
    
    def copy_special_element(self, source_elem, target_doc, target_style_name,
                             source_part=None, page_width_emu=None, available_width_emu=None,
                             warning_callback=None):
        """复制特殊元素（OLE对象、Visio图等）
        OLE 对象：提取其预览图转为 PNG 后作为普通图片插入；
        独立 VML 形状（非 OLE 预览）无法安全复制关系 ID，跳过其 XML 结构。
        :param warning_callback: 警告回调函数 callback(message)
        """
        try:
            # 创建一个新的段落来容纳特殊对象
            new_para = target_doc.add_paragraph()
            try:
                new_para.style = target_style_name
            except KeyError:
                new_para.style = target_doc.styles['Normal']
            
            # 检查是否包含OLE对象或形状（使用正确的命名空间）
            objects = source_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object')
            shapes = source_elem.findall('.//{urn:schemas-microsoft-com:vml}shape')
            
            if objects or shapes:
                inserted = False
                if source_part is not None and page_width_emu is not None and available_width_emu is not None:
                    for obj in objects:
                        result = self._extract_ole_preview(source_part, obj)
                        if result is not None:
                            png_bytes, emu_w, emu_h = result
                            pic_run = new_para.add_run()
                            self.add_picture(pic_run, png_bytes, page_width_emu, available_width_emu, emu_w, emu_h)
                            inserted = True
                            break
                if not inserted:
                    # 无法提取预览图时回退占位提示
                    new_para.add_run("[OLE对象]")
                    if warning_callback:
                        try:
                            warning_callback(f"[WARNING] 存在无法自动提取预览图的 OLE 对象，已跳过")
                        except Exception:
                            pass
                return new_para
            else:
                # 如果没有特殊对象，返回空段落
                return new_para
        except Exception as e:
            print(f"警告：复制特殊元素时出错: {e}")
            # 出错时返回一个空段落
            new_para = target_doc.add_paragraph()
            try:
                new_para.style = target_style_name
            except KeyError:
                new_para.style = target_doc.styles['Normal']
            return new_para
    
    def extract_and_add_images(self, source_para, new_para, page_width_emu, available_width_emu):
        """从源段落提取图片并添加到新段落（DRY原则：避免代码重复）"""
        for run in source_para.runs:
            blips = run._element.findall('.//' + qn('a:blip'))
            for blip in blips:
                rId = blip.get(qn('r:embed'))
                if rId:
                    try:
                        img_part = source_para.part.related_parts[rId]
                        img_bytes = img_part.blob
                        emu_w, emu_h = self.get_image_extent(blip)
                        pic_run = new_para.add_run()
                        self.add_picture(pic_run, img_bytes, page_width_emu, available_width_emu, emu_w, emu_h)
                    except Exception:
                        pass
    
    def copy_paragraph_with_images(self, source_para, target_doc, target_style_name,
                                   page_width_emu, available_width_emu, para_idx, source_file="",
                                   warning_callback=None, image_style_override=None, enable_image_style=False,
                                   remove_chapter_label=False,
                                   list_method='bullet', list_style='Body Text',
                                   enable_list_style=True, resolved_numbering_text=None):
        """复制段落（包含图片、Visio图、OLE对象等）
        :param warning_callback: 警告回调函数 callback(message)
        :param image_style_override: 图片样式覆盖（当enable_image_style=True时使用）
        :param enable_image_style: 是否启用图片样式覆盖
        :param remove_chapter_label: 是否清除"第X章/第X节"等章节标记
        :param list_method: 列表段落处理方式 'bullet'（符号）或 'style'（样式）
        :param list_style: 列表段落兜底样式名（当list_method='style'时使用）
        :param enable_list_style: 是否启用列表样式处理
        """
        # 调试：检查大纲级别
        outline_level = self.get_outline_level(source_para)
        
        # 检查是否为目录段落，如果是则保持原样式
        if self.is_toc_paragraph(source_para):
            new_para = target_doc.add_paragraph()
            # 保持原始样式或应用目标样式
            try:
                new_para.style = target_style_name
            except KeyError:
                new_para.style = target_doc.styles['Normal']
            
            # 复制内容但不修改
            for run in source_para.runs:
                new_run = new_para.add_run(run.text)
                # 复制格式
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                
                # 复制图片
                blips = run._element.findall('.//' + qn('a:blip'))
                for blip in blips:
                    rId = blip.get(qn('r:embed'))
                    if rId:
                        try:
                            img_part = source_para.part.related_parts[rId]
                            img_bytes = img_part.blob
                            emu_w, emu_h = self.get_image_extent(blip)
                            pic_run = new_para.add_run()
                            self.add_picture(pic_run, img_bytes, page_width_emu, available_width_emu, emu_w, emu_h)
                        except Exception:
                            pass
            return new_para
        
        # 检查是否包含 OLE 对象或 VML 形状
        has_ole_objects = source_para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object')
        has_vml_shapes = source_para._element.findall('.//{urn:schemas-microsoft-com:vml}shape')
        
        # 包含 OLE/VML 对象的段落，按源文档顺序重建内容：
        # OLE 对象提取其预览图作为普通图片插入；独立 VML 形状跳过其 XML；
        # 文本内容保持原位置。不 deepcopy OLE 的 XML（rId 在新文档中无效，会导致文档损坏）。
        if has_ole_objects or has_vml_shapes:
            # 创建新段落，设置目标样式（OLE 图片段落使用图片兜底样式）
            ole_para_style = target_style_name
            if enable_image_style and image_style_override:
                try:
                    target_doc.styles[image_style_override]
                    ole_para_style = image_style_override
                except KeyError:
                    pass
            new_para = target_doc.add_paragraph()
            try:
                new_para.style = ole_para_style
            except KeyError:
                new_para.style = target_doc.styles['Normal']
            
            ole_fallback = False
            for run in source_para.runs:
                # 检查这个run是否包含OLE对象
                run_has_ole = bool(run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object'))
                run_has_vml = bool(run._element.findall('.//{urn:schemas-microsoft-com:vml}shape'))
                
                if run_has_ole:
                    # 提取 OLE 预览图作为普通图片插入
                    inserted = False
                    for obj in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object'):
                        result = self._extract_ole_preview(source_para.part, obj)
                        if result is not None:
                            png_bytes, emu_w, emu_h = result
                            pic_run = new_para.add_run()
                            self.add_picture(pic_run, png_bytes, page_width_emu, available_width_emu, emu_w, emu_h)
                            inserted = True
                            break
                    if not inserted:
                        ole_fallback = True
                elif run_has_vml:
                    # 独立 VML 形状（非 OLE 预览）：跳过其 XML，不复制
                    pass
                elif run.text:
                    # 普通文本：复制文本及格式
                    new_run = new_para.add_run(run.text)
                    # 复制基本格式
                    if run.font:
                        new_run.font.bold = run.font.bold
                        new_run.font.italic = run.font.italic
                        new_run.font.underline = run.font.underline
                        if run.font.size:
                            new_run.font.size = run.font.size
                        if run.font.color and run.font.color.rgb:
                            try:
                                new_run.font.color.rgb = run.font.color.rgb
                            except:
                                pass
            
            if ole_fallback:
                warning_msg = f"[WARNING] 段落 {para_idx} 存在无法自动提取预览图的 OLE 对象，已跳过该对象"
                print(warning_msg)
                if warning_callback:
                    try:
                        warning_callback(warning_msg)
                    except:
                        pass
            
            return new_para
        
        # 普通段落处理（原有逻辑）
        has_image = any(run._element.findall('.//' + qn('a:blip')) for run in source_para.runs)
        new_para = target_doc.add_paragraph()
        src_style_name = source_para.style.name
        
        if outline_level > 0:
            style_map = getattr(self, 'current_style_map', STYLE_MAP)
            # 生成虚拟大纲样式名，用于查找用户映射
            virtual_style_name = f'[大纲级别 {outline_level}]'
            mapped_style = style_map.get(virtual_style_name)
            if mapped_style is None:
                # 回退：用原始样式名查找
                mapped_style = style_map.get(src_style_name)
            
            if mapped_style is not None:
                # 使用用户配置的映射样式
                final_style = mapped_style
            else:
                # 未配置映射，使用大纲级别对应的默认样式
                final_style = OUTLINE_STYLE_MAP.get(outline_level)
                if final_style is None:
                    final_style = f"Heading {outline_level}"
            
            # 检查目标文档中是否存在该样式
            try:
                target_doc.styles[final_style]
            except KeyError:
                # 样式不存在，尝试从源文档复制
                print(f"[WARNING] 模板中缺少样式 '{final_style}'，尝试从源文档复制...")
                try:
                    source_style = source_para.part.document.styles[final_style]
                    # 复制样式到目标文档
                    new_style = target_doc.styles.add_style(final_style, WD_STYLE_TYPE.PARAGRAPH)
                    new_style.base_style = target_doc.styles['Normal']
                    # 复制基本格式
                    if hasattr(source_style, 'font'):
                        if source_style.font.bold:
                            new_style.font.bold = True
                        if source_style.font.size:
                            new_style.font.size = source_style.font.size
                        if source_style.font.color and source_style.font.color.rgb:
                            new_style.font.color.rgb = source_style.font.color.rgb
                    print(f"[INFO] 已从源文档复制样式 '{final_style}'")
                except Exception as e:
                    print(f"[WARNING] 无法复制样式 '{final_style}': {e}，将尝试创建基本样式")
                    try:
                        # 创建基本的标题样式
                        new_style = target_doc.styles.add_style(final_style, WD_STYLE_TYPE.PARAGRAPH)
                        new_style.base_style = target_doc.styles['Normal']
                        new_style.font.bold = True
                        # 根据大纲级别设置字体大小
                        size_map = {1: 24, 2: 20, 3: 16, 4: 14, 5: 13, 6: 12, 7: 12, 8: 12, 9: 12}
                        from docx.shared import Pt
                        new_style.font.size = Pt(size_map.get(outline_level, 12))
                        print(f"[INFO] 已创建基本样式 '{final_style}'")
                    except Exception as e2:
                        print(f"[ERROR] 无法创建样式 '{final_style}': {e2}，使用默认样式")
                        final_style = DEFAULT_TARGET
        else:
            if has_image:
                # 图片段落样式：不受样式映射影响，只按单独的图片样式定义处理
                # 1. enable_image_style=True → 使用image_style_override指定的样式
                # 2. 未启用 → 保留源样式名（模板中存在则用，否则DEFAULT_TARGET）
                if enable_image_style and image_style_override:
                    # 级别1：复选框选中，使用覆盖样式
                    final_style = image_style_override
                    try:
                        target_doc.styles[final_style]
                    except KeyError:
                        final_style = DEFAULT_TARGET
                else:
                    # 级别2：保留源样式名
                    try:
                        target_doc.styles[src_style_name]
                        final_style = src_style_name
                    except KeyError:
                        final_style = DEFAULT_TARGET
            else:
                final_style = target_style_name
        
        try:
            new_para.style = final_style
        except Exception:
            new_para.style = target_doc.styles['Normal']
        
        is_heading_by_outline = outline_level > 0
        is_heading_by_style = src_style_name in HEADING_STYLES
        is_heading_by_target = final_style in HEADING_STYLES
        is_custom_src_heading = '标题' in src_style_name or src_style_name.startswith('Heading')
        is_custom_tgt_heading = '标题' in final_style or final_style.startswith('Heading')
        
        if is_heading_by_outline or is_heading_by_style or is_heading_by_target or is_custom_src_heading or is_custom_tgt_heading:
            # 使用 para.text 而非手动拼接 runs，因为 runs 不包含超链接(w:hyperlink)内部的 run 文本
            # 例如"第六章  图纸（如有）"中，"第六章"和"图纸"在 hyperlink 内部，runs 无法获取
            full_text = source_para.text
            has_auto_numbering = self.has_numbering(source_para)
            
            if remove_chapter_label:
                # 勾选"清除第X章/第X节"：
                # 1. 移除自动编号（章节标记可能来自自动编号，如numId=17→"第二节"）
                self.remove_auto_numbering(new_para)
                # 2. 清理文本中的"第X章/第X节/第X篇"
                cleaned_text = self.remove_chapter_section_marking(full_text)
                # 3. 清理常规手动编号（"一、"、"3.1"等）
                cleaned_text = self.remove_manual_numbering(cleaned_text)
                # 4. 如果有自动编号，清理残留的手动编号痕迹（如".总则"中的前导点）
                if has_auto_numbering:
                    cleaned_text = self._clean_residual_numbering_artifacts(cleaned_text)
            else:
                # 未勾选"清除第X章/第X节"：
                # 保留章节标记，只清理常规手动编号（"一、"、"3.1"、"（1）"等）
                cleaned_text = self.remove_manual_numbering(full_text)
                # 如果有自动编号，清理残留的手动编号痕迹
                if has_auto_numbering:
                    cleaned_text = self._clean_residual_numbering_artifacts(cleaned_text)
                # ★ 修复：不直接复制 numPr（不同文档的 numId 映射不同，会导致错误编号），
                # 改为选择性解析自动编号：
                # 1. 如果编号是章节样式（如"第%1节"→"第二节"），解析并拼接到文本前
                # 2. 普通数字编号（如"%1"→"1"、"%1.%2"→"1.1"）不解析，避免文本中出现冗余编号
                if has_auto_numbering and self._is_chapter_style_numbering(source_para):
                    numbering_text = self._resolve_auto_numbering_text(source_para)
                    if numbering_text:
                        cleaned_text = numbering_text + cleaned_text
            new_para.clear()
            new_para.add_run(cleaned_text)
            # [HIGH_VOLTAGE] 使用统一的图片处理方法
            self.extract_and_add_images(source_para, new_para, page_width_emu, available_width_emu)
            return new_para
        
        if self.has_numbering(source_para) and enable_list_style:
            if list_method == 'style':
                # "样式"模式：使用 Step 4 的 list_style 作为列表段落的样式。
                # 列表段落的样式由 Step 4 的"列表段落"配置区独立控制，不使用 Step 3 的样式映射结果。
                # 不加符号，不清除自动编号，只复制文本内容（保留原始编号）
                try:
                    target_doc.styles[list_style]
                    new_para.style = list_style
                except Exception:
                    try:
                        new_para.style = target_doc.styles['Normal']
                    except Exception:
                        pass
                # 检查目标样式本身是否已经包含 numPr 定义（如 BN_原文引用列表项目符号
                # 样式自带 numId=4 的项目符号编号）。如果样式自带编号，则不再从源段落
                # 拷贝 numPr，以免覆盖样式中定义的项目符号/编号格式。
                style_has_numPr = False
                try:
                    style_xml = target_doc.styles[list_style]._element.xml
                    if '<w:numPr>' in style_xml or '<w:numPr ' in style_xml:
                        style_has_numPr = True
                except Exception:
                    pass
                
                if not style_has_numPr:
                    # ★ 修复：如果目标样式没有自带编号定义（如 BN_正文），
                    # 说明用户希望将列表段落转为普通正文段落。
                    # 此时不应该从源段落复制自动编号，而应该移除原有编号。
                    # 清除源段落原有的自动编号，使其变为普通正文段落
                    self.remove_auto_numbering(new_para)
                # 使用 para.text 获取完整文本（包括超链接内部的run）
                list_runs_text = ''.join(run.text for run in source_para.runs)
                list_full_text = source_para.text
                if list_full_text and len(list_full_text) > len(list_runs_text):
                    # 有超链接内部文本，使用完整文本
                    # 但仍需处理图片
                    for run_idx, run in enumerate(source_para.runs):
                        blips = run._element.findall('.//' + qn('a:blip'))
                        if blips:
                            for blip in blips:
                                rId = blip.get(qn('r:embed'))
                                if rId:
                                    try:
                                        img_part = source_para.part.related_parts[rId]
                                        img_bytes = img_part.blob
                                        emu_w, emu_h = self.get_image_extent(blip)
                                        pic_run = new_para.add_run()
                                        self.add_picture(pic_run, img_bytes, page_width_emu, available_width_emu, emu_w, emu_h)
                                    except Exception:
                                        pass
                    if list_full_text.strip():
                        new_para.add_run(list_full_text.strip())
                else:
                    for run_idx, run in enumerate(source_para.runs):
                        blips = run._element.findall('.//' + qn('a:blip'))
                        if blips:
                            for blip in blips:
                                rId = blip.get(qn('r:embed'))
                                if rId:
                                    try:
                                        img_part = source_para.part.related_parts[rId]
                                        img_bytes = img_part.blob
                                        emu_w, emu_h = self.get_image_extent(blip)
                                        pic_run = new_para.add_run()
                                        self.add_picture(pic_run, img_bytes, page_width_emu, available_width_emu, emu_w, emu_h)
                                    except Exception:
                                        pass
                        else:
                            if run.text:
                                new_para.add_run(run.text)
                return new_para
            else:
                # "符号"模式：保留原有逻辑（添加 bullet 符号，清除编号）
                new_para.add_run(self.list_bullet)
                self.remove_auto_numbering(new_para)
                # 对于列表段落，使用专门的编号清理函数
                # 使用 para.text 获取完整文本（包括超链接内部的run）
                symbol_runs_text = ''.join(run.text for run in source_para.runs)
                symbol_full_text = source_para.text
                if symbol_full_text and len(symbol_full_text) > len(symbol_runs_text):
                    full_text = symbol_full_text
                else:
                    full_text = symbol_runs_text
                cleaned_text = clean_list_numbering(full_text)
                if cleaned_text:
                    new_para.add_run(cleaned_text)
                for run_idx, run in enumerate(source_para.runs):
                    blips = run._element.findall('.//' + qn('a:blip'))
                    for blip in blips:
                        rId = blip.get(qn('r:embed'))
                        if rId:
                            try:
                                img_part = source_para.part.related_parts[rId]
                                img_bytes = img_part.blob
                                emu_w, emu_h = self.get_image_extent(blip)
                                pic_run = new_para.add_run()
                                self.add_picture(pic_run, img_bytes, page_width_emu, available_width_emu, emu_w, emu_h)
                            except Exception:
                                pass
                return new_para
        
        # ★ 修复：如果传入了 resolved_numbering_text（虚拟样式映射到非列表样式时，
        # 需要保留原始编号文本），将其作为第一个 run 添加到段落最前面
        if resolved_numbering_text:
            new_para.add_run(resolved_numbering_text)
        
        # 获取段落完整文本（包括超链接内部的run文本）
        # python-docx 的 para.runs 不返回超链接(w:hyperlink)内部的run，需要用 para.text
        runs_text = ''.join(run.text for run in source_para.runs)
        full_para_text = source_para.text
        
        if full_para_text and len(full_para_text) > len(runs_text):
            # 段落中有超链接内部的隐藏文本，使用完整文本
            # 但仍需处理图片（图片在直接子级run中）
            for run_idx, run in enumerate(source_para.runs):
                blips = run._element.findall('.//' + qn('a:blip'))
                if blips:
                    for blip in blips:
                        rId = blip.get(qn('r:embed'))
                        if rId:
                            try:
                                img_part = source_para.part.related_parts[rId]
                                img_bytes = img_part.blob
                                emu_w, emu_h = self.get_image_extent(blip)
                                pic_run = new_para.add_run()
                                self.add_picture(pic_run, img_bytes, page_width_emu, available_width_emu, emu_w, emu_h)
                            except Exception:
                                pass
            # 使用完整段落文本
            cleaned_text = full_para_text.strip()
            if cleaned_text:
                new_para.add_run(cleaned_text)
        else:
            for run_idx, run in enumerate(source_para.runs):
                blips = run._element.findall('.//' + qn('a:blip'))
                if blips:
                    for blip in blips:
                        rId = blip.get(qn('r:embed'))
                        if rId:
                            try:
                                img_part = source_para.part.related_parts[rId]
                                img_bytes = img_part.blob
                                emu_w, emu_h = self.get_image_extent(blip)
                                pic_run = new_para.add_run()
                                self.add_picture(pic_run, img_bytes, page_width_emu, available_width_emu, emu_w, emu_h)
                            except Exception:
                                pass
                else:
                    if run.text:
                        new_para.add_run(run.text)
        
        return new_para
    
    def _copy_table_grid(self, new_table, source_table):
        """复制源表格的列宽网格（w:tblGrid，纯数值，不含任何样式定义）。"""
        src_grid = source_table._tbl.find(qn('w:tblGrid'))
        new_grid = new_table._tbl.find(qn('w:tblGrid'))
        if src_grid is None or new_grid is None:
            return
        # 清空目标表格默认的等宽 gridCol
        for gc in list(new_grid.findall(qn('w:gridCol'))):
            new_grid.remove(gc)
        # 复制源表格的列宽（deepcopy 仅复制列宽数值，安全）
        for gc in src_grid.findall(qn('w:gridCol')):
            new_grid.append(deepcopy(gc))
    
    def _collect_merge_regions(self, source_table):
        """收集表格中所有合并单元格区域。
        
        返回 [(top, left, bottom, right), ...]，坐标为网格坐标，bottom/right 为开区间（结束行/列索引）。
        横向合并通过 gridSpan 识别，纵向合并通过 vMerge=restart 及后续 continue 链识别。
        """
        tbl_el = source_table._tbl
        tr_lst = tbl_el.tr_lst
        regions = []
        for ri, tr in enumerate(tr_lst):
            for tc in tr.tc_lst:
                vm = tc.vMerge
                if vm == 'continue':
                    # continue 单元格不是合并区域起点，跳过
                    continue
                span = tc.grid_span
                top = tc.top
                left = tc.left
                right = left + span
                bottom = top + 1
                if vm == 'restart':
                    # 向下找纵向 continue 链，计算合并高度
                    r = ri + 1
                    while r < len(tr_lst):
                        found = False
                        for tcc in tr_lst[r].tc_lst:
                            if tcc.vMerge == 'continue' and tcc.left == left:
                                bottom = r + 1
                                found = True
                                break
                        if not found:
                            break
                        r += 1
                if span > 1 or vm == 'restart':
                    regions.append((top, left, bottom, right))
        return regions
    
    def _copy_cell_content(self, source_cell, new_cell, target_doc, table_idx, cell_pos,
                           available_width_emu, warning_callback=None,
                           table_style_override=None, enable_table_style=False):
        """复制单个单元格内容（段落文本、图片、OLE/VML 对象）。
        
        样式决策保持与原来一致：表格内段落只按表格样式定义处理，不参与正文样式映射。
        """
        def _get_table_para_style(src_style_name):
            if enable_table_style and table_style_override:
                try:
                    target_doc.styles[table_style_override]
                    return table_style_override
                except KeyError:
                    return DEFAULT_TARGET
            else:
                try:
                    target_doc.styles[src_style_name]
                    return src_style_name
                except KeyError:
                    return DEFAULT_TARGET
        
        # 清空单元格内容
        new_cell._element.clear_content()
        
        for para in source_cell.paragraphs:
            new_para = new_cell.add_paragraph()
            src_para_style = para.style.name
            new_para.style = _get_table_para_style(src_para_style)
            
            if self.has_numbering(para):
                new_para.add_run(self.list_bullet)
                self.remove_auto_numbering(new_para)
                full_text = ''.join(run.text for run in para.runs)
                cleaned_text = clean_list_numbering(full_text)
                if cleaned_text:
                    new_para.add_run(cleaned_text)
                for run in para.runs:
                    blips = run._element.findall('.//' + qn('a:blip'))
                    for blip in blips:
                        rId = blip.get(qn('r:embed'))
                        if rId:
                            try:
                                img_part = para.part.related_parts[rId]
                                img_bytes = img_part.blob
                                emu_w, emu_h = self.get_image_extent(blip)
                                pic_run = new_para.add_run()
                                self.add_picture(pic_run, img_bytes, available_width_emu, available_width_emu, emu_w, emu_h)
                            except Exception:
                                pass
                continue
            
            # 检查是否包含特殊对象（Visio图、OLE对象等）
            objects = para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object')
            shapes = para._element.findall('.//{urn:schemas-microsoft-com:vml}shape')
            
            if objects or shapes:
                # OLE 对象：提取预览图作为普通图片插入；独立 VML 形状跳过其 XML（不 deepcopy，避免损坏文档）
                inserted = False
                for obj in objects:
                    result = self._extract_ole_preview(para.part, obj)
                    if result is not None:
                        png_bytes, emu_w, emu_h = result
                        pic_run = new_para.add_run()
                        self.add_picture(pic_run, png_bytes, available_width_emu, available_width_emu, emu_w, emu_h)
                        inserted = True
                        break
                if not inserted:
                    if warning_callback:
                        try:
                            warning_msg = f"表格 {table_idx} 单元格 {cell_pos} 存在无法自动提取预览图的 OLE/VML 对象"
                            warning_callback(warning_msg)
                        except Exception:
                            pass
            else:
                for run in para.runs:
                    blips = run._element.findall('.//' + qn('a:blip'))
                    if blips:
                        for blip in blips:
                            rId = blip.get(qn('r:embed'))
                            if rId:
                                try:
                                    img_part = para.part.related_parts[rId]
                                    img_bytes = img_part.blob
                                    emu_w, emu_h = self.get_image_extent(blip)
                                    pic_run = new_para.add_run()
                                    self.add_picture(pic_run, img_bytes, available_width_emu, available_width_emu, emu_w, emu_h)
                                except Exception:
                                    pass
                    else:
                        if run.text:
                            new_para.add_run(run.text)
    
    def copy_table_with_images(self, source_table, target_doc, table_idx, available_width_emu, source_file="",
                               warning_callback=None, table_style_override=None, enable_table_style=False):
        """
        复制表格（包含图片、边框、合并单元格结构）
        合并单元格通过底层网格坐标 + 官方 merge() 重建，仅搬移合并结构（gridSpan/vMerge），
        不 deepcopy 源表格 XML，因此不会引入模板之外的样式。
        :param source_table: 源表格
        :param target_doc: 目标文档
        :param table_idx: 表格索引
        :param available_width_emu: 可用宽度
        :param source_file: 源文件名
        :param warning_callback: 警告回调函数
        :param table_style_override: 表格样式覆盖（当enable_table_style=True时使用）
        :param enable_table_style: 是否启用表格样式覆盖
        """
        # 获取源表格的网格维度（用底层网格，避免 row.cells 展开导致维度失真）
        tbl_el = source_table._tbl
        rows = len(tbl_el.tr_lst)
        cols = tbl_el.col_count
        
        # 创建规则网格表格
        new_table = target_doc.add_table(rows=rows, cols=cols)
        try:
            new_table.style = source_table.style
        except (KeyError, ValueError):
            # 模板中不存在该表格样式时保持默认样式，避免引入模板外样式
            pass
        
        # 复制列宽、设置宽度和边框
        self._copy_table_grid(new_table, source_table)
        self.set_table_width(new_table, available_width_emu)
        self.set_table_borders(new_table)
        
        # 先建立合并结构（在空表上 merge，仅重建 gridSpan/vMerge）
        for (top, left, bottom, right) in self._collect_merge_regions(source_table):
            try:
                new_table.cell(top, left).merge(new_table.cell(bottom - 1, right - 1))
            except Exception:
                # 个别非法/重叠区域跳过，不影响整体转换
                pass
        
        # 复制内容：遍历底层 tc，跳过 vMerge=continue，用网格坐标定位目标单元格。
        # 先一次性构建目标表格 grid 坐标 -> _Cell 映射，避免反复调用 table.cell()
        # （每次 cell() 都会重算整个 _cells 列表，导致 O(n²) 性能问题）。
        cell_map = {}
        for tr in new_table._tbl.tr_lst:
            for tc in tr.tc_lst:
                # ★ 修复：纵向合并的 continue cell 的 top 会继承 restart 的 top 值，
                # 若不加过滤，continue 会用相同 (top,left) 覆盖 restart 的映射，
                # 导致内容被复制到 continue 位置而 restart 位置为空。
                if tc.vMerge == 'continue':
                    continue
                cell_map[(tc.top, tc.left)] = _Cell(tc, new_table)
        
        for tr in tbl_el.tr_lst:
            for tc in tr.tc_lst:
                if tc.vMerge == 'continue':
                    continue
                top = tc.top
                left = tc.left
                new_cell = cell_map.get((top, left))
                if new_cell is None:
                    continue
                source_cell = _Cell(tc, source_table)
                cell_pos = f"[{top},{left}]"
                self._copy_cell_content(source_cell, new_cell, target_doc, table_idx, cell_pos,
                                        available_width_emu, warning_callback,
                                        table_style_override=table_style_override,
                                        enable_table_style=enable_table_style)
        
        return new_table
    
    def convert_styles(self, source_file, template_file, output_file, custom_style_map=None, list_bullet=None,
                       warning_callback=None, source_styles_cache=None,
                       table_style_override=None, enable_table_style=False,
                       image_style_override=None, enable_image_style=False,
                       remove_chapter_label=False,
                       list_method='bullet', list_style='Body Text',
                       enable_list_style=True):
        """
        样式转换主函数
        :param source_file: 源文件路径
        :param template_file: 模板文件路径
        :param output_file: 输出文件路径
        :param custom_style_map: 自定义样式映射表（可选）
        :param list_bullet: 列表段落符号（可选，默认为配置常量）
        :param warning_callback: 警告回调函数 callback(message)
        :param source_styles_cache: 缓存的源文件样式列表（可选，避免重复分析）
        :param table_style_override: 表格样式覆盖（当enable_table_style=True时使用）
        :param enable_table_style: 是否启用表格样式覆盖
        :param image_style_override: 图片样式覆盖（当enable_image_style=True时使用）
        :param enable_image_style: 是否启用图片样式覆盖
        :param remove_chapter_label: 是否清除"第X章/第X节"等章节标记
        :param list_method: 列表段落处理方式 'bullet'（符号）或 'style'（样式）
        :param list_style: 列表段落兜底样式名（当list_method='style'时使用）
        :param enable_list_style: 是否启用列表样式处理
        :return: (success, actual_file, message)
        """
        # 使用局部样式映射副本，避免修改全局变量
        style_map = STYLE_MAP.copy()
        if custom_style_map:
            style_map.update(custom_style_map)
        
        # 将样式映射存储为实例变量，供get_target_style使用
        self.current_style_map = style_map
        
        # 设置列表符号
        if list_bullet is not None:
            self.list_bullet = list_bullet
        
        logger = self.setup_logger(source_file)
        
        try:
            template_doc = Document(template_file)
        except Exception as e:
            return False, f"加载模板文档失败: {e}"
        
        # 获取模板和源文档的样式
        self.template_styles = self.get_template_styles(template_doc)
        
        try:
            source_doc = Document(source_file)
        except Exception as e:
            return False, f"加载源文档失败: {e}"
        
        # [HIGH_VOLTAGE] 性能优化：使用缓存的样式列表，避免重复分析
        if source_styles_cache:
            self.source_styles = source_styles_cache
        else:
            # 如果没有缓存，重新分析（兜底逻辑）
            self.source_styles = self.get_all_styles_from_doc(source_doc)
        
        new_doc = Document(template_file)
        self.clear_document_content(new_doc)
        
        section = new_doc.sections[0]
        page_width = section.page_width
        left_margin = section.left_margin
        right_margin = section.right_margin
        available_width = page_width - left_margin - right_margin
        
        body = source_doc.element.body
        para_idx = 0
        table_idx = 0
        self.stats = {"para": 0, "table": 0, "heading": 0}
        
        for child in body:
            if child.tag == qn('w:p'):
                if para_idx < len(source_doc.paragraphs):
                    para = source_doc.paragraphs[para_idx]
                    src_style = para.style.name
                    
                    # ★ 修复：对于有编号的列表段落，使用虚拟样式名（如 "1 列表段落"）进行样式映射，
                    # 而不是用真实的样式名（通常是 "Normal"）。这样 Step 3 中用户配置的
                    # "1 列表段落" → "BN_原文引用列表项目符号" 映射才能生效。
                    # ★ 修复：标题段落（有outlineLevel或样式为Heading）即使有编号也不视为列表段落，
                    # 应走正常的标题样式映射路径。
                    is_heading_by_outline = self.get_outline_level(para) > 0
                    is_heading_by_style = src_style in HEADING_STYLES
                    style_map = getattr(self, 'current_style_map', STYLE_MAP)
                    is_heading_by_mapped = style_map.get(src_style) in HEADING_STYLES
                    is_custom_heading = '标题' in src_style or src_style.startswith('Heading')
                    if is_heading_by_outline or is_heading_by_style or is_heading_by_mapped or is_custom_heading:
                        # 标题段落：走标题样式映射
                        target_style = self.get_target_style(src_style, new_doc, source_file)
                        _enable_list_fallback = enable_list_style
                        _list_style = list_style
                        _resolve_num_text = None
                    elif self.has_numbering(para):
                        # 非标题的列表段落：使用虚拟样式名进行映射
                        virtual_style = self._detect_numbering_format(para)
                        target_style = self.get_target_style(virtual_style, new_doc, source_file)
                        _resolve_num_text = None
                        if virtual_style in style_map:
                            # ★ 修复：虚拟样式已在 Step 3 映射
                            # 判断目标样式是否为列表段落样式（模板中自带 w:numPr 定义）
                            target_has_numPr = False
                            try:
                                tpl_style_xml = new_doc.styles[target_style]._element.xml
                                if '<w:numPr>' in tpl_style_xml or '<w:numPr ' in tpl_style_xml:
                                    target_has_numPr = True
                            except KeyError:
                                pass
                            
                            if target_has_numPr:
                                # 目标样式是列表段落样式 → 走样式模式兜底，用映射目标样式替代 list_style
                                # 这样 copy_paragraph_with_images 会应用目标样式的 numPr 编号
                                _enable_list_fallback = True
                                _list_style = target_style
                            else:
                                # 目标样式是正文/普通段落 → 保留原编号文本，应用目标样式
                                _enable_list_fallback = False
                                _resolve_num_text = self._resolve_auto_numbering_text(para)
                        else:
                            _enable_list_fallback = enable_list_style
                            _list_style = list_style
                    else:
                        target_style = self.get_target_style(src_style, new_doc, source_file)
                        _enable_list_fallback = enable_list_style
                        _list_style = list_style
                        _resolve_num_text = None
                    
                    new_para = self.copy_paragraph_with_images(
                        para, new_doc, target_style,
                        page_width, available_width,
                        para_idx, source_file,
                        warning_callback,
                        image_style_override=image_style_override,
                        enable_image_style=enable_image_style,
                        remove_chapter_label=remove_chapter_label,
                        list_method=list_method,
                        list_style=_list_style,
                        enable_list_style=_enable_list_fallback,
                        resolved_numbering_text=_resolve_num_text
                    )
                    
                    if self.get_outline_level(para) > 0 or src_style in HEADING_STYLES or style_map.get(src_style) in HEADING_STYLES or '标题' in src_style or src_style.startswith('Heading'):
                        self.stats["heading"] += 1
                    self.stats["para"] += 1
                    para_idx += 1
            
            elif child.tag == qn('w:tbl'):
                if table_idx < len(source_doc.tables):
                    table = source_doc.tables[table_idx]
                    self.copy_table_with_images(table, new_doc, table_idx, available_width, source_file,
                                               warning_callback,
                                               table_style_override=table_style_override,
                                               enable_table_style=enable_table_style)
                    self.stats["table"] += 1
                    table_idx += 1
            
            else:
                # 处理其他类型的元素（如OLE对象、Visio图等）
                # 尝试获取样式名称，如果无法获取则使用默认样式
                try:
                    # 对于非段落/表格元素，尝试查找其所属段落的样式
                    parent_para = child.getparent()
                    while parent_para is not None and parent_para.tag != qn('w:p'):
                        parent_para = parent_para.getparent()
                    
                    if parent_para is not None:
                        # 找到父段落，尝试获取其样式
                        pPr = parent_para.find(qn('w:pPr'))
                        if pPr is not None:
                            pStyle = pPr.find(qn('w:pStyle'))
                            if pStyle is not None:
                                style_id = pStyle.get(qn('w:val'))
                                if style_id:
                                    target_style = self.get_target_style(style_id, new_doc, source_file)
                                else:
                                    target_style = DEFAULT_TARGET
                            else:
                                target_style = DEFAULT_TARGET
                        else:
                            target_style = DEFAULT_TARGET
                    else:
                        target_style = DEFAULT_TARGET
                except:
                    target_style = DEFAULT_TARGET
                
                # 复制特殊元素
                special_para = self.copy_special_element(child, new_doc, target_style, warning_callback)  # [FIX] 传递warning_callback参数
                if special_para is not None:
                    self.stats["para"] += 1  # 计入统计
        
        # 使用重试机制保存文档
        success, actual_file, msg = self.save_with_retry(new_doc, output_file)
        if success:
            return True, f"转换完成！段落: {self.stats['para']}, 表格: {self.stats['table']}, 标题: {self.stats['heading']}。{msg}"
        else:
            return False, msg
    
    def is_part_of_exception(self, full_text, match_start, match_end, word):
        """判断单字词是否属于例外词
        
        特殊处理"应对"：复合词"应对"(yìng duì = 处理/对付) 作为例外保留"应"，
        但情态动词"应"+介词"对"(yīng + duì = 应该 + 对于) 不是例外，需去掉"应"。
        区分方法：检查"对"之后 ~25 字范围是否存在 "负责/进行/予以..." 等标志动词，
        且动词距"对" > 2 字 → 分离结构，不视为例外。
        """
        if word == "应":
            exceptions = EXCEPTION_WORDS_YING
        elif word == "须":
            exceptions = EXCEPTION_WORDS_XU
        else:
            return False
        
        start = max(0, match_start - 20)
        end = min(len(full_text), match_end + 20)
        substr = full_text[start:end]
        for exc in exceptions:
            if exc in substr:
                pos = substr.find(exc)
                while pos != -1:
                    exc_start = start + pos
                    exc_end = exc_start + len(exc)
                    if exc_start <= match_start < exc_end:
                        # 特殊处理"应对"：区分复合词 vs "应+对"分离结构
                        if exc == "应对":
                            if self._is_ying_dui_separable(full_text, match_end):
                                return False  # 类型B：应+对分离，不视为例外，正常去掉"应"
                        return True
                    pos = substr.find(exc, pos+1)
        return False
    
    def _is_ying_dui_separable(self, full_text, dui_pos):
        """判断"应+对"是否属于分离结构（应 yīng + 对 duì = 应当 + 对于）
        
        参数 dui_pos 是"对"字在 full_text 中的位置（即 match_end）。
        检查"对"之后 ~25 字范围内是否存在标志动词（负责/进行/予以...），
        且动词距"对" > 2 字，防止误杀"应对措施/应对方案"等复合词。
        
        返回 True 表示是分离结构（类型B），"应"不应受"应对"例外保护。
        """
        after_dui = full_text[dui_pos:dui_pos + 25]
        for m in _YING_DUI_VERBS.finditer(after_dui):
            verb_dist = m.start()  # 动词距"对"的字数
            if verb_dist > 2:
                return True
        return False
    
    def is_multi_exception(self, full_text, match_start, match_end, word):
        """判断多字祈使词是否属于例外词"""
        start = max(0, match_start - 20)
        end = min(len(full_text), match_end + 20)
        substr = full_text[start:end]
        for exc in MULTI_EXCEPTIONS:
            if exc in substr:
                pos = substr.find(exc)
                while pos != -1:
                    exc_start = start + pos
                    exc_end = exc_start + len(exc)
                    if exc_start <= match_start < exc_end:
                        return True
                    pos = substr.find(exc, pos+1)
        return False
    
    def replace_multiple_imperative(self, run_text, full_text, run_start_offset):
        """替换多字祈使词"""
        if not run_text:
            return run_text
        def repl(match):
            word = match.group(0)
            abs_start = run_start_offset + match.start()
            abs_end = run_start_offset + match.end()
            if self.is_multi_exception(full_text, abs_start, abs_end, word):
                return word
            return MULTI_IMPERATIVE_TO_STATEMENT.get(word, word)
        return MULTI_IMPERATIVE_REGEX.sub(repl, run_text)
    
    def replace_single_imperative(self, run_text, full_text, run_start_offset):
        """替换单字祈使词"""
        if not run_text:
            return run_text
        def repl(match):
            word = match.group(0)
            abs_start = run_start_offset + match.start()
            abs_end = run_start_offset + match.end()
            if self.is_part_of_exception(full_text, abs_start, abs_end, word):
                return word
            return SINGLE_REPLACE.get(word, word)
        return SINGLE_IMPERATIVE_REGEX.sub(repl, run_text)
    
    _ZW_CHARS_RE = re.compile(r'[\u200b\u200c\u200d\ufeff]')

    def process_paragraph_mood(self, para):
        """处理段落语气转换 - 基于段落全文匹配，支持跨run和零宽字符
        
        修复两个遗留问题：
        A. Word 拼写检查/语言检测/Track Changes 导致"投标人"被拆分到多个 run
        B. 零宽空格(U+200B)等不可见字符夹在"投标人"中间导致正则断裂
        
        策略：在段落全文(para.text)上匹配，清理零宽字符后再做正则搜索，
        匹配结果映射回原始 run 边界执行替换。
        """
        runs = para.runs
        if not runs:
            return False

        # 收集所有 run 的原始文本
        run_texts_orig = [run.text for run in runs]
        full_text = ''.join(run_texts_orig)
        if not full_text.strip():
            return False

        # 清理零宽字符，用于正则匹配
        clean_text = self._ZW_CHARS_RE.sub('', full_text)
        if not clean_text.strip():
            return False

        # 构建 clean_text -> full_text 索引映射
        clean_to_full = []
        for fi, ch in enumerate(full_text):
            if ch not in '\u200b\u200c\u200d\ufeff':
                clean_to_full.append(fi)

        if not clean_to_full:
            return False

        # 收集所有替换（在 full_text 坐标中）
        replacements = []

        def _add_match(clean_start, clean_end, repl_text):
            """将 clean_text 中的匹配转换为 full_text 坐标并加入替换列表"""
            if clean_start >= len(clean_to_full):
                return
            fs = clean_to_full[clean_start]
            if clean_end - 1 < len(clean_to_full):
                fe = clean_to_full[clean_end - 1] + 1
            else:
                fe = len(full_text)
            orig = full_text[fs:fe]
            if repl_text and repl_text != orig:
                replacements.append((fs, fe, repl_text))

        # 1. REPLACE_REGEX: "投标人" -> "本投标人" 等固定替换
        if REPLACE_REGEX:
            for m in REPLACE_REGEX.finditer(clean_text):
                _add_match(m.start(), m.end(), REPLACE_MAP.get(m.group(0), ''))

        # 2. MULTI_IMPERATIVE_REGEX: 多字祈使词
        for m in MULTI_IMPERATIVE_REGEX.finditer(clean_text):
            word = m.group(0)
            cs, ce = m.start(), m.end()
            if cs >= len(clean_to_full) or ce - 1 >= len(clean_to_full):
                continue
            fs = clean_to_full[cs]
            fe = clean_to_full[ce - 1] + 1
            if self.is_multi_exception(full_text, fs, fe, word):
                continue
            repl = MULTI_IMPERATIVE_TO_STATEMENT.get(word)
            if repl:
                replacements.append((fs, fe, repl))

        # 3. SINGLE_IMPERATIVE_REGEX: 单字祈使词
        for m in SINGLE_IMPERATIVE_REGEX.finditer(clean_text):
            word = m.group(0)
            cs, ce = m.start(), m.end()
            if cs >= len(clean_to_full) or ce - 1 >= len(clean_to_full):
                continue
            fs = clean_to_full[cs]
            fe = clean_to_full[ce - 1] + 1
            if self.is_part_of_exception(full_text, fs, fe, word):
                continue
            repl = SINGLE_REPLACE.get(word)
            if repl:
                replacements.append((fs, fe, repl))

        if not replacements:
            return False

        # 按起始位置降序排列，从后往前执行替换避免偏移问题
        replacements.sort(key=lambda x: x[0], reverse=True)

        # 在 full_text 上执行所有替换
        result = full_text
        for fs, fe, repl in replacements:
            result = result[:fs] + repl + result[fe:]

        # "将将" -> "将把"（后处理，全局替换）
        result = result.replace('将将', '将把')

        # 按原始字符数比例将 result 分配回各个 run
        total_orig = len(full_text)
        total_result = len(result)

        if total_orig == 0:
            return False

        new_run_texts = []
        result_pos = 0
        cum_orig = 0

        if len(run_texts_orig) == 1:
            new_run_texts.append(result)
        else:
            for i, orig_text in enumerate(run_texts_orig[:-1]):
                cum_orig += len(orig_text)
                target_end = int(total_result * cum_orig / total_orig)
                target_end = max(result_pos, min(target_end, total_result))
                new_text = result[result_pos:target_end]
                new_run_texts.append(new_text)
                result_pos = target_end
            new_run_texts.append(result[result_pos:])

        # 写回 run 文本
        modified = False
        for i, run in enumerate(runs):
            if run.text != new_run_texts[i]:
                run.text = new_run_texts[i]
                modified = True

        return modified
    
    def convert_mood(self, input_file, output_file=None):
        """
        祈使语气转换
        :param input_file: 输入文件
        :param output_file: 输出文件（如果为None则覆盖原文件）
        :return: (success, actual_output_file, message)
        """
        if output_file is None:
            output_file = input_file
        
        try:
            doc = Document(input_file)
        except Exception as e:
            return False, f"加载文档失败: {e}"
        
        modified_count = 0
        para_count = 0
        
        for para in doc.paragraphs:
            para_count += 1
            # 跳过标记为 keepOriginal 的段落（copy_chapter 模式的第一份副本）
            if self._is_keep_original_paragraph(para._element):
                continue
            # 跳过标题段落（标题中的"投标人"不应转换）
            if self.is_heading_paragraph(para._element, doc):
                continue
            if self.process_paragraph_mood(para):
                modified_count += 1
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para_count += 1
                        if self._is_keep_original_paragraph(para._element):
                            continue
                        # 跳过标题段落（表格内通常无标题，但保持一致）
                        if self.is_heading_paragraph(para._element, doc):
                            continue
                        if self.process_paragraph_mood(para):
                            modified_count += 1
        
        # 使用重试机制保存文档
        success, actual_file, msg = self.save_with_retry(doc, output_file)
        if success:
            return True, f"语气转换完成！处理段落: {para_count}, 修改: {modified_count}。{msg}"
        else:
            return False, msg
    
    def ensure_style_exists(self, doc, style_name):
        """确保文档中存在指定样式"""
        try:
            doc.styles[style_name]
        except KeyError:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = doc.styles['Normal']
    
    def get_style_id(self, elem):
        """获取段落元素的样式ID"""
        pPr = elem.find(qn('w:pPr'))
        if pPr is None:
            return None
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is None:
            return None
        return pStyle.get(qn('w:val'))
    
    def is_plain_paragraph(self, elem):
        """判断是否为无样式的普通正文段落"""
        if not hasattr(elem, 'tag'):
            return False
        if elem.tag != qn('w:p'):
            return False
        pPr = elem.find(qn('w:pPr'))
        if pPr is None:
            return True
        pStyle = pPr.find(qn('w:pStyle'))
        return pStyle is None
    
    def contains_image(self, elem):
        """检查段落是否包含图片"""
        if not hasattr(elem, 'tag'):
            return False
        if elem.tag != qn('w:p'):
            return False
        blips = elem.findall('.//' + qn('a:blip'))
        return len(blips) > 0
    
    def is_table_elem(self, elem):
        """判断是否为表格"""
        if not hasattr(elem, 'tag'):
            return False
        return elem.tag == qn('w:tbl')
    
    def create_answer_paragraph_element(self, doc, answer_text, answer_style):
        """创建应答句段落XML元素"""
        temp_para = doc.add_paragraph(answer_text)
        temp_para.style = answer_style
        para_elem = deepcopy(temp_para._element)
        temp_para._element.getparent().remove(temp_para._element)
        return para_elem
    
    def is_heading_paragraph(self, elem, doc=None):
        """判断段落是否为标题（通过大纲级别 + 样式名/ID判断）
        
        支持标准 Heading 1-9 样式、自定义"标题"样式（如 BN_标题0），
        以及具有大纲级别的段落。
        
        兼容两种调用场景：
        - python-docx Paragraph 对象：.style 返回 Style 对象，有 .name 属性
        - lxml/CT_P XML 元素：.style 返回样式 ID 字符串（如 '1', 'a3', 'BN0'）
        """
        if self.get_outline_level(elem, doc) > 0:
            return True
        
        if hasattr(elem, 'style') and elem.style:
            s = elem.style
            if isinstance(s, str):
                # XML元素 (CT_P等)：.style 返回样式ID字符串
                style_id = s
                # 检查是否为内置标题样式ID (1-9)
                if style_id in HEADING_STYLE_IDS:
                    return True
                # 尝试通过doc解析样式名称，识别自定义标题样式
                if doc is not None:
                    try:
                        style_obj = doc.styles[style_id]
                        style_name = style_obj.name
                        if style_name in HEADING_STYLES:
                            return True
                        if '标题' in style_name or style_name.startswith('Heading'):
                            return True
                    except (KeyError, AttributeError):
                        pass
            elif hasattr(s, 'name'):
                # Paragraph对象：.style 返回Style对象
                style_name = s.name
                if style_name in HEADING_STYLES:
                    return True
                if '标题' in style_name or style_name.startswith('Heading'):
                    return True
        
        return False
    
    def insert_response_after_headings(self, input_file, output_file=None, 
                                       answer_text=None, answer_style=None,
                                       mode='before_heading',
                                       answer_source_style=None,
                                       answer_copy_style=None,
                                       table_answer_style=None,
                                       list_method='bullet',
                                       list_style='Body Text',
                                       list_answer_method='bullet',
                                       list_answer_style='Body Text',
                                       list_answer_bullet='● ',
                                       enable_image_style=False,
                                       image_style_override=None):
        """
        在章节或段落前后插入应答句（改进版：使用大纲级别识别标题）
        :param input_file: 输入文件
        :param output_file: 输出文件（如果为None则覆盖原文件）
        :param answer_text: 应答文本
        :param answer_style: 应答样式
        :param mode: 插入模式
            - 'before_heading': 章节前插入（默认，原有逻辑）
            - 'after_heading': 章节末插入（在章节正文最后一段后插入）
            - 'copy_chapter': 复制章节内容（未来扩展）
            - 'before_paragraph': 逐段前应答（未来扩展）
            - 'after_paragraph': 逐段后应答（未来扩展）
        :return: 是否成功，消息
        """
        if output_file is None:
            output_file = input_file
        if answer_text is None:
            answer_text = ANSWER_TEXT
        if answer_style is None:
            answer_style = ANSWER_STYLE
        
        try:
            doc = Document(input_file)
        except Exception as e:
            return False, f"加载文档失败: {e}"
        
        self.ensure_style_exists(doc, answer_style)
        
        # 预创建应答段落模板
        temp_para = doc.add_paragraph(answer_text)
        temp_para.style = answer_style
        answer_template = deepcopy(temp_para._element)
        temp_para._element.getparent().remove(temp_para._element)
        
        body = doc.element.body
        children = list(body)
        new_children = []
        insert_count = 0
        total_heading_count = 0
        para_index = 0
        i = 0
        
        # 根据模式选择不同的处理逻辑
        if mode == 'before_heading':
            # 原有逻辑：标题前插入（下一个不是标题时插入）
            insert_count, total_heading_count = self._insert_before_headings(
                children, new_children, answer_template, doc
            )
        elif mode == 'after_heading':
            # 新需求：章节后插入（在章节正文最后一段后插入）
            insert_count, total_heading_count = self._insert_after_headings(
                children, new_children, answer_template, doc
            )
        elif mode == 'copy_chapter':
            # 复制章节插入（在应答句后复制章节内容）
            insert_count, total_heading_count = self._insert_with_copy_chapter(
                children, new_children, answer_template, doc
            )
        elif mode == 'before_paragraph':
            # 逐段前应答（在非标题段落前插入应答句）
            insert_count, total_heading_count = self._insert_before_paragraphs(
                children, new_children, answer_template, doc
            )
        elif mode == 'after_paragraph':
            # 逐段后应答（在非标题段落后插入应答句）
            insert_count, total_heading_count = self._insert_after_paragraphs(
                children, new_children, answer_template, doc
            )
        else:
            return False, output_file, f"不支持的插入模式: {mode}"
        
        # 清空并重组body
        for child in list(body):
            body.remove(child)
        for elem in new_children:
            body.append(elem)
        
        # 使用重试机制保存文档
        success, actual_file, msg = self.save_with_retry(doc, output_file)
        if success:
            return True, actual_file, f"已插入 {insert_count} 个应答句，共发现标题 {total_heading_count} 个。{msg}"
        else:
            return False, output_file, msg
    
    def _insert_before_headings(self, children, new_children, answer_template, doc):
        """
        在章节前插入应答句（原有逻辑）
        判断条件：如果标题后下一个元素不是标题，则在该标题前插入
        :return: (insert_count, total_heading_count)
        """
        insert_count = 0
        total_heading_count = 0
        i = 0
        
        while i < len(children):
            child = children[i]
            
            # 安全检查
            if not hasattr(child, 'tag'):
                i += 1
                continue
            
            new_children.append(child)
            
            # 检查是否为标题
            if child.tag == qn('w:p') and self.is_heading_paragraph(child, doc):
                total_heading_count += 1
                
                # 检查下一个元素
                if i + 1 < len(children):
                    next_elem = children[i + 1]
                    
                    if hasattr(next_elem, 'tag'):
                        # 如果下一个不是标题，则插入应答句
                        if not self.is_heading_paragraph(next_elem, doc):
                            answer_elem = deepcopy(answer_template)
                            new_children.append(answer_elem)
                            insert_count += 1
            
            i += 1
        
        return insert_count, total_heading_count
    
    def _insert_after_headings(self, children, new_children, answer_template, doc):
        """
        在章节末插入应答句（需求2：章节后插入）
        判断条件：如果标题前是正文段落（不是标题），就在这个正文后插入应答句
        特殊情况：
        - 全文档第一个标题前不插入（因为前面没有章节）
        - 两个标题之间不插入
        - 文章最后一段如果是正文，在其后插入
        :return: (insert_count, total_heading_count)
        """
        insert_count = 0
        total_heading_count = 0
        
        # 第一步：遍历所有元素，添加元素并统计标题
        for i, child in enumerate(children):
            # 安全检查
            if not hasattr(child, 'tag'):
                new_children.append(child)
                continue
            
            # 统计标题数量
            if child.tag == qn('w:p') and self.is_heading_paragraph(child, doc):
                total_heading_count += 1
            
            # 添加当前元素
            new_children.append(child)
            
            # 第二步：判断是否需要在当前元素后插入应答句
            # 条件：当前不是标题 + 下一个元素是标题
            if i + 1 < len(children):
                next_elem = children[i + 1]
                
                # 检查下一个元素是否为标题
                if hasattr(next_elem, 'tag') and next_elem.tag == qn('w:p'):
                    if self.is_heading_paragraph(next_elem, doc):
                        # 下一个是标题，检查当前元素是否不是标题
                        is_not_heading = True
                        
                        # 检查当前是否为标题
                        if hasattr(child, 'tag') and child.tag == qn('w:p'):
                            if self.is_heading_paragraph(child, doc):
                                is_not_heading = False
                        
                        # 如果当前不是标题（可以是正文、表格、图片等），则在其后插入应答句
                        if is_not_heading:
                            answer_elem = deepcopy(answer_template)
                            new_children.append(answer_elem)
                            insert_count += 1
            else:
                # 当前是最后一个元素
                # 如果当前不是标题，在其后插入应答句
                is_not_heading = True
                if hasattr(child, 'tag') and child.tag == qn('w:p'):
                    if self.is_heading_paragraph(child, doc):
                        is_not_heading = False
                
                if is_not_heading:
                    answer_elem = deepcopy(answer_template)
                    new_children.append(answer_elem)
                    insert_count += 1
        
        return insert_count, total_heading_count
    
    def _insert_with_copy_chapter(self, children, new_children, answer_template, doc,
                                   answer_source_style=None,
                                   answer_copy_style=None,
                                   table_answer_style=None,
                                   list_method='bullet',
                                   list_style='Body Text',
                                   list_answer_method='bullet',
                                   list_answer_style='Body Text',
                                   list_answer_bullet='● ',
                                   enable_image_style=False,
                                   image_style_override=None,
                                   enable_list_style=True):
        """
        原文+应答句+应答原文（需求3：copy_chapter）
        最终效果：标题 → 提示语 → 原文（未转换，标记为 keepOriginal）→ 应答句 → 原文（语气转换后）
        注意：此模式在 full_convert 中会调换流水线顺序（先插入应答句，后语气转换），
              因此原始正文在插入应答句时仍是未转换状态，加上 keepOriginal 标记后会被跳过。
        
        增强功能：
        - answer_source_style: 原文（第一份副本）的样式
        - answer_copy_style: 应答原文（第二份副本）的样式
        - table_answer_style: 表格应答原文样式
        - list_method/list_style: 原文列表段落处理
        - list_answer_method/list_answer_style/list_answer_bullet: 应答原文列表段落处理
        - enable_image_style/image_style_override: 图片兜底样式
        :return: (insert_count, total_heading_count)
        """
        from copy import deepcopy as deep_copy
        
        insert_count = 0
        total_heading_count = 0
        bookmark_id = 0  # 书签 ID 计数器
        
        # 创建 source_template（原文样式模板）
        source_template = deep_copy(answer_template)
        if answer_source_style:
            try:
                src_style = doc.styles[answer_source_style]
                source_pPr = source_template.find(qn('w:pPr'))
                if source_pPr is None:
                    source_pPr = OxmlElement('w:pPr')
                    source_template.insert(0, source_pPr)
                source_pStyle = source_pPr.find(qn('w:pStyle'))
                if source_pStyle is None:
                    source_pStyle = OxmlElement('w:pStyle')
                    source_pPr.append(source_pStyle)
                source_pStyle.set(qn('w:val'), src_style.style_id)
            except Exception:
                pass
        
        # 创建 copy_template（应答原文副本样式模板）
        copy_template = deep_copy(answer_template)
        if answer_copy_style:
            try:
                copy_style = doc.styles[answer_copy_style]
                copy_pPr = copy_template.find(qn('w:pPr'))
                if copy_pPr is None:
                    copy_pPr = OxmlElement('w:pPr')
                    copy_template.insert(0, copy_pPr)
                copy_pStyle = copy_pPr.find(qn('w:pStyle'))
                if copy_pStyle is None:
                    copy_pStyle = OxmlElement('w:pStyle')
                    copy_pPr.append(copy_pStyle)
                copy_pStyle.set(qn('w:val'), copy_style.style_id)
            except Exception:
                pass
        
        def is_heading(elem):
            """判断元素是否为标题段落"""
            if not hasattr(elem, 'tag'):
                return False
            if elem.tag != qn('w:p'):
                return False
            return self.is_heading_paragraph(elem, doc)
        
        def remove_keep_original_from_element(elem):
            """移除元素中的 keepOriginal 书签标记"""
            if not hasattr(elem, 'tag') or elem.tag != qn('w:p'):
                return
            bookmark_ids = set()
            starts_to_remove = []
            ends_to_remove = []
            for child in elem:
                if child.tag == qn('w:bookmarkStart'):
                    if child.get(qn('w:name')) == '_keepOriginal_':
                        bookmark_ids.add(child.get(qn('w:id')))
                        starts_to_remove.append(child)
                elif child.tag == qn('w:bookmarkEnd'):
                    if child.get(qn('w:id')) in bookmark_ids:
                        ends_to_remove.append(child)
            for start in starts_to_remove:
                elem.remove(start)
            for end in ends_to_remove:
                elem.remove(end)
        
        def _apply_paragraph_style(elem, style_id):
            """将段落样式ID写入段落的 pStyle，保留已有 numPr/numId 等编号信息。"""
            if not hasattr(elem, 'tag') or elem.tag != qn('w:p') or not style_id:
                return
            pPr = elem.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                elem.insert(0, pPr)
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is None:
                pStyle = OxmlElement('w:pStyle')
                pPr.append(pStyle)
            pStyle.set(qn('w:val'), style_id)
        
        # 当前章节标题索引，None 表示尚未进入任何章节
        current_chapter_heading = None
        # 章节内容缓冲区：保存当前章节内的非提示语元素，用于生成第二份副本
        chapter_buffer = []
        
        i = 0
        while i < len(children):
            child = children[i]
            
            # 安全检查：非 XML 元素直接输出
            if not hasattr(child, 'tag'):
                new_children.append(child)
                i += 1
                continue
            
            # 遇到标题：先刷新前一个章节，再输出当前标题
            if is_heading(child):
                # 刷新前一个章节的内容
                if current_chapter_heading is not None and chapter_buffer:
                    # 检查是否有非提示语内容
                    has_real_content = any(not self._is_hint_paragraph(elem) for elem in chapter_buffer)
                    if has_real_content:
                        # 插入应答句
                        answer_elem = deep_copy(answer_template)
                        new_children.append(answer_elem)
                        insert_count += 1
                        
                        # 复制应答原文副本（用 copy_template 样式，对应应答原文格式），跳过提示语
                        for elem in chapter_buffer:
                            if self._is_hint_paragraph(elem):
                                continue
                            # 表格元素：深拷贝原始表格，并将所有段落的样式改为 table_answer_style（优先）或 answer_copy_style
                            if elem.tag == qn('w:tbl'):
                                source_elem = deep_copy(elem)
                                tbl_style_id = None
                                if table_answer_style:
                                    tbl_style_id = self.get_style_id_by_name(doc, table_answer_style)
                                if not tbl_style_id:
                                    _copy_sid = self.get_style_id_by_name(doc, answer_copy_style) if answer_copy_style else None
                                    if not _copy_sid:
                                        _copy_sid = self.get_style_id(copy_template) if hasattr(self, 'get_style_id') else None
                                    tbl_style_id = _copy_sid
                                if tbl_style_id:
                                    for p_elem in source_elem.iter(qn('w:p')):
                                        pPr = p_elem.find(qn('w:pPr'))
                                        if pPr is None:
                                            pPr = OxmlElement('w:pPr')
                                            p_elem.insert(0, pPr)
                                        pStyle = pPr.find(qn('w:pStyle'))
                                        if pStyle is None:
                                            pStyle = OxmlElement('w:pStyle')
                                            pPr.append(pStyle)
                                        pStyle.set(qn('w:val'), tbl_style_id)
                                new_children.append(source_elem)
                                continue
                            # 列表段落用 deepcopy(elem) 保留原始列表结构
                            if DocumentConverter._elem_has_numbering(elem, doc):
                                source_elem = deep_copy(elem)
                                remove_keep_original_from_element(source_elem)
                                if list_answer_method == 'style' and list_answer_style:
                                    list_sid = self.get_style_id_by_name(doc, list_answer_style)
                                    if list_sid:
                                        _apply_paragraph_style(source_elem, list_sid)
                                new_children.append(source_elem)
                            else:
                                # ★ 修复：OLE占位段落应用图片兜底样式
                                is_ole_placeholder = self._is_ole_placeholder_paragraph(elem)
                                if is_ole_placeholder and enable_image_style and image_style_override:
                                    img_sid = self.get_style_id_by_name(doc, image_style_override)
                                    if img_sid:
                                        source_elem = deep_copy(elem)
                                        _apply_paragraph_style(source_elem, img_sid)
                                        remove_keep_original_from_element(source_elem)
                                    else:
                                        source_elem = deep_copy(copy_template)
                                        source_runs = source_elem.findall('.//' + qn('w:r'))
                                        orig_runs = elem.findall('.//' + qn('w:r'))
                                        for r in source_runs:
                                            source_elem.remove(r)
                                        for r in orig_runs:
                                            source_elem.append(deep_copy(r))
                                        remove_keep_original_from_element(source_elem)
                                else:
                                    # 用 copy_template 替换内容，保持应答原文副本样式（answer_copy_style）
                                    source_elem = deep_copy(copy_template)
                                    source_runs = source_elem.findall('.//' + qn('w:r'))
                                    orig_runs = elem.findall('.//' + qn('w:r'))
                                    for r in source_runs:
                                        source_elem.remove(r)
                                    for r in orig_runs:
                                        source_elem.append(deep_copy(r))
                                    remove_keep_original_from_element(source_elem)
                                    # 应答原文副本图片段落应用图片兜底样式
                                    if enable_image_style and image_style_override:
                                        has_img = (elem.find('.//' + qn('w:drawing')) is not None or
                                                   elem.find('.//' + qn('w:pict')) is not None or
                                                   elem.find('.//' + qn('pic:pic')) is not None or
                                                   elem.find('.//' + qn('wp:inline')) is not None or
                                                   elem.find('.//' + qn('wp:anchor')) is not None)
                                        if has_img:
                                            img_sid = self.get_style_id_by_name(doc, image_style_override)
                                            if img_sid:
                                                _apply_paragraph_style(source_elem, img_sid)
                                new_children.append(source_elem)
                
                chapter_buffer.clear()
                new_children.append(child)
                current_chapter_heading = i
                total_heading_count += 1
                i += 1
                continue
            
            # 非标题元素
            if current_chapter_heading is not None:
                # 在章节内：提示语直接输出，其他内容标记为 keepOriginal 后输出，并加入缓冲区
                if self._is_hint_paragraph(child):
                    new_children.append(child)
                else:
                    # 用 source_template 替换内容，应用原文格式（answer_source_style）
                    if child.tag == qn('w:p'):
                        child_pPr = child.find(qn('w:pPr'))
                        child_numPr = child_pPr.find(qn('w:numPr')) if child_pPr is not None else None
                        
                        _pStyle_val = None
                        if child_pPr is not None:
                            _pStyle_elem = child_pPr.find(qn('w:pStyle'))
                            if _pStyle_elem is not None:
                                _pStyle_val = _pStyle_elem.get(qn('w:val'))
                        
                        # 标题段落（包括带编号的标题）不视为列表段落
                        if _pStyle_val and self.is_heading_paragraph(child, doc):
                            _is_list_para = False
                        else:
                            _is_list_para = DocumentConverter._elem_has_numbering(child, doc)
                            if not _is_list_para:
                                if _pStyle_val and ('列表段落' in _pStyle_val):
                                    _is_list_para = True
                                if not _is_list_para:
                                    if _pStyle_val and list_method == 'style' and list_style:
                                        _list_sid = self.get_style_id_by_name(doc, list_style)
                                        if _list_sid and _pStyle_val == _list_sid:
                                            _is_list_para = True
                        
                        if _is_list_para and enable_list_style:
                            source_elem = deep_copy(child)
                            bookmark_start = OxmlElement('w:bookmarkStart')
                            bookmark_start.set(qn('w:id'), str(bookmark_id))
                            bookmark_start.set(qn('w:name'), '_keepOriginal_')
                            source_elem.insert(0, bookmark_start)
                            bookmark_end = OxmlElement('w:bookmarkEnd')
                            bookmark_end.set(qn('w:id'), str(bookmark_id))
                            source_elem.append(bookmark_end)
                            bookmark_id += 1
                            if list_method == 'style' and list_style:
                                list_sid = self.get_style_id_by_name(doc, list_style)
                                if list_sid:
                                    _apply_paragraph_style(source_elem, list_sid)
                                try:
                                    _style_xml2 = doc.styles[list_style]._element.xml
                                    if not ('<w:numPr>' in _style_xml2 or '<w:numPr ' in _style_xml2):
                                        self.remove_auto_numbering(source_elem)
                                except Exception:
                                    pass
                            if enable_image_style and image_style_override:
                                has_img = (child.find('.//' + qn('w:drawing')) is not None or
                                           child.find('.//' + qn('w:pict')) is not None or
                                           child.find('.//' + qn('pic:pic')) is not None or
                                           child.find('.//' + qn('wp:inline')) is not None or
                                           child.find('.//' + qn('wp:anchor')) is not None)
                                if has_img:
                                    img_sid = self.get_style_id_by_name(doc, image_style_override)
                                    if img_sid:
                                        _apply_paragraph_style(source_elem, img_sid)
                            new_children.append(source_elem)
                            chapter_buffer.append(source_elem)
                            i += 1
                            continue
                        
                        # ★ 修复：OLE占位段落使用图片兜底样式，而非source_template样式
                        is_ole_ph = self._is_ole_placeholder_paragraph(child)
                        if is_ole_ph and enable_image_style and image_style_override:
                            img_sid = self.get_style_id_by_name(doc, image_style_override)
                            if img_sid:
                                source_elem = deep_copy(child)
                                _apply_paragraph_style(source_elem, img_sid)
                            else:
                                source_elem = deep_copy(source_template)
                                source_runs = source_elem.findall('.//' + qn('w:r'))
                                orig_runs = child.findall('.//' + qn('w:r'))
                                for r in source_runs:
                                    source_elem.remove(r)
                                for r in orig_runs:
                                    source_elem.append(deep_copy(r))
                        else:
                            source_elem = deep_copy(source_template)
                            source_runs = source_elem.findall('.//' + qn('w:r'))
                            orig_runs = child.findall('.//' + qn('w:r'))
                            for r in source_runs:
                                source_elem.remove(r)
                            for r in orig_runs:
                                source_elem.append(deep_copy(r))
                        bookmark_start = OxmlElement('w:bookmarkStart')
                        bookmark_start.set(qn('w:id'), str(bookmark_id))
                        bookmark_start.set(qn('w:name'), '_keepOriginal_')
                        source_elem.insert(0, bookmark_start)
                        bookmark_end = OxmlElement('w:bookmarkEnd')
                        bookmark_end.set(qn('w:id'), str(bookmark_id))
                        source_elem.append(bookmark_end)
                        bookmark_id += 1
                        # 检查是否有 numPr（列表编号），如果有则在 source_elem 中处理
                        if child_numPr is not None:
                            _style_has_numPr = False
                            if list_method == 'style' and list_style:
                                try:
                                    _style_xml = doc.styles[list_style]._element.xml
                                    if '<w:numPr>' in _style_xml or '<w:numPr ' in _style_xml:
                                        _style_has_numPr = True
                                except Exception:
                                    pass
                            if not _style_has_numPr:
                                src_pPr2 = source_elem.find(qn('w:pPr'))
                                if src_pPr2 is not None:
                                    old_numPr = src_pPr2.find(qn('w:numPr'))
                                    if old_numPr is not None:
                                        src_pPr2.remove(old_numPr)
                            else:
                                src_pPr = source_elem.find(qn('w:pPr'))
                                if src_pPr is None:
                                    src_pPr = OxmlElement('w:pPr')
                                    source_elem.insert(0, src_pPr)
                                src_numPr = src_pPr.find(qn('w:numPr'))
                                if src_numPr is None:
                                    src_numPr = deep_copy(child_numPr)
                                    src_pPr.append(src_numPr)
                        if child_numPr is not None and list_method == 'style' and list_style and _is_list_para:
                            list_sid = self.get_style_id_by_name(doc, list_style)
                            if list_sid:
                                _apply_paragraph_style(source_elem, list_sid)
                        if enable_image_style and image_style_override:
                            has_img = (child.find('.//' + qn('w:drawing')) is not None or
                                       child.find('.//' + qn('w:pict')) is not None or
                                       child.find('.//' + qn('pic:pic')) is not None or
                                       child.find('.//' + qn('wp:inline')) is not None or
                                       child.find('.//' + qn('wp:anchor')) is not None)
                            if has_img:
                                img_sid = self.get_style_id_by_name(doc, image_style_override)
                                if img_sid:
                                    _apply_paragraph_style(source_elem, img_sid)
                        new_children.append(source_elem)
                        chapter_buffer.append(source_elem)
                    else:
                        # 表格等非段落元素：需要给原文副本的表格添加 keepOriginal 保护，
                        # 避免语气转换时把表格中的"投标人"转为"本投标人"
                        if child.tag == qn('w:tbl'):
                            source_elem = deepcopy(child)
                            bookmark_id = self._add_keep_original_to_table(source_elem, bookmark_id)
                            new_children.append(source_elem)
                            chapter_buffer.append(child)  # 保留原始引用，供应答副本使用
                        else:
                            new_children.append(child)
                            chapter_buffer.append(child)
            else:
                # 不在任何章节内（文档开头无标题），直接输出
                new_children.append(child)
            
            i += 1
        
        # 处理最后一个章节
        if current_chapter_heading is not None and chapter_buffer:
            has_real_content = any(not self._is_hint_paragraph(elem) for elem in chapter_buffer)
            if has_real_content:
                # 插入应答句
                answer_elem = deep_copy(answer_template)
                new_children.append(answer_elem)
                insert_count += 1
                
                # 复制第二份副本（应答原文，用 copy_template 样式），跳过提示语
                for elem in chapter_buffer:
                    if self._is_hint_paragraph(elem):
                        continue
                    if elem.tag == qn('w:tbl'):
                        source_elem = deep_copy(elem)
                        tbl_style_id = None
                        if table_answer_style:
                            tbl_style_id = self.get_style_id_by_name(doc, table_answer_style)
                        if not tbl_style_id:
                            _sid = self.get_style_id_by_name(doc, answer_copy_style) if answer_copy_style else None
                            tbl_style_id = _sid
                        if tbl_style_id:
                            for p_elem in source_elem.iter(qn('w:p')):
                                pPr = p_elem.find(qn('w:pPr'))
                                if pPr is None:
                                    pPr = OxmlElement('w:pPr')
                                    p_elem.insert(0, pPr)
                                pStyle = pPr.find(qn('w:pStyle'))
                                if pStyle is None:
                                    pStyle = OxmlElement('w:pStyle')
                                    pPr.append(pStyle)
                                pStyle.set(qn('w:val'), tbl_style_id)
                        new_children.append(source_elem)
                        continue
                    _has_numpr = DocumentConverter._elem_has_numbering(elem, doc)
                    _has_list_style = False
                    _is_style_match = False
                    if not _has_numpr:
                        _pPr_el = elem.find(qn('w:pPr'))
                        if _pPr_el is not None:
                            _pS_el = _pPr_el.find(qn('w:pStyle'))
                            if _pS_el is not None:
                                _ps_val = _pS_el.get(qn('w:val'))
                                if _ps_val and '列表段落' in str(_ps_val):
                                    _has_list_style = True
                                if not _has_list_style and _ps_val and list_answer_method == 'style' and list_answer_style:
                                    _ans_sid = self.get_style_id_by_name(doc, list_answer_style)
                                    if _ans_sid and _ps_val == _ans_sid:
                                        _is_style_match = True
                    is_list_para = ((_has_numpr or _has_list_style or _is_style_match) and 
                                    list_answer_method == 'style' and 
                                    list_answer_style)
                    if is_list_para:
                        list_sid = self.get_style_id_by_name(doc, list_answer_style)
                        if list_sid:
                            source_elem = deep_copy(elem)
                            _apply_paragraph_style(source_elem, list_sid)
                            remove_keep_original_from_element(source_elem)
                            new_children.append(source_elem)
                        else:
                            source_elem = deep_copy(copy_template)
                            source_runs = source_elem.findall('.//' + qn('w:r'))
                            orig_runs = elem.findall('.//' + qn('w:r'))
                            for r in source_runs:
                                source_elem.remove(r)
                            for r in orig_runs:
                                source_elem.append(deep_copy(r))
                            remove_keep_original_from_element(source_elem)
                            if enable_image_style and image_style_override:
                                has_img = (elem.find('.//' + qn('w:drawing')) is not None or
                                           elem.find('.//' + qn('w:pict')) is not None or
                                           elem.find('.//' + qn('pic:pic')) is not None or
                                           elem.find('.//' + qn('wp:inline')) is not None or
                                           elem.find('.//' + qn('wp:anchor')) is not None)
                                if has_img:
                                    img_sid = self.get_style_id_by_name(doc, image_style_override)
                                    if img_sid:
                                        _apply_paragraph_style(source_elem, img_sid)
                            new_children.append(source_elem)
                    else:
                        source_elem = deep_copy(copy_template)
                        source_runs = source_elem.findall('.//' + qn('w:r'))
                        orig_runs = elem.findall('.//' + qn('w:r'))
                        for r in source_runs:
                            source_elem.remove(r)
                        for r in orig_runs:
                            source_elem.append(deep_copy(r))
                        remove_keep_original_from_element(source_elem)
                        if enable_image_style and image_style_override:
                            has_img = (elem.find('.//' + qn('w:drawing')) is not None or
                                       elem.find('.//' + qn('w:pict')) is not None or
                                       elem.find('.//' + qn('pic:pic')) is not None or
                                       elem.find('.//' + qn('wp:inline')) is not None or
                                       elem.find('.//' + qn('wp:anchor')) is not None)
                            if has_img:
                                img_sid = self.get_style_id_by_name(doc, image_style_override)
                                if img_sid:
                                    _apply_paragraph_style(source_elem, img_sid)
                        new_children.append(source_elem)
        
        return insert_count, total_heading_count
    
    def _insert_before_paragraphs(self, children, new_children, answer_template, doc):
        """
        逐段前应答（需求4）- 改进版：支持语义段落分组
        逻辑：
        1. 将连续的语义相关段落分组（短句、引号上下文、列表）
        2. 在每个语义单元前插入一个应答句
        :return: (insert_count, total_heading_count)
        """
        from copy import deepcopy as deep_copy
        
        insert_count = 0
        total_heading_count = 0
        
        # 第一步：将元素分组为语义单元
        semantic_groups = self._group_semantic_units(children, doc)
        
        # 第二步：遍历每个语义单元，在单元前插入应答句
        for group in semantic_groups:
            if not group:
                continue
            
            first_elem = group[0]
            
            # 统计标题数量
            if hasattr(first_elem, 'tag') and first_elem.tag == qn('w:p'):
                if self.is_heading_paragraph(first_elem, doc):
                    total_heading_count += len([e for e in group if hasattr(e, 'tag') and e.tag == qn('w:p') and self.is_heading_paragraph(e, doc)])
            
            # 判断是否为需要插入应答句的语义单元
            should_insert = self._should_insert_answer_for_group(group, doc)
            
            if should_insert:
                # 在语义单元前插入应答句
                answer_elem = deep_copy(answer_template)
                new_children.append(answer_elem)
                insert_count += 1
            
            # 添加语义单元中的所有元素
            for elem in group:
                new_children.append(elem)
        
        return insert_count, total_heading_count
    
    def _is_list_paragraph(self, elem):
        """判断段落是否是列表（有编号或项目符号）
        numId=0 表示无编号，应视为没有列表。
        """
        if not hasattr(elem, 'tag') or elem.tag != qn('w:p'):
            return False
        
        pPr = elem.find(qn('w:pPr'))
        if pPr is not None:
            # 检查是否有编号（numPr）
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                numId = numPr.find(qn('w:numId'))
                if numId is not None:
                    val = numId.get(qn('w:val'))
                    if val == '0':
                        return False
                return True
        
        return False
    
    def _get_paragraph_text(self, elem):
        """
        获取段落的文本内容
        :param elem: 段落元素
        :return: 文本字符串
        """
        if not hasattr(elem, 'tag') or elem.tag != qn('w:p'):
            return ""
        
        text_elems = elem.findall('.//' + qn('w:t'))
        return ''.join([t.text for t in text_elems if t.text])
    
    def _ends_with_colon_or_quote(self, text):
        """
        判断文本是否以冒号或引号结尾（需要与下一段合并）
        :param text: 文本内容
        :return: True 如果需要合并
        """
        if not text:
            return False
        
        # 去除末尾空白
        text = text.rstrip()
        
        # 检查是否以冒号、冒号+引号结尾
        if text.endswith('：') or text.endswith(':'):
            return True
        if text.endswith('”') or text.endswith('"'):
            # 检查前面是否有冒号
            if len(text) > 1 and (text[-2] == '：' or text[-2] == ':'):
                return True
        
        return False
    
    def _is_short_paragraph(self, text, threshold=20):
        """
        判断是否为短段落
        :param text: 文本内容
        :param threshold: 字数阈值
        :return: True 如果是短段落
        """
        if not text:
            return True
        return len(text.strip()) < threshold
    
    def _is_manual_numbered_paragraph(self, text):
        """
        判断段落是否是手动编号（如"1、"、"2）"、"a."等）
        :param text: 段落文本
        :return: True 如果是手动编号段落
        """
        if not text:
            return False
        
        text = text.strip()
        
        # 匹配常见的中文编号格式
        import re
        patterns = [
            r'^\d+[、\.．]',      # 1、 或 1. 或 1．
            r'^\d+）',             # 1）
            r'^\d+\)',             # 1)
            r'^[（(]\d+[）)]',     # （1） 或 (1)
            r'^[一二三四五六七八九十]+[、\.．]',  # 一、 或 一.
            r'^[a-zA-Z][、\.．]',  # a. 或 A、
            r'^[a-zA-Z]）',        # a）
            r'^[a-zA-Z]\)',        # a)
        ]
        
        for pattern in patterns:
            if re.match(pattern, text):
                return True
        
        return False
    
    def _is_bullet_point_paragraph(self, text):
        """
        判断段落是否是项目符号列表（如●、■、◆等）
        :param text: 段落文本
        :return: True 如果是项目符号段落
        """
        if not text:
            return False
        
        text = text.strip()
        
        # 常见的项目符号
        bullet_symbols = ['●', '○', '■', '□', '◆', '◇', '▲', '△', '►', '▶', '•', '-', '*']
        
        for symbol in bullet_symbols:
            if text.startswith(symbol):
                return True
        
        return False
    
    def _is_empty_paragraph(self, text):
        """
        判断段落是否为空行
        :param text: 段落文本
        :return: True 如果是空行
        """
        if not text:
            return True
        return len(text.strip()) == 0
    
    def _group_semantic_units(self, children, doc):
        """
        将元素分组为语义单元
        规则：
        1. 标题单独成组
        2. 连续的列表项合并为一个组
        3. 以冒号/引号结尾的段落与下一段合并
        4. 连续的短段落合并为一个组
        :param children: 所有元素列表
        :param doc: 文档对象
        :return: 分组后的列表 [[elem1, elem2], [elem3], ...]
        """
        groups = []
        current_group = []
        
        for i, child in enumerate(children):
            if not hasattr(child, 'tag'):
                # 非段落元素（如表格），单独成组
                if current_group:
                    groups.append(current_group)
                    current_group = []
                groups.append([child])
                continue
            
            # 检查是否为标题
            is_heading = False
            if child.tag == qn('w:p') and self.is_heading_paragraph(child, doc):
                is_heading = True
            
            # 获取文本内容（提前获取，供后续判断使用）
            text = self._get_paragraph_text(child) if child.tag == qn('w:p') else ""
            
            # 检查是否为空行
            is_empty = False
            if child.tag == qn('w:p'):
                is_empty = self._is_empty_paragraph(text)
            
            # 如果是空行，跳过不处理
            if is_empty:
                continue
            
            # 检查是否为列表
            is_list = False
            if child.tag == qn('w:p'):
                is_list = self._is_list_paragraph(child)
            
            # 检查是否为手动编号段落
            is_manual_numbered = False
            if child.tag == qn('w:p'):
                is_manual_numbered = self._is_manual_numbered_paragraph(text)
            
            # 检查是否为项目符号段落
            is_bullet_point = False
            if child.tag == qn('w:p'):
                is_bullet_point = self._is_bullet_point_paragraph(text)
            
            # 决策逻辑
            if is_heading:
                # 标题：结束当前组，标题单独成组
                if current_group:
                    groups.append(current_group)
                    current_group = []
                groups.append([child])
            elif is_list or is_manual_numbered or is_bullet_point:
                # 列表项（包括自动列表、手动编号、项目符号）：如果前一个也是列表，则合并；否则新起一组
                should_merge = False
                
                if current_group:
                    # 检查前一段是否是列表/编号/项目符号
                    if (self._is_last_group_list(current_group, doc) or 
                        self._is_last_group_manual_numbered(current_group) or
                        self._is_last_group_bullet_point(current_group)):
                        should_merge = True
                    else:
                        # 检查前一段是否以冒号/引号结尾
                        prev_text = self._get_last_paragraph_text(current_group)
                        if self._ends_with_colon_or_quote(prev_text):
                            should_merge = True
                
                if should_merge:
                    current_group.append(child)
                else:
                    if current_group:
                        groups.append(current_group)
                    current_group = [child]
            elif self._ends_with_colon_or_quote(text):
                # 以冒号/引号结尾：与下一段合并
                current_group.append(child)
                # 标记需要与下一段合并（通过保持 current_group 不结束）
            elif self._is_short_paragraph(text) and current_group:
                # 短段落：尝试与前一段合并
                prev_text = self._get_last_paragraph_text(current_group)
                if self._is_short_paragraph(prev_text) or self._ends_with_colon_or_quote(prev_text):
                    # 前一段也是短句或以冒号结尾，合并
                    current_group.append(child)
                else:
                    # 否则新起一组
                    groups.append(current_group)
                    current_group = [child]
            else:
                # 普通段落：检查是否需要与前一段合并
                if current_group:
                    prev_text = self._get_last_paragraph_text(current_group)
                    prev_is_numbered_or_bullet = self._is_last_group_manual_numbered(current_group) or self._is_last_group_bullet_point(current_group)
                    
                    if self._ends_with_colon_or_quote(prev_text):
                        # 前一段以冒号结尾，合并
                        current_group.append(child)
                    elif prev_is_numbered_or_bullet:
                        # 前一段是手动编号或项目符号，当前普通段落是其内容，合并
                        current_group.append(child)
                    else:
                        # 否则新起一组
                        groups.append(current_group)
                        current_group = [child]
                else:
                    current_group = [child]
        
        # 处理最后一组
        if current_group:
            groups.append(current_group)
        
        return groups
    
    def _is_last_group_list(self, group, doc):
        """检查组中最后一个元素是否是列表"""
        if not group:
            return False
        last_elem = group[-1]
        if hasattr(last_elem, 'tag') and last_elem.tag == qn('w:p'):
            return self._is_list_paragraph(last_elem)
        return False
    
    def _is_last_group_manual_numbered(self, group):
        """检查组中最后一个元素是否是手动编号段落"""
        if not group:
            return False
        last_elem = group[-1]
        if hasattr(last_elem, 'tag') and last_elem.tag == qn('w:p'):
            text = self._get_paragraph_text(last_elem)
            return self._is_manual_numbered_paragraph(text)
        return False
    
    def _is_last_group_bullet_point(self, group):
        """检查组中最后一个元素是否是项目符号段落"""
        if not group:
            return False
        last_elem = group[-1]
        if hasattr(last_elem, 'tag') and last_elem.tag == qn('w:p'):
            text = self._get_paragraph_text(last_elem)
            return self._is_bullet_point_paragraph(text)
        return False
    
    def _get_last_paragraph_text(self, group):
        """获取组中最后一个段落的文本"""
        for elem in reversed(group):
            if hasattr(elem, 'tag') and elem.tag == qn('w:p'):
                return self._get_paragraph_text(elem)
        return ""
    
    def _should_insert_answer_for_group(self, group, doc):
        """
        判断是否应该为该语义单元插入应答句
        :param group: 语义单元（元素列表）
        :param doc: 文档对象
        :return: True 如果需要插入
        """
        if not group:
            return False
        
        first_elem = group[0]
        
        # 排除标题
        if hasattr(first_elem, 'tag') and first_elem.tag == qn('w:p'):
            if self.is_heading_paragraph(first_elem, doc):
                return False
        
        # 排除图片
        if hasattr(first_elem, 'tag') and first_elem.tag == qn('w:p'):
            if len(first_elem.findall('.//' + qn('a:blip'))) > 0:
                return False
        
        # 其他情况都插入（包括列表、普通段落等）
        return True
    
    def _insert_after_paragraphs(self, children, new_children, answer_template, doc):
        """
        逐段后应答（需求5）- 改进版：支持语义段落分组
        逻辑：
        1. 将连续的语义相关段落分组（短句、引号上下文、列表）
        2. 在每个语义单元后插入一个应答句
        :return: (insert_count, total_heading_count)
        """
        from copy import deepcopy as deep_copy
        
        insert_count = 0
        total_heading_count = 0
        
        # 第一步：将元素分组为语义单元
        semantic_groups = self._group_semantic_units(children, doc)
        
        # 第二步：遍历每个语义单元，在单元后插入应答句
        for group in semantic_groups:
            if not group:
                continue
            
            first_elem = group[0]
            
            # 统计标题数量
            if hasattr(first_elem, 'tag') and first_elem.tag == qn('w:p'):
                if self.is_heading_paragraph(first_elem, doc):
                    total_heading_count += len([e for e in group if hasattr(e, 'tag') and e.tag == qn('w:p') and self.is_heading_paragraph(e, doc)])
            
            # 先添加语义单元中的所有元素
            for elem in group:
                new_children.append(elem)
            
            # 判断是否为需要插入应答句的语义单元
            should_insert = self._should_insert_answer_for_group(group, doc)
            
            if should_insert:
                # 在语义单元后插入应答句
                answer_elem = deep_copy(answer_template)
                new_children.append(answer_elem)
                insert_count += 1
        
        return insert_count, total_heading_count
    
    def full_convert(self, source_file, template_file, output_file, 
                     custom_style_map=None, do_mood=True, 
                     answer_text=None, answer_style=None,
                     answer_source_style=None,
                     answer_copy_style=None,
                     table_answer_style=None,
                     list_bullet=None, do_answer_insertion=True,
                     answer_mode='before_heading',
                     do_hint_insertion=False, hint_type='text',
                     hint_text='招标文件原文', hint_image_path=None,
                     hint_style='Normal',
                     progress_callback=None, warning_callback=None,
                     source_styles_cache=None,
                     table_style_override=None, enable_table_style=False,
                     image_style_override=None, enable_image_style=False,
                     remove_chapter_label=False,
                     list_method='bullet', list_style='Body Text',
                     list_answer_method='bullet', list_answer_style='Body Text',
                     list_answer_bullet='● ',
                     enable_list_style=True):
        """
        完整转换流程：样式转换 -> 语气转换 -> 插入应答句
        [HIGH_VOLTAGE] 性能优化：合并为一次性流水线，避免多次加载/保存文档
        :param source_file: 源文件
        :param template_file: 模板文件
        :param output_file: 最终输出文件
        :param custom_style_map: 自定义样式映射
        :param do_mood: 是否进行语气转换
        :param answer_text: 应答文本
        :param answer_style: 应答样式
        :param list_bullet: 列表段落符号
        :param do_answer_insertion: 是否插入应答句
        :param answer_mode: 应答句插入模式
            - 'before_heading': 标题前插入（默认）
            - 'after_heading': 章节后插入（在章节正文最后一段后插入）
        :param progress_callback: 进度回调函数 callback(step, message)
        :param warning_callback: 警告回调函数 callback(message)
        :param source_styles_cache: 缓存的源文件样式列表（可选，避免重复分析）
        :param table_style_override: 表格样式覆盖（当enable_table_style=True时使用）
        :param enable_table_style: 是否启用表格样式覆盖
        :param image_style_override: 图片样式覆盖（当enable_image_style=True时使用）
        :param enable_image_style: 是否启用图片样式覆盖
        :return: (success, actual_output_file, message)
        """
        import time
        start_time = time.time()
        
        # 固定7个步骤，确保进度条能正确填满
        if progress_callback:
            progress_callback(1, "正在进行转换...")
        
        # ========== [HIGH_VOLTAGE] 性能优化：一次性流水线处理 ==========
        # 原来：Load → StyleConv → Save → Load → MoodConv → Save → Load → AnswerInsert → Save
        # 现在：  Load → StyleConv → MoodConv → AnswerInsert → Save（一次加载，一次保存）
        
        # 特殊处理：copy_chapter 模式时，调换语气转换和应答句插入的顺序，
        # 使第一份副本保留原文（祈使语气），第二份副本完成语气转换。
        
        # 步骤1：样式转换（返回Document对象，不保存）
        doc = self._convert_styles_in_memory(source_file, template_file, custom_style_map, list_bullet,
                                              warning_callback, source_styles_cache,
                                              table_style_override, enable_table_style,
                                              image_style_override, enable_image_style,
                                              remove_chapter_label=remove_chapter_label,
                                              list_method=list_method,
                                              list_style=list_style,
                                              enable_list_style=enable_list_style)
        if doc is None:
            elapsed = time.time() - start_time
            return False, output_file, f"样式转换失败（耗时{elapsed:.1f}秒）"
        
        if progress_callback:
            progress_callback(2, "正在进行转换...")
        
        # ========== 章节提示语插入（在语气转换之前） ==========
        if do_hint_insertion:
            hint_result = self._insert_hint_in_memory(
                doc, hint_type, hint_text, hint_image_path, hint_style
            )
            if not hint_result:
                elapsed = time.time() - start_time
                return False, output_file, f"插入提示语失败（耗时{elapsed:.1f}秒）"
        
        actual_output_file = output_file
        
        # ========== 根据 answer_mode 决定流水线顺序 ==========
        if answer_mode == 'copy_chapter' and do_answer_insertion and do_mood:
            # ===== copy_chapter 模式专用流水线 =====
            # 步骤2-3：插入应答句（在语气转换之前，此时原文未转换）
            if progress_callback:
                progress_callback(3, "正在插入应答句（保留原文模式）...")
            insert_result = self._insert_response_in_memory(
                doc, answer_text, answer_style, mode=answer_mode,
                answer_source_style=answer_source_style,
                answer_copy_style=answer_copy_style,
                table_answer_style=table_answer_style,
                list_method=list_method,
                list_style=list_style,
                list_answer_method=list_answer_method,
                list_answer_style=list_answer_style,
                list_answer_bullet=list_answer_bullet,
                enable_image_style=enable_image_style,
                image_style_override=image_style_override,
                enable_list_style=enable_list_style
            )
            if not insert_result:
                elapsed = time.time() - start_time
                return False, output_file, f"插入应答句失败（耗时{elapsed:.1f}秒）"
            
            if progress_callback:
                progress_callback(4, "正在插入应答句...")
            
            # 步骤4-5：语气转换（跳过标记为 keepOriginal 的第一份副本段落）
            if progress_callback:
                progress_callback(5, "正在语气转换（跳过原文副本）...")
            mood_result = self._convert_mood_in_memory(doc)
            if not mood_result:
                elapsed = time.time() - start_time
                return False, output_file, f"语气转换失败（耗时{elapsed:.1f}秒）"
            
            if progress_callback:
                progress_callback(6, "正在转换...")
        
        else:
            # ===== 标准流水线（其他模式） =====
            # 步骤2-3：语气转换（直接在内存中的Document对象上操作）
            if do_mood:
                if progress_callback:
                    progress_callback(3, "正在进行转换...")
                mood_result = self._convert_mood_in_memory(doc)
                if not mood_result:
                    elapsed = time.time() - start_time
                    return False, output_file, f"语气转换失败（耗时{elapsed:.1f}秒）"
                if progress_callback:
                    progress_callback(4, "正在进行转换...")
            else:
                # 跳过语气转换，但仍然占用步骤3和4
                if progress_callback:
                    progress_callback(3, "正在进行转换...")
                    progress_callback(4, "正在进行转换...")
            
            # 步骤5-6：插入应答句（直接在内存中的Document对象上操作）
            if do_answer_insertion:
                if progress_callback:
                    progress_callback(5, "正在进行转换...")
                insert_result = self._insert_response_in_memory(
                    doc, answer_text, answer_style, mode=answer_mode,
                    answer_source_style=answer_source_style,
                    answer_copy_style=answer_copy_style,
                    table_answer_style=table_answer_style,
                    list_method=list_method,
                    list_style=list_style,
                    list_answer_method=list_answer_method,
                    list_answer_style=list_answer_style,
                    list_answer_bullet=list_answer_bullet,
                    enable_image_style=enable_image_style,
                    image_style_override=image_style_override,
                    enable_list_style=enable_list_style
                )
                if not insert_result:
                    elapsed = time.time() - start_time
                    return False, output_file, f"插入应答句失败（耗时{elapsed:.1f}秒）"
                
                if progress_callback:
                    progress_callback(6, "正在进行转换...")
            else:
                # 不插入应答句，但仍然占用步骤5和6
                if progress_callback:
                    progress_callback(5, "正在进行转换...")
                    progress_callback(6, "正在进行转换...")
        
        # 步骤7：保存文档（只保存一次！）
        if progress_callback:
            progress_callback(7, "正在保存...")
        
        success, actual_file, msg = self.save_with_retry(doc, output_file)
        elapsed = time.time() - start_time
        
        if success:
            return True, actual_file, f"转换成功完成！（耗时{elapsed:.1f}秒）"
        else:
            return False, output_file, f"保存失败: {msg}（耗时{elapsed:.1f}秒）"
    
    def _generate_unique_filename(self, output_file):
        """
        生成不重复的文件名：若原始文件名已存在，则追加 _HHMMSS 时间戳；
        若同一秒内仍有冲突，再追加三位序号（_001）。
        :param output_file: 原始输出文件路径
        :return: 可用的文件路径
        """
        import os

        if not os.path.exists(output_file):
            return output_file

        base, ext = os.path.splitext(output_file)
        time_suffix = datetime.now().strftime("_%H%M%S")
        candidate = f"{base}{time_suffix}{ext}"

        if not os.path.exists(candidate):
            return candidate

        # 极端情况：同一秒内仍有冲突，追加序号
        for i in range(1, 1000):
            candidate = f"{base}{time_suffix}_{i:03d}{ext}"
            if not os.path.exists(candidate):
                return candidate

        # 兜底，理论上不会走到这里
        return output_file

    def save_with_retry(self, doc, output_file, max_retries=10):
        """
        智能保存文档：若目标文件已存在（重名），自动追加 _HHMMSS 时间戳；
        保存过程中如遇占用，也会自动生成新的备用文件名并重试。
        :param doc: Document对象
        :param output_file: 原始输出文件路径
        :param max_retries: 最大重试次数
        :return: (success, actual_output_file, message)
        """
        import os
        import time

        # 首次保存：若存在重名文件则生成带时间戳的备用文件名
        current_file = self._generate_unique_filename(output_file)
        if current_file != output_file:
            print(f"  检测到重名文件，使用备用文件名: {current_file}")

        for attempt in range(max_retries):
            try:
                doc.save(current_file)
                if current_file == output_file:
                    return True, current_file, f"文档已保存到 {current_file}"
                else:
                    return True, current_file, f"检测到重名，文档已保存到: {current_file}"
            except (PermissionError, OSError, IOError) as e:
                # 文件在保存过程中被占用（罕见情况）
                if attempt == 0:
                    print(f"  警告：保存文档失败（文件可能被占用）: {e}")

                # 生成新的文件名
                base, ext = os.path.splitext(output_file)
                time_suffix = datetime.now().strftime("_%H%M%S")
                current_file = f"{base}{time_suffix}{ext}"

                # 如果新文件名仍冲突，追加序号避免覆盖
                idx = 1
                while os.path.exists(current_file) and idx < 1000:
                    current_file = f"{base}{time_suffix}_{idx:03d}{ext}"
                    idx += 1

                print(f"  尝试备用文件名: {current_file}")

                # 稍等片刻再重试
                time.sleep(0.3)
            except Exception as e:
                # 其他异常直接返回失败
                return False, output_file, f"保存文档失败: {e}"

        # 重试次数用尽
        return False, output_file, f"无法保存文档，已尝试 {max_retries} 次"
    
    def _convert_styles_in_memory(self, source_file, template_file, custom_style_map=None, list_bullet=None,
                                   warning_callback=None, source_styles_cache=None,
                                   table_style_override=None, enable_table_style=False,
                                   image_style_override=None, enable_image_style=False,
                                   remove_chapter_label=False,
                                   list_method='bullet', list_style='Body Text',
                                   enable_list_style=True):
        """
        [HIGH_VOLTAGE] 性能优化：在内存中进行样式转换，不保存中间文件
        :param table_style_override: 表格样式覆盖（当enable_table_style=True时使用）
        :param enable_table_style: 是否启用表格样式覆盖
        :param image_style_override: 图片样式覆盖（当enable_image_style=True时使用）
        :param enable_image_style: 是否启用图片样式覆盖
        :return: Document对象或None（失败时）
        """
        try:
            from docx import Document
            from copy import deepcopy
            from lxml import etree
            from docx.oxml.ns import qn
            
            # 加载源文档和模板文档
            source_doc = Document(source_file)
            new_doc = Document(template_file)
            self.clear_document_content(new_doc)
            
            # 设置样式映射
            style_map = STYLE_MAP.copy()
            if custom_style_map:
                style_map.update(custom_style_map)
            self.current_style_map = style_map
            
            # 使用缓存的样式列表或重新分析
            if source_styles_cache:
                self.source_styles = source_styles_cache
            else:
                self.source_styles = self.get_all_styles_from_doc(source_doc)
            
            # 获取页面宽度信息
            section = new_doc.sections[0]
            page_width = section.page_width
            left_margin = section.left_margin
            right_margin = section.right_margin
            available_width = page_width - left_margin - right_margin
            
            # 处理源文档的所有元素
            body = source_doc.element.body
            para_idx = 0
            table_idx = 0
            
            for child in body:
                if child.tag == qn('w:p'):
                    if para_idx < len(source_doc.paragraphs):
                        para = source_doc.paragraphs[para_idx]
                        src_style = para.style.name
                        
                        # ★ 修复：对于有编号的列表段落，使用虚拟样式名（如 "1 列表段落"）进行样式映射，
                        # 而不是用真实的样式名（通常是 "Normal"）。这样 Step 3 中用户配置的
                        # "1 列表段落" → "BN_原文引用列表项目符号" 映射才能生效。
                        # ★ 修复：标题段落（有outlineLevel或样式为Heading）即使有编号也不视为列表段落，
                        # 应走正常的标题样式映射路径。
                        is_heading_by_outline = self.get_outline_level(para) > 0
                        is_heading_by_style = src_style in HEADING_STYLES
                        style_map = getattr(self, 'current_style_map', STYLE_MAP)
                        is_heading_by_mapped = style_map.get(src_style) in HEADING_STYLES
                        is_custom_heading = '标题' in src_style or src_style.startswith('Heading')
                        if is_heading_by_outline or is_heading_by_style or is_heading_by_mapped or is_custom_heading:
                            # 标题段落：走标题样式映射
                            target_style = self.get_target_style(src_style, new_doc, source_file)
                            _enable_list_fallback = enable_list_style
                            _list_style = list_style
                            _resolve_num_text = None
                        elif self.has_numbering(para):
                            # 非标题的列表段落：使用虚拟样式名进行映射
                            virtual_style = self._detect_numbering_format(para)
                            target_style = self.get_target_style(virtual_style, new_doc, source_file)
                            _resolve_num_text = None
                            if virtual_style in style_map:
                                # ★ 修复：虚拟样式已在 Step 3 映射
                                target_has_numPr = False
                                try:
                                    tpl_style_xml = new_doc.styles[target_style]._element.xml
                                    if '<w:numPr>' in tpl_style_xml or '<w:numPr ' in tpl_style_xml:
                                        target_has_numPr = True
                                except KeyError:
                                    pass
                                if target_has_numPr:
                                    _enable_list_fallback = True
                                    _list_style = target_style
                                else:
                                    _enable_list_fallback = False
                                    _resolve_num_text = self._resolve_auto_numbering_text(para)
                            else:
                                _enable_list_fallback = enable_list_style
                                _list_style = list_style
                        else:
                            target_style = self.get_target_style(src_style, new_doc, source_file)
                            _enable_list_fallback = enable_list_style
                            _list_style = list_style
                            _resolve_num_text = None
                        
                        # 使用copy_paragraph_with_images方法复制段落
                        self.copy_paragraph_with_images(
                            para, new_doc, target_style,
                            page_width, available_width,
                            para_idx, source_file,
                            warning_callback,
                            image_style_override=image_style_override,
                            enable_image_style=enable_image_style,
                            remove_chapter_label=remove_chapter_label,
                            list_method=list_method,
                            list_style=_list_style,
                            enable_list_style=_enable_list_fallback,
                            resolved_numbering_text=_resolve_num_text
                        )
                        para_idx += 1
                elif child.tag == qn('w:tbl'):
                    if table_idx < len(source_doc.tables):
                        table = source_doc.tables[table_idx]
                        self.copy_table_with_images(
                            table, new_doc, table_idx, available_width,
                            source_file, warning_callback,
                            table_style_override=table_style_override,
                            enable_table_style=enable_table_style
                        )
                        table_idx += 1
            
            return new_doc
        except Exception as e:
            print(f"样式转换失败: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _convert_mood_in_memory(self, doc):
        """
        [HIGH_VOLTAGE] 性能优化：在内存中进行语气转换，不保存中间文件
        跳过标记为 _keepOriginal_ 的段落（copy_chapter 模式的第一份副本），转换完成后清除标记
        :param doc: Document对象
        :return: True/False
        """
        try:
            modified_count = 0
            para_count = 0
            
            for para in doc.paragraphs:
                para_count += 1
                # 跳过标记为 keepOriginal 的段落
                if self._is_keep_original_paragraph(para._element):
                    continue
                # 跳过标题段落（标题中的"投标人"不应转换）
                if self.is_heading_paragraph(para._element, doc):
                    continue
                if self.process_paragraph_mood(para):
                    modified_count += 1
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            para_count += 1
                            if self._is_keep_original_paragraph(para._element):
                                continue
                            # 跳过标题段落（表格内通常无标题，但保持一致）
                            if self.is_heading_paragraph(para._element, doc):
                                continue
                            if self.process_paragraph_mood(para):
                                modified_count += 1
            
            # 清除所有 _keepOriginal_ 书签标记
            self._remove_keep_original_markers(doc)
            
            print(f"语气转换完成！处理段落: {para_count}, 修改: {modified_count}")
            return True
        except Exception as e:
            print(f"语气转换失败: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _is_keep_original_paragraph(self, elem):
        """检查段落是否标记为 keepOriginal（不做语气转换）"""
        if not hasattr(elem, 'tag') or elem.tag != qn('w:p'):
            return False
        for child in elem:
            if child.tag == qn('w:bookmarkStart'):
                if child.get(qn('w:name')) == '_keepOriginal_':
                    return True
        return False
    
    def _remove_keep_original_markers(self, doc):
        """清除文档中所有 _keepOriginal_ 书签标记"""
        body = doc.element.body
        for elem in body.iter(qn('w:p')):
            bookmark_ids_to_remove = set()
            starts_to_remove = []
            ends_to_remove = []
            
            for child in elem:
                if child.tag == qn('w:bookmarkStart'):
                    if child.get(qn('w:name')) == '_keepOriginal_':
                        bookmark_ids_to_remove.add(child.get(qn('w:id')))
                        starts_to_remove.append(child)
                elif child.tag == qn('w:bookmarkEnd'):
                    if child.get(qn('w:id')) in bookmark_ids_to_remove:
                        ends_to_remove.append(child)
            
            for start in starts_to_remove:
                elem.remove(start)
            for end in ends_to_remove:
                elem.remove(end)

    def _add_keep_original_to_table(self, tbl_elem, start_bookmark_id):
        """为表格内所有段落添加 _keepOriginal_ 标记，返回下一个可用的 bookmark_id
        
        用于 copy_chapter 模式：表格属于原文副本的一部分，其内容（如"投标人"）
        不应被语气转换修改。
        """
        bid = start_bookmark_id
        for p_elem in tbl_elem.iter(qn('w:p')):
            bookmark_start = OxmlElement('w:bookmarkStart')
            bookmark_start.set(qn('w:id'), str(bid))
            bookmark_start.set(qn('w:name'), '_keepOriginal_')
            p_elem.insert(0, bookmark_start)
            bookmark_end = OxmlElement('w:bookmarkEnd')
            bookmark_end.set(qn('w:id'), str(bid))
            p_elem.append(bookmark_end)
            bid += 1
        return bid

    def _is_hint_paragraph(self, elem):
        """检查段落是否标记为提示语（hint）"""
        if not hasattr(elem, 'tag') or elem.tag != qn('w:p'):
            return False
        for child in elem:
            if child.tag == qn('w:bookmarkStart'):
                if child.get(qn('w:name')) == '_hint_':
                    return True
        return False

    def _is_ole_placeholder_paragraph(self, elem):
        """检查段落是否包含OLE占位提示文本 [OLE对象，请手动复制]"""
        if not hasattr(elem, 'tag') or elem.tag != qn('w:p'):
            return False
        for run in elem.findall('.//' + qn('w:r')):
            text_elem = run.find(qn('w:t'))
            if text_elem is not None and text_elem.text and '[OLE对象，请手动复制]' in text_elem.text:
                return True
        return False

    def _remove_hint_markers(self, doc):
        """清除文档中所有 _hint_ 书签标记"""
        body = doc.element.body
        for elem in body.iter(qn('w:p')):
            bookmark_ids_to_remove = set()
            starts_to_remove = []
            ends_to_remove = []
            for child in elem:
                if child.tag == qn('w:bookmarkStart'):
                    if child.get(qn('w:name')) == '_hint_':
                        bookmark_ids_to_remove.add(child.get(qn('w:id')))
                        starts_to_remove.append(child)
                elif child.tag == qn('w:bookmarkEnd'):
                    if child.get(qn('w:id')) in bookmark_ids_to_remove:
                        ends_to_remove.append(child)
            for start in starts_to_remove:
                elem.remove(start)
            for end in ends_to_remove:
                elem.remove(end)
    
    def _insert_hint_in_memory(self, doc, hint_type='text', hint_text='招标文件原文',
                               hint_image_path=None, hint_style='Normal'):
        """
        在内存中插入章节提示语（在每个章节标题后、正文开始前插入提示语）
        :param doc: Document对象
        :param hint_type: 提示语类型 'text' 或 'image'
        :param hint_text: 提示语文本内容
        :param hint_image_path: 提示语图片文件路径
        :param hint_style: 提示语段落样式
        :return: True/False
        """
        try:
            from copy import deepcopy
            
            self.ensure_style_exists(doc, hint_style)

            # 计算可用页面宽度（用于图片提示语）
            section = doc.sections[0]
            available_width = section.page_width - section.left_margin - section.right_margin
            
            body = doc.element.body
            children = list(body)
            new_children = []
            insert_count = 0
            total_heading_count = 0
            hint_bookmark_id = 0
            
            i = 0
            while i < len(children):
                child = children[i]
                
                if not hasattr(child, 'tag'):
                    new_children.append(child)
                    i += 1
                    continue
                
                new_children.append(child)
                
                # 检查是否为标题
                if child.tag == qn('w:p') and self.is_heading_paragraph(child, doc):
                    total_heading_count += 1
                    
                    # 检查下一个元素：如果不是标题，在标题后插入提示语
                    if i + 1 < len(children):
                        next_elem = children[i + 1]
                        if hasattr(next_elem, 'tag') and not self.is_heading_paragraph(next_elem, doc):
                            if hint_type == 'text':
                                # 文本提示语
                                hint_para = doc.add_paragraph(hint_text)
                                hint_para.style = hint_style
                                hint_elem = deepcopy(hint_para._element)
                                hint_para._element.getparent().remove(hint_para._element)
                                # 标记为提示语，避免 copy_chapter 模式重复复制
                                self._add_hint_marker(hint_elem, hint_bookmark_id)
                                hint_bookmark_id += 1
                                new_children.append(hint_elem)
                                insert_count += 1
                            elif hint_type == 'image' and hint_image_path:
                                # 图片提示语：宽度设为版心宽度，高度按比例缩放
                                hint_para = doc.add_paragraph()
                                hint_para.style = hint_style
                                try:
                                    picture = hint_para.add_run().add_picture(hint_image_path)
                                    picture.width = available_width
                                    # 使用 PIL 按原始宽高比显式计算高度，避免比例异常
                                    try:
                                        from PIL import Image
                                        with Image.open(hint_image_path) as img:
                                            img_w, img_h = img.size
                                        if img_w and img_w > 0:
                                            picture.height = int(int(picture.width) * img_h / img_w)
                                    except Exception:
                                        pass
                                except Exception as e:
                                    print(f"插入提示语图片失败: {e}")
                                    hint_para.text = hint_text  # 回退为文本
                                hint_elem = deepcopy(hint_para._element)
                                hint_para._element.getparent().remove(hint_para._element)
                                # 标记为提示语，避免 copy_chapter 模式重复复制
                                self._add_hint_marker(hint_elem, hint_bookmark_id)
                                hint_bookmark_id += 1
                                new_children.append(hint_elem)
                                insert_count += 1
                
                i += 1
            
            # 清空并重组body
            for child in list(body):
                body.remove(child)
            for elem in new_children:
                body.append(elem)
            
            print(f"章节提示语插入完成！插入: {insert_count}个，标题: {total_heading_count}个")
            return True
        except Exception as e:
            print(f"插入提示语失败: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _add_hint_marker(self, elem, bookmark_id):
        """为提示语段落添加 _hint_ 书签标记"""
        if not hasattr(elem, 'tag') or elem.tag != qn('w:p'):
            return
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), str(bookmark_id))
        bookmark_start.set(qn('w:name'), '_hint_')
        elem.insert(0, bookmark_start)
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), str(bookmark_id))
        elem.append(bookmark_end)
    
    def _insert_response_in_memory(self, doc, answer_text=None, answer_style=None, mode='before_heading',
                                    answer_source_style=None,
                                    answer_copy_style=None,
                                    table_answer_style=None,
                                    list_method='bullet', list_style='Body Text',
                                    list_answer_method='bullet', list_answer_style='Body Text',
                                    list_answer_bullet='● ',
                                    enable_image_style=False, image_style_override=None,
                                    enable_list_style=True):
        """
        [HIGH_VOLTAGE] 性能优化：在内存中插入应答句，不保存中间文件
        :param doc: Document对象
        :param answer_text: 应答文本
        :param answer_style: 应答样式
        :param mode: 插入模式
        :return: True/False
        """
        try:
            from copy import deepcopy
            from docx.oxml.ns import qn
            
            if answer_text is None:
                answer_text = ANSWER_TEXT
            if answer_style is None:
                answer_style = ANSWER_STYLE
            
            self.ensure_style_exists(doc, answer_style)
            
            # 预创建应答段落模板
            temp_para = doc.add_paragraph(answer_text)
            temp_para.style = answer_style
            answer_template = deepcopy(temp_para._element)
            temp_para._element.getparent().remove(temp_para._element)
            
            body = doc.element.body
            children = list(body)
            new_children = []
            
            # 根据模式选择不同的处理逻辑（与 insert_response_after_headings 保持一致）
            if mode == 'before_heading':
                insert_count, total_heading_count = self._insert_before_headings(
                    children, new_children, answer_template, doc
                )
            elif mode == 'after_heading':
                insert_count, total_heading_count = self._insert_after_headings(
                    children, new_children, answer_template, doc
                )
            elif mode == 'copy_chapter':
                insert_count, total_heading_count = self._insert_with_copy_chapter(
                    children, new_children, answer_template, doc,
                    answer_source_style=answer_source_style,
                    answer_copy_style=answer_copy_style,
                    table_answer_style=table_answer_style,
                    list_method=list_method,
                    list_style=list_style,
                    list_answer_method=list_answer_method,
                    list_answer_style=list_answer_style,
                    list_answer_bullet=list_answer_bullet,
                    enable_image_style=enable_image_style,
                    image_style_override=image_style_override,
                    enable_list_style=enable_list_style
                )
            elif mode == 'before_paragraph':
                insert_count, total_heading_count = self._insert_before_paragraphs(
                    children, new_children, answer_template, doc
                )
            elif mode == 'after_paragraph':
                insert_count, total_heading_count = self._insert_after_paragraphs(
                    children, new_children, answer_template, doc
                )
            else:
                # 默认使用标题前插入
                insert_count, total_heading_count = self._insert_before_headings(
                    children, new_children, answer_template, doc
                )
            
            # 清空并重组body
            for child in list(body):
                body.remove(child)
            for elem in new_children:
                body.append(elem)
            
            # 清除提示语标记（避免残留到最终文档）
            self._remove_hint_markers(doc)
            
            print(f"插入应答句完成！插入: {insert_count}个，标题: {total_heading_count}个")
            return True
        except Exception as e:
            print(f"插入应答句失败: {e}")
            import traceback
            traceback.print_exc()
            return False


if __name__ == "__main__":
    # 测试代码
    converter = DocumentConverter()
    print("文档转换器模块加载成功")
    print(f"Pillow可用: {PIL_AVAILABLE}")
