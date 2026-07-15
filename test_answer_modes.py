# -*- coding: utf-8 -*-
"""
测试应答句插入功能的4种模式
验证命名修改后的功能完整性
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import sys

# 添加项目路径
sys.path.insert(0, os.path.dirname(__file__))

from doc_converter import DocumentConverter


def create_test_document():
    """创建测试文档"""
    doc = Document()
    
    # 添加一些样式
    styles = doc.styles
    
    # 创建应答句样式
    if '应答句' not in styles:
        answer_style = styles.add_style('应答句', 1)  # 1 = paragraph style
        answer_style.font.size = Pt(12)
        answer_style.font.bold = True
        answer_style.font.color.rgb = None  # 自动颜色
        answer_style.paragraph_format.space_after = Pt(6)
    
    # 标题1
    heading1 = doc.add_heading('第一章 项目概述', level=1)
    
    # 正文段落1
    para1 = doc.add_paragraph('这是一个正文段落，描述项目的基本情况和背景信息。这个段落比较长，用来测试应答句插入功能。')
    
    # 正文段落2
    para2 = doc.add_paragraph('这是第二个正文段落，继续描述项目的详细内容。包括技术方案、实施计划等关键信息。')
    
    # 标题2
    heading2 = doc.add_heading('第二章 技术方案', level=1)
    
    # 正文段落3
    para3 = doc.add_paragraph('技术方案的详细描述，包括系统架构、技术选型等内容。')
    
    # 表格
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    for i in range(2):
        for j in range(3):
            table.cell(i, j).text = f'单元格{i+1},{j+1}'
    
    # 标题3
    heading3 = doc.add_heading('第三章 实施计划', level=1)
    
    # 正文段落4
    para4 = doc.add_paragraph('实施计划的详细说明，包括时间安排、人员配置等。')
    
    # 保存测试文档
    test_file = 'test_answer_modes_source.docx'
    doc.save(test_file)
    print(f"✅ 测试文档已创建: {test_file}")
    return test_file


def test_all_modes(source_file):
    """测试所有5种模式"""
    converter = DocumentConverter()
    
    modes = [
        ('before_heading', '章节前插入'),
        ('after_heading', '章节末插入'),
        ('before_paragraph', '逐段前应答'),
        ('after_paragraph', '逐段后应答'),
        ('copy_chapter', '原文+应答句+应答原文')
    ]
    
    results = []
    
    for mode_key, mode_name in modes:
        output_file = f'test_answer_{mode_key}_result.docx'
        
        print(f"\n{'='*60}")
        print(f"测试模式: {mode_name} ({mode_key})")
        print(f"{'='*60}")
        
        try:
            success, actual_file, message = converter.insert_response_after_headings(
                input_file=source_file,
                output_file=output_file,
                answer_text='应答：本投标人理解并满足要求。',
                answer_style='应答句',
                mode=mode_key
            )
            
            if success:
                print(f"✅ 转换成功: {message}")
                
                # 验证输出文件
                if os.path.exists(output_file):
                    result_doc = Document(output_file)
                    para_count = len(result_doc.paragraphs)
                    table_count = len(result_doc.tables)
                    
                    print(f"   - 段落数: {para_count}")
                    print(f"   - 表格数: {table_count}")
                    print(f"   - 输出文件: {output_file}")
                    
                    results.append({
                        'mode': mode_name,
                        'success': True,
                        'file': output_file,
                        'paragraphs': para_count,
                        'tables': table_count
                    })
                else:
                    print(f"❌ 输出文件不存在: {output_file}")
                    results.append({
                        'mode': mode_name,
                        'success': False,
                        'error': '输出文件不存在'
                    })
            else:
                print(f"❌ 转换失败: {message}")
                results.append({
                    'mode': mode_name,
                    'success': False,
                    'error': message
                })
        
        except Exception as e:
            print(f"❌ 异常: {e}")
            import traceback
            traceback.print_exc()
            results.append({
                'mode': mode_name,
                'success': False,
                'error': str(e)
            })
    
    return results


def print_summary(results):
    """打印测试总结"""
    print(f"\n\n{'='*60}")
    print("测试总结")
    print(f"{'='*60}")
    
    success_count = sum(1 for r in results if r['success'])
    total_count = len(results)
    
    print(f"\n总测试数: {total_count}")
    print(f"成功: {success_count}")
    print(f"失败: {total_count - success_count}")
    
    print(f"\n详细结果:")
    for r in results:
        status = "✅ 成功" if r['success'] else f"❌ 失败: {r.get('error', '未知错误')}"
        print(f"  - {r['mode']}: {status}")
        if r['success']:
            print(f"    文件: {r['file']}")
            print(f"    段落数: {r['paragraphs']}, 表格数: {r['tables']}")
    
    print(f"\n{'='*60}")
    if success_count == total_count:
        print("🎉 所有测试通过！")
    else:
        print(f"⚠️  有 {total_count - success_count} 个测试失败，请检查错误信息")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    print("="*60)
    print("应答句插入模式功能测试")
    print("="*60)
    print("\n测试内容：")
    print("1. 章节前插入 (before_heading)")
    print("2. 章节末插入 (after_heading)")
    print("3. 逐段前应答 (before_paragraph)")
    print("4. 逐段后应答 (after_paragraph)")
    print("5. 原文+应答句+应答原文 (copy_chapter)")
    print("="*60)
    
    # 创建测试文档
    source_file = create_test_document()
    
    # 测试所有模式
    results = test_all_modes(source_file)
    
    # 打印总结
    print_summary(results)
    
    print("\n提示：请手动打开生成的docx文件验证应答句插入位置是否正确")
