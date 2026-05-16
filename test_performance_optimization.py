"""
性能优化验证测试脚本
测试一次性流水线优化的正确性和性能提升
"""
import os
import sys
import time
from docx import Document
from doc_converter import DocumentConverter

# 测试配置
TEST_DIR = "test_performance"
SOURCE_FILE = os.path.join(TEST_DIR, "test_source.docx")
TEMPLATE_FILE = os.path.join(TEST_DIR, "test_template.docx")
OUTPUT_FILE_OLD = os.path.join(TEST_DIR, "output_old.docx")
OUTPUT_FILE_NEW = os.path.join(TEST_DIR, "output_new.docx")


def create_test_documents():
    """创建测试文档"""
    print("📝 创建测试文档...")
    
    if not os.path.exists(TEST_DIR):
        os.makedirs(TEST_DIR)
    
    # 创建源文档（包含标题、正文、表格）
    source_doc = Document()
    
    # 添加标题
    source_doc.add_heading('第一章 项目概述', level=1)
    source_doc.add_paragraph('这是第一段正文内容，需要进行样式转换和语气转换。')
    source_doc.add_paragraph('第二段正文内容，继续测试。')
    
    source_doc.add_heading('1.1 项目背景', level=2)
    source_doc.add_paragraph('项目背景描述，包含多个段落。')
    source_doc.add_paragraph('更多背景信息。')
    
    source_doc.add_heading('1.2 项目目标', level=2)
    source_doc.add_paragraph('项目目标说明。')
    
    # 添加表格
    table = source_doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = '列1'
    table.cell(0, 1).text = '列2'
    table.cell(1, 0).text = '数据1'
    table.cell(1, 1).text = '数据2'
    table.cell(2, 0).text = '数据3'
    table.cell(2, 1).text = '数据4'
    
    source_doc.add_heading('第二章 技术方案', level=1)
    source_doc.add_paragraph('技术方案详细描述。')
    
    source_doc.save(SOURCE_FILE)
    
    # 创建模板文档
    template_doc = Document()
    template_doc.add_heading('模板标题', level=1)
    template_doc.add_paragraph('模板正文')
    template_doc.save(TEMPLATE_FILE)
    
    print(f"✅ 测试文档创建完成")
    print(f"   源文档: {SOURCE_FILE}")
    print(f"   模板文档: {TEMPLATE_FILE}")


def test_old_conversion(converter):
    """测试旧的三阶段转换方式（模拟）"""
    print("\n⏱️  测试旧方式（三阶段流水线）...")
    
    start_time = time.time()
    
    # 模拟旧方式的三次加载/保存
    # 第1阶段：样式转换
    temp_file_1 = os.path.join(TEST_DIR, "temp_stage1.docx")
    converter.convert_styles(SOURCE_FILE, TEMPLATE_FILE, temp_file_1)
    
    # 第2阶段：语气转换
    temp_file_2 = os.path.join(TEST_DIR, "temp_stage2.docx")
    converter.convert_mood(temp_file_1, temp_file_2)
    
    # 第3阶段：插入应答句
    converter.insert_response_after_headings(
        temp_file_2, 
        OUTPUT_FILE_OLD,
        answer_text="【应答】",
        answer_style='Normal',
        mode='before_heading'
    )
    
    # 清理临时文件
    for temp_file in [temp_file_1, temp_file_2]:
        if os.path.exists(temp_file):
            os.remove(temp_file)
    
    elapsed = time.time() - start_time
    print(f"   耗时: {elapsed:.2f}秒")
    
    return elapsed, OUTPUT_FILE_OLD


def test_new_conversion(converter):
    """测试新的一次性流水线方式"""
    print("\n⚡ 测试新方式（一次性流水线）...")
    
    start_time = time.time()
    
    converter.full_convert(
        source_file=SOURCE_FILE,
        template_file=TEMPLATE_FILE,
        output_file=OUTPUT_FILE_NEW,
        custom_style_map=None,
        do_mood=True,
        answer_text="【应答】",
        answer_style='Normal',
        list_bullet=None,
        do_answer_insertion=True,
        answer_mode='before_heading',
        progress_callback=None,
        warning_callback=None,
        source_styles_cache=None
    )
    
    elapsed = time.time() - start_time
    print(f"   耗时: {elapsed:.2f}秒")
    
    return elapsed, OUTPUT_FILE_NEW


def verify_output_consistency(old_file, new_file):
    """验证两种方式的输出一致性"""
    print("\n🔍 验证输出一致性...")
    
    old_doc = Document(old_file)
    new_doc = Document(new_file)
    
    # 比较段落数
    old_para_count = len(old_doc.paragraphs)
    new_para_count = len(new_doc.paragraphs)
    
    print(f"   旧方式段落数: {old_para_count}")
    print(f"   新方式段落数: {new_para_count}")
    
    if old_para_count != new_para_count:
        print(f"   ⚠️  警告：段落数不一致！")
        return False
    
    # 比较表格数
    old_table_count = len(old_doc.tables)
    new_table_count = len(new_doc.tables)
    
    print(f"   旧方式表格数: {old_table_count}")
    print(f"   新方式表格数: {new_table_count}")
    
    if old_table_count != new_table_count:
        print(f"   ⚠️  警告：表格数不一致！")
        return False
    
    # 比较标题数
    old_headings = [p for p in old_doc.paragraphs if p.style.name.startswith('Heading')]
    new_headings = [p for p in new_doc.paragraphs if p.style.name.startswith('Heading')]
    
    print(f"   旧方式标题数: {len(old_headings)}")
    print(f"   新方式标题数: {len(new_headings)}")
    
    if len(old_headings) != len(new_headings):
        print(f"   ⚠️  警告：标题数不一致！")
        return False
    
    print(f"   ✅ 输出一致性验证通过")
    return True


def main():
    """主测试函数"""
    print("=" * 60)
    print("性能优化验证测试")
    print("=" * 60)
    
    # 步骤1：创建测试文档
    create_test_documents()
    
    # 步骤2：初始化转换器
    print("\n🔧 初始化转换器...")
    converter = DocumentConverter()
    
    # 步骤3：测试旧方式
    try:
        old_time, old_file = test_old_conversion(converter)
    except Exception as e:
        print(f"❌ 旧方式测试失败: {e}")
        import traceback
        traceback.print_exc()
        old_time = None
        old_file = None
    
    # 步骤4：测试新方式
    try:
        new_time, new_file = test_new_conversion(converter)
    except Exception as e:
        print(f"❌ 新方式测试失败: {e}")
        import traceback
        traceback.print_exc()
        new_time = None
        new_file = None
    
    # 步骤5：比较性能
    print("\n" + "=" * 60)
    print("性能对比结果")
    print("=" * 60)
    
    if old_time and new_time:
        improvement = ((old_time - new_time) / old_time) * 100
        print(f"旧方式耗时: {old_time:.2f}秒")
        print(f"新方式耗时: {new_time:.2f}秒")
        print(f"性能提升: {improvement:.1f}%")
        print(f"时间节省: {old_time - new_time:.2f}秒")
        
        if improvement > 30:
            print(f"✅ 性能优化显著（提升超过30%）")
        elif improvement > 10:
            print(f"⚠️  性能有提升，但未达到预期目标（60%）")
        else:
            print(f"❌ 性能提升不明显")
    
    # 步骤6：验证输出一致性
    if old_file and new_file and os.path.exists(old_file) and os.path.exists(new_file):
        consistency_ok = verify_output_consistency(old_file, new_file)
        
        if consistency_ok:
            print("\n✅ 测试通过：新方式与旧方式输出一致")
        else:
            print("\n❌ 测试失败：新方式与旧方式输出不一致")
    else:
        print("\n⚠️  无法验证输出一致性（文件不存在）")
    
    print("\n" + "=" * 60)
    print("测试完成")
    print("=" * 60)


if __name__ == "__main__":
    main()
