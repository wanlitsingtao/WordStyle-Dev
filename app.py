# -*- coding: utf-8 -*-
"""
文档转换工具 - Web 版本 (MVP)
基于 Streamlit 快速搭建
"""
import streamlit as st

# [WARN] set_page_config必须在所有Streamlit命令之前调用
# [OK] 支持UptimeRobot健康检查：通过URL参数检测
import sys
from urllib.parse import urlparse, parse_qs
try:
    # 获取当前URL参数
    query_params = st.query_params
    if 'health' in query_params:
        # 返回health检查响应
        import json
        st.json({"status": "healthy", "service": "user-page", "version": "1.0.0"})
        st.stop()
except Exception:
    # 如果query_params不可用（旧版本Streamlit），忽略
    pass
st.set_page_config(
    page_title="标书抄写神器",
    page_icon="📄",
    layout="wide",  # 使用宽屏布局
    initial_sidebar_state="expanded"
)

import os
import sys
import json
import time
import threading
import logging
from datetime import datetime, timedelta
from pathlib import Path
from contextlib import contextmanager  # 添加缺失的导入

# ==================== 组件导入 ====================
from components.dialogs.feedback import show_feedback_dialog
from components.dialogs.history import show_history_dialog
from components.dialogs.style_mapping import show_style_mapping_dialog
from state import app_state  # 统一状态管理器
from components.config_panel import render_conversion_config
from components.upload import count_paragraphs, get_template_styles_list, analyze_source_styles, count_pages

# 配置日志系统
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger('WordStyle')

# [OK] 服务级单例：Streamlit rerun 时复用清理任务和后台线程
@st.cache_resource
def _start_file_cleanup_scheduler():
    try:
        from file_manager import cleanup_on_startup, schedule_daily_cleanup
        cleanup_on_startup()

        from apscheduler.schedulers.background import BackgroundScheduler
        scheduler = BackgroundScheduler()
        scheduler.add_job(
            func=schedule_daily_cleanup,
            trigger='cron',
            hour=0,
            minute=0,
            id='daily_file_cleanup',
            name='每日文件清理任务',
            replace_existing=True
        )
        scheduler.start()
        logger.info("[OK] 每日文件清理任务已启动（每天零点执行）")
        return scheduler
    except ImportError:
        logger.warning("[WARN] APScheduler未安装，跳过定时清理任务")
    except Exception as e:
        logger.warning(f"定时清理任务启动失败: {e}")
    return None


_start_file_cleanup_scheduler()

# 添加当前目录到路径，以便导入其他模块
sys.path.insert(0, os.path.dirname(__file__))

# 导入配置
from config import (
    DEFAULT_ANSWER_TEXT, DEFAULT_ANSWER_STYLE, DEFAULT_ANSWER_MODE,
    ANSWER_MODE_OPTIONS, DEFAULT_LIST_BULLET, PAGE_TITLE, PAGE_ICON,
    LAYOUT, SIDEBAR_STATE, FREE_PARAGRAPHS_DAILY, DATA_SOURCE  # [OK] 修复：添加DATA_SOURCE导入
)


# 导入工具函数
from utils import (
    sanitize_html, sanitize_filename, validate_docx_file, convert_server_time_to_local
)

# 导入用户管理
from data_manager import (
    load_user_data, save_user_data, claim_free_paragraphs, register_or_login_user,
    get_config
)


@st.cache_data(ttl=30, show_spinner=False)
def _get_cached_config(key):
    """短期缓存不随页面 rerun 变化的系统配置。"""
    return get_config(key)


@st.cache_data(show_spinner=False)
def _get_image_base64(image_path):
    with open(image_path, 'rb') as image_file:
        import base64
        return base64.b64encode(image_file.read()).decode()

# 导入评论管理
from comments_manager import (
    load_comments, save_comments, add_comment, like_comment,
    get_comments, get_comment_stats, validate_comment_content,
    add_feedback, get_feedbacks, get_feedback_stats
)

# 导入临时文件清理模块
# from temp_file_cleanup import cleanup_on_startup  # [WARN] 已移动到archive目录

# 导入转换器
from doc_converter import DocumentConverter

# [OK] 所有配置已从 config.py 和 utils.py 导入，不再重复定义
# 参见：config.py, utils.py, user_manager.py, comments_manager.py
# ==================== 初始化会话状态 ====================
# [OK] 基于设备指纹的用户识别系统
# 设计原则：简单、可靠、99.99%成功率

import hashlib
from data_manager import generate_device_fingerprint, get_or_create_user_by_device

# 标记：用户初始化是否成功
user_init_success = False

try:
    # 第一步：获取客户端User-Agent并生成设备指纹
    try:
        headers = st.context.headers if hasattr(st, 'context') and hasattr(st.context, 'headers') else {}
        user_agent = headers.get('User-Agent', 'unknown')
        device_fingerprint = generate_device_fingerprint(user_agent)
        logger.info(f"设备指纹生成成功: {device_fingerprint[:16]}...")
    except Exception as e:
        logger.warning(f"[WARN] User-Agent获取失败，使用备用方案: {e}")
        device_fingerprint = generate_device_fingerprint(f"fallback_{id(st.session_state)}")
    
    # 第二步：通过设备指纹从数据库获取或创建用户
    # 但是如果用户已通过账号登录，则保持登录身份不变
    _logged_in_uid = st.session_state.get('logged_in_user_id', None)
    _user_cache_key = f"account:{_logged_in_uid}" if _logged_in_uid else f"device:{device_fingerprint}"
    _cached_user_data = st.session_state.get('user_data')
    _use_cached_user_data = (
        _cached_user_data is not None
        and st.session_state.get('user_data_cache_key') == _user_cache_key
        and not st.session_state.get('user_data_stale', False)
    )
    if _logged_in_uid:
        # 已登录：使用账号身份，加载账号对应的用户数据（而非设备指纹数据）
        app_state.set_user_id(_logged_in_uid)
        app_state.set_device_fingerprint(device_fingerprint)
        user_data = _cached_user_data if _use_cached_user_data else load_user_data(_logged_in_uid)
        if not user_data:
            # 降级：如果账号数据加载失败，回退到设备指纹
            user_data = get_or_create_user_by_device(device_fingerprint, user_agent)
        logger.info(f"[OK] 已登录用户 - ID: {_logged_in_uid}, 额度: {user_data.get('paragraphs_remaining', 0)}")
    else:
        user_data = _cached_user_data if _use_cached_user_data else get_or_create_user_by_device(device_fingerprint, user_agent)
        app_state.set_user_id(user_data['user_id'])
        app_state.set_device_fingerprint(device_fingerprint)
        logger.info(f"[OK] 用户初始化成功 - ID: {app_state.get_user_id()}")
    user_init_success = True
    st.session_state.user_data = user_data
    st.session_state.user_data_cache_key = _user_cache_key
    st.session_state.user_data_stale = False
    
except Exception as e:
    logger.error(f"[ERROR] 用户初始化失败: {e}", exc_info=True)
    
    # 最终降级方案：生成一个本地可用的临时ID
    try:
        fallback_id = hashlib.md5(f"temp_{id(st.session_state)}_{datetime.now().timestamp()}".encode()).hexdigest()[:12]
    except:
        fallback_id = f"temp_error_{id(st.session_state)}"
    
    app_state.set_user_id(fallback_id)
    app_state.set_device_fingerprint(None)
    app_state.set_user_init_failed(True)  # 标记初始化失败
    
    user_data = {
        'user_id': fallback_id,
        'balance': 0.0,
        'paragraphs_remaining': 0,  # [WARN] 失败时额度为0
        'total_paragraphs_used': 0,
        'total_converted': 0,
        'is_active': False,
        'created_at': datetime.now().isoformat(),
        'last_login': datetime.now().isoformat(),
        'conversion_history': [],  # [OK] 添加转换历史字段
    }
    logger.warning(f"[WARN] 使用临时用户ID（无额度）: {fallback_id}")

# 第三步：只有在初始化成功时才尝试领取免费额度（仅首次加载时执行）
if user_init_success and 'free_claimed_today' not in st.session_state:
    try:
        free_paragraphs = claim_free_paragraphs(app_state.get_user_id())
        if free_paragraphs > 0:
            st.toast(f"🎉 欢迎！今日免费额度已重置为 {free_paragraphs:,} 段", icon="🎁")
            user_data['paragraphs_remaining'] = free_paragraphs
            logger.info(f"[OK] 免费额度领取成功: {free_paragraphs}")
            # 标记今日已领取，避免重复显示toast
            app_state.set_free_claimed_today(True)
        else:
            logger.info(f"ℹ️ 无需领取额度或已领取过，当前额度: {user_data.get('paragraphs_remaining', 0)}")
            # 即使没有新领取，也标记已检查过
            app_state.set_free_claimed_today(True)
    except Exception as e:
        logger.warning(f"[WARN] 领取免费额度失败: {e}，但不影响用户使用")
        app_state.set_free_claimed_today(True)
else:
    if not user_init_success:
        logger.warning("[WARN] 用户初始化失败，跳过额度领取")

st.session_state.user_data = user_data
logger.info(f"用户 {app_state.get_user_id()} 初始化完成，剩余额度: {user_data['paragraphs_remaining']}")


# 新手引导标志
if 'has_seen_guide' not in st.session_state:
    app_state.set_has_seen_guide(False)

# 每日免费额度机制，不再需要 free_paragraphs_claimed 标记
# if 'free_paragraphs_claimed' not in st.session_state:
#     user_data = load_user_data()
#     has_used = (...)
#     app_state.set_free_paragraphs_claimed(has_used)

# ==================== 评论区功能 ====================

COMMENTS_FILE = Path("comments_data.json")

def load_comments():
    """加载评论数据（优先从API获取）"""
    # [OK] 修复：使用 API 加载评论（兼容多实例部署）
    from config import BACKEND_URL
    
    if BACKEND_URL and DATA_SOURCE == 'api':
        # API 模式：通过后端 API 获取
        try:
            import requests
            api_url = f"{BACKEND_URL.rstrip('/')}/api/comments/comments/list?limit=100"  # 修复：后端路由有双重 /comments
            response = requests.get(api_url, timeout=10)
            response.raise_for_status()
            comments = response.json()
            # 转换UUID为字符串，保持兼容性
            for c in comments:
                if isinstance(c.get('id'), str) and len(c['id']) > 20:
                    # UUID格式，截取前8位作为显示ID
                    c['display_id'] = c['id'][:8]
            return comments
        except Exception as e:
            logger.error(f"[ERROR] API加载评论失败: {e}，降级到本地文件")
            # 降级到本地文件
    
    # 本地/Supabase 模式：使用本地文件（兜底逻辑）
    if COMMENTS_FILE.exists():
        with open(COMMENTS_FILE, 'r', encoding='utf-8') as f:
            try:
                return json.load(f)
            except:
                return []
    return []

def save_comments(comments):
    """保存评论数据"""
    with open(COMMENTS_FILE, 'w', encoding='utf-8') as f:
        json.dump(comments, f, ensure_ascii=False, indent=2)

def add_comment(username, content, rating=5):
    """添加新评论（使用API提交到数据库）"""
    # [OK] 修复：使用 API 提交评论（兼容多实例部署）
    from config import BACKEND_URL  # 添加BACKEND_URL的局部导入
    import requests  # [OK] 修复：在函数开头导入requests，确保except子句可用
    
    if BACKEND_URL and DATA_SOURCE == 'api':
        # API 模式：通过后端 API 提交
        try:
            api_url = f"{BACKEND_URL.rstrip('/')}/api/comments/comments/submit"  # 修复：后端路由有双重 /comments
            logger.info(f"[INFO] 尝试调用API: {api_url}")
            response = requests.post(
                api_url,
                json={
                    'username': username or f'用户{app_state.get_user_id()[:6]}',
                    'content': content,
                    'rating': rating,
                    'user_id': app_state.get_user_id()
                },
                timeout=10
            )
            logger.info(f"[INFO] API响应状态码: {response.status_code}")
            
            # [OK] 修复：检查HTTP状态码，确保数据库写入成功
            if response.status_code == 200:
                result = response.json()
                logger.info(f"[INFO] API返回结果: {result}")
                
                # [OK] 修复：API成功后同步写入本地文件，确保数据一致性
                new_comment = {
                    'id': result.get('id'),
                    'username': result.get('username'),
                    'content': result.get('content'),
                    'rating': result.get('rating'),
                    'timestamp': result.get('timestamp'),
                    'likes': result.get('likes', 0),
                    'user_id': result.get('user_id')
                }
                
                # 同步到本地文件（作为缓存和降级备份）
                comments = load_comments()
                comments.append(new_comment)
                save_comments(comments)
                
                logger.info(f"[SUCCESS] 评论已成功写入数据库并同步到本地")
                return new_comment
            else:
                # HTTP状态码不是200，说明写入失败
                error_detail = response.json().get('detail', '未知错误')
                logger.error(f"[ERROR] API返回错误状态码 {response.status_code}: {error_detail}")
                raise Exception(f"数据库写入失败: {error_detail}")
                
        except requests.exceptions.Timeout:
            logger.error(f"[ERROR] API请求超时（10秒）")
            # 降级到本地存储
        except requests.exceptions.ConnectionError:
            logger.error(f"[ERROR] 无法连接到后端API服务器")
            # 降级到本地存储
        except Exception as e:
            logger.error(f"[ERROR] API提交评论失败: {e}，降级到本地存储")
            # 降级到本地存储，继续执行下面的本地存储逻辑
    
    # 本地/Supabase 模式：使用本地存储（兜底逻辑）
    logger.info(f"[INFO] 使用本地存储模式保存评论")
    comments = load_comments()
    
    new_comment = {
        'id': len(comments) + 1,
        'username': username or f'用户{app_state.get_user_id()[:6]}',
        'content': content,
        'rating': rating,
        'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'likes': 0,
        'user_id': app_state.get_user_id()
    }
    
    comments.append(new_comment)
    save_comments(comments)
    logger.info(f"[SUCCESS] 评论已保存到本地文件")
    return new_comment

def like_comment(comment_id):
    """点赞评论"""
    comments = load_comments()
    for comment in comments:
        if comment['id'] == comment_id:
            comment['likes'] += 1
            break
    save_comments(comments)

@st.fragment
def show_comments_section():
    """显示评论区"""
    # 检查是否需要刷新评论
    if app_state.get_comment_refresh_needed():
        app_state.set_comment_refresh_needed(False)  # 清除标记
        st.rerun()  # 刷新fragment以加载最新评论
    
    # 会话内复用评论列表，发表评论或点赞后由上方逻辑主动失效
    comments = st.session_state.get('comments_cache')
    if comments is None:
        comments = load_comments()
        st.session_state.comments_cache = comments
    
    # 显示统计信息
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("总评论数", len(comments))
    with col2:
        if comments:
            avg_rating = sum(c.get('rating', 5) for c in comments) / len(comments)
            st.metric("平均评分", f"{avg_rating:.1f} ⭐")
        else:
            st.metric("平均评分", "暂无")
    with col3:
        total_likes = sum(c.get('likes', 0) for c in comments)
        st.metric("总点赞数", total_likes)
    
    st.markdown("---")
    
    # 发表评论表单
    with st.expander("✍️ 发表评论", expanded=False):
        with st.form("comment_form", clear_on_submit=True):  # 修复：添加clear_on_submit清空表单
            rating = st.slider("评分", 1, 5, 5, help="请为工具打分")
            
            content = st.text_area(
                "评论内容",
                placeholder="分享您的使用体验、建议或问题...",
                height=100,
                max_chars=500
            )
            
            col_submit, col_cancel = st.columns([1, 3])
            with col_submit:
                submit_comment = st.form_submit_button("📤 发表", type="primary")
            
            if submit_comment:
                if not content.strip():
                    st.error("❌ 请输入评论内容")
                else:
                    # 显示加载状态
                    with st.spinner('正在提交评论...'):
                        new_comment = add_comment(None, content, rating)  # 匿名评论
                    
                    if new_comment:
                        st.success("✅ 评论发表成功！")
                        st.session_state.comments_cache = None
                        # [OK] 优化：设置标记，告诉侧边栏使用缓存数据
                        st.session_state.comment_refresh_only = True
                        # 使用session_state标记，通知fragment刷新
                        app_state.set_comment_refresh_needed(True)
                        # [OK] 修复：立即触发fragment刷新，确保评论立即显示
                        st.rerun()
                    else:
                        st.error("❌ 评论发表失败，请稍后重试")
    
    st.markdown("---")
    
    # 显示评论列表
    if not comments:
        st.info("💭 暂无评论，快来发表第一条评论吧！")
    else:
        # 按时间倒序显示
        comments_sorted = sorted(comments, key=lambda x: x['timestamp'], reverse=True)
        
        for comment in comments_sorted[:20]:  # 最多显示20条
            with st.container():
                col_header, col_like = st.columns([4, 1])
                
                with col_header:
                    # 显示评分（使用emoji星号）和时间
                    rating = comment.get('rating', 5)
                    stars = "⭐" * rating
                    st.markdown(f"{stars}")
                    st.caption(f"🕒 {comment.get('timestamp', '')}")
                
                with col_like:
                    # 点赞按钮
                    likes = comment.get('likes', 0)
                    if st.button(f"👍 {likes}", key=f"like_{comment['id']}"):
                        like_comment(comment['id'])
                        st.session_state.comments_cache = None
                        # Streamlit会在按钮点击后自动重新运行脚本
                
                # 显示评论内容
                st.markdown(f"<div style='padding: 10px; background-color: #f0f2f6; border-radius: 5px; margin: 5px 0;'>{sanitize_html(comment.get('content', ''))}</div>", unsafe_allow_html=True)
                
                st.markdown("---")
        
        if len(comments) > 20:
            st.caption(f"显示最近20条评论，共 {len(comments)} 条")



# ==================== 页面配置 ====================
# ==================== 应用启动时清理临时文件 ====================
# [WARN] 已禁用：cleanup_on_startup 函数已被移除
# 临时文件清理功能已整合到其他模块中

# ==================== 维护模式检查（每次页面渲染时执行）====================
try:
    maintenance_mode = _get_cached_config('maintenance_mode')
    
    # 处理多种可能的值类型：字符串'true'、布尔值True、字符串'1'等
    is_maintenance = False
    if maintenance_mode is not None:
        if isinstance(maintenance_mode, bool):
            is_maintenance = maintenance_mode
        elif isinstance(maintenance_mode, str):
            is_maintenance = maintenance_mode.lower() in ('true', '1', 'yes', 'on')
        else:
            is_maintenance = bool(maintenance_mode)
    
    if is_maintenance:
        # 如果启用维护模式，显示维护页面并停止执行
        # 获取当前脚本所在目录
        script_dir = Path(__file__).parent
        logo_path = script_dir / "resource" / "wh.jpg"
        
        # 自定义CSS - 黑色背景，移除所有空白
        st.markdown("""
<style>
    /* 全局样式 */
    .stApp {
        background-color: #000000 !important;
        margin: 0 !important;
        padding: 0 !important;
    }
    
    /* 隐藏默认Streamlit元素 */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    [data-testid="stHeader"] {display: none !important;}
    [data-testid="stToolbar"] {display: none !important;}
    
    /* 移除所有默认内边距 */
    .block-container {
        padding-top: 0 !important;
        padding-bottom: 0 !important;
        padding-left: 0 !important;
        padding-right: 0 !important;
        margin-top: 0 !important;
        margin-bottom: 0 !important;
        max-width: 100% !important;
    }
    
    /* main容器 */
    main {
        padding-top: 0 !important;
        margin-top: 0 !important;
        display: block !important;
        visibility: visible !important;
    }
    
    /* 确保内容可见 */
    div[data-testid="stVerticalBlock"] {
        display: block !important;
        visibility: visible !important;
    }
    
    /* 呼吸动画文字 */
    @keyframes breathe {
        0%, 100% { opacity: 1; transform: scale(1); }
        50% { opacity: 0.6; transform: scale(1.05); }
    }
    
    .breathe-text {
        animation: breathe 3s ease-in-out infinite;
        color: #ffffff !important;
        font-size: 2rem;
        font-weight: bold;
        text-align: center;
        margin-top: 2rem;
        text-shadow: 0 0 20px rgba(255, 255, 255, 0.8);
        display: block !important;
        visibility: visible !important;
    }
</style>
""", unsafe_allow_html=True)
        
        # 显示Logo图片（铺满全屏，完整显示不裁剪）
        if logo_path.exists():
            try:
                encoded_image = _get_image_base64(str(logo_path))
                
                st.markdown(f'''
                <div style="width: 100vw; height: 100vh; margin: 0; padding: 0; position: fixed; top: 0; left: 0; z-index: 0; overflow: hidden;">
                    <img src="data:image/jpeg;base64,{encoded_image}" 
                         style="position: absolute; top: 0; left: 50%; transform: translateX(-50%); height: 100vh; width: auto; min-width: 100vw; display: block;">
                </div>
                ''', unsafe_allow_html=True)
            except Exception as e:
                logger.error(f"[维护模式] Logo图片加载失败: {e}")
        
        # 显示呼吸文字（图片下方底部）
        st.markdown('''
<div class="breathe-text" style="position: fixed; bottom: 5vh; left: 0; right: 0; text-align: center; z-index: 1;">
    我会回来的！
</div>
''', unsafe_allow_html=True)
        
        st.stop()  # 停止执行后续代码
except Exception as e:
    logger.warning(f"维护模式检查失败（不影响服务）: {e}")

# ==================== 主界面 ====================
# 使用 resource/logo.png 替换 emoji 图标
_logo_path = Path(__file__).parent / "resource" / "logo.png"
if _logo_path.exists():
    _logo_b64 = _get_image_base64(str(_logo_path))
    st.markdown(f'''
    <div style="display: flex; align-items: center; gap: 12px; margin-bottom: 0.5rem;">
        <img src="data:image/png;base64,{_logo_b64}" 
             style="height: 2.5rem; vertical-align: middle;">
        <span style="font-size: 2rem; font-weight: 700; color: #262730;">标书抄写神器</span>
    </div>
    ''', unsafe_allow_html=True)
else:
    st.title("📄 标书抄写神器")

# 全屏提示
st.markdown("""
<div style='background-color: #e3f2fd; padding: 10px; border-radius: 5px; margin-bottom: 10px;'>
💡 <strong>提示：</strong>按 <kbd>F11</kbd> 键可以让浏览器全屏显示，获得更好的体验
</div>
""", unsafe_allow_html=True)

# 说明信息
st.markdown("""
<div style='background-color: #fff3cd; padding: 10px; border-radius: 5px; margin-bottom: 10px;'>
 <strong>说明：</strong>文档转换完成后，请及时下载，源文档和结果文档会被自动清理。自行做好标书检查，Good luck！
</div>
""", unsafe_allow_html=True)

# 自定义CSS，优化页面显示（简化版，让Streamlit自动处理布局）
st.markdown("""
<style>
    /* 隐藏页脚 */
    footer {visibility: hidden;}
    
    /* 强制主要内容区域使用最大宽度 */
    .block-container {
        max-width: 100% !important;
        padding-top: 4rem !important;  /* 增加顶部间距，避免被Streamlit Cloud工具栏遮挡 */
        padding-bottom: 1rem !important;
    }
    
    /* 优化文件上传器大小 */
    .stFileUploader > div {
        min-height: 80px;
    }
    
    /* 增大按钮 */
    .stButton > button {
        height: 3em;
        font-size: 1.1em;
        width: 100%;
    }
    
    /* 优化指标显示 */
    [data-testid="stMetric"] {
        background-color: #f0f2f6;
        padding: 15px;
        border-radius: 5px;
    }
    
    /* 修复侧边栏隐藏后的布局问题 */
    div[data-testid="stAppViewContainer"] {
        width: 100% !important;
    }
    
    /* 确保主内容区域的父容器正确响应侧边栏变化 */
    section[data-testid="stSidebar"] + div {
        flex-grow: 1 !important;
        width: auto !important;
    }
    
    /* 转换历史对话框 - 设置较大的默认尺寸 */
    [data-testid="stDialog"]:has([data-testid="stMarkdownContainer"] h2:first-child),
    div[role="dialog"] {
        min-width: 900px !important;
        min-height: 600px !important;
        max-width: 95vw !important;
        max-height: 90vh !important;
    }
</style>

<script>
// 监听侧边栏按钮点击并强制重新布局
setTimeout(function() {
    // 查找侧边栏切换按钮
    const toggleButtons = document.querySelectorAll('button[title*="sidebar"], button[aria-label*="sidebar"], [data-testid="stSidebarCollapsedControl"]');
    
    toggleButtons.forEach(function(btn) {
        btn.addEventListener('click', function() {
            // 延迟执行以确保DOM已更新
            setTimeout(function() {
                // 触发窗口resize事件
                window.dispatchEvent(new Event('resize'));
                
                // 强制重新计算布局
                const mainContainer = document.querySelector('.main');
                if (mainContainer) {
                    mainContainer.style.display = 'none';
                    setTimeout(function() {
                        mainContainer.style.display = '';
                    }, 10);
                }
            }, 300);
        });
    });
}, 2000);
</script>
""", unsafe_allow_html=True)

# 侧边栏：用户信息
with st.sidebar:
    st.header("👤 用户信息")
    
    # 🔍 调试信息：显示当前user_id
    # [OK] 显示用户ID或错误提示；登录后显示用户名
    _logged_in_name = st.session_state.get('logged_in_username', None)
    if st.session_state.get('user_init_failed', False):
        st.error("❌ 获取用户ID失败")
        st.caption("用户服务暂时不可用，请稍后刷新页面重试")
    else:
        if _logged_in_name:
            st.caption(f"👤 {_logged_in_name}")
        else:
            st.caption(f"用户ID: {app_state.get_user_id()[:12]}...")
    
    # [OK] 只有初始化成功才从 API 加载数据
    if not st.session_state.get('user_init_failed', False):
        # 复用页面初始化阶段加载的数据，写操作通过 user_data_stale 触发刷新
        user_data = st.session_state.get('user_data', user_data)
        st.session_state.sidebar_user_data = user_data
    else:
        # 初始化失败：使用本地默认数据（额度为0）
        user_data = {
            'user_id': app_state.get_user_id(),
            'balance': 0.0,
            'paragraphs_remaining': 0,  # [WARN] 失败时额度为0
            'total_paragraphs_used': 0,
            'total_converted': 0,
            'is_active': False,
            'created_at': '',
            'last_login': '',
        }
        logger.warning(f"[WARN] 用户初始化失败，使用本地默认数据（额度=0）")
    
    # 🔧 容错处理：如果用户数据为空，尝试重新初始化
    if user_data is None:
        logger.warning(f"[WARN] 用户数据加载失败: {app_state.get_user_id()}，尝试重新初始化")
        try:
            # 通过设备指纹重新获取用户
            device_fingerprint = st.session_state.get('device_fingerprint', '')
            if device_fingerprint:
                from data_manager import get_or_create_user_by_device
                user_data = get_or_create_user_by_device(device_fingerprint)
                app_state.set_user_id(user_data['user_id'])
                logger.info(f"[OK] 重新初始化用户成功: {app_state.get_user_id()}")
            else:
                # 降级方案：创建临时用户数据
                user_data = {
                    'user_id': app_state.get_user_id(),
                    'balance': 0.0,
                    'paragraphs_remaining': 0,
                    'total_paragraphs_used': 0,
                    'total_converted': 0,
                    'is_active': False,
                    'created_at': '',
                    'last_login': '',
                    'conversion_history': [],  # [OK] 添加转换历史字段
                }
                logger.warning(f"[WARN] 使用临时用户数据")
        except Exception as e:
            logger.error(f"[ERROR] 重新初始化用户失败: {e}")
            user_data = {
                'user_id': app_state.get_user_id(),
                'balance': 0.0,
                'paragraphs_remaining': 0,
                'total_paragraphs_used': 0,
                'total_converted': 0,
                'is_active': False,
                'created_at': '',
                'last_login': '',
                'conversion_history': [],  # [OK] 添加转换历史字段
            }
    
    # 显示段落数和统计信息
    st.metric("剩余段落数", f"{user_data['paragraphs_remaining']:,}")
    st.metric("累计转换文档", user_data['total_converted'])
    
    # [OK] 暂时隐藏查看转换历史按钮
    # if st.button("📋 查看转换历史", use_container_width=True, key="view_history_btn"):
    #     show_history_dialog()
    
    # 管理后台入口（隐藏链接，通过URL访问）
    # st.markdown("[[TOOL] 管理后台](/?page=admin)")
    
    # ==================== 账号绑定/登录 ====================

    from account_manager import create_account_manager

    # 检查当前设备是否已绑定账号
    _device_fp = st.session_state.get('device_fingerprint', '')
    _logged_in_user = st.session_state.get('logged_in_username', None)
    _logged_in_uid = st.session_state.get('logged_in_user_id', None)

    if _logged_in_user:
        # 已登录状态：显示用户名 + 解绑 + 退出按钮
        st.markdown(f"👤 **{_logged_in_user}**")

        # 解绑确认状态
        if st.session_state.get('show_unbind_confirm', False):
            st.warning("确定要解绑账号吗？解绑后用户名和密码将被清除，恢复设备指纹身份。")
            col_ub1, col_ub2 = st.columns(2)
            with col_ub1:
                if st.button("✅ 确认解绑", key="confirm_unbind_btn", use_container_width=True):
                    mgr = create_account_manager()
                    success, msg = mgr.unbind_account(_device_fp)
                    if success:
                        st.session_state.logged_in_username = None
                        st.session_state.logged_in_user_id = None
                        st.session_state.sidebar_user_data = None
                        st.session_state.user_data_stale = True
                        st.session_state.show_unbind_confirm = False
                        st.success(msg + " 已恢复设备指纹身份。")
                        st.rerun()
                    else:
                        st.error(msg)
            with col_ub2:
                if st.button("取消", key="cancel_unbind_btn", use_container_width=True):
                    st.session_state.show_unbind_confirm = False
                    st.rerun()
        else:
            if st.button("🔓 解绑用户", key="unbind_account_btn", use_container_width=True):
                st.session_state.show_unbind_confirm = True
                st.rerun()

        if st.button("🚪 退出登录", key="logout_btn", use_container_width=True):
            # 退出：恢复设备指纹身份
            st.session_state.logged_in_username = None
            st.session_state.logged_in_user_id = None
            st.session_state.sidebar_user_data = None  # 强制刷新用户数据
            st.session_state.user_data_stale = True
            st.rerun()
    else:
        # 未登录状态：显示绑定/登录按钮
        col_a, col_b = st.columns(2)
        with col_a:
            if st.button("🔗 绑定账号", key="bind_account_btn", use_container_width=True):
                st.session_state.show_bind_dialog = True
                st.session_state.show_login_dialog = False

        with col_b:
            if st.button("🔑 账号登录", key="login_account_btn", use_container_width=True):
                st.session_state.show_login_dialog = True
                st.session_state.show_bind_dialog = False

    # ==================== 绑定账号对话框 ====================
    if st.session_state.get('show_bind_dialog', False):
        with st.expander("🔗 绑定账号", expanded=True):
            st.markdown("将当前设备与一个易记的用户名绑定，方便以后跨设备登录。")
            # 使用 st.form 保证所有输入原子提交，避免 password 字段值丢失
            with st.form("bind_account_form"):
                bind_username = st.text_input(
                    "设置用户名",
                    placeholder="字母、数字或中文，不区分大小写",
                    key="bind_username_input"
                )
                bind_password = st.text_input(
                    "设置密码",
                    type="password",
                    placeholder="设置登录密码",
                    key="bind_password_input"
                )
                bind_password2 = st.text_input(
                    "确认密码",
                    type="password",
                    placeholder="再次输入密码",
                    key="bind_password2_input"
                )

                col_b1, col_b2 = st.columns(2)
                with col_b1:
                    submitted_bind = st.form_submit_button("✅ 确定", use_container_width=True)
                with col_b2:
                    cancelled_bind = st.form_submit_button("取消", use_container_width=True)

            if cancelled_bind:
                st.session_state.show_bind_dialog = False
                # 清理 form 中的 key 避免残留
                for k in ('bind_username_input', 'bind_password_input', 'bind_password2_input'):
                    st.session_state.pop(k, None)
                st.rerun()

            if submitted_bind:
                if bind_password != bind_password2:
                    st.error("两次输入的密码不一致")
                elif not bind_username or not bind_password:
                    st.error("用户名和密码不能为空")
                else:
                    mgr = create_account_manager()
                    success, msg = mgr.bind_account(
                        _device_fp, bind_username, bind_password
                    )
                    if success:
                        # 绑定成功 → 自动登录，侧边栏显示用户名
                        st.session_state.logged_in_username = bind_username.strip()
                        st.session_state.logged_in_user_id = app_state.get_user_id()
                        st.session_state.sidebar_user_data = None
                        st.session_state.user_data_stale = True
                        st.session_state.show_bind_dialog = False
                        for k in ('bind_username_input', 'bind_password_input', 'bind_password2_input'):
                            st.session_state.pop(k, None)
                        st.success(msg + " 已自动登录。")
                        st.rerun()
                    else:
                        st.error(msg)

    # ==================== 账号登录对话框 ====================
    if st.session_state.get('show_login_dialog', False):
        with st.expander("🔑 账号登录", expanded=True):
            st.markdown("使用已绑定的用户名和密码登录。")
            with st.form("login_account_form"):
                login_username = st.text_input(
                    "用户名",
                    placeholder="输入已绑定的用户名",
                    key="login_username_input"
                )
                login_password = st.text_input(
                    "密码",
                    type="password",
                    placeholder="输入登录密码",
                    key="login_password_input"
                )

                col_l1, col_l2 = st.columns(2)
                with col_l1:
                    submitted_login = st.form_submit_button("✅ 登录", use_container_width=True)
                with col_l2:
                    cancelled_login = st.form_submit_button("取消", use_container_width=True)

            if cancelled_login:
                st.session_state.show_login_dialog = False
                for k in ('login_username_input', 'login_password_input'):
                    st.session_state.pop(k, None)
                st.rerun()

            if submitted_login:
                if not login_username or not login_password:
                    st.error("用户名和密码不能为空")
                else:
                    mgr = create_account_manager()
                    success, msg, user_id = mgr.login_account(login_username, login_password)
                    if success and user_id:
                        st.session_state.logged_in_username = login_username.strip()
                        st.session_state.logged_in_user_id = user_id
                        # 切换身份后刷新用户数据
                        app_state.set_user_id(user_id)
                        st.session_state.sidebar_user_data = None
                        st.session_state.user_data_stale = True
                        st.session_state.show_login_dialog = False
                        for k in ('login_username_input', 'login_password_input'):
                            st.session_state.pop(k, None)
                        st.success(msg)
                        st.rerun()
                    else:
                        st.error(msg)

    st.markdown("---")

    # 显示提示文字
    st.markdown('<div style="text-align: center; margin-bottom: 1rem;">', unsafe_allow_html=True)
    st.markdown('**更好的体验，需要你的支持！**')
    st.markdown('</div>', unsafe_allow_html=True)

    # 显示ds.jpg图片
    try:
        ds_image_path = Path("resource/ds.jpg")
        if ds_image_path.exists():
            st.image(str(ds_image_path), use_container_width=True)
    except Exception as e:
        logger.warning(f"加载ds.jpg失败: {e}")

    st.markdown("---")

    # 需求提交入口
    if st.button("💡 提交需求/反馈", use_container_width=True, key="feedback_btn"):
        show_feedback_dialog()

    # 居中显示版本号和版权信息
    col1, col2, col3 = st.columns([1, 6, 1])
    with col2:
        st.markdown('<div style="text-align: center; white-space: nowrap;">', unsafe_allow_html=True)
        st.markdown('<p style="text-align: center; margin: 0.5rem 0 0 0; color: #666; font-size: 0.875rem;">标书抄写神器2.0.0</p>', unsafe_allow_html=True)
        st.markdown('<p style="text-align: center; margin: 0.25rem 0 0 0; color: #666; font-size: 0.75rem; white-space: nowrap;">© 2026 文档转换工具 保留所有权利</p>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

# ==================== 主功能区 ====================

# 文件上传区（修复版：上下排列，避免st.columns导致的布局震荡）
# 使用 session_state 保持上传器状态，避免页面刷新时消失
if 'source_files_uploaded' not in st.session_state:
    app_state.set_source_files_uploaded(False)

st.subheader("📄 上传源文档")

st.info("⚠️ 仅支持 **.docx** 格式（Word 2007 及以上）。如果您的文档是 **.doc** 格式（Word 97-2003），请先用 Word 打开后「另存为」**.docx** 格式再上传。", icon="💡")

source_files = st.file_uploader(
    "选择要转换的 Word 文档（可多选）",
    type=['docx'],
    help="仅支持 .docx 格式，可同时选择多个文件",
    accept_multiple_files=True,
    key="source_uploader"
)

# 标记已上传状态
if source_files and not app_state.get_source_files_uploaded():
    app_state.set_source_files_uploaded(True)

# [OK] 修复：优先使用session_state中的文件，如果为空则使用file_uploader返回的文件
current_source_files = st.session_state.get('current_source_files', None)
if source_files:
    # 如果有新上传的文件，更新session_state
    current_source_files = source_files
    app_state.set_current_source_files(source_files)

if current_source_files:
    # [OK] 修复：不再重复设置，已经在第830-831行设置过了
    
    # 检查是否需要重新分析（文件变化或尚未分析）
    need_analyze = False
    current_file_names = [sf.name for sf in current_source_files]
    analyzed_file_names = list(st.session_state.get('file_styles_map', {}).keys())
    
    if not analyzed_file_names or set(current_file_names) != set(analyzed_file_names):
        need_analyze = True
    
    # 始终创建进度条组件（避免作用域问题）
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    # 如果需要分析，显示进度条（基于段落数量更新）
    if need_analyze:
        # 初始化进度条为0
        progress_bar.progress(0)
        status_text.text(" 正在分析源文档...")
        
        # [HIGH_VOLTAGE] 性能优化：记录开始时间
        start_time = time.time()
        
        # 分析源文档样式（基于段落数量更新进度条）
        file_styles_map = {}
        
        # [HIGH_VOLTAGE] 性能优化：单次遍历完成段落计数和样式分析
        file_styles_map = {}
        file_paragraph_counts = {}
        total_paragraphs = 0
        total_files = len(source_files)  # [WARN] 修复：定义total_files变量
        
        for idx, source_file in enumerate(source_files, 1):
            temp_source = f"temp_source_{app_state.get_user_id()}_{source_file.name}"
            with open(temp_source, 'wb') as f:
                f.write(source_file.getbuffer())
            
            from docx import Document
            doc = Document(temp_source)  # ← 只加载1次
            
            current_file_total = len(doc.paragraphs)
            file_paragraph_counts[source_file.name] = current_file_total
            total_paragraphs += current_file_total
            
            styles = set()
            
            # ★ 收集列表段落虚拟样式（数字编号/符号编号），用于 Step 3 样式映射
            converter_temp = DocumentConverter()
            list_virtual_styles = converter_temp.get_list_virtual_styles(doc)
            styles.update(list_virtual_styles)
            
            status_text.text(f"🔍 正在分析文件 {idx}/{total_files}: {source_file.name}...")
            
            for para_idx, para in enumerate(doc.paragraphs):
                if para.style and para.style.name:
                    styles.add(para.style.name)
                    # 检测大纲级别（outlineLvl）并生成虚拟样式名
                    para_style_lower = para.style.name.lower()
                    if not (para_style_lower.startswith('heading') or para_style_lower.startswith('head')):
                        elem = para._element
                        pPr = elem.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
                        if pPr is not None:
                            outline = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl')
                            if outline is not None:
                                val = outline.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                                if val is not None:
                                    try:
                                        level = int(val) + 1
                                        if 1 <= level <= 9:
                                            styles.add(f'[大纲级别 {level}]')
                                    except ValueError:
                                        pass
                
                # 每处理10个段落或最后一个段落时更新进度
                if (para_idx + 1) % 10 == 0 or para_idx == len(doc.paragraphs) - 1:
                    completed_files_progress = (idx - 1) * (100 / total_files)
                    current_file_progress = ((para_idx + 1) / current_file_total) * (100 / total_files)
                    total_progress = completed_files_progress + current_file_progress

                    now = time.monotonic()
                    last_progress_update = st.session_state.get('_analysis_progress_time', 0.0)
                    if now - last_progress_update >= 0.2 or para_idx == len(doc.paragraphs) - 1:
                        progress_bar.progress(min(total_progress / 100, 1.0))
                        st.session_state._analysis_progress_time = now
                
            # 保存该文件的样式和段落数
            file_styles_map[source_file.name] = sorted(list(styles))
            
            # 确保进度至少增加（处理空文件）
            if current_file_total == 0:
                completed_files_progress = idx * (100 / total_files)
                progress_bar.progress(min(completed_files_progress / 100, 1.0))
        
        # 分析完成
        elapsed = time.time() - start_time
        progress_bar.progress(1.0)
        status_text.text(f"[OK] 分析完成！耗时: {elapsed:.1f}秒")
        
        app_state.set_file_styles_map(file_styles_map)
        app_state.set_file_paragraph_counts(file_paragraph_counts) # [HIGH_VOLTAGE] 保存段落数供后续使用
        
        # 合并所有文件的样式用于显示
        all_styles = set()
        for styles in file_styles_map.values():
            all_styles.update(styles)
        all_styles = sorted(list(all_styles))
        app_state.set_source_styles(all_styles)
    else:
        # 使用已缓存的样式，显示进度条（直接100%）
        file_styles_map = app_state.get_file_styles_map()
        file_paragraph_counts = st.session_state.get('file_paragraph_counts', {})  # [WARN] 修复：从缓存中恢复
        all_styles = app_state.get_source_styles()
        progress_bar.progress(1.0)
        status_text.text("[OK] 已分析完成（使用缓存）")
    
    # [HIGH_VOLTAGE] 性能优化：使用分析阶段已计算的段落数，避免重复读取
    file_info = [(sf.name, file_paragraph_counts[sf.name]) for sf in current_source_files]
    total_paragraphs = sum(file_paragraph_counts.values())
    
    # 将所有信息整合到一个expander中
    with st.expander(f"📄 源文档信息：{len(source_files)}个文件 | {len(all_styles)}种样式 | {total_paragraphs:,}段落", expanded=True):
        # 第一行：基本信息
        st.markdown(f"**✅ 已上传:** {len(source_files)} 个文件")
        st.markdown(f"**📋 检测到样式:** {len(all_styles)} 种 - {', '.join(all_styles[:10])}{'...' if len(all_styles) > 10 else ''}")
        
        # 第二行：文件详情
        st.markdown("**ℹ️ 文件详情：**")
        for fname, fpara in file_info:
            st.markdown(f"  • {fname}: {fpara:,} 个段落")
        
        # 第三行：段落数
        st.markdown(f"**📊 总段落数:** {total_paragraphs:,}")
        
        # [OK] 只在非转换完成状态下检查余额（防止重渲染时误报）
        if not st.session_state.get('show_download_buttons', False):
            if total_paragraphs > user_data['paragraphs_remaining']:
                st.error(f"❌ 余额不足！需要 {total_paragraphs:,}，剩余 {user_data['paragraphs_remaining']:,}")

# 模板文档上传（上下排列）
# 使用 session_state 保持上传器状态
if 'template_file_uploaded' not in st.session_state:
    app_state.set_template_file_uploaded(False)

st.subheader("📋 上传模板文档")
template_file = st.file_uploader(
    "选择模板文档（仅支持 .docx）",
    type=['docx'],
    help="用于定义目标样式的 Word 文档",
    key="template_uploader"
)

# 标记已上传状态
if template_file and not app_state.get_template_file_uploaded():
    app_state.set_template_file_uploaded(True)

# [OK] 修复：优先使用session_state中的模板文件路径，如果为空则使用file_uploader返回的文件
current_temp_template = st.session_state.get('current_temp_template', None)
last_template_name = st.session_state.get('last_template_name', None)

if template_file:
    # 如果有新上传的模板文件，保存并更新session_state
    
    # [OK] 修复：清除旧的样式缓存，强制重新解析
    # 防止用户上传新模板后仍使用旧模板的样式缓存
    if 'template_styles' in st.session_state:
        app_state.delete_key('template_styles')
        logger.info(f"[REFRESH] 清除旧模板样式缓存，准备重新解析")
    
    temp_template = f"temp_template_{app_state.get_user_id()}.docx"
    with open(temp_template, 'wb') as f:
        f.write(template_file.getbuffer())
    current_temp_template = temp_template
    last_template_name = template_file.name
    app_state.set_current_temp_template(temp_template)
    app_state.set_last_template_name(last_template_name)

if current_temp_template:
    # [OK] 修复：不再重复保存文件，直接使用current_temp_template
    
    # 检查是否需要重新分析模板样式
    need_analyze_template = ('template_styles' not in st.session_state or 
                             st.session_state.get('last_template_name') != last_template_name)
    
    # 始终创建进度条组件（避免作用域问题）
    template_progress_bar = st.progress(0)
    template_status_text = st.empty()
    
    if need_analyze_template:
        # 初始化进度条为0
        template_progress_bar.progress(0)
        template_status_text.text("[SEARCH] 正在分析模板样式...")
        
        # 修复：提取模板文档中所有定义的段落样式（不是只提取使用的）
        template_progress_bar.progress(0.5)
        template_status_text.text("正在提取所有段落样式...")
        
        # 使用正确的函数：从doc.styles中提取所有段落样式
        template_styles_list = get_template_styles_list(current_temp_template)
        
        # 分析完成
        template_progress_bar.progress(1.0)
        template_status_text.text(f"[OK] 已提取 {len(template_styles_list)} 种样式！")
        
        app_state.set_template_styles(template_styles_list)
        app_state.set_last_template_name(last_template_name)
    else:
        # 使用已缓存的样式，显示进度条（直接100%）
        template_styles = app_state.get_template_styles()
        template_progress_bar.progress(1.0)
        template_status_text.text("[OK] 已分析完成（使用缓存）")
    
    # 将模板信息整合到一个expander中
    with st.expander(f"📋 模板文档信息：{os.path.basename(current_temp_template)} | {len(app_state.get_template_styles())}种样式", expanded=True):
        st.markdown(f"**✅ 已上传:** {os.path.basename(current_temp_template)}")
        st.markdown(f"**📋 检测到样式:** {len(app_state.get_template_styles())} 种 - {', '.join(app_state.get_template_styles()[:10])}{'...' if len(app_state.get_template_styles()) > 10 else ''}")

# 转换配置
st.markdown("---")
st.subheader("⚙️ 转换配置")

# ====================================================================
# 加载用户持久化的默认配置（跨会话保持，多用户隔离）
# 所有"设为默认"按钮保存的配置在此统一恢复
# ====================================================================
_user_defaults = {}
_sm = {}  # ★ 修复：初始化，供下方 file_style_mappings 使用

# ★ 修复：将 style_mappings 完整加载到 session_state
# 桌面版/Web版"设为默认"保存的样式映射、表格/图片/列表兜底配置，
# 页面刷新/重启后通过此处恢复到 session_state，供转换时回退使用
if 'file_style_mappings' not in st.session_state:
    # 仅首次会话初始化时读取持久化默认配置
    try:
        _uid = app_state.get_user_id()
        if _uid:
            _ud = load_user_data(_uid)
            if _ud and 'style_mappings' in _ud:
                _sm = _ud['style_mappings']
                _user_defaults = {
                    'hint':      _sm.get('_default_hint_settings', {}) or {},
                    'answer':    _sm.get('_default_answer_config', {}) or {},
                    'list':      _sm.get('_default_list_config', {}) or {},
                    'tbl_img':   _sm.get('_default_tbl_img_config', {}) or {},
                    'rm_chapter': _sm.get('_default_remove_chapter_label', None),
                }
    except Exception:
        pass
    st.session_state.file_style_mappings = _sm if isinstance(_sm, dict) else {}

_h  = _user_defaults.get('hint', {})
_a  = _user_defaults.get('answer', {})
_l  = _user_defaults.get('list', {})
_t  = _user_defaults.get('tbl_img', {})

if 'do_mood_config' not in st.session_state:
    app_state.set_do_mood_config(True)

# ── 应答句配置 ──
if 'do_answer_config' not in st.session_state:
    st.session_state.do_answer_config = _a.get('do_answer', False)
if 'answer_text_config' not in st.session_state:
    app_state.set_answer_text_config(_a.get('answer_text', '应答：本投标人理解并满足要求。'))
if 'answer_style_config' not in st.session_state:
    app_state.set_answer_style_config(_a.get('answer_style', 'Normal'))
if 'answer_mode_config' not in st.session_state:
    app_state.set_answer_mode_config(_a.get('answer_mode', 'copy_chapter'))
if 'answer_source_style_config' not in st.session_state:
    app_state.set_answer_source_style_config(_a.get('answer_source_style', ''))
if 'answer_copy_style_config' not in st.session_state:
    app_state.set_answer_copy_style_config(_a.get('answer_copy_style', ''))

# ── 列表段落兜底配置 ──
if 'enable_list_style_config' not in st.session_state:
    app_state.set_enable_list_style_config(_l.get('enable_list', True))
if 'list_method_config' not in st.session_state:
    app_state.set_list_method_config(_l.get('method', 'bullet'))
if 'list_bullet_config' not in st.session_state:
    app_state.set_list_bullet_config(_l.get('bullet', '•'))
if 'list_style_config' not in st.session_state:
    app_state.set_list_style_config(_l.get('style', 'Body Text'))
if 'list_answer_method_config' not in st.session_state:
    app_state.set_list_answer_method_config(_l.get('answer_method', 'bullet'))
if 'list_answer_bullet_config' not in st.session_state:
    app_state.set_list_answer_bullet_config(_l.get('answer_bullet', '•'))
if 'list_answer_style_config' not in st.session_state:
    app_state.set_list_answer_style_config(_l.get('answer_style', 'Body Text'))

# ── 表格/图片兜底配置 ──
if 'enable_table_style_config' not in st.session_state:
    app_state.set_enable_table_style_config(_t.get('enable_table_style', False))
if 'table_style_config' not in st.session_state:
    app_state.set_table_style_config(_t.get('table_style', 'Body Text'))
if 'table_answer_style_config' not in st.session_state:
    app_state.set_table_answer_style_config(_t.get('table_answer_style', ''))
if 'enable_image_style_config' not in st.session_state:
    app_state.set_enable_image_style_config(_t.get('enable_image_style', False))
if 'image_style_config' not in st.session_state:
    app_state.set_image_style_config(_t.get('image_style', 'Body Text'))
if 'image_answer_style_config' not in st.session_state:
    app_state.set_image_answer_style_config(_t.get('image_answer_style', ''))

# ── 清除章节标签 ──
if 'remove_chapter_label_config' not in st.session_state:
    _rm = _user_defaults.get('rm_chapter', None)
    app_state.set_remove_chapter_label_config(False if _rm is None else bool(_rm))

# ── 提示语配置 ──
if 'do_hint_config' not in st.session_state:
    app_state.set_do_hint_config(_h.get('do_hint', False))
if 'hint_type_config' not in st.session_state:
    app_state.set_hint_type_config(_h.get('hint_type', 'text'))
if 'hint_text_config' not in st.session_state:
    app_state.set_hint_text_config(_h.get('hint_text', '招标文件原文'))
if 'hint_style_config' not in st.session_state:
    app_state.set_hint_style_config(_h.get('hint_style', 'Normal'))
if 'hint_image_config' not in st.session_state:
    app_state.set_hint_image_config(None)  # 图片路径不跨会话恢复（临时文件）


# 在“⚙️ 转换配置”下放置“配置样式映射”按钮
st.markdown("---")
map_col1, map_col2 = st.columns([2, 8])
with map_col1:
    if st.button("📊 配置样式映射", key="open_style_mapping_btn", use_container_width=True,
                 help="完整的四步样式配置（标题映射、应答句、正文映射、表格/图片/列表兜底）"):
        from components.dialogs.style_mapping import show_style_mapping_dialog
        show_style_mapping_dialog()
with map_col2:
    if current_source_files and st.session_state.get('template_styles'):
        st.caption("点击按钮配置当前文件的样式映射")
    else:
        st.caption("上传源文档和模板文档后即可配置样式映射")

# ==================== [FIX] 调用配置区组件渲染实际的UI控件 ====================
# render_conversion_config() 来自 components/config_panel.py
# 完全参照桌面版布局：
# 转换选项区：祈使语气转换checkbox
# 章节提示语区：插入提示语checkbox + 类型/样式/内容配置
# 所有其他配置（应答句、列表段落、表格/图片兜底、清除章节标签）均在样式映射对话框中管理
result = render_conversion_config()
do_mood, do_answer, list_bullet, answer_text, answer_style, answer_mode = result[0:6]
do_hint, hint_type, hint_text, hint_image_path, hint_style = result[6:11]
answer_source_style, answer_copy_style = result[11:13]
list_method, list_style, list_answer_method, list_answer_style, list_answer_bullet = result[13:18]
remove_chapter_label = result[18]
enable_list_style = result[19] if len(result) > 19 else True

# 不插入应答句时使用默认值（确保变量存在）
if not do_answer:
    answer_text = app_state.get_answer_text_config()
    answer_style = app_state.get_answer_style_config()
    answer_mode = app_state.get_answer_mode_config()
    answer_source_style = app_state.get_answer_source_style_config()
    answer_copy_style = app_state.get_answer_copy_style_config()


# 开始转换按钀
st.markdown("---")

# 检查是否正在前台转换中
is_converting = st.session_state.get('is_converting', False)

if is_converting:
    # 如果正在转换，显示提示信恀
    st.warning("⏳ **正在进行前台转换，请稍后...**\n\n转换期间无法进行其他操作，请耐心等待转换完成。")
    st.info("💡 转换完成后将自动恢复操作权限")
else:
    # 正常状态，显示开始转换按钀
    if st.button("🚀 开始转换", type="primary", use_container_width=True):
            # [OK] 修复：清空之前的转换结果，避免新旧结果混合
            st.session_state.conversion_file_results = []
            st.session_state.recent_results = []
            st.session_state.show_download_buttons = False
            
            # [OK] 提前定义进度条和状态文本（必须在所有使用之前）
            progress_placeholder = st.empty()
            status_placeholder = st.empty()
            progress_bar = progress_placeholder.progress(0)
            
            # 验证输入（从session_state中恢复文件变量）
            # [OK] 修复：从session_state获取文件，而不是依赖局部变量（页面刷新后会丢失（
            current_source_files = st.session_state.get('current_source_files', None)
            current_temp_template = st.session_state.get('current_temp_template', None)
                        
            if not current_source_files or not current_temp_template:
                st.error("❌ 请上传源文档和模板文档")
                status_placeholder.text("[ERROR] 验证失败：缺少文件")
                progress_bar.progress(0)
                st.stop()
            elif not os.path.exists(current_temp_template):
                st.error("❌ 文件上传失败，请重试")
                status_placeholder.text("[ERROR] 验证失败：文件上传错误")
                progress_bar.progress(0)
                st.stop()
            else:
                # 设置转换标志，禁用后续操作
                st.session_state.is_converting = True
                
                # [HIGH_VOLTAGE] 性能优化：立即更新进度条，不要等验证完成
                status_placeholder.text("⏳ 正在验证输入...")
                progress_bar.progress(5)
            
            # [HIGH_VOLTAGE] 性能优化：使用分析阶段已计算的段落数（file_paragraph_counts已在笀86-899行计算）
            # 如果file_paragraph_counts不存在（异常情况），使用兜底逻辑
            if 'file_paragraph_counts' in st.session_state and st.session_state.file_paragraph_counts:
                file_paragraph_counts = st.session_state.file_paragraph_counts
                file_info = [(sf.name, file_paragraph_counts[sf.name]) for sf in current_source_files]
                total_paragraphs = sum(file_paragraph_counts.values())
            else:
                # 兜底逻辑：重新计算（不应该发生）
                logger.warning("file_paragraph_counts 不存在，使用兜底逻辑重新计算")
                total_paragraphs = 0
                file_info = []
                for sf in current_source_files:
                    temp_source = f"temp_source_{st.session_state.user_id}_{sf.name}"
                    paragraphs = count_paragraphs(temp_source)
                    total_paragraphs += paragraphs
                    file_info.append((sf.name, paragraphs))
            
            
            progress_bar.progress(10)
            status_placeholder.text("⏳ 准备转换...")
            
            # [HIGH_VOLTAGE] 性能优化：使用缓存的文件信息，避免重复读叀
            source_files_info = []
            for fname, fpara in file_info:
                temp_source = f"temp_source_{st.session_state.user_id}_{fname}"
                source_files_info.append((fname, temp_source, fpara))
            
            # 配置字典
            config = {
                'do_mood': do_mood,
                'answer_text': answer_text,
                'answer_style': answer_style,
                'list_bullet': list_bullet if list_bullet else "—",
                'do_answer_insertion': do_answer,
                'answer_mode': answer_mode,
                'answer_source_style': answer_source_style,
                'answer_copy_style': answer_copy_style,
                'list_method': list_method,
                'list_style': list_style,
                'list_answer_method': list_answer_method,
                'list_answer_style': list_answer_style,
                'list_answer_bullet': list_answer_bullet,
                'remove_chapter_label': remove_chapter_label,
                'custom_style_map': st.session_state.get('style_mapping', None)  # 用户配置的样式映尀
            }
                        
            # ========== 前台转换模式 ==========
            # 进度条已在按钮点击时创建
                        
            # 添加"转为后台"按钮（使用session_state标记（
            if 'switch_to_background' not in st.session_state:
                st.session_state.switch_to_background = False
                        
            try:
                # 更新进度提示
                status_placeholder.text("⏳ 正在初始化转换器...")
                progress_bar.progress(10)
                                
                # 创建转换噀
                converter = DocumentConverter()
                progress_bar.progress(10)
                
                # 处理每个文件
                output_files = []
                success_count = 0
                fail_count = 0
                total_success_paragraphs = 0  # 成功转换的段落数
                
                # [OK] 初始化文件级结果列表（用于持久化保存）
                # [FIX] 只在首次转换时初始化，避免rerun后清空已保存的结果
                if 'conversion_file_results' not in st.session_state:
                    st.session_state.conversion_file_results = []
                
                for idx, source_file_obj in enumerate(current_source_files):
                    # 输出文件路径 - 保存到conversion_results目录
                    base_name = os.path.splitext(source_file_obj.name)[0]
                    output_filename = f"result_{base_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                    output_file = os.path.join("conversion_results", output_filename)
                    temp_source = f"temp_source_{st.session_state.user_id}_{source_file_obj.name}"
                    
                    # [FIX] 确保临时文件存在（防止Streamlit热重载或清理后丢失）
                    if not os.path.exists(temp_source):
                        import logging
                        logger = logging.getLogger(__name__)
                        logger.warning(f"临时文件 {temp_source} 不存在，从 UploadedFile 重新创建")
                        for sf in st.session_state.current_source_files:
                            if sf.name == source_file_obj.name:
                                with open(temp_source, 'wb') as f:
                                    f.write(sf.getbuffer())
                                break
                    
                    # [HIGH_VOLTAGE] 性能优化：从缓存中获取段落数，避免重复读叀
                    file_paragraphs = 0
                    for fname, fpara in file_info:
                        if fname == source_file_obj.name:
                            file_paragraphs = fpara
                            break
                    
                    status_placeholder.text(f" 正在转换第 {idx+1}/{len(current_source_files)} 个文件 {source_file_obj.name} ({file_paragraphs:,} 段落)")
                    
                    # [OK] 修复：使用每个文件各自的样式映射配置（与桌面版一致）
                    # ★ 如果没有文件级配置，回退到用户设定的默认配置
                    file_mapping = None
                    file_tbl_img_config = {}
                    file_list_config = {}
                    if 'file_style_mappings' in st.session_state and source_file_obj.name in st.session_state.file_style_mappings:
                        file_mapping_data = st.session_state.file_style_mappings[source_file_obj.name]
                        # 样式映射部分（排除特殊键）
                        file_mapping = {k: v for k, v in file_mapping_data.items() if not k.startswith('_')}
                        # 表格/图片样式配置（从映射数据中的特殊键取出）
                        file_tbl_img_config = file_mapping_data.get('_table_image_style', {})
                        # 列表段落兜底配置（从映射数据中的特殊键取出）
                        file_list_config = file_mapping_data.get('_list_config', {})
                    
                    # ★ 修复：文件无单独配置时，回退到"设为默认"保存的默认配置
                    if not file_mapping:
                        default_style_map = st.session_state.file_style_mappings.get('_default_style_map', {})
                        if default_style_map:
                            # 排除嵌套的特殊键，只取样式映射部分
                            file_mapping = {k: v for k, v in default_style_map.items() if not k.startswith('_')}
                            if file_mapping:
                                st.info(f"📋 {source_file_obj.name}: 使用默认样式映射 ({len(file_mapping)} 个样式)")
                    
                    if not file_tbl_img_config:
                        default_tbl_img = st.session_state.file_style_mappings.get('_default_tbl_img_config', {})
                        if default_tbl_img:
                            file_tbl_img_config = default_tbl_img
                    
                    if not file_list_config:
                        default_list = st.session_state.file_style_mappings.get('_default_list_config', {})
                        if default_list:
                            file_list_config = default_list
                    
                    # 警告收集
                    warnings_list = []
                    def warning_callback(msg):
                        warnings_list.append(msg)
                    
                    # 进度回调函数 - 实时更新进度杀
                    def make_progress_callback(file_idx, total_files):
                        last_update = [0.0]

                        def callback(step, message):
                            # 计算总体进度 (10% - 80%)
                            base_progress = 10 + int((file_idx / total_files) * 70)
                            step_progress = int((step / 7) * (70 / total_files))
                            current_progress = min(base_progress + step_progress, 80)
                            now = time.monotonic()
                            if current_progress >= 80 or now - last_update[0] >= 0.2:
                                progress_bar.progress(current_progress)
                                status_placeholder.text(f"⏀{message}")
                                last_update[0] = now
                        return callback
                    
                    # [HIGH_VOLTAGE] 性能优化：传递缓存的样式列表，避免重复分枀
                    source_styles_for_file = st.session_state.file_styles_map.get(source_file_obj.name, None)
                    
                    # 获取该文件的表格/图片应答样式配置（★ 修复：回退到全局默认配置）
                    file_table_answer_style = file_tbl_img_config.get('table_answer_style') or st.session_state.get('table_answer_style_config', '')
                    file_image_answer_style = file_tbl_img_config.get('image_answer_style') or st.session_state.get('image_answer_style_config', '')
                    
                    # 文件级列表段落配置覆盖全局设置
                    _list_bullet = file_list_config.get('bullet', list_bullet if list_bullet else "—")
                    _list_method = file_list_config.get('method', list_method)
                    _list_style = file_list_config.get('style', list_style)
                    _list_answer_method = file_list_config.get('answer_method', list_answer_method)
                    _list_answer_style = file_list_config.get('answer_style', list_answer_style)
                    _list_answer_bullet = file_list_config.get('answer_bullet', list_answer_bullet)
                    _enable_list_style = file_list_config.get('enable_list', enable_list_style)
                    
                    # 执行转换
                    success, actual_file, msg = converter.full_convert(
                        source_file=temp_source,
                        template_file=current_temp_template,
                        output_file=output_file,
                        custom_style_map=file_mapping,  # [OK] 修复：使用每个文件各自的映射配置
                        do_mood=do_mood,
                        answer_text=answer_text,
                        answer_style=answer_style,
                        answer_source_style=answer_source_style,
                        answer_copy_style=answer_copy_style,
                        table_answer_style=file_table_answer_style,
                        list_bullet=_list_bullet,
                        list_method=_list_method,
                        list_style=_list_style,
                        list_answer_method=_list_answer_method,
                        list_answer_style=_list_answer_style,
                        list_answer_bullet=_list_answer_bullet,
                        do_answer_insertion=do_answer,
                        answer_mode=answer_mode,
                        do_hint_insertion=do_hint,
                        hint_type=hint_type,
                        hint_text=hint_text,
                        hint_image_path=hint_image_path,
                        hint_style=hint_style,
                        progress_callback=make_progress_callback(idx, len(current_source_files)),
                        warning_callback=warning_callback,
                        source_styles_cache=source_styles_for_file,  # [HIGH_VOLTAGE] 传递缓存的样式列表
                        table_style_override=file_tbl_img_config.get('table_style') or st.session_state.get('table_style_config', 'Body Text'),
                        enable_table_style=file_tbl_img_config.get('enable_table_style', st.session_state.get('enable_table_style_config', False)),
                        image_style_override=file_tbl_img_config.get('image_style') or st.session_state.get('image_style_config', 'Body Text'),
                        enable_image_style=file_tbl_img_config.get('enable_image_style', st.session_state.get('enable_image_style_config', False)),
                        remove_chapter_label=remove_chapter_label,
                        enable_list_style=_enable_list_style
                    )
                    
                    if success:
                        output_files.append(actual_file)
                        success_count += 1
                        total_success_paragraphs += file_paragraphs
                        
                        # [OK] 保存文件级结果到 session_state（防止重渲染后丢失）
                        st.session_state.conversion_file_results.append({
                            'name': source_file_obj.name,
                            'status': 'success',
                            'paragraphs': file_paragraphs,
                            'warnings': warnings_list.copy()  # 复制列表，避免后续修改影哀
                        })
                    else:
                        fail_count += 1
                        
                        # [OK] 保存文件级失败结果到 session_state
                        st.session_state.conversion_file_results.append({
                            'name': source_file_obj.name,
                            'status': 'fail',
                            'msg': msg
                        })
                
                progress_bar.progress(90)
                
                if success_count > 0:
                    progress_bar.progress(100)
                    
                    # [OK] 扣除段落脚额（只扣段落数，不涉及费用（
                    if user_data['paragraphs_remaining'] >= total_success_paragraphs:
                        user_data['paragraphs_remaining'] -= total_success_paragraphs
                    else:
                        # 如果余额不足，只扣除剩余部分，最低为0
                        user_data['paragraphs_remaining'] = 0
                    
                    # 更新用户统计
                    user_data['total_converted'] += success_count
                    user_data['total_paragraphs_used'] += total_success_paragraphs
                    
                    # 记录转换历史（不包含费用（
                    conversion_record = {
                        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                        'files': len(current_source_files),
                        'success': success_count,
                        'failed': fail_count,
                        'paragraphs_charged': total_success_paragraphs,
                        'mode': 'foreground'
                    }
                    
                    # [OK] 防御性编程：确保conversion_history字段存在
                    if 'conversion_history' not in user_data:
                        user_data['conversion_history'] = []
                    
                    user_data['conversion_history'].append(conversion_record)
                    
                    # [OK] 修复：调用add_conversion_record写入conversion_tasks表（API模式（
                    from data_manager import add_conversion_record
                    add_conversion_record(
                        files_count=len(current_source_files),
                        success_count=success_count,
                        failed_count=fail_count,
                        user_id=st.session_state.user_id,
                        paragraphs=total_success_paragraphs  # [OK] 新增：传递段落数
                    )
                    
                    # 保存用户数据（使用统一数据接口（
                    from data_manager import save_user_data
                    save_user_data(user_data, st.session_state.user_id)
                    st.session_state.user_data_stale = True
                    
                    # [OK] 修复：将转换结果文件路径保存刀session_state，防止刷新后丢失
                    if 'recent_results' not in st.session_state:
                        st.session_state.recent_results = []
                    
                    # 添加本次转换的结果文什
                    for output_file in output_files:
                        if os.path.exists(output_file):
                            file_info = {
                                'path': output_file,
                                'name': os.path.basename(output_file),
                                'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                            }
                            st.session_state.recent_results.append(file_info)
                    
                    # 重置转换标志
                    st.session_state.is_converting = False
                    
                    # [OK] 保存转换总结信息到session_state（在下载区域统一显示，防止重复）
                    st.session_state.conversion_summary = {
                        'success_count': success_count,
                        'fail_count': fail_count,
                        'total_paragraphs': total_success_paragraphs
                    }
                    
                    # [OK] 清理临时文件（源文件和模板文件）
                    try:
                        from file_manager import get_file_manager
                        fm = get_file_manager()
                        cleanup_stats = fm.cleanup_temp_files(st.session_state.user_id)
                        logger.info(f"临时文件清理完成: {cleanup_stats}")
                    except Exception as cleanup_error:
                        logger.warning(f"临时文件清理失败（不影响转换结果（ {cleanup_error}")
                    
                    # [OK] 标记显示下载按钮（用于页面刷新后保持状态）
                    st.session_state.show_download_buttons = True
                    
                    # [OK] 强制重新渲染，避免在同一轮渲染中重复显示转换总结
                    st.rerun()
                else:
                    # 所有文件都转换失败
                    status_placeholder.text("[ERROR] 转换失败")
                    progress_bar.progress(100)
                    st.session_state.is_converting = False
                    
                    # [OK] 保存转换总结信息（即使全部失败也要保存，以便显示详细错误）
                    st.session_state.conversion_summary = {
                        'success_count': success_count,
                        'fail_count': fail_count,
                        'total_paragraphs': total_success_paragraphs
                    }
                    
                    # [OK] 标记显示下载按钮（用于显示失败详情）
                    st.session_state.show_download_buttons = True
                    
                    st.error("❌ 所有文件转换失败，请检查错误信息")
                    st.info("💡 请查看下方的具体错误提示，修正后重试")
                    
                    # [OK] 强制重新渲染，显示失败详情
                    st.rerun()
        
            except Exception as e:
                # 重置转换标志
                st.session_state.is_converting = False
                
                error_msg = f"发生错误: {str(e)}"
                st.error(error_msg)
                
                # [FIX] 如果发生未预期的异常，也要保存错误信息
                if 'conversion_file_results' not in st.session_state or not st.session_state.conversion_file_results:
                    # 如果没有文件级结果，创建一个通用的错误记录
                    st.session_state.conversion_file_results = []
                    current_source_files = st.session_state.get('current_source_files', [])
                    for source_file_obj in current_source_files:
                        st.session_state.conversion_file_results.append({
                            'name': source_file_obj.name,
                            'status': 'fail',
                            'msg': f"系统异常: {error_msg}"
                        })
                
                # 保存转换总结信息
                st.session_state.conversion_summary = {
                    'success_count': 0,
                    'fail_count': len(st.session_state.get('current_source_files', [])),
                    'total_paragraphs': 0
                }
                
                # 标记显示下载按钮
                st.session_state.show_download_buttons = True
                
                import traceback
                with st.expander("📋 查看详细错误堆栈"):
                    st.code(traceback.format_exc())
                
                # 强制重新渲染，显示错误详情
                st.rerun()

# [OK] 转换完成后显示下载按钮和转换总结信息（在按钮之后，从session_state读取（
if 'show_download_buttons' in st.session_state and st.session_state.show_download_buttons:
    # [OK] 显示转换总结信息（从session_state读取，防止刷新后丢失（
    if 'conversion_summary' in st.session_state and st.session_state.conversion_summary:
        summary = st.session_state.conversion_summary
        st.success(f"🎉 转换完成！成功 {summary['success_count']} 个，失败: {summary['fail_count']} 个")
        if summary['fail_count'] > 0:
            st.warning(f"⚠️ 有 {summary['fail_count']} 个文件转换失败")
        st.info(f"处理 {summary['total_paragraphs']:,} 个段落")
    
    # [OK] 恢复每个文件的转换结果（从session_state读取，防止重渲染后丢失）
    if 'conversion_file_results' in st.session_state:
        for result in st.session_state.conversion_file_results:
            if result['status'] == 'success':
                st.success(f"✅ {result['name']} 转换成功")
                # 显示警告汇总信息（如果有）
                if result.get('warnings'):
                    # 统计无法自动提取预览图的 OLE 对象数（成功提取预览图的不会产生 warning）
                    ole_fail_count = sum(1 for w in result['warnings'] if '无法自动提取预览图' in w)
                    
                    # 只有 OLE 转换失败时才提示手动处理（成功转换的不提示）
                    if ole_fail_count > 0:
                        st.warning(
                            f"⚠️ **{result['name']}** 文件中有 {ole_fail_count} 个 OLE 对象无法自动转换，"
                            f"已在文档相应位置标注「[OLE对象，请手动复制]」，请手动处理。"
                        )
            else:
                st.error(f"❌ {result['name']} 转换失败: {result.get('msg', '')}")
    
    st.subheader("📥 下载转换结果")
    
    # 显示所有转换结果文什
    if 'recent_results' in st.session_state and st.session_state.recent_results:
        for idx, file_info in enumerate(st.session_state.recent_results):
            if os.path.exists(file_info['path']):
                with open(file_info['path'], 'rb') as f:
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        st.download_button(
                            label=f"[DOWNLOAD] 下载: {file_info['name']}",
                            data=f.read(),
                            file_name=file_info['name'],
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                            key=f"download_recent_{file_info['name']}"
                        )
                    with col2:
                        st.caption(f"转换时间: {file_info['time']}")
            else:
                st.warning(f"⚠️ 文件已过期或不存在: {file_info["name"]}")
    
    st.markdown("---")


st.subheader("📖 使用说明")

# 添加自定义CSS增强使用说明的视觉效果
st.markdown("""
<style>
    .usage-section {
        background-color: #f8f9fa;
        border-left: 4px solid #4CAF50;
        padding: 15px;
        margin: 10px 0;
        border-radius: 5px;
    }
    .usage-section h3 {
        color: #2c3e50;
        border-bottom: 2px solid #e0e0e0;
        padding-bottom: 8px;
        margin-top: 20px;
    }
    .usage-note {
        background-color: #e7f3ff;
        border-left: 3px solid #2196F3;
        padding: 10px;
        margin: 10px 0;
    }
    /* [WARN] 强制显示 expander 箭头 */
    .streamlit-expanderHeader {
        cursor: pointer;
    }
    .streamlit-expanderHeader:hover {
        background-color: rgba(0, 0, 0, 0.05);
    }
</style>
""", unsafe_allow_html=True)

with st.expander("📖 使用说明", expanded=False):
    # 动态获取免费额度配置
    try:
        free_paragraphs_value = _get_cached_config('free_paragraphs_daily')
        if free_paragraphs_value:
            free_paragraphs_display = f"{int(free_paragraphs_value):,}"
        else:
            free_paragraphs_display = "10,000"  # 默认值
    except Exception:
        free_paragraphs_display = "10,000"  # 降级方案
    
    st.markdown(f"""

### 🎯 本工具能帮你解决什么

如果你也是一名苦逼的售前，是否也被抄写标书这种低级的牛马工作折磨过？要把样式乱七八糟的需求文档或厂家方案，按照公司要求的标书样式重新复制粘贴一遍。为了不引入新格式、确保合稿顺利，还必须粘贴为纯文本，再一点一点调整格式……

**现在一键搞定**——过去数天的工作，现在只需几分钟：配置好源文档与模板的样式映射，工具自动完成全部转换。同时支持：
- 将招标文件的祈使语气自动转为投标人口吻（"应""须""必须"→统统消失）
- 五种应答句插入模式，自动批量生成应答内容
- 图片/文本提示语插入，标注原文位置
剩下的就是舒心检查、把标书打磨得更完美。

---

### ℹ️ 操作步骤

1. **上传源文档**：选择需要转换样式的 Word 文档（支持同时上传多个 `.docx` 文件）
2. **上传模板**：选择定义了目标样式的 Word 模板文档
3. **配置样式映射**（重要‼️）：点击「📊 配置样式映射」按钮，完成四步配置：
   - **Step 1 - 标题映射**：将源文档标题样式映射到模板标题样式
   - **Step 2 - 应答句配置**：是否插入应答句、应答文本、插入位置（5种模式）、原文/应答原文样式
   - **Step 3 - 正文+列表段落映射**：将正文和列表段落样式映射到目标样式
   - **Step 4 - 兜底配置**：表格样式、图片样式 + 列表段落兜底设置（未映射列表的处理方式）+ 清除章/节/篇编号
4. **配置转换选项**：
   - ☑️ 祈使语气转换：将"投标人需""应""须"等转为投标人承诺语气
   - ☑️ 插入提示语：在每个章节标题后插入指定提示语（支持文本或图片），样式可自定义
5. **点击「🚀 开始转换」**：系统自动处理并生成结果
6. **下载结果**：转换完成后下载文档

---

### 🎁 段落额度说明

**每日免费额度**：每位用户每天可获得 **{free_paragraphs_display}** 段落的免费转换额度。

**额度规则**：
- 免费额度按日计算，每天自动重置
- 只统计非标题的正文段落
- 转换失败的文件不计入额度消耗

---

### 📊 段落定义

**什么是段落？**
- 段落指 Word 文档中的**正文内容段落**
- **不包括**标题（Heading 1-9、标题 1-9 等样式）
- **包括**普通文本段落、列表项、表格外文字等

**举例：**
```
标题 1：项目概述          ← 不计段落（标题）
这是一个项目...           ← 计段落（正文）

标题 2：技术方案          ← 不计段落（标题）
我们采用...               ← 计段落（正文）
- 第一点                  ← 计段落（列表）
- 第二点                  ← 计段落（列表）
```

---

### 👤 账号绑定（可选）

系统默认使用设备指纹识别用户身份，**无需注册登录即可使用**。

**绑定账号的优势：**
- 更换设备后可登录账号恢复额度
- 用户名易记易识别，方便区分多设备使用场景
- 绑定后每日额度与账号关联，不受设备限制

**操作方式：**
- 点击左侧栏「🔗 绑定账号」→ 设置用户名和密码
- 点击「🔑 账号登录」→ 输入已绑定的用户名密码
- 解绑：登录后点击「🔓 解绑用户」即可恢复设备指纹身份

> ⚠️ 一个设备指纹只能绑定一个账号，用户名不区分大小写且不可重复。

---

**好消息：**
- 工具会自动检测以上问题，并在转换结束后明确提示
- 即使有这些限制，也已经至少帮你自动完成了 90%以上 的工作！
    """)
# ==================== 评论区 ====================
st.markdown("---")
st.subheader("💬 用户评论")
show_comments_section()

st.markdown("---")
st.markdown(
    "<div style='text-align: center; color: gray;'>Powered by Streamlit | MVP Version</div>",
    unsafe_allow_html=True
)
