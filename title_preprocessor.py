# -*- coding: utf-8 -*-
"""
源文档标题预处理引擎（T04 / 工具箱 Tab A）
无 Streamlit 依赖，可独立测试。

职责：
1. 检测以正文格式出现的编号标题（数字编号 + 制表符 + 单列标题，如 "1.1\t线路"）
2. 按编号层级推断大纲级别（1→H1, 1.1→H2, 1.1.1→H3）
3. 将选中段落应用对应 Heading N 样式（并设置大纲级别）
4. 判断文档是否已使用标题样式（供转换页引导提示）
"""
import re
from typing import List, Dict, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 编号标题检测：数字编号 + 制表符 + 单列标题文本（如 "1.1\t线路"）。
# 排除列表项（"1、"、"1）"、"2)"）、多列表格行（"1\t小于 300\t1.2\t65"）等。
# 捕获组 1 = 纯数字编号（用于推断层级），组 2 = 标题正文。
HEADING_PATTERN = re.compile(r'^\s*(\d+(?:\.\d+)*)\t([^\t]*)$')

# 章节标题检测："第一章" / "第1章" / "第一章 概述" 等，默认大纲级别 1。
# 捕获组 1 = 章节编号（如 "第一章"），组 2 = 标题正文。
CHAPTER_PATTERN = re.compile(r'^\s*(第[一二三四五六七八九十百\d]+章)\s*(.*)$')

# 已有标题样式名 -> 大纲级别
HEADING_STYLE_LEVELS = {
    'heading 1': 1, 'heading 2': 2, 'heading 3': 3, 'heading 4': 4, 'heading 5': 5,
    'heading 6': 6, 'heading 7': 7, 'heading 8': 8, 'heading 9': 9,
    '标题 1': 1, '标题 2': 2, '标题 3': 3, '标题 4': 4, '标题 5': 5,
    '标题 6': 6, '标题 7': 7, '标题 8': 8, '标题 9': 9,
}


class TitlePreprocessor:
    """源文档标题预处理引擎（无状态工具类）"""

    MAX_LEVEL = 9

    @staticmethod
    def infer_level(number_part: str) -> int:
        """根据编号推断大纲级别：层级深度 = 点号数量 + 1，上限 9。

        Examples:
            "1"     -> 1
            "1.1"   -> 2
            "1.1.1" -> 3
        """
        depth = number_part.count('.') + 1
        return min(depth, TitlePreprocessor.MAX_LEVEL)

    @staticmethod
    def detect_headings(docx_file) -> List[Dict]:
        """检测文档中的编号标题段落。

        Args:
            docx_file: docx 文件路径。

        Returns:
            List[Dict]，每个元素结构：
                {
                    "index": int,             # 段落索引（doc.paragraphs 中的位置）
                    "text": str,              # 段落原文
                    "number": str,            # 编号部分（如 "1.1.1"）
                    "detected_level": int,    # 自动推断的级别（1-9）
                }
        """
        doc = Document(docx_file)
        candidates = []
        for idx, para in enumerate(doc.paragraphs):
            # 用原始文本匹配（不 strip），以便区分"单列标题"与"多列表格行/列表项"。
            raw_text = para.text
            if not raw_text.strip():
                continue
            m = HEADING_PATTERN.match(raw_text)
            if not m:
                continue
            number_part = m.group(1)
            content = m.group(2).strip()
            # 排除纯数字段落（如单纯的 "1" 后面无内容）
            if not content:
                continue
            # 排除以句号/分号结尾的正文条款（标题通常不以这些标点结尾）
            if content.endswith('。') or content.endswith('；'):
                continue
            # 排除过长的段落（标题通常较短）
            text = raw_text.strip()
            if len(text) > 60:
                continue
            candidates.append({
                "index": idx,
                "text": text,
                "number": number_part,
                "content": content,
                "detected_level": TitlePreprocessor.infer_level(number_part),
            })

        # 先过滤掉内容以数字/符号开头的段落（多为数值、表格行，如 "1294.53～"、"100℃×…"、">10P 250"）。
        filtered = []
        for c in candidates:
            if c["content"][0] in '0123456789±><≥≤~～＋－':
                continue
            filtered.append(c)

        # 单级编号（无 "."）列表：按文档顺序记录其在 filtered 中的位置。
        single_positions = [k for k, c in enumerate(filtered) if '.' not in c["number"]]

        def _has_child(pos):
            """单级编号 N 到下一个单级编号之间，是否存在 N.x 子级。"""
            num = filtered[pos]["number"]
            nxt = None
            for sp in single_positions:
                if sp > pos:
                    nxt = sp
                    break
            end = nxt if nxt is not None else len(filtered)
            return any(
                '.' in filtered[j]["number"] and filtered[j]["number"].startswith(num + '.')
                for j in range(pos + 1, end)
            )

        # 连续性判断：把连续递增的单级编号归为一段（如 3、4、5 或 1、2、3、4）。
        # 段内只要有一个编号存在子级，整段都视为标题（用户规则：3 无子级但 4 有 4.1，则 3、4 都是）。
        single_runs = []  # 每项为位置列表
        for sp in single_positions:
            if single_runs and int(filtered[single_runs[-1][-1]]["number"]) + 1 == int(filtered[sp]["number"]):
                single_runs[-1].append(sp)
            else:
                single_runs.append([sp])

        heading_positions = set()
        for run in single_runs:
            if any(_has_child(p) for p in run):
                heading_positions.update(run)

        headings = []
        for k, c in enumerate(filtered):
            # 多级编号（如 1.1 / 1.1.1）直接视为标题；单级编号按连续性判断。
            if '.' in c["number"] or k in heading_positions:
                headings.append({
                    "index": c["index"],
                    "text": c["text"],
                    "number": c["number"],
                    "detected_level": c["detected_level"],
                })

        # 章节标题（"第一章" / "第1章"）检测：默认大纲级别 1。
        for idx, para in enumerate(doc.paragraphs):
            raw_text = para.text
            if not raw_text.strip():
                continue
            m = CHAPTER_PATTERN.match(raw_text)
            if not m:
                continue
            content = m.group(2).strip()
            # 排除章节号引用（如 "第二章 4.10.3 节 ..."）及句末标点结尾的正文
            if content and (
                content[0].isdigit()
                or content.endswith('。')
                or content.endswith('；')
            ):
                continue
            text = raw_text.strip()
            if len(text) > 60:
                continue
            headings.append({
                "index": idx,
                "text": text,
                "number": m.group(1),
                "detected_level": 1,
            })

        # 已有标题样式 / 大纲级别的段落：全文识别，统一处理成大纲级别标题。
        # 已出现在 headings 中的段落（编号/章节标题）跳过，避免重复。
        seen_idx = {h["index"] for h in headings}
        for idx, para in enumerate(doc.paragraphs):
            if idx in seen_idx:
                continue
            level = TitlePreprocessor._existing_heading_level(para)
            if not level:
                continue
            text = para.text.strip()
            if not text or len(text) > 60:
                continue
            headings.append({
                "index": idx,
                "text": text,
                "number": "",
                "detected_level": level,
            })

        # 按段落索引排序，保持文档顺序
        headings.sort(key=lambda h: h["index"])
        return headings

    @staticmethod
    def _existing_heading_level(paragraph) -> Optional[int]:
        """返回段落已有的大纲级别（1-9）；非标题段落返回 None。

        优先按样式名（Heading N / 标题 N），其次按 w:outlineLvl。
        """
        if paragraph.style is not None and paragraph.style.name:
            level = HEADING_STYLE_LEVELS.get(paragraph.style.name.lower())
            if level:
                return level
        pPr = paragraph._p.find(qn('w:pPr'))
        if pPr is not None:
            outline = pPr.find(qn('w:outlineLvl'))
            if outline is not None:
                val = outline.get(qn('w:val'))
                if val is not None:
                    try:
                        level = int(val) + 1
                        if 1 <= level <= 9:
                            return level
                    except ValueError:
                        pass
        return None

    @staticmethod
    def _set_outline_level(paragraph, level: int) -> None:
        """直接设置段落的大纲级别（w:outlineLvl）。"""
        pPr = paragraph._p.get_or_add_pPr()
        outline = pPr.find(qn('w:outlineLvl'))
        if outline is None:
            outline = OxmlElement('w:outlineLvl')
            pPr.append(outline)
        outline.set(qn('w:val'), str(level - 1))  # outlineLvl 是 0-based

    @staticmethod
    def _apply_heading_style(doc, paragraph, level: int) -> None:
        """将段落应用 Heading N 样式，并设置大纲级别。

        优先使用 python-docx 的样式赋值（按样式名解析 styleId）；
        若文档 styles.xml 未定义该样式，则回退到直接写 XML（pStyle + outlineLvl），
        保证 Word 打开时能识别为内置标题样式。
        """
        style_name = f"Heading {level}"
        applied = False
        try:
            paragraph.style = doc.styles[style_name]
            applied = True
        except Exception:
            applied = False

        if not applied:
            # 回退：直接写 pStyle（内置标题 styleId 为 "Heading1"、"Heading2"...）
            pPr = paragraph._p.get_or_add_pPr()
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is None:
                pStyle = OxmlElement('w:pStyle')
                pPr.insert(0, pStyle)
            pStyle.set(qn('w:val'), f'Heading{level}')

        # 无论样式是否成功，都设置大纲级别（确保被大纲级别检测识别）
        TitlePreprocessor._set_outline_level(paragraph, level)

    @staticmethod
    def apply_headings(docx_file, output_file, selections: List[Dict], progress_callback=None) -> None:
        """将选中段落应用对应 Heading N 样式，保存为新文档。

        Args:
            docx_file: 输入 docx 路径。
            output_file: 输出 docx 路径。
            selections: List[Dict]，每个元素结构：
                {"index": int, "target_level": int}
                target_level 为 1-9；不在 selections 中的段落保持不变。
            progress_callback: 可选，每处理一个选中段落后调用
                progress_callback(done: int, total: int)。
        """
        doc = Document(docx_file)
        # 按 index 建立查找表
        sel_map = {s["index"]: s["target_level"] for s in selections if s.get("target_level", 0) >= 1}
        total = len(sel_map)
        done = 0
        for idx, para in enumerate(doc.paragraphs):
            if idx in sel_map:
                TitlePreprocessor._apply_heading_style(doc, para, sel_map[idx])
                done += 1
                if progress_callback:
                    progress_callback(done, total)
        doc.save(output_file)
        if progress_callback:
            progress_callback(total, total)

    @staticmethod
    def has_heading_styles(docx_file) -> bool:
        """判断文档是否已使用标题样式（Heading N / 标题 N / 大纲级别）。"""
        doc = Document(docx_file)
        heading_style_names = {
            'heading 1', 'heading 2', 'heading 3', 'heading 4', 'heading 5',
            'heading 6', 'heading 7', 'heading 8', 'heading 9',
            '标题 1', '标题 2', '标题 3', '标题 4', '标题 5',
            '标题 6', '标题 7', '标题 8', '标题 9',
        }
        for para in doc.paragraphs:
            if para.style and para.style.name:
                if para.style.name.lower() in heading_style_names:
                    return True
            # 检查大纲级别
            pPr = para._p.find(qn('w:pPr'))
            if pPr is not None and pPr.find(qn('w:outlineLvl')) is not None:
                return True
        return False
