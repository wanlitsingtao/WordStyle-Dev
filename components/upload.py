# -*- coding: utf-8 -*-
"""
文件上传与样式分析组件
从 app.py 提取的工具函数
"""
import streamlit as st
from pathlib import Path
import logging
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger('WordStyle')


def count_paragraphs(docx_file):
    """统计文档段落数（不包括标题）"""
    try:
        doc = Document(docx_file)
        paragraph_count = 0
        
        for para in doc.paragraphs:
            # 检查是否为标题样式
            style_name = para.style.name.lower() if para.style else ''
            
            # 排除所有标题样式（Heading 1-9）
            is_heading = (
                'heading' in style_name or
                '标题' in style_name or
                para.style.type == WD_STYLE_TYPE.PARAGRAPH and hasattr(para, 'outline_level') and para.outline_level is not None
            )
            
            # 只统计非标题段落
            if not is_heading:
                paragraph_count += 1
        
        return paragraph_count
    except:
        return 0


def get_template_styles_list(template_file):
    """获取模板文档中的所有段落样式"""
    try:
        doc = Document(template_file)
        styles = []
        for style in doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                styles.append(style.name)
        return sorted(styles)
    except:
        return ["Normal"]  # 默认返回Normal样式


def analyze_source_styles(source_files, user_id):
    """
    分析源文档样式（不显示进度条，避免布局问题）
    :param source_files: 上传的文件对象列表
    :param user_id: 用户ID
    :return: {filename: [styles]} 字典，每个文件对应其样式列表
    """
    import os
    
    file_styles_map = {}  # {filename: [styles]}
    total_files = len(source_files)
    
    for idx, source_file in enumerate(source_files, 1):
        # 保存临时文件
        temp_source = f"temp_source_{user_id}_{source_file.name}"
        try:
            with open(temp_source, 'wb') as f:
                f.write(source_file.getbuffer())
            
            # 读取样式
            doc = Document(temp_source)
            styles = set()
            
            for para in doc.paragraphs:
                if para.style and para.style.name:
                    styles.add(para.style.name)
            
            # 保存该文件的样式
            file_styles_map[source_file.name] = sorted(list(styles))
            
        except Exception as e:
            st.error(f"❌ 分析文件 {source_file.name} 失败: {e}")
            continue
    
    return file_styles_map


def count_pages(docx_file):
    """估算文档页数（基于段落数）"""
    try:
        doc = Document(docx_file)
        # 粗略估算：每50个段落约1页
        paragraphs = len(doc.paragraphs)
        estimated_pages = max(1, paragraphs // 50)
        return estimated_pages
    except:
        return 0  # 无法计算时返回0
