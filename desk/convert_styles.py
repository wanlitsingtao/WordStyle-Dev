# convert_styles_bodytext_with_resize.py (最终修正版：保留源图片尺寸，超出时缩放)
try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml, OxmlElement
    from docx.shared import Emu
    from docx.image.exceptions import UnrecognizedImageError
except ImportError:
    print("错误：未安装 python-docx 库，请运行: pip install python-docx")
    sys.exit(1)

import re
import io
import os
import sys
import argparse
import logging
import time
from datetime import datetime
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml, OxmlElement
from docx.shared import Emu
from docx.image.exceptions import UnrecognizedImageError

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    print("警告：未安装 Pillow 库，图片将无法自动调整大小。如需调整，请运行: pip install Pillow")

# ==================== 配置区域 ====================
DEFAULT_SOURCE = "source.docx"
DEFAULT_TEMPLATE = "mb.docx"
DEFAULT_OUTPUT_SUFFIX = "_styled"

DEFAULT_TARGET = "Normal"
TARGET_STYLE_FOR_TABLES = "Body Text"
TARGET_STYLE_FOR_IMAGES = "Body Text"
IMAGE_SCALE_RATIO = 2 / 3

LIST_BULLET_SYMBOL = "● "

TABLE_BORDER_SIZE = '4'
TABLE_BORDER_COLOR = '000000'

STYLE_MAP = {
    "Heading 1": "Heading 1",
    "Heading 2": "Heading 2",
    "Heading 3": "Heading 3",
    "Heading 4": "Heading 4",
    "Heading 5": "Heading 5",
    "Heading 6": "Heading 6",
    "BN_列表项目符号": "List Paragraph",
    "BN_标书应答": "应答句",
}

OUTLINE_STYLE_MAP = {
    1: "Heading 1",
    2: "Heading 2",
    3: "Heading 3",
    4: "Heading 4",
    5: "Heading 5",
    6: "Heading 6",
    7: "Heading 7",
    8: "Heading 8",
    9: "Heading 9",
}
# =================================================

HEADING_STYLES = {f"Heading {i}" for i in range(1, 10)}
LOG_SUFFIX = "_err.log"

def setup_logger(source_file):
    log_filename = os.path.splitext(source_file)[0] + LOG_SUFFIX
    logger = logging.getLogger(f"converter_{os.path.basename(source_file)}")
    logger.setLevel(logging.WARNING)
    if not logger.handlers:
        handler = logging.FileHandler(log_filename, encoding='utf-8')
        formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(filename)s - %(message)s')
        handler.setFormatter(formatter)
        logger.addHandler(handler)
    return logger

def clear_document_content(doc):
    for table in doc.tables:
        table._element.getparent().remove(table._element)
    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)
    doc.add_paragraph()

def get_outline_level(paragraph):
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        outline = pPr.find(qn('w:outlineLvl'))
        if outline is not None:
            val = outline.get(qn('w:val'))
            if val is not None:
                try:
                    return int(val) + 1
                except ValueError:
                    pass
    return 0

def has_numbering(paragraph):
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            return True
    return False

def remove_auto_numbering(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    numPr = pPr.find(qn('w:numPr'))
    if numPr is not None:
        pPr.remove(numPr)

def remove_manual_numbering(text):
    fragment_patterns = [
        r'\d+(?:\.\d+)*\.?',
        r'\([0-9]+\)',
        r'[①-⑩]',
        r'[A-Za-z]\.',
    ]
    pattern = r'^\s*(' + '|'.join(fragment_patterns) + r')[\s、，]*'
    compiled = re.compile(pattern)
    cleaned = text
    while True:
        m = compiled.match(cleaned)
        if m:
            cleaned = cleaned[m.end():]
        else:
            break
    return cleaned

def clean_list_numbering(text):
    pattern = r'^\s*(?:\d+[、\)）]|[（\(]\d+[）\)])\s*'
    cleaned = re.sub(pattern, '', text, count=1)
    return cleaned

def get_target_style(original_style_name, template_doc, logger, source_file):
    target = STYLE_MAP.get(original_style_name)
    if target is not None:
        try:
            template_doc.styles[target]
            return target
        except KeyError:
            msg = f"样式缺失: 映射样式 '{target}' 不存在，将使用 '{DEFAULT_TARGET}'"
            print(f"  警告：{msg}")
            logger.warning(f"{msg} (源文件: {source_file})")
            return DEFAULT_TARGET
    else:
        return DEFAULT_TARGET

def get_safe_style(style_name, template_doc, fallback=DEFAULT_TARGET):
    try:
        template_doc.styles[style_name]
        return style_name
    except KeyError:
        print(f"  警告：模板中缺少样式 '{style_name}'，将使用 '{fallback}'")
        return fallback

def get_image_extent(blip_element):
    """从 blip 元素向上查找 wp:inline，并返回图片的原始显示尺寸 (cx, cy) 单位 EMU，若失败则返回 (None, None)"""
    parent = blip_element.getparent()
    while parent is not None:
        if parent.tag == qn('wp:inline'):
            extent = parent.find(qn('wp:extent'))
            if extent is not None:
                cx = int(extent.get('cx', '0'))
                cy = int(extent.get('cy', '0'))
                return (cx, cy)
            break
        parent = parent.getparent()
    return (None, None)

def _add_picture(run, img_bytes, page_width_emu, available_width_emu, emu_width=None, emu_height=None):
    """
    添加图片。
    若提供了源文档的显示尺寸 (emu_width, emu_height)，则优先使用；
    当图片宽度超出可用宽度时，缩放到页面宽度的 IMAGE_SCALE_RATIO，高度等比缩放。
    如果未提供尺寸，则使用图片的像素尺寸（96 DPI 计算）并做相同处理。
    """
    if not PIL_AVAILABLE:
        run.add_picture(io.BytesIO(img_bytes))
        return

    # 如果有源尺寸，直接使用；否则从像素计算
    if emu_width is not None and emu_height is not None:
        w_emu = emu_width
        h_emu = emu_height
    else:
        try:
            img = Image.open(io.BytesIO(img_bytes))
            w_px, h_px = img.size
            w_emu = int(w_px / 96 * 914400)
            h_emu = int(h_px / 96 * 914400)
        except:
            run.add_picture(io.BytesIO(img_bytes))
            return

    # 如果宽度超出可用宽度，则按页面宽度的 IMAGE_SCALE_RATIO 缩放
    if w_emu > available_width_emu:
        target_w = int(page_width_emu * IMAGE_SCALE_RATIO)
        scale = target_w / w_emu
        new_w = int(w_emu * scale)
        new_h = int(h_emu * scale)
        run.add_picture(io.BytesIO(img_bytes), width=Emu(new_w), height=Emu(new_h))
    else:
        run.add_picture(io.BytesIO(img_bytes), width=Emu(w_emu), height=Emu(h_emu))

def set_table_width(table, width_emu):
    width_dxa = int(width_emu / 635)
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml('<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = parse_xml('<w:tblW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(width_dxa))
    tblW.set(qn('w:type'), 'dxa')

def set_table_borders(table, border_size=TABLE_BORDER_SIZE, border_color=TABLE_BORDER_COLOR):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for border_name in ['top', 'left', 'bottom', 'right']:
                existing = tcPr.find(qn(f'w:{border_name}'))
                if existing is not None:
                    tcPr.remove(existing)
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), border_size)
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), border_color)
                tcPr.append(border)

def copy_paragraph_with_images(source_para, target_doc, target_style_name,
                               page_width_emu, available_width_emu,
                               para_idx, logger, source_file):
    has_image = any(run._element.findall('.//' + qn('a:blip')) for run in source_para.runs)
    new_para = target_doc.add_paragraph()
    outline_level = get_outline_level(source_para)
    src_style_name = source_para.style.name

    if outline_level > 0:
        final_style = OUTLINE_STYLE_MAP.get(outline_level)
        if final_style is None:
            final_style = f"Heading {outline_level}"
            msg = f"未定义映射: 大纲级别 {outline_level} 未在 OUTLINE_STYLE_MAP 中配置，使用默认样式 '{final_style}'"
            print(f"  警告：{msg}")
            logger.warning(f"段落 {para_idx+1}: {msg} (源文件: {source_file})")
        try:
            target_doc.styles[final_style]
        except KeyError:
            msg = f"样式缺失: 大纲级别 {outline_level} -> 目标样式 '{final_style}' 不存在，使用 '{DEFAULT_TARGET}'"
            print(f"  警告：{msg}")
            logger.warning(f"段落 {para_idx+1}: {msg} (源文件: {source_file})")
            final_style = DEFAULT_TARGET
    else:
        if has_image:
            final_style = get_safe_style(TARGET_STYLE_FOR_IMAGES, target_doc)
        else:
            final_style = target_style_name
    try:
        new_para.style = final_style
    except Exception as e:
        msg = f"应用样式失败: '{final_style}' - {e}"
        print(f"  警告：{msg}")
        logger.warning(f"段落 {para_idx+1}: {msg} (源文件: {source_file})")
        new_para.style = target_doc.styles['Normal']

    is_heading_by_outline = outline_level > 0
    is_heading_by_style = src_style_name in HEADING_STYLES

    # ---------- 处理标题 ----------
    if is_heading_by_outline or is_heading_by_style:
        if not is_heading_by_outline:
            remove_auto_numbering(new_para)
        full_text = ''.join(run.text for run in source_para.runs)
        cleaned_text = remove_manual_numbering(full_text)
        new_para.clear()
        new_para.add_run(cleaned_text)
        for run_idx, run in enumerate(source_para.runs):
            blips = run._element.findall('.//' + qn('a:blip'))
            for blip in blips:
                rId = blip.get(qn('r:embed'))
                if rId:
                    try:
                        img_part = source_para.part.related_parts[rId]
                        img_bytes = img_part.blob
                        emu_w, emu_h = get_image_extent(blip)
                        pic_run = new_para.add_run()
                        _add_picture(pic_run, img_bytes, page_width_emu, available_width_emu, emu_w, emu_h)
                    except UnrecognizedImageError:
                        msg = f"无法识别的图片格式"
                        print(f"  警告：{msg}")
                        logger.warning(f"段落 {para_idx+1}, run {run_idx+1}: {msg}")
                    except Exception as e:
                        msg = f"图片添加失败: {e}"
                        print(f"  警告：{msg}")
                        logger.warning(f"段落 {para_idx+1}, run {run_idx+1}: {msg}")
        return new_para

    # ---------- 列表段落 ----------
    if has_numbering(source_para):
        remove_auto_numbering(new_para)
        if target_style_name == DEFAULT_TARGET:
            new_para.add_run(LIST_BULLET_SYMBOL)
            first_text_run = True
            for run in source_para.runs:
                blips = run._element.findall('.//' + qn('a:blip'))
                if blips:
                    for blip in blips:
                        rId = blip.get(qn('r:embed'))
                        if rId:
                            try:
                                img_part = source_para.part.related_parts[rId]
                                img_bytes = img_part.blob
                                emu_w, emu_h = get_image_extent(blip)
                                pic_run = new_para.add_run()
                                _add_picture(pic_run, img_bytes, page_width_emu, available_width_emu, emu_w, emu_h)
                            except UnrecognizedImageError:
                                msg = f"列表段落图片格式无法识别"
                                print(f"  警告：{msg}")
                                logger.warning(f"段落 {para_idx+1}: {msg}")
                            except Exception as e:
                                msg = f"列表段落图片添加失败: {e}"
                                print(f"  警告：{msg}")
                                logger.warning(f"段落 {para_idx+1}: {msg}")
                else:
                    if run.text:
                        text_to_add = run.text
                        if first_text_run:
                            text_to_add = clean_list_numbering(text_to_add)
                            first_text_run = False
                        if text_to_add:
                            new_para.add_run(text_to_add)
        else:
            for run in source_para.runs:
                blips = run._element.findall('.//' + qn('a:blip'))
                if blips:
                    for blip in blips:
                        rId = blip.get(qn('r:embed'))
                        if rId:
                            try:
                                img_part = source_para.part.related_parts[rId]
                                img_bytes = img_part.blob
                                emu_w, emu_h = get_image_extent(blip)
                                pic_run = new_para.add_run()
                                _add_picture(pic_run, img_bytes, page_width_emu, available_width_emu, emu_w, emu_h)
                            except UnrecognizedImageError:
                                msg = f"列表段落图片格式无法识别"
                                print(f"  警告：{msg}")
                                logger.warning(f"段落 {para_idx+1}: {msg}")
                            except Exception as e:
                                msg = f"列表段落图片添加失败: {e}"
                                print(f"  警告：{msg}")
                                logger.warning(f"段落 {para_idx+1}: {msg}")
                else:
                    if run.text:
                        new_para.add_run(run.text)
        return new_para

    # ---------- 普通段落 ----------
    for run in source_para.runs:
        blips = run._element.findall('.//' + qn('a:blip'))
        if blips:
            for blip in blips:
                rId = blip.get(qn('r:embed'))
                if rId:
                    try:
                        img_part = source_para.part.related_parts[rId]
                        img_bytes = img_part.blob
                        emu_w, emu_h = get_image_extent(blip)
                        pic_run = new_para.add_run()
                        _add_picture(pic_run, img_bytes, page_width_emu, available_width_emu, emu_w, emu_h)
                    except UnrecognizedImageError:
                        msg = f"普通段落图片格式无法识别"
                        print(f"  警告：{msg}")
                        logger.warning(f"段落 {para_idx+1}: {msg}")
                    except Exception as e:
                        msg = f"普通段落图片添加失败: {e}"
                        print(f"  警告：{msg}")
                        logger.warning(f"段落 {para_idx+1}: {msg}")
        else:
            if run.text:
                new_para.add_run(run.text)
    return new_para

def copy_table_with_images(source_table, target_doc, default_style, table_idx, logger, source_file, available_width_emu, page_width_emu):
    rows = len(source_table.rows)
    cols = len(source_table.columns)
    new_table = target_doc.add_table(rows=rows, cols=cols)
    new_table.style = source_table.style
    table_style = get_safe_style(TARGET_STYLE_FOR_TABLES, target_doc)

    set_table_width(new_table, available_width_emu)
    set_table_borders(new_table)

    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            new_cell = new_table.cell(i, j)
            new_cell._element.clear_content()
            for para_idx, para in enumerate(cell.paragraphs):
                new_para = new_cell.add_paragraph()
                new_para.style = table_style
                if has_numbering(para):
                    remove_auto_numbering(new_para)
                    new_para.add_run(LIST_BULLET_SYMBOL)
                    first_text_run = True
                    for run in para.runs:
                        blips = run._element.findall('.//' + qn('a:blip'))
                        if blips:
                            for blip in blips:
                                rId = blip.get(qn('r:embed'))
                                if rId:
                                    try:
                                        img_part = para.part.related_parts[rId]
                                        img_bytes = img_part.blob
                                        emu_w, emu_h = get_image_extent(blip)
                                        pic_run = new_para.add_run()
                                        _add_picture(pic_run, img_bytes, page_width_emu, available_width_emu, emu_w, emu_h)
                                    except Exception as e:
                                        msg = f"表格图片添加失败: {e}"
                                        print(f"  警告：{msg}")
                                        logger.warning(f"表格 {table_idx+1}, 单元格({i+1},{j+1}), 段落 {para_idx+1}: {msg}")
                        else:
                            if run.text:
                                text_to_add = run.text
                                if first_text_run:
                                    text_to_add = clean_list_numbering(text_to_add)
                                    first_text_run = False
                                if text_to_add:
                                    new_para.add_run(text_to_add)
                    continue
                for run in para.runs:
                    blips = run._element.findall('.//' + qn('a:blip'))
                    if blips:
                        for blip in blips:
                            rId = blip.get(qn('r:embed'))
                            if rId:
                                try:
                                    img_part = para.part.related_parts[rId]
                                    img_bytes = img_part.blob
                                    emu_w, emu_h = get_image_extent(blip)
                                    pic_run = new_para.add_run()
                                    _add_picture(pic_run, img_bytes, page_width_emu, available_width_emu, emu_w, emu_h)
                                except Exception as e:
                                    msg = f"表格图片添加失败: {e}"
                                    print(f"  警告：{msg}")
                                    logger.warning(f"表格 {table_idx+1}, 单元格({i+1},{j+1}), 段落 {para_idx+1}: {msg}")
                    else:
                        if run.text:
                            new_para.add_run(run.text)
    return new_table

def save_document_with_retry(doc, base_output_file, logger):
    max_retries = 10
    output_file = base_output_file
    for attempt in range(max_retries):
        try:
            doc.save(output_file)
            print(f"保存新文档到 {output_file}")
            return output_file
        except (PermissionError, OSError, IOError) as e:
            if attempt == 0:
                msg = f"保存文档失败（文件可能被占用）: {e}"
                print(f"  警告：{msg}")
                logger.warning(msg)
            base, ext = os.path.splitext(base_output_file)
            time_suffix = datetime.now().strftime("_%H%M")
            output_file = f"{base}{time_suffix}{ext}"
            print(f"  尝试备用文件名: {output_file}")
            time.sleep(0.5)
        except Exception as e:
            raise e
    raise RuntimeError(f"无法保存文档，已尝试 {max_retries} 次")

def convert_one_file(source_file, template_file, output_file):
    logger = setup_logger(source_file)
    print(f"\n正在处理: {source_file}")
    print("1. 加载模板文档...")
    try:
        template_doc = Document(template_file)
    except Exception as e:
        msg = f"加载模板文档失败: {e}"
        print(f"错误：{msg}")
        logger.error(msg)
        return
    new_doc = Document(template_file)
    clear_document_content(new_doc)

    section = new_doc.sections[0]
    page_width = section.page_width
    left_margin = section.left_margin
    right_margin = section.right_margin
    available_width = page_width - left_margin - right_margin
    print(f"页面总宽度: {page_width} EMU (约 {page_width / 914400:.2f} 英寸)")
    print(f"页面可用宽度: {available_width} EMU (约 {available_width / 914400:.2f} 英寸)")
    print(f"图片超出可用宽度时将缩放至页面宽度的 {IMAGE_SCALE_RATIO*100:.0f}%")

    print("2. 加载源文档...")
    try:
        source_doc = Document(source_file)
    except Exception as e:
        msg = f"加载源文档失败: {e}"
        print(f"错误：{msg}")
        logger.error(msg)
        return

    print("3. 开始迁移内容...")
    body = source_doc.element.body
    para_idx = 0
    table_idx = 0
    stats = {"para": 0, "table": 0, "heading": 0}

    for child in body:
        if child.tag == qn('w:p'):
            if para_idx < len(source_doc.paragraphs):
                para = source_doc.paragraphs[para_idx]
                src_style = para.style.name
                target_style = get_target_style(src_style, new_doc, logger, source_file)
                print(f"段落 {para_idx+1}: '{src_style}' -> '{target_style}'")

                new_para = copy_paragraph_with_images(
                    para, new_doc, target_style,
                    page_width, available_width,
                    para_idx, logger, source_file
                )

                if get_outline_level(para) > 0 or src_style in HEADING_STYLES:
                    stats["heading"] += 1
                stats["para"] += 1
                para_idx += 1

        elif child.tag == qn('w:tbl'):
            if table_idx < len(source_doc.tables):
                table = source_doc.tables[table_idx]
                print(f"表格 {table_idx+1}: 复制中...")
                copy_table_with_images(table, new_doc, DEFAULT_TARGET, table_idx, logger, source_file, available_width, page_width)
                stats["table"] += 1
                table_idx += 1

    print(f"统计：共处理段落 {stats['para']} 个，表格 {stats['table']} 个，标题 {stats['heading']} 个")
    
    try:
        final_output = save_document_with_retry(new_doc, output_file, logger)
        print(f"完成！最终保存为: {final_output}")
    except Exception as e:
        msg = f"保存文档最终失败: {e}"
        print(f"错误：{msg}")
        logger.error(msg)

def main():
    parser = argparse.ArgumentParser(description='将 Word 文档的样式转换为模板样式，并处理图片缩放。支持多文件。')
    parser.add_argument('inputs', nargs='*', default=None,
                        help='输入文件路径，可指定多个（若不提供则使用默认文件 source.docx）')
    parser.add_argument('-t', '--template', default=DEFAULT_TEMPLATE,
                        help=f'模板文档路径（默认: {DEFAULT_TEMPLATE}）')
    parser.add_argument('-o', '--output', default=None,
                        help='输出文档路径（仅当处理单个输入文件时有效）')
    args = parser.parse_args()

    if args.inputs:
        source_files = args.inputs
    else:
        source_files = [DEFAULT_SOURCE]

    template_file = args.template
    if not os.path.exists(template_file):
        print(f"错误：模板文件不存在 {template_file}")
        sys.exit(1)

    for sf in source_files:
        if not os.path.exists(sf):
            print(f"错误：源文件不存在 {sf}")
            sys.exit(1)

    if len(source_files) == 1:
        source_file = source_files[0]
        if args.output:
            output_file = args.output
        else:
            base, ext = os.path.splitext(source_file)
            output_file = f"{base}{DEFAULT_OUTPUT_SUFFIX}{ext}"
        convert_one_file(source_file, template_file, output_file)
    else:
        if args.output:
            print("警告：处理多个文件时 -o 参数无效，将自动生成输出文件名。")
        for source_file in source_files:
            base, ext = os.path.splitext(source_file)
            output_file = f"{base}{DEFAULT_OUTPUT_SUFFIX}{ext}"
            convert_one_file(source_file, template_file, output_file)

if __name__ == "__main__":
    main()