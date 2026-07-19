# -*- coding: utf-8 -*-
"""
文档格式转换程序 - Tkinter图形界面（优化版）
修复问题：
1. 添加独立的样式映射配置按钮
2. 支持默认映射+用户自定义
3. 界面自适应分辨率和任务栏
4. 优化性能避免卡顿
5. 自动生成输出文件，显示文件列表
"""
import tkinter as tk
from tkinter import (Tk, Frame, Label, LabelFrame, Button, Entry, Listbox, Scrollbar, 
                     Checkbutton, IntVar, StringVar, Text, messagebox, Canvas,
                     Radiobutton, filedialog, ttk, VERTICAL, HORIZONTAL, END, LEFT, RIGHT, 
                     TOP, BOTTOM, X, Y, BOTH, W, E, N, S, CENTER)
from datetime import datetime
import os
import json
import subprocess
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

# 默认配置文件路径（保存在desk目录下）
DEFAULT_CONFIG_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "default_config.json")

from doc_converter import DocumentConverter


class StyleMappingDialog:
    """样式映射对话框（新版：四步配置流程）

    步骤1: 标题样式映射（统一配置，不区分原文/应答句）
    步骤2: 应答句配置（根据模式决定后续单列/双列）
    步骤3: 样式映射（根据应答句模式动态切换单/双列）
    步骤4: 表格图片兜底 + 列表段落兜底（也分两套）
    """

    def __init__(self, parent, source_styles, template_styles, current_mapping=None,
                 saved_default_mapping=None, save_default_callback=None,
                 current_tbl_img_config=None, current_remove_chapter_label=0,
                 current_answer_config=None, current_list_config=None):
        self.parent = parent
        self.all_source_styles = sorted(source_styles)
        self.template_styles = sorted(template_styles)
        self.current_mapping = current_mapping or {}
        self.saved_default_mapping = saved_default_mapping or {}
        self.current_tbl_img_config = current_tbl_img_config or {}
        self.current_remove_chapter_label = current_remove_chapter_label
        self.current_answer_config = current_answer_config or {}
        self.current_list_config = current_list_config or {}
        self.save_default_callback = save_default_callback
        self.result = None

        # 分离标题样式和正文样式
        self.heading_styles = []
        self.body_styles = []
        for s in self.all_source_styles:
            is_heading = False
            if s.startswith('[大纲级别') or s.startswith('[Outline'):
                is_heading = True
            for i in range(1, 10):
                if s == f'Heading {i}' or s == f'heading {i}' or s == f'Heading{i}':
                    is_heading = True
                    break
            if is_heading:
                self.heading_styles.append(s)
            else:
                self.body_styles.append(s)

        self.dialog = None
        self.body_mapping_widgets = []
        self.answer_mapping_widgets = []
        # 表格/图片双列控件
        self.table_style_combo2 = None
        self.image_style_combo2 = None
        # 列表段落双列控件
        self.list_style_combo2 = None

        self.create_dialog()

    def create_dialog(self):
        """创建对话框界面"""
        self.dialog = tk.Toplevel(self.parent)
        self.dialog.title("样式映射配置")
        screen_width = self.dialog.winfo_screenwidth()
        screen_height = self.dialog.winfo_screenheight()
        # 固定大小，完整显示所有步骤 + 底部按钮（步骤1/3内部滚动）
        dialog_width = min(1280, screen_width - 20)
        dialog_height = min(int(screen_height * 0.85), screen_height - 40)
        dialog_height = max(dialog_height, 600)
        self.dialog.minsize(1050, 550)
        x = max(0, (screen_width - dialog_width) // 2)
        y = max(0, (screen_height - dialog_height) // 2 - 15)
        self.dialog.geometry(f'{dialog_width}x{dialog_height}+{x}+{y}')
        self.dialog.transient(self.parent)
        self.dialog.grab_set()
        # 确保对话框在显示前完成布局计算
        self.dialog.update_idletasks()

        # ===== 主容器：使用普通 Frame + pack fill，不再整体滚动 =====
        # 步骤1/3各自内部有独立滚动，主窗口固定显示所有步骤和按钮
        main_container = Frame(self.dialog)
        main_container.pack(fill=BOTH, expand=True)

        # 使用普通 Frame 作为主内容容器（无需滚动）
        scrollable_main = Frame(main_container, borderwidth=0)
        scrollable_main.pack(fill=BOTH, expand=True)

        # ===== 步骤1: 标题样式映射（内部可滚动） =====
        step1_frame = LabelFrame(scrollable_main, text="1. 标题样式映射（统一，不分原文/应答句）",
                                 font=("微软雅黑", 10, "bold"), padx=3, pady=2, fg="#1565C0")
        step1_frame.pack(fill=X, padx=5, pady=(1, 1))
        if not self.heading_styles:
            Label(step1_frame, text="（当前源文档未检测到标题样式）",
                  font=("微软雅黑", 9), fg="black").pack(anchor=W, pady=1)
        else:
            # 步骤1内部使用可滚动容器：固定高度，内容多时滚动
            hdr = Frame(step1_frame)
            hdr.pack(fill=X, pady=(1, 1))
            Label(hdr, text="源标题样式", width=26, font=("微软雅黑", 9, "bold"), anchor=W).pack(side=LEFT, padx=2)
            Label(hdr, text="→", width=2, anchor=CENTER).pack(side=LEFT)
            Label(hdr, text="目标样式", width=26, font=("微软雅黑", 9, "bold"), anchor=W).pack(side=LEFT, padx=2)

            # 标题列表滚动容器
            s1_canvas = Canvas(step1_frame, highlightthickness=0, borderwidth=0, height=70)
            s1_scrollbar = Scrollbar(step1_frame, orient=VERTICAL, command=s1_canvas.yview)
            s1_canvas.configure(yscrollcommand=s1_scrollbar.set)
            s1_canvas.pack(side=LEFT, fill=BOTH, expand=True)
            s1_scrollbar.pack(side=RIGHT, fill=Y)
            s1_inner = Frame(s1_canvas, borderwidth=0)
            s1_window_id = s1_canvas.create_window((0, 0), window=s1_inner, anchor="nw")
            def _s1_configure(event):
                s1_canvas.configure(scrollregion=s1_canvas.bbox("all"))
                cw = s1_canvas.winfo_width()
                if cw > 20:
                    s1_canvas.itemconfig(s1_window_id, width=cw)
            s1_inner.bind("<Configure>", _s1_configure)
            # 鼠标滚轮
            def _s1_scroll(event):
                s1_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
            s1_canvas.bind("<MouseWheel>", _s1_scroll)
            s1_inner.bind("<MouseWheel>", _s1_scroll)

            self.heading_widgets = []
            for src_style in self.heading_styles:
                row = Frame(s1_inner)
                row.pack(fill=X, pady=0)
                Label(row, text=src_style, width=26, anchor=W, font=("微软雅黑", 9)).pack(side=LEFT, padx=2)
                Label(row, text="→", width=2, anchor=CENTER).pack(side=LEFT)
                default_val = self.current_mapping.get(src_style, self.saved_default_mapping.get(src_style, src_style if src_style in self.template_styles else "Normal"))
                if default_val not in self.template_styles:
                    default_val = src_style if src_style in self.template_styles else "Normal"
                var = StringVar(value=default_val)
                combo = ttk.Combobox(row, textvariable=var, width=26, state="readonly")
                combo['values'] = self.template_styles
                combo.pack(side=LEFT, padx=3)
                self.heading_widgets.append((src_style, var))

        # ===== 步骤2: 应答句配置 =====
        step2_frame = LabelFrame(scrollable_main, text="2. 应答句配置",
                                 font=("微软雅黑", 10, "bold"), padx=5, pady=2, fg="#1565C0")
        step2_frame.pack(fill=X, padx=5, pady=1)
        # 第1行：启用 + 文本（一行）
        r0 = Frame(step2_frame)
        r0.pack(fill=X, pady=0)
        self.do_answer_var = IntVar(value=self.current_answer_config.get('do_answer', 0))
        Checkbutton(r0, text="插入模式", variable=self.do_answer_var,
                    font=("微软雅黑", 9), command=self.toggle_answer_section).pack(side=LEFT, padx=(0, 3))
        Label(r0, text="文本:", width=4, anchor=W, font=("微软雅黑", 9)).pack(side=LEFT)
        saved_answer_text = self.current_answer_config.get('answer_text', '应答：本投标人理解并满足要求。')
        self.answer_text_var = StringVar(value=saved_answer_text)
        self.answer_text_entry = Entry(r0, textvariable=self.answer_text_var, font=("微软雅黑", 9))
        self.answer_text_entry.pack(side=LEFT, fill=X, expand=True, padx=3)

        # 第2行：答样式 + 插入模式 + 原文格式 + 应答原文格式（一行）
        r1 = Frame(step2_frame)
        r1.pack(fill=X, pady=0)
        Label(r1, text="答样式:", width=7, anchor=W, font=("微软雅黑", 9)).pack(side=LEFT)
        saved_answer_style = self.current_answer_config.get('answer_style', '应答句')
        self.answer_style_var = StringVar(value=saved_answer_style)
        self.answer_style_combo = ttk.Combobox(r1, textvariable=self.answer_style_var, width=22, state="readonly")
        self.answer_style_combo['values'] = self.template_styles
        self.answer_style_combo.pack(side=LEFT, padx=1)

        Label(r1, text="插入模式:", width=7, anchor=W, font=("微软雅黑", 9)).pack(side=LEFT, padx=(5, 0))
        self.answer_mode_labels = ["章节标题后插入", "章节末尾插入", "原文+应答句+应答原文", "逐段前插入", "逐段后插入"]
        saved_answer_mode = self.current_answer_config.get('answer_mode', self.answer_mode_labels[2])
        if saved_answer_mode not in self.answer_mode_labels:
            saved_answer_mode = self.answer_mode_labels[2]
        self.answer_mode_var = StringVar(value=saved_answer_mode)
        self.answer_mode_combo = ttk.Combobox(r1, textvariable=self.answer_mode_var, width=22, state="readonly")
        self.answer_mode_combo['values'] = self.answer_mode_labels
        self.answer_mode_combo.bind('<<ComboboxSelected>>', lambda e: self.on_answer_mode_change())
        self.answer_mode_combo.pack(side=LEFT, padx=1)

        Label(r1, text="原文:", width=5, anchor=W, font=("微软雅黑", 9)).pack(side=LEFT, padx=(5, 0))
        saved_answer_source = self.current_answer_config.get('answer_source_style', '')
        self.answer_source_var = StringVar(value=saved_answer_source)
        self.answer_source_combo = ttk.Combobox(r1, textvariable=self.answer_source_var, width=22, state="readonly")
        self.answer_source_combo['values'] = self.template_styles
        self.answer_source_combo.pack(side=LEFT, padx=1)

        Label(r1, text="答原文:", width=6, anchor=W, font=("微软雅黑", 9)).pack(side=LEFT, padx=(3, 0))
        saved_answer_copy = self.current_answer_config.get('answer_copy_style', '')
        self.answer_copy_var = StringVar(value=saved_answer_copy)
        self.answer_copy_combo = ttk.Combobox(r1, textvariable=self.answer_copy_var, width=22, state="readonly")
        self.answer_copy_combo['values'] = self.template_styles
        self.answer_copy_combo.pack(side=LEFT, padx=1)

        # ===== 步骤3: 样式映射（内部可滚动） =====
        self.step3_frame = LabelFrame(scrollable_main, text="3. 样式映射（正文+列表段落）",
                                      font=("微软雅黑", 10, "bold"), padx=3, pady=2, fg="#1565C0")
        self.step3_frame.pack(fill=X, padx=5, pady=1)
        # 步骤3内部使用可滚动容器
        self.s3_canvas = Canvas(self.step3_frame, highlightthickness=0, borderwidth=0, height=100)
        self.s3_scrollbar = Scrollbar(self.step3_frame, orient=VERTICAL, command=self.s3_canvas.yview)
        self.s3_canvas.configure(yscrollcommand=self.s3_scrollbar.set)
        self.s3_canvas.pack(side=LEFT, fill=BOTH, expand=True)
        self.s3_scrollbar.pack(side=RIGHT, fill=Y)
        self.step3_container = Frame(self.s3_canvas, borderwidth=0)
        self.s3_window_id = self.s3_canvas.create_window((0, 0), window=self.step3_container, anchor="nw")
        def _s3_configure(event):
            self.s3_canvas.configure(scrollregion=self.s3_canvas.bbox("all"))
            cw = self.s3_canvas.winfo_width()
            if cw > 20:
                self.s3_canvas.itemconfig(self.s3_window_id, width=cw)
        self.step3_container.bind("<Configure>", _s3_configure)
        def _s3_scroll(event):
            self.s3_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        self.s3_canvas.bind("<MouseWheel>", _s3_scroll)
        self.step3_container.bind("<MouseWheel>", _s3_scroll)
        # 应答句配置完成后再重建样式映射
        self.toggle_answer_section()

        # ===== 步骤4: 表格/图片/列表兜底 =====
        step4_frame = LabelFrame(scrollable_main, text="4. 表格/图片/列表兜底配置",
                                 font=("微软雅黑", 10, "bold"), padx=3, pady=2, fg="#1565C0")
        step4_frame.pack(fill=X, padx=5, pady=1)

        # --- 第0行：列标题（原文 / 应答原文） ---
        thf = Frame(step4_frame)
        thf.pack(fill=X, pady=0)
        Label(thf, text="", font=("微软雅黑", 9), width=8, anchor=W).pack(side=LEFT)
        Label(thf, text="原文", font=("微软雅黑", 9, "bold"), fg="black", width=15, anchor=CENTER).pack(side=LEFT)
        Label(thf, text="应答原文", font=("微软雅黑", 9, "bold"), fg="black", width=15, anchor=CENTER).pack(side=LEFT)
        Label(thf, text="  图片", font=("微软雅黑", 9, "bold"), fg="black", anchor=W).pack(side=LEFT, padx=(18, 1))
        Label(thf, text="原文", font=("微软雅黑", 8), width=15, anchor=CENTER).pack(side=LEFT)
        Label(thf, text="应答原文", font=("微软雅黑", 8), width=15, anchor=CENTER).pack(side=LEFT)

        # --- 第1行：表格 + 图片（同一行） ---
        tif = Frame(step4_frame)
        tif.pack(fill=X, pady=0)
        # 表格
        self.enable_table_style_var = IntVar(value=self.current_tbl_img_config.get('enable_table_style', 0))
        Checkbutton(tif, text="表格", variable=self.enable_table_style_var,
                    font=("微软雅黑", 9), command=self.toggle_table_style).pack(side=LEFT, padx=(0, 1))
        tsd = self.current_tbl_img_config.get('table_style', 'Body Text')
        self.table_style_var = StringVar(value=tsd)
        self.table_style_combo = ttk.Combobox(tif, textvariable=self.table_style_var, width=18, state="readonly")
        self.table_style_combo['values'] = self.template_styles
        self.table_style_combo.pack(side=LEFT, padx=1)
        tad = self.current_tbl_img_config.get('table_answer_style', tsd)
        self.table_answer_var = StringVar(value=tad)
        self.table_style_combo2 = ttk.Combobox(tif, textvariable=self.table_answer_var, width=18, state="readonly")
        self.table_style_combo2['values'] = self.template_styles
        self.table_style_combo2.pack(side=LEFT, padx=1)
        Label(tif, text="  ", font=("微软雅黑", 9)).pack(side=LEFT)
        self.enable_image_style_var = IntVar(value=self.current_tbl_img_config.get('enable_image_style', 0))
        Checkbutton(tif, text="图片", variable=self.enable_image_style_var,
                    font=("微软雅黑", 9), command=self.toggle_image_style).pack(side=LEFT, padx=(0, 1))
        isd = self.current_tbl_img_config.get('image_style', 'Body Text')
        self.image_style_var = StringVar(value=isd)
        self.image_style_combo = ttk.Combobox(tif, textvariable=self.image_style_var, width=18, state="readonly")
        self.image_style_combo['values'] = self.template_styles
        self.image_style_combo.pack(side=LEFT, padx=1)
        iad = self.current_tbl_img_config.get('image_answer_style', isd)
        self.image_answer_var = StringVar(value=iad)
        self.image_style_combo2 = ttk.Combobox(tif, textvariable=self.image_answer_var, width=18, state="readonly")
        self.image_style_combo2['values'] = self.template_styles
        self.image_style_combo2.pack(side=LEFT, padx=1)
        self.toggle_table_style()
        self.toggle_image_style()

        # --- 第2行：列表段落（原文+应答 同一行，标签+控件全部在同一行） ---
        lf = Frame(step4_frame)
        lf.pack(fill=X, pady=(1, 0))
        # 标签：列表段落（未映射）
        Label(lf, text="列表段落（未映射）", font=("微软雅黑", 9, "bold"), fg="black", width=14, anchor=W).pack(side=LEFT)
        # 原文控件（符号/样式在同一行）
        self.list_method_var = StringVar(value=self.current_list_config.get('method', 'bullet'))
        Radiobutton(lf, text="符号", variable=self.list_method_var, value="bullet",
                    font=("微软雅黑", 9), command=self.toggle_list_method).pack(side=LEFT, padx=0)
        sld = self.current_list_config.get('bullet', '* ')
        self.list_bullet_var = StringVar(value=sld)
        self.list_bullet_entry = Entry(lf, textvariable=self.list_bullet_var, width=4, font=("微软雅黑", 9))
        self.list_bullet_entry.pack(side=LEFT, padx=1)
        Radiobutton(lf, text="样式", variable=self.list_method_var, value="style",
                    font=("微软雅黑", 9), command=self.toggle_list_method).pack(side=LEFT, padx=(3, 0))
        slsd = self.current_list_config.get('style', 'Body Text')
        self.list_style_var = StringVar(value=slsd)
        self.list_style_combo = ttk.Combobox(lf, textvariable=self.list_style_var, width=18, state="readonly")
        self.list_style_combo['values'] = self.template_styles
        self.list_style_combo.pack(side=LEFT, padx=1)
        self.toggle_list_method()
        # 应答原文控件
        Label(lf, text="", width=2).pack(side=LEFT)
        saved_list_method2 = self.current_list_config.get('answer_method', self.list_method_var.get())
        self.list_method_var2 = StringVar(value=saved_list_method2)
        Radiobutton(lf, text="符号", variable=self.list_method_var2, value="bullet",
                    font=("微软雅黑", 9), command=self.toggle_list_method2).pack(side=LEFT, padx=0)
        sld2 = self.current_list_config.get('answer_bullet', self.list_bullet_var.get())
        self.list_bullet_var2 = StringVar(value=sld2)
        self.list_bullet_entry2 = Entry(lf, textvariable=self.list_bullet_var2, width=4, font=("微软雅黑", 9))
        self.list_bullet_entry2.pack(side=LEFT, padx=1)
        Radiobutton(lf, text="样式", variable=self.list_method_var2, value="style",
                    font=("微软雅黑", 9), command=self.toggle_list_method2).pack(side=LEFT, padx=(3, 0))
        slsd2 = self.current_list_config.get('answer_style', self.list_style_var.get())
        self.list_style_var2 = StringVar(value=slsd2)
        self.list_style_combo2 = ttk.Combobox(lf, textvariable=self.list_style_var2, width=18, state="readonly")
        self.list_style_combo2['values'] = self.template_styles
        self.list_style_combo2.pack(side=LEFT, padx=1)
        self.toggle_list_method2()

        # 分隔线 + 清除章节编号（放在列表段落之后）
        sep = Frame(step4_frame, height=1, bg="#DDDDDD")
        sep.pack(fill=X, pady=(1, 0))
        cr = Frame(step4_frame)
        cr.pack(fill=X, pady=0)
        self.remove_chapter_label_var = IntVar(value=self.current_remove_chapter_label)
        Checkbutton(cr, text='清除标题中的"第X章/第X节"等字样',
                    variable=self.remove_chapter_label_var, font=("微软雅黑", 9)).pack(side=LEFT)

        # 对话框初始状态同步：根据 do_answer 和 answer_mode 设置 Step 4 应答列状态
        do_answer = (self.do_answer_var.get() == 1)
        is_dual = (do_answer and self.answer_mode_var.get() == "原文+应答句+应答原文")
        self._update_dual_mode_widgets(is_dual)

        # ===== 底部按钮 =====
        bf = Frame(scrollable_main)
        bf.pack(fill=X, padx=5, pady=(2, 4))
        bf.columnconfigure(0, weight=1)
        bf.columnconfigure(1, weight=1)
        bf.columnconfigure(2, weight=1)
        Button(bf, text="恢复默认", width=10, font=("微软雅黑", 9), command=self.reset_defaults).grid(row=0, column=0, padx=2)
        Button(bf, text="设为默认", width=10, font=("微软雅黑", 9), command=self.save_as_default).grid(row=0, column=1, padx=2)
        Button(bf, text="确认", width=10, font=("微软雅黑", 9, "bold"), command=self.confirm).grid(row=0, column=2, padx=2)

        # 强制刷新布局
        self.dialog.update_idletasks()

    # ==================== 回调方法 ====================

    def toggle_answer_section(self):
        """根据应答句启用状态切换相关控件的可用性"""
        enabled = self.do_answer_var.get() == 1
        state = "normal" if enabled else "disabled"
        self.answer_text_entry.config(state=state)
        self.answer_style_combo.config(state=state)
        self.answer_mode_combo.config(state=state)
        self.answer_source_combo.config(state=state)
        self.answer_copy_combo.config(state=state)
        # Step 4 应答原文列控件状态
        is_dual = (enabled and self.answer_mode_var.get() == "原文+应答句+应答原文")
        answer_state = "normal" if (enabled and is_dual) else "disabled"
        # 表格应答列
        if hasattr(self, 'table_style_combo2') and self.table_style_combo2:
            if self.enable_table_style_var.get() == 1 and enabled and is_dual:
                self.table_style_combo2.config(state="normal")
            else:
                self.table_style_combo2.config(state="disabled")
        # 图片应答列
        if hasattr(self, 'image_style_combo2') and self.image_style_combo2:
            if self.enable_image_style_var.get() == 1 and enabled and is_dual:
                self.image_style_combo2.config(state="normal")
            else:
                self.image_style_combo2.config(state="disabled")
        # 列表段落应答控件
        self._set_list_answer_state(answer_state)
        self.rebuild_style_mapping()

    def on_answer_mode_change(self):
        """应答句模式变化时重建样式映射"""
        mode = self.answer_mode_var.get()
        is_dual = (mode == "原文+应答句+应答原文" and self.do_answer_var.get() == 1)
        state = "normal" if is_dual else "disabled"
        self.answer_source_combo.config(state=state)
        self.answer_copy_combo.config(state=state)
        self._update_dual_mode_widgets(is_dual)
        self.rebuild_style_mapping()

    def _update_dual_mode_widgets(self, is_dual):
        """更新双列模式下第二列控件的状态"""
        state = "normal" if is_dual else "disabled"
        if self.table_style_combo2:
            tbl_enabled = self.enable_table_style_var.get() == 1
            self.table_style_combo2.config(state="normal" if (tbl_enabled and is_dual) else "disabled")
        if self.image_style_combo2:
            img_enabled = self.enable_image_style_var.get() == 1
            self.image_style_combo2.config(state="normal" if (img_enabled and is_dual) else "disabled")
        self._set_list_answer_state(state)

    def _set_list_answer_state(self, state):
        """设置列表段落应答控件的可用状态"""
        if hasattr(self, 'list_style_combo2') and self.list_style_combo2:
            is_style = self.list_method_var2.get() == "style"
            self.list_style_combo2.config(state="normal" if (state == "normal" and is_style) else "disabled")
        if hasattr(self, 'list_bullet_entry2') and self.list_bullet_entry2:
            is_bullet = self.list_method_var2.get() == "bullet"
            self.list_bullet_entry2.config(state="normal" if (state == "normal" and is_bullet) else "disabled")
        # RadioButton 本身不做 disabled，因为 value 仍然需要可读，但它们的容器状态由外部控制

    def rebuild_style_mapping(self):
        """重建样式映射部分（根据应答句模式切换单列/双列）"""
        for widget in self.step3_container.winfo_children():
            widget.destroy()
        self.body_mapping_widgets = []
        self.answer_mapping_widgets = []
        is_dual = (self.do_answer_var.get() == 1 and self.answer_mode_var.get() == "原文+应答句+应答原文")
        if not self.body_styles:
            # 显示更醒目的提示，避免用户误以为"完全空白"
            frame = Frame(self.step3_container, bg="#FFF3E0")
            frame.pack(fill=X, pady=5, padx=5)
            Label(frame, text="⚠ 当前源文档未检测到正文样式",
                  font=("微软雅黑", 9, "bold"), fg="#E65100", bg="#FFF3E0").pack(anchor=W, pady=2)
            Label(frame, text="请确认已加载源文档且文档中包含除标题外的段落样式（如 Normal、Body Text 等）。",
                  font=("微软雅黑", 8), fg="#BF360C", bg="#FFF3E0").pack(anchor=W, pady=1)
            return
        hdr = Frame(self.step3_container)
        hdr.pack(fill=X, pady=(2, 1))
        cw = 24 if is_dual else 28
        Label(hdr, text="源样式", width=cw, font=("微软雅黑", 9, "bold"), anchor=W).pack(side=LEFT, padx=3)
        Label(hdr, text=" > ", width=3, anchor=CENTER).pack(side=LEFT)
        if is_dual:
            Label(hdr, text="原文目标样式", width=cw, font=("微软雅黑", 9, "bold"), anchor=W).pack(side=LEFT, padx=3)
            Label(hdr, text="应答目标样式", width=cw, font=("微软雅黑", 9, "bold"), anchor=W).pack(side=LEFT, padx=3)
        else:
            Label(hdr, text="目标样式", width=cw, font=("微软雅黑", 9, "bold"), anchor=W).pack(side=LEFT, padx=3)
        for src_style in self.body_styles:
            row = Frame(self.step3_container)
            row.pack(fill=X, pady=0)
            Label(row, text=src_style, width=cw, anchor=W, font=("微软雅黑", 9)).pack(side=LEFT, padx=3)
            Label(row, text=" > ", width=3, anchor=CENTER).pack(side=LEFT)
            if is_dual:
                default_val = self.current_mapping.get(src_style, self.saved_default_mapping.get(src_style, "Body Text"))
                if default_val not in self.template_styles:
                    default_val = "Body Text"
                var = StringVar(value=default_val)
                combo = ttk.Combobox(row, textvariable=var, width=cw, state="readonly")
                combo['values'] = self.template_styles
                combo.pack(side=LEFT, padx=3)
                self.body_mapping_widgets.append((src_style, var))
                answer_key = f"answer_{src_style}"
                answer_default = self.current_mapping.get(answer_key, self.saved_default_mapping.get(answer_key, default_val))
                if answer_default not in self.template_styles:
                    answer_default = default_val
                answer_var = StringVar(value=answer_default)
                answer_combo = ttk.Combobox(row, textvariable=answer_var, width=cw, state="readonly")
                answer_combo['values'] = self.template_styles
                answer_combo.pack(side=LEFT, padx=3)
                self.answer_mapping_widgets.append((src_style, answer_var))
            else:
                default_val = self.current_mapping.get(src_style, self.saved_default_mapping.get(src_style, src_style if src_style in self.template_styles else "Body Text"))
                if default_val not in self.template_styles:
                    default_val = src_style if src_style in self.template_styles else "Body Text"
                var = StringVar(value=default_val)
                combo = ttk.Combobox(row, textvariable=var, width=cw, state="readonly")
                combo['values'] = self.template_styles
                combo.pack(side=LEFT, padx=3)
                self.body_mapping_widgets.append((src_style, var))

    def toggle_table_style(self):
        """切换表格样式控件的启用状态"""
        enabled = self.enable_table_style_var.get() == 1
        self.table_style_combo.config(state="normal" if enabled else "disabled")
        if self.table_style_combo2:
            is_dual = (self.do_answer_var.get() == 1 and self.answer_mode_var.get() == "原文+应答句+应答原文")
            self.table_style_combo2.config(state="normal" if (enabled and is_dual) else "disabled")

    def toggle_image_style(self):
        """切换图片样式控件的启用状态"""
        enabled = self.enable_image_style_var.get() == 1
        self.image_style_combo.config(state="normal" if enabled else "disabled")
        if self.image_style_combo2:
            is_dual = (self.do_answer_var.get() == 1 and self.answer_mode_var.get() == "原文+应答句+应答原文")
            self.image_style_combo2.config(state="normal" if (enabled and is_dual) else "disabled")

    def toggle_list_method(self):
        """切换原文列表段落处理方式（默认符号/指定样式）"""
        is_style = self.list_method_var.get() == "style"
        self.list_style_combo.config(state="normal" if is_style else "disabled")
        self.list_bullet_entry.config(state="normal" if not is_style else "disabled")

    def toggle_list_method2(self):
        """切换应答句列表段落处理方式（默认符号/指定样式）"""
        is_style = self.list_method_var2.get() == "style"
        is_dual = (self.do_answer_var.get() == 1 and self.answer_mode_var.get() == "原文+应答句+应答原文")
        self.list_style_combo2.config(state="normal" if (is_style and is_dual) else "disabled")
        self.list_bullet_entry2.config(state="normal" if (not is_style and is_dual) else "disabled")

    def save_as_default(self):
        """保存当前配置为默认"""
        mapping = self._collect_mapping()
        answer_config = self._collect_answer_config()
        tbl_img_config = self._collect_tbl_img_config()
        list_config = self._collect_list_config()
        remove_chapter_label = self.remove_chapter_label_var.get()
        if self.save_default_callback:
            self.save_default_callback(mapping, answer_config, tbl_img_config, list_config, remove_chapter_label)

    def reset_defaults(self):
        """恢复默认配置"""
        if self.saved_default_mapping:
            self.current_mapping = dict(self.saved_default_mapping)
        # 重新创建所有控件的值
        # 简化处理：关闭对话框并让调用者重新打开
        messagebox.showinfo("提示", "已恢复默认配置，请关闭对话框重新打开以生效。")

    def confirm(self):
        """确认并返回所有配置"""
        mapping = self._collect_mapping()
        answer_config = self._collect_answer_config()
        tbl_img_config = self._collect_tbl_img_config()
        list_config = self._collect_list_config()
        remove_chapter_label = self.remove_chapter_label_var.get()
        self.result = (mapping, answer_config, tbl_img_config, list_config, remove_chapter_label)
        self.dialog.destroy()

    def show(self):
        """显示对话框并等待结果"""
        self.dialog.wait_window()
        if self.result:
            return self.result
        return None, None, None, None, None

    def _collect_mapping(self):
        """收集样式映射结果"""
        mapping = {}
        # 标题
        if hasattr(self, 'heading_widgets'):
            for src, var in self.heading_widgets:
                if var.get():
                    mapping[src] = var.get()
        # 正文
        for src, var in self.body_mapping_widgets:
            if var.get():
                mapping[src] = var.get()
        # 应答句正文（双列模式）
        for src, var in self.answer_mapping_widgets:
            if var.get():
                mapping[f"answer_{src}"] = var.get()
        return mapping

    def _collect_answer_config(self):
        """收集应答句配置"""
        return {
            'do_answer': self.do_answer_var.get(),
            'answer_text': self.answer_text_var.get(),
            'answer_style': self.answer_style_var.get(),
            'answer_mode': self.answer_mode_var.get(),
            'answer_source_style': self.answer_source_var.get(),
            'answer_copy_style': self.answer_copy_var.get(),
        }

    def _collect_tbl_img_config(self):
        """收集表格/图片兜底配置"""
        do_answer = (self.do_answer_var.get() == 1 and self.answer_mode_var.get() == "原文+应答句+应答原文")
        cfg = {
            'enable_table_style': self.enable_table_style_var.get(),
            'table_style': self.table_style_var.get(),
            'enable_image_style': self.enable_image_style_var.get(),
            'image_style': self.image_style_var.get(),
        }
        if self.table_answer_var and do_answer:
            cfg['table_answer_style'] = self.table_answer_var.get()
        if self.image_answer_var and do_answer:
            cfg['image_answer_style'] = self.image_answer_var.get()
        return cfg

    def _collect_list_config(self):
        """收集列表段落兜底配置"""
        do_answer = (self.do_answer_var.get() == 1 and self.answer_mode_var.get() == "原文+应答句+应答原文")
        cfg = {
            'method': self.list_method_var.get(),
            'bullet': self.list_bullet_var.get(),
            'style': self.list_style_var.get(),
        }
        if do_answer:
            cfg['answer_method'] = self.list_method_var2.get()
            cfg['answer_bullet'] = self.list_bullet_var2.get()
            cfg['answer_style'] = self.list_style_var2.get()
        return cfg


class DocumentConverterGUI:
    """文档转换器主界面（优化版）"""
    
    def __init__(self, root):
        self.root = root
        self.root.title("Word文档格式转换工具")
        self._apply_ui_theme()
        
        # 设置窗口为最大化状态，自动适应屏幕并考虑任务栏
        try:
            # Windows下使用wm_state设置最大化
            self.root.state('zoomed')
        except Exception:
            # 其他系统使用geometry最大化
            screen_width = root.winfo_screenwidth()
            screen_height = root.winfo_screenheight()
            self.root.geometry(f'{screen_width}x{screen_height}+0+0')
        
        # 设置最小窗口大小，兼容本机分辨率与 1920×1080
        self.root.minsize(1180, 720)
        
        self.converter = DocumentConverter()
        
        # 加载默认配置
        self.default_config = self._load_default_config()
        
        # 变量
        self.source_files = []  # 源文件列表
        self.template_file = StringVar()  # 模板文件路径
        self.do_mood_conversion = IntVar(value=1)  # 是否进行语气转换
        self.use_word_com = IntVar(value=0)  # 是否使用 Word COM 转换（保留 Visio 图）
        # 应答句样式和列表段落符号等配置已迁移到样式映射对话框中管理
        # 以下变量仅用于 full_convert 调用时的参数传递（从 file_style_maps 获取）
        self.list_bullet = StringVar(value="● ")  # 保留默认值，但实际从配置中读取
        
        # 应答句插入模式（显示标签 -> 模式值映射）
        self.answer_mode_options = {
            "章节标题后插入": "before_heading",
            "章节末尾插入": "after_heading",
            "原文+应答句+应答原文": "copy_chapter",
            "逐段前插入": "before_paragraph",
            "逐段后插入": "after_paragraph"
        }
        self.answer_mode_labels = list(self.answer_mode_options.keys())
        
        # 章节提示语配置（从 default_config.json 加载完整配置）
        saved_do_hint = self.default_config.get("do_hint_insertion", 0)
        saved_hint_type = self.default_config.get("hint_type", "text")
        saved_hint_text = self.default_config.get("hint_text", "招标文件原文")
        saved_hint_image_path = self.default_config.get("hint_image_path", "")
        saved_hint_style = self.default_config.get("hint_style", "Normal")
        self.do_hint_insertion = IntVar(value=saved_do_hint)  # 是否插入章节提示语
        self.hint_type = StringVar(value=saved_hint_type)  # 提示语类型：text 或 image
        self.hint_text = StringVar(value=saved_hint_text)  # 提示语文本
        self.hint_image_path = StringVar(value=saved_hint_image_path)  # 提示语图片路径
        self.hint_style = StringVar(value=saved_hint_style)  # 提示语样式
        
        # 样式信息
        self.source_styles = set()
        self.template_styles = set()
        self.custom_style_map = {}  # 用户自定义的样式映射（全局）
        self.custom_tbl_img_config = {}  # 用户自定义的表格/图片样式配置（全局，当file_style_maps中没有时使用）
        self.custom_remove_chapter_label = 0  # 用户自定义的清除章节编号配置（全局）
        self.custom_answer_config = {}  # 用户自定义的应答句配置（全局）
        self.custom_list_config = {}  # 用户自定义的列表配置（全局）
        self.remove_chapter_label = self.default_config.get('remove_chapter_label', 0)  # 当前清除章节编号配置
        self.default_style_map = self.default_config.get("style_map", {})  # 用户保存的默认样式映射
        self.file_style_maps = {}  # 每个文件的独立映射配置 {file_path: {style_map}}
        self.selected_source_file = None  # 当前选中的源文件
        
        # 输出文件列表
        self.output_files = []
        
        # 转换控制
        self.conversion_thread = None  # 转换线程
        self.stop_conversion_flag = False  # 停止转换标志
        
        self.create_widgets()

    def _apply_ui_theme(self):
        """统一桌面端界面主题，提升本机分辨率与 1920×1080 的视觉适配。"""
        self.root.configure(bg="#f4f7fb")
        self.root.option_add("*Font", "微软雅黑 10")
        self.root.option_add("*Background", "#f4f7fb")
        self.root.option_add("*Foreground", "#23395d")

        style = ttk.Style(self.root)
        try:
            style.theme_use("clam")
        except Exception:
            pass
        style.configure("TFrame", background="#f4f7fb")
        style.configure("TLabelframe", background="#f4f7fb", bordercolor="#d6e2f0")
        style.configure("TLabelframe.Label", background="#f4f7fb", foreground="#1f4e79", font=("微软雅黑", 10, "bold"))
        style.configure("TButton", padding=(10, 7), font=("微软雅黑", 10), background="#2563eb", foreground="white")
        style.map("TButton", background=[("active", "#1d4ed8")], foreground=[("active", "white")])
        style.configure("TEntry", padding=(6, 4), fieldbackground="#ffffff")
        style.configure("TCombobox", padding=(6, 4), fieldbackground="#ffffff")
        style.configure("Horizontal.TProgressbar", background="#34d399", troughcolor="#e5e7eb")

    def create_widgets(self):
        """创建界面组件"""
        # 创建主容器 - 使用Canvas实现可滚动内容
        main_canvas = tk.Canvas(self.root, highlightthickness=0, bg="#f4f7fb", borderwidth=0)
        main_scrollbar = Scrollbar(self.root, orient=VERTICAL, command=main_canvas.yview)
        scrollable_main = Frame(main_canvas, bg="#f4f7fb")
        
        # 配置滚动
        scrollable_main.bind(
            "<Configure>",
            lambda e: main_canvas.configure(scrollregion=main_canvas.bbox("all"))
        )
        
        # 初始创建窗口，宽度稍后更新
        canvas_window = main_canvas.create_window((0, 0), window=scrollable_main, anchor="nw")
        main_canvas.configure(yscrollcommand=main_scrollbar.set)
        
        # Canvas大小变化时更新内部窗口宽度
        def _on_canvas_resize(event):
            if event.width > 0:
                main_canvas.itemconfig(canvas_window, width=event.width)
        main_canvas.bind("<Configure>", _on_canvas_resize)
        
        # 绑定鼠标滚轮事件 - 智能判断焦点
        def _on_mousewheel(event):
            # 获取当前获得焦点的控件
            focused_widget = self.root.focus_get()
            
            # 如果焦点在Listbox或Text控件上，让它们自己处理滚动
            if focused_widget:
                widget_class = str(focused_widget.winfo_class())
                if widget_class in ('Listbox', 'Text'):
                    # 让控件自己处理滚动，不干预
                    return
            
            # 否则滚动主Canvas
            main_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        
        # 绑定到根窗口，所有控件都能捕获
        self.root.bind_all("<MouseWheel>", _on_mousewheel)
        
        # 布局Canvas和滚动条
        main_canvas.pack(side=LEFT, fill=BOTH, expand=True)
        main_scrollbar.pack(side=RIGHT, fill=Y)
        
        # 在可滚动容器中创建内容
        # ========== 文件选择区域 ==========
        file_frame = LabelFrame(scrollable_main, text="文件选择", font=("微软雅黑", 10, "bold"),
                                padx=12, pady=10, fg="#1f4e79", bg="#ffffff", borderwidth=1, relief="groove")
        file_frame.pack(fill=X, pady=(10, 6), padx=10)
        
        # 添加 Visio 图提示（居中显示）
        visio_hint_frame = Frame(file_frame, bg="#fff7ed")
        visio_hint_frame.pack(fill=X, pady=(0, 8))
        
        visio_hint = Label(visio_hint_frame, 
                          text="[TIP] 提示：建议提前将源文档中的 Visio 图提前转换为 JPG/PNG 等图片格式，为保证转换后文档最大可用性，本软件暂不支持viso等OLE对象的复制。",
                          font=("微软雅黑", 9), fg="#c2410c", bg="#fff7ed", justify=CENTER)
        visio_hint.pack(anchor=CENTER)
        
        # 源文件
        source_row = Frame(file_frame)
        source_row.pack(fill=X, pady=3)
        
        Label(source_row, text="源文档:", width=10, anchor=W,
              font=("微软雅黑", 9)).pack(side=LEFT)
        
        self.source_entry = Entry(source_row, font=("微软雅黑", 9))
        self.source_entry.pack(side=LEFT, fill=X, expand=True, padx=5)
        
        Button(source_row, text="浏览...", command=self.browse_source,
               width=8, font=("微软雅黑", 9)).pack(side=LEFT, padx=2)
        Button(source_row, text="多选...", command=self.browse_multiple_sources,
               width=8, font=("微软雅黑", 9)).pack(side=LEFT, padx=2)
        
        # 模板文件
        template_row = Frame(file_frame)
        template_row.pack(fill=X, pady=3)
        
        Label(template_row, text="模板文档:", width=10, anchor=W,
              font=("微软雅黑", 9)).pack(side=LEFT)
        
        Entry(template_row, textvariable=self.template_file,
              font=("微软雅黑", 9)).pack(side=LEFT, fill=X, expand=True, padx=5)
        
        Button(template_row, text="浏览...", command=self.browse_template,
               width=8, font=("微软雅黑", 9)).pack(side=LEFT, padx=2)
        
        # 进度条区域
        progress_frame = Frame(file_frame)
        progress_frame.pack(fill=X, pady=(5, 0))
        
        Label(progress_frame, text="加载进度:", width=10, anchor=W,
              font=("微软雅黑", 9)).pack(side=LEFT)
        
        self.loading_progress = ttk.Progressbar(progress_frame, mode='determinate')
        self.loading_progress.pack(side=LEFT, fill=X, expand=True, padx=5)
        self.loading_progress['value'] = 0
        
        # ========== 样式信息区域（包含三个列表）==========
        styles_frame = LabelFrame(scrollable_main, text="样式信息", font=("微软雅黑", 10, "bold"),
                                  padx=12, pady=10, fg="#1f4e79", bg="#ffffff", borderwidth=1, relief="groove")
        styles_frame.pack(fill=BOTH, expand=False, pady=(0, 6), padx=10)
        
        # 样式映射按钮和提示
        btn_row = Frame(styles_frame)
        btn_row.pack(fill=X, pady=(0, 5))
        
        Button(btn_row, text="配置样式映射", command=self.open_style_mapping,
               width=15, font=("微软雅黑", 9, "bold"), bg="#2196F3", fg="white").pack(side=LEFT)
        
        Label(btn_row, text="  （选择源文件后点击按钮配置该文件的样式映射）", 
              font=("微软雅黑", 8), fg="black").pack(side=LEFT)
        
        # 三栏布局：源文件列表、源样式列表、模板样式列表
        styles_container = Frame(styles_frame)
        styles_container.pack(fill=BOTH, expand=True)
        
        # 第一栏：源文件列表
        file_list_frame = Frame(styles_container)
        file_list_frame.pack(side=LEFT, fill=BOTH, expand=True, padx=5)
        
        Label(file_list_frame, text="源文件列表", font=("微软雅黑", 9, "bold")).pack(anchor=W)
        Label(file_list_frame, text="（点击选择）", font=("微软雅黑", 8), fg="black").pack(anchor=W)
        
        file_scroll = Scrollbar(file_list_frame)
        file_scroll.pack(side=RIGHT, fill=Y)
        
        self.file_listbox = Listbox(file_list_frame, yscrollcommand=file_scroll.set,
                                    font=("微软雅黑", 9), height=6, bg="#ffffff", selectbackground="#93c5fd", selectforeground="#0f172a")
        self.file_listbox.pack(fill=BOTH, expand=True)
        file_scroll.config(command=self.file_listbox.yview)
        
        # 绑定选择事件
        self.file_listbox.bind('<<ListboxSelect>>', self.on_source_file_select)
        
        # 第二栏：源文档样式
        source_style_frame = Frame(styles_container)
        source_style_frame.pack(side=LEFT, fill=BOTH, expand=True, padx=5)
        
        Label(source_style_frame, text="源文档样式", font=("微软雅黑", 9, "bold")).pack(anchor=W)
        Label(source_style_frame, text="（选中文件的样式）", font=("微软雅黑", 8), fg="black").pack(anchor=W)
        
        source_scroll = Scrollbar(source_style_frame)
        source_scroll.pack(side=RIGHT, fill=Y)
        
        self.source_listbox = Listbox(source_style_frame, yscrollcommand=source_scroll.set,
                                      font=("微软雅黑", 9), height=6, bg="#ffffff", selectbackground="#93c5fd", selectforeground="#0f172a")
        self.source_listbox.pack(fill=BOTH, expand=True)
        source_scroll.config(command=self.source_listbox.yview)
        
        # 第三栏：模板文档样式
        right_frame = Frame(styles_container)
        right_frame.pack(side=LEFT, fill=BOTH, expand=True, padx=5)
        
        Label(right_frame, text="模板文档样式", font=("微软雅黑", 9, "bold")).pack(anchor=W)
        Label(right_frame, text="（可用的目标样式）", font=("微软雅黑", 8), fg="black").pack(anchor=W)
        
        template_scroll = Scrollbar(right_frame)
        template_scroll.pack(side=RIGHT, fill=Y)
        
        self.template_listbox = Listbox(right_frame, yscrollcommand=template_scroll.set,
                                        font=("微软雅黑", 9), height=6, bg="#ffffff", selectbackground="#93c5fd", selectforeground="#0f172a")
        self.template_listbox.pack(fill=BOTH, expand=True)
        template_scroll.config(command=self.template_listbox.yview)
        
        # ========== 转换选项区域 ==========
        options_frame = LabelFrame(scrollable_main, text="转换选项", font=("微软雅黑", 10, "bold"),
                                   padx=12, pady=10, fg="#1f4e79", bg="#ffffff", borderwidth=1, relief="groove")
        options_frame.pack(fill=X, pady=(0, 6), padx=10)
        
        # 复选框（并排显示）
        check_row = Frame(options_frame)
        check_row.pack(fill=X, pady=3)
        
        Checkbutton(check_row, text="进行祈使语气转换",
                    variable=self.do_mood_conversion, font=("微软雅黑", 9)).pack(side=LEFT, padx=5)
        
        # 应答句配置已迁移到"样式映射"对话框（步骤2）中管理
        
        # ========== 章节提示语配置区域 ==========
        hint_frame = LabelFrame(options_frame, text="章节提示语", font=("微软雅黑", 10),
                                padx=10, pady=5)
        hint_frame.pack(fill=X, pady=5)
        
        # 第一行：启用开关 + 提示语类型选择
        hint_check_row = Frame(hint_frame)
        hint_check_row.pack(fill=X, pady=2)
        
        self.hint_check = Checkbutton(hint_check_row, text="插入章节提示语",
                                       variable=self.do_hint_insertion, font=("微软雅黑", 9),
                                       command=self.toggle_hint_controls)
        self.hint_check.pack(side=LEFT, padx=5)
        
        # 提示语类型单选按钮
        self.hint_type_text_rb = tk.Radiobutton(hint_check_row, text="文本提示语",
                                                 variable=self.hint_type, value="text",
                                                 font=("微软雅黑", 9),
                                                 command=self.toggle_hint_type_controls)
        self.hint_type_text_rb.pack(side=LEFT, padx=5)
        
        self.hint_type_image_rb = tk.Radiobutton(hint_check_row, text="图片提示语",
                                                  variable=self.hint_type, value="image",
                                                  font=("微软雅黑", 9),
                                                  command=self.toggle_hint_type_controls)
        self.hint_type_image_rb.pack(side=LEFT, padx=5)
        
        # 第二行：提示语文本输入框（始终显示，图片模式灰显）
        hint_text_row = Frame(hint_frame)
        hint_text_row.pack(fill=X, pady=2)
        
        Label(hint_text_row, text="提示语文本:", width=10, anchor=W,
              font=("微软雅黑", 9)).pack(side=LEFT)
        self.hint_text_entry = Entry(hint_text_row, textvariable=self.hint_text,
                                     font=("微软雅黑", 9))
        self.hint_text_entry.pack(side=LEFT, fill=X, expand=True, padx=5)
        
        # 第三行：图片文件选择（始终显示，文本模式灰显）
        self.hint_image_row = Frame(hint_frame)
        self.hint_image_row.pack(fill=X, pady=2)
        
        Label(self.hint_image_row, text="图片文件:", width=10, anchor=W,
              font=("微软雅黑", 9)).pack(side=LEFT)
        self.hint_image_entry = Entry(self.hint_image_row, textvariable=self.hint_image_path,
                                       font=("微软雅黑", 9), state='readonly')
        self.hint_image_entry.pack(side=LEFT, fill=X, expand=True, padx=5)
        self.hint_image_select_btn = Button(self.hint_image_row, text="选择图片", width=10,
                                             font=("微软雅黑", 9),
                                             command=self.select_hint_image)
        self.hint_image_select_btn.pack(side=LEFT, padx=(0, 5))
        
        # 第四行：提示语样式下拉框 + 设为默认按钮（始终显示，在图片文件下面）
        hint_style_row = Frame(hint_frame)
        hint_style_row.pack(fill=X, pady=2)
        
        Label(hint_style_row, text="提示语样式:", width=10, anchor=W,
              font=("微软雅黑", 9)).pack(side=LEFT)
        self.hint_style_combo = ttk.Combobox(hint_style_row, textvariable=self.hint_style,
                                               width=20, state="readonly", font=("微软雅黑", 9))
        self.hint_style_combo.pack(side=LEFT, padx=5)
        self.hint_style_combo['values'] = []  # 等待模板加载后填充
        
        self.set_hint_default_btn = Button(hint_style_row, text="设为默认", width=8,
                                            font=("微软雅黑", 9),
                                            command=self.save_default_hint_settings)
        self.set_hint_default_btn.pack(side=LEFT, padx=(5, 0))
        
        # 初始化提示语控件状态
        self.toggle_hint_controls()
        
        # ========== 输出文件列表区域 ==========
        output_frame = LabelFrame(scrollable_main, text="输出文件列表", font=("微软雅黑", 10, "bold"),
                                  padx=12, pady=10, fg="#1f4e79", bg="#ffffff", borderwidth=1, relief="groove")
        output_frame.pack(fill=BOTH, expand=True, pady=(0, 6), padx=10)
        
        output_scroll = Scrollbar(output_frame)
        output_scroll.pack(side=RIGHT, fill=Y)
        
        self.output_listbox = Listbox(output_frame, yscrollcommand=output_scroll.set,
                                      font=("微软雅黑", 9), height=6, bg="#ffffff", selectbackground="#93c5fd", selectforeground="#0f172a")
        self.output_listbox.pack(fill=BOTH, expand=True)
        output_scroll.config(command=self.output_listbox.yview)
        
        # 双击打开文件
        self.output_listbox.bind('<Double-Button-1>', self.open_selected_file)
        
        # ========== 日志输出区域 ==========
        log_frame = LabelFrame(scrollable_main, text="处理日志", font=("微软雅黑", 10, "bold"),
                               padx=12, pady=10, fg="#1f4e79", bg="#ffffff", borderwidth=1, relief="groove")
        log_frame.pack(fill=BOTH, expand=True, pady=(0, 6), padx=10)
        
        # 日志控制区域
        log_control_frame = Frame(log_frame)
        log_control_frame.pack(fill=X, pady=(0, 5))
        
        self.detailed_log_var = IntVar(value=0)  # 默认关闭详细日志
        Checkbutton(log_control_frame, text="详细日志",
                    variable=self.detailed_log_var, font=("微软雅黑", 9)).pack(side=LEFT)
        
        log_scroll = Scrollbar(log_frame)
        log_scroll.pack(side=RIGHT, fill=Y)
        
        self.log_text = Text(log_frame, height=7, font=("Consolas", 9),
                             yscrollcommand=log_scroll.set, bg="#f8fafc", padx=8, pady=6)
        self.log_text.pack(fill=BOTH, expand=True)
        log_scroll.config(command=self.log_text.yview)
        
        # ========== 按钮区域（固定在底部）==========
        btn_frame = Frame(scrollable_main, bg="#f4f7fb", pady=8)
        btn_frame.pack(fill=X, side=BOTTOM, pady=(4, 10), padx=10)
        
        self.convert_btn = Button(btn_frame, text="开始转换", command=self.start_conversion,
                                  width=15, font=("微软雅黑", 11, "bold"), bg="#16a34a",
                                  fg="white", relief="raised")
        self.convert_btn.pack(side=LEFT, padx=5)
        
        self.stop_btn = Button(btn_frame, text="停止转换", command=self.stop_conversion,
                               width=10, font=("微软雅黑", 10), bg="#ef4444",
                               fg="white", state='disabled', relief="raised")
        self.stop_btn.pack(side=LEFT, padx=5)
        
        self.clear_log_btn = Button(btn_frame, text="清空日志", command=self.clear_log,
                                    width=10, font=("微软雅黑", 10))
        self.clear_log_btn.pack(side=LEFT, padx=5)
        
        self.open_folder_btn = Button(btn_frame, text="打开输出文件夹", command=self.open_output_folder,
                                      width=12, font=("微软雅黑", 10), bg="#f8fafc", relief="raised")
        self.open_folder_btn.pack(side=LEFT, padx=5)
        
        Button(btn_frame, text="退出", command=self.root.quit,
               width=10, font=("微软雅黑", 10), bg="#f8fafc", relief="raised").pack(side=RIGHT, padx=5)
    
    def _start_loading_progress(self, total_steps):
        """启动进度条"""
        self.loading_progress['maximum'] = total_steps
        self.loading_progress['value'] = 0
        # 灰化按钮，防止用户在加载过程中操作
        self.set_buttons_state('disabled')
        self.root.update_idletasks()
    
    def _update_loading_progress(self, step):
        """更新进度条"""
        self.loading_progress['value'] = step
        self.root.update_idletasks()
    
    def _stop_loading_progress(self):
        """停止进度条"""
        self.loading_progress['value'] = 0
        # 恢复按钮状态
        self.set_buttons_state('normal')
        self.root.update_idletasks()
    
    def browse_source(self):
        """浏览选择单个源文件"""
        filename = filedialog.askopenfilename(
            title="选择源文档",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if filename:
            self.source_files = [filename]
            self.source_entry.delete(0, END)
            self.source_entry.insert(0, filename)
            
            # 只清空源文件相关数据，保留模板信息
            self.clear_source_data()
            
            # 注意：不在这里添加文件到列表，等待分析完成后再添加
            # self.file_listbox.insert(END, os.path.basename(filename))
            
            # 分析样式（带进度条），分析完成后会添加文件到列表
            self.analyze_source_styles_with_progress([filename])
    
    def analyze_source_styles_with_progress(self, files):
        """分析源文档样式（使用后台线程，按段落更新进度）"""
        import threading
        
        # 先计算总段落数和每个文件的段落数
        total_paragraphs = 0
        file_paragraph_counts = {}
        for file in files:
            try:
                doc = Document(file)
                count = len(doc.paragraphs)
                file_paragraph_counts[file] = count
                total_paragraphs += count
            except:
                file_paragraph_counts[file] = 0
        
        # 设置进度条 maximum = 100（固定）
        self._start_loading_progress(100)
        
        # 在后台线程中执行分析
        def analyze_in_background():
            try:
                temp_styles = set()
                total_files = len(files)
                
                for idx, file in enumerate(files, 1):
                    self.log(f"  正在分析: {os.path.basename(file)}...")
                    
                    doc = Document(file)
                    para_count = file_paragraph_counts.get(file, 0)
                    current_file_total = len(doc.paragraphs)  # 当前文件的总段落数

                    # 先收集列表段落虚拟样式（数字编号/符号编号）
                    list_virtual_styles = self.converter.get_list_virtual_styles(doc)
                    temp_styles.update(list_virtual_styles)

                    for para_idx, para in enumerate(doc.paragraphs):
                        if para.style and para.style.name:
                            temp_styles.add(para.style.name)
                            # 检测大纲级别（outlineLvl）并生成虚拟样式名
                            # 仅当段落样式不是已知的标题样式时，才检测直接 outlineLvl
                            para_style_lower = para.style.name.lower()
                            if not (para_style_lower.startswith('heading') or para_style_lower.startswith('head')):
                                elem = para._element
                                pPr = elem.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
                                if pPr is not None:
                                    outline = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl')
                                    if outline is not None:
                                        val = outline.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                                        if val is not None:
                                            try:
                                                level = int(val) + 1
                                                if 1 <= level <= 9:
                                                    temp_styles.add(f'[大纲级别 {level}]')
                                            except ValueError:
                                                pass
                        
                        # 每处理10个段落或最后一个段落时更新进度
                        if (para_idx + 1) % 10 == 0 or para_idx == len(doc.paragraphs) - 1:
                            # 两层进度计算：
                            # 第一层：已完成 (idx-1) 个文件，每个占 100/total_files
                            # 第二层：当前文件已处理 (para_idx+1)/current_file_total
                            completed_files_progress = (idx - 1) * (100 / total_files)
                            current_file_progress = ((para_idx + 1) / current_file_total) * (100 / total_files)
                            total_progress = completed_files_progress + current_file_progress
                            
                            self.root.after(0, lambda p=total_progress: self._update_loading_progress(p))
                    
                    # 确保进度至少增加（处理空文件）
                    if para_count == 0:
                        completed_files_progress = idx * (100 / total_files)
                        self.root.after(0, lambda p=completed_files_progress: self._update_loading_progress(p))
                
                # 分析完成，在主线程中更新UI
                self.root.after(0, lambda: self._finish_style_analysis(temp_styles))
            except Exception as e:
                self.root.after(0, lambda: self._handle_analysis_error(e))
        
        thread = threading.Thread(target=analyze_in_background)
        thread.daemon = True
        thread.start()
    
    def _finish_style_analysis(self, styles):
        """完成样式分析的UI更新"""
        self.source_styles = styles
        
        # 更新列表框
        self.source_listbox.delete(0, END)
        for style in sorted(self.source_styles):
            self.source_listbox.insert(END, style)
        
        # 分析完成后，添加文件到源文件列表
        if self.source_files:
            for f in self.source_files:
                self.file_listbox.insert(END, os.path.basename(f))
        
        self.log(f"✓ 源文档样式分析完成，共 {len(self.source_styles)} 种样式")
        self._stop_loading_progress()
    
    def _handle_analysis_error(self, error):
        """处理分析错误"""
        messagebox.showerror("错误", f"分析源文档样式失败: {error}")
        self._stop_loading_progress()
    
    def browse_multiple_sources(self):
        """浏览选择多个源文件"""
        filenames = filedialog.askopenfilenames(
            title="选择多个源文档",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if filenames:
            self.source_files = list(filenames)
            self.source_entry.delete(0, END)
            
            # 只清空源文件相关数据，保留模板信息
            self.clear_source_data()
            
            if len(filenames) == 1:
                self.source_entry.insert(0, filenames[0])
                # 不在这里添加文件到列表，等待分析完成
                self.analyze_source_styles_with_progress(list(filenames))
            else:
                self.source_entry.insert(0, f"已选择 {len(filenames)} 个文件")
                self.log(f"已选择 {len(filenames)} 个源文件")
                
                # 不在这里填充源文件列表，等待分析完成
                # for f in filenames:
                #     self.file_listbox.insert(END, os.path.basename(f))
                
                # 分析所有文件的样式（带进度条），分析完成后会添加文件到列表
                self.analyze_multiple_source_styles_with_progress(list(filenames))
    
    def analyze_multiple_source_styles_with_progress(self, files):
        """分析多个源文件的样式（使用后台线程，按段落更新进度）"""
        import threading
        
        # 先计算总段落数和每个文件的段落数
        total_paragraphs = 0
        file_paragraph_counts = {}
        for file in files:
            try:
                doc = Document(file)
                count = len(doc.paragraphs)
                file_paragraph_counts[file] = count
                total_paragraphs += count
            except:
                file_paragraph_counts[file] = 0
        
        # 设置进度条 maximum = 100（固定）
        self._start_loading_progress(100)
        
        # 在后台线程中执行分析
        def analyze_in_background():
            try:
                temp_file_maps = {}
                total_files = len(files)
                
                for idx, file in enumerate(files, 1):
                    self.log(f"  [{idx}/{len(files)}] 分析: {os.path.basename(file)}...")
                    
                    doc = Document(file)
                    styles = set()
                    para_count = file_paragraph_counts.get(file, 0)
                    current_file_total = len(doc.paragraphs)  # 当前文件的总段落数

                    # 先收集列表段落虚拟样式（数字编号/符号编号）并合并到 styles 中
                    list_virtual_styles = self.converter.get_list_virtual_styles(doc)
                    styles.update(list_virtual_styles)

                    for para_idx, para in enumerate(doc.paragraphs):
                        if para.style and para.style.name:
                            styles.add(para.style.name)
                            # 检测大纲级别（outlineLvl）并生成虚拟样式名
                            para_style_lower = para.style.name.lower()
                            if not (para_style_lower.startswith('heading') or para_style_lower.startswith('head')):
                                elem = para._element
                                pPr = elem.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
                                if pPr is not None:
                                    outline = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl')
                                    if outline is not None:
                                        val = outline.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                                        if val is not None:
                                            try:
                                                level = int(val) + 1
                                                if 1 <= level <= 9:
                                                    styles.add(f'[大纲级别 {level}]')
                                            except ValueError:
                                                pass
                        
                        # 每处理10个段落或最后一个段落时更新进度
                        if (para_idx + 1) % 10 == 0 or para_idx == len(doc.paragraphs) - 1:
                            # 两层进度计算：
                            # 第一层：已完成 (idx-1) 个文件，每个占 100/total_files
                            # 第二层：当前文件已处理 (para_idx+1)/current_file_total
                            completed_files_progress = (idx - 1) * (100 / total_files)
                            current_file_progress = ((para_idx + 1) / current_file_total) * (100 / total_files)
                            total_progress = completed_files_progress + current_file_progress
                            
                            self.root.after(0, lambda p=total_progress: self._update_loading_progress(p))
                    
                    # 保存该文件的样式
                    temp_file_maps[file] = {
                        'styles': styles,
                        'mapping': {}
                    }
                    
                    # 确保进度至少增加（处理空文件）
                    if para_count == 0:
                        completed_files_progress = idx * (100 / total_files)
                        self.root.after(0, lambda p=completed_files_progress: self._update_loading_progress(p))
                
                # 分析完成，在主线程中更新UI
                self.root.after(0, lambda: self._finish_multi_file_analysis(temp_file_maps))
            except Exception as e:
                self.root.after(0, lambda: self._handle_analysis_error(e))
        
        thread = threading.Thread(target=analyze_in_background)
        thread.daemon = True
        thread.start()
    
    def _finish_multi_file_analysis(self, file_maps):
        """完成多文件分析的UI更新"""
        self.file_style_maps = file_maps
        
        # 分析完成后，添加所有文件到源文件列表
        for f in self.source_files:
            self.file_listbox.insert(END, os.path.basename(f))
        
        self.log(f"✓ 已分析 {len(file_maps)} 个文件的样式")
        self._stop_loading_progress()
    
    def on_source_file_select(self, event=None):
        """源文件列表选择事件"""
        selection = self.file_listbox.curselection()
        if selection and len(self.source_files) > 0:
            index = selection[0]
            if 0 <= index < len(self.source_files):
                selected_file = self.source_files[index]
                self.selected_source_file = selected_file
                
                # 更新源样式列表
                if selected_file in self.file_style_maps:
                    styles = self.file_style_maps[selected_file]['styles']
                    self.source_listbox.delete(0, END)
                    for style in sorted(styles):
                        self.source_listbox.insert(END, style)
                    
                    self.log(f"已选择文件: {os.path.basename(selected_file)} (共{len(styles)}种样式)")
    
    def clear_source_data(self):
        """清空源文件相关数据（保留模板信息）"""
        # 清空源文件样式和映射
        self.source_styles = set()
        self.file_style_maps = {}
        self.selected_source_file = None
        
        # 清空源文件列表和源样式列表
        self.file_listbox.delete(0, END)
        self.source_listbox.delete(0, END)
        
        # 清空输出文件列表
        self.output_files = []
        self.output_listbox.delete(0, END)
        
        # 清空日志
        self.log_text.delete(1.0, END)
        
        # 重置自定义映射
        self.custom_style_map = {}
        self.custom_tbl_img_config = {}
        self.custom_remove_chapter_label = 0
        self.custom_answer_config = {}
        self.custom_list_config = {}
        self.remove_chapter_label = self.default_config.get('remove_chapter_label', 0)
    
    def clear_template_data(self):
        """清空模板相关数据"""
        # 清空模板样式
        self.template_styles = set()
        
        # 清空模板样式列表
        self.template_listbox.delete(0, END)
    
    def clear_all_data(self):
        """清空所有数据（源文件和模板）"""
        self.clear_source_data()
        self.clear_template_data()
    
    def _load_default_config(self):
        """从配置文件加载默认设置"""
        try:
            if os.path.exists(DEFAULT_CONFIG_FILE):
                with open(DEFAULT_CONFIG_FILE, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                return config
        except Exception as e:
            print(f"加载默认配置失败: {e}")
        return {}
    
    def _save_default_config(self, config):
        """将默认设置保存到配置文件"""
        try:
            with open(DEFAULT_CONFIG_FILE, 'w', encoding='utf-8') as f:
                json.dump(config, f, ensure_ascii=False, indent=2)
            return True
        except Exception as e:
            print(f"保存默认配置失败: {e}")
            return False
    
    def save_default_style_map(self, style_map, answer_config=None, tbl_img_config=None, list_config=None, remove_chapter_label=None):
        """将当前样式映射保存为默认（包含样式映射、应答句配置、表格/图片/列表样式定义和清除章节编号配置）"""
        self.default_config["style_map"] = style_map
        # 保存应答句配置
        if answer_config:
            self.default_config["do_answer"] = answer_config.get('do_answer', 0)
            self.default_config["answer_text"] = answer_config.get('answer_text', '应答：本投标人理解并满足要求。')
            self.default_config["answer_style"] = answer_config.get('answer_style', '应答句')
            self.default_config["answer_mode"] = answer_config.get('answer_mode', '章节标题后插入')
            self.default_config["answer_source_style"] = answer_config.get('answer_source_style', '')
            self.default_config["answer_copy_style"] = answer_config.get('answer_copy_style', '')
        # 保存表格/图片样式定义
        if tbl_img_config:
            self.default_config["enable_table_style"] = tbl_img_config.get('enable_table_style', 0)
            self.default_config["table_style"] = tbl_img_config.get('table_style', 'Body Text')
            self.default_config["table_answer_style"] = tbl_img_config.get('table_answer_style', 'Body Text')
            self.default_config["enable_image_style"] = tbl_img_config.get('enable_image_style', 0)
            self.default_config["image_style"] = tbl_img_config.get('image_style', 'Body Text')
            self.default_config["image_answer_style"] = tbl_img_config.get('image_answer_style', 'Body Text')
        # 保存列表段落兜底配置
        if list_config:
            self.default_config["list_method"] = list_config.get('method', 'bullet')
            self.default_config["list_bullet"] = list_config.get('bullet', '● ')
            self.default_config["list_style"] = list_config.get('style', 'Body Text')
            self.default_config["list_answer_method"] = list_config.get('answer_method', 'bullet')
            self.default_config["list_answer_bullet"] = list_config.get('answer_bullet', '● ')
            self.default_config["list_answer_style"] = list_config.get('answer_style', 'Body Text')
        # 保存清除章节标题编号配置
        if remove_chapter_label is not None:
            self.default_config["remove_chapter_label"] = remove_chapter_label
        if self._save_default_config(self.default_config):
            configured_count = sum(1 for v in style_map.values() if v)
            tbl_info = ""
            if tbl_img_config:
                tbl_on = tbl_img_config.get('enable_table_style', 0)
                img_on = tbl_img_config.get('enable_image_style', 0)
                tbl_info = f"\n表格样式覆盖: {'启用 (' + tbl_img_config.get('table_style', '') + ')' if tbl_on else '未启用'}"
                tbl_info += f"\n图片样式覆盖: {'启用 (' + tbl_img_config.get('image_style', '') + ')' if img_on else '未启用'}"
            ch_label_info = ""
            if remove_chapter_label is not None:
                ch_label_info = f"\n清除章节标题编号: {'是' if remove_chapter_label else '否'}"
            self.log(f"✓ 已将样式映射设为默认，共 {configured_count} 个映射关系")
            messagebox.showinfo("成功", f"已将样式映射设为默认！\n\n共 {configured_count} 个样式映射关系{tbl_info}{ch_label_info}\n下次打开样式映射对话框时将自动恢复此配置。")
        else:
            messagebox.showerror("错误", "保存默认配置失败，请检查文件权限。")
    
    def toggle_hint_controls(self):
        """切换章节提示语控件的可用状态（灰显/可用，不隐藏）"""
        if self.do_hint_insertion.get():
            # 选中：启用控件
            self.hint_type_text_rb.config(state='normal')
            self.hint_type_image_rb.config(state='normal')
            self.hint_style_combo.config(state='readonly')
            self.set_hint_default_btn.config(state='normal')
            self.toggle_hint_type_controls()  # 根据类型切换文本/图片控件灰显
        else:
            # 未选中：灰化所有控件
            self.hint_type_text_rb.config(state='disabled')
            self.hint_type_image_rb.config(state='disabled')
            self.hint_text_entry.config(state='disabled')
            self.set_hint_default_btn.config(state='disabled')
            self.hint_style_combo.config(state='disabled')
            self.hint_image_entry.config(state='disabled')
            self.hint_image_select_btn.config(state='disabled')
    
    def toggle_hint_type_controls(self):
        """根据提示语类型切换文本/图片控件的灰显状态（所有控件始终可见）"""
        if not self.do_hint_insertion.get():
            return  # 未启用时不操作
        
        if self.hint_type.get() == "text":
            # 文本模式：文本输入可用，图片选择灰显
            self.hint_text_entry.config(state='normal')
            self.hint_image_entry.config(state='disabled')
            self.hint_image_select_btn.config(state='disabled')
        else:
            # 图片模式：文本输入灰显，图片选择可用
            self.hint_text_entry.config(state='disabled')
            self.hint_image_entry.config(state='readonly')
            self.hint_image_select_btn.config(state='normal')
        # 设为默认按钮始终可用（保存全部提示语配置）
    
    def select_hint_image(self):
        """选择提示语图片文件"""
        file_path = filedialog.askopenfilename(
            title="选择提示语图片",
            filetypes=[
                ("图片文件", "*.png *.jpg *.jpeg *.bmp *.gif *.tiff *.tif"),
                ("所有文件", "*.*")
            ]
        )
        if file_path:
            self.hint_image_path.set(file_path)
    
    def save_default_hint_settings(self):
        """保存当前提示语全部配置为默认（是否启用、类型、文本、图片路径、样式）"""
        self.default_config["do_hint_insertion"] = self.do_hint_insertion.get()
        self.default_config["hint_type"] = self.hint_type.get()
        self.default_config["hint_text"] = self.hint_text.get()
        self.default_config["hint_image_path"] = self.hint_image_path.get()
        self.default_config["hint_style"] = self.hint_style.get()
        if self._save_default_config(self.default_config):
            hint_enabled = "启用" if self.do_hint_insertion.get() else "未启用"
            hint_type_str = "文本" if self.hint_type.get() == "text" else "图片"
            messagebox.showinfo("提示", f"已将当前提示语配置设为默认！\n\n状态: {hint_enabled}\n类型: {hint_type_str}")
        else:
            messagebox.showerror("错误", "保存默认配置失败，请检查文件权限。")

    def stop_conversion(self):
        """停止转换"""
        if self.conversion_thread and self.conversion_thread.is_alive():
            self.stop_conversion_flag = True
            self.log("\n⚠ 正在停止转换，请稍候...")
            self.stop_btn.config(state='disabled')
            self.convert_btn.config(state='disabled')
    
    def set_buttons_state(self, state='normal'):
        """设置按钮的可用状态"""
        # 转换相关按钮
        self.convert_btn.config(state=state)
        self.clear_log_btn.config(state=state)
        self.open_folder_btn.config(state=state)
        
        # 文件选择按钮（浏览、多选）
        # 需要找到这些按钮并设置状态
        # 由于它们是局部变量，我们需要通过遍历来查找
        for widget in self.root.winfo_children():
            if isinstance(widget, Frame):
                for child in widget.winfo_children():
                    if isinstance(child, LabelFrame):
                        for grandchild in child.winfo_children():
                            if isinstance(grandchild, Frame):
                                for btn in grandchild.winfo_children():
                                    if isinstance(btn, Button) and btn.cget('text') in ['浏览...', '多选...']:
                                        btn.config(state=state)
    
    def browse_template(self):
        """浏览选择模板文件"""
        filename = filedialog.askopenfilename(
            title="选择模板文档",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if filename:
            self.template_file.set(filename)
            
            # 清空之前的模板数据
            self.clear_template_data()
            
            self.analyze_template_styles_with_progress(filename)
    
    def analyze_template_styles_with_progress(self, template_file):
        """分析模板文档样式（使用后台线程）"""
        import threading
        
        self.log(f"正在分析模板: {os.path.basename(template_file)}...")
        
        # 在后台线程中执行分析
        def analyze_in_background():
            try:
                doc = Document(template_file)
                
                styles_list = list(doc.styles)
                total = len(styles_list)
                
                # 使用after在主线程中更新进度条最大值
                self.root.after(0, lambda t=total: self._start_loading_progress(t))
                
                temp_styles = set()
                for idx, style in enumerate(styles_list, 1):
                    if style.type == WD_STYLE_TYPE.PARAGRAPH:
                        temp_styles.add(style.name)
                    
                    # 每10个样式更新一次进度
                    if idx % 10 == 0:
                        self.root.after(0, lambda i=idx: self._update_loading_progress(i))
                
                # 完成，使用最终进度
                self.root.after(0, lambda t=total: self._update_loading_progress(t))
                
                # 在主线程中更新UI
                self.root.after(0, lambda: self._finish_template_analysis(temp_styles))
            except Exception as e:
                self.root.after(0, lambda: self._handle_analysis_error(e))
        
        thread = threading.Thread(target=analyze_in_background)
        thread.daemon = True
        thread.start()
    
    def _finish_template_analysis(self, styles):
        """完成模板分析的UI更新"""
        self.template_styles = styles
        
        # 更新列表框
        self.template_listbox.delete(0, END)
        for style in sorted(self.template_styles):
            self.template_listbox.insert(END, style)
        
        # 更新提示语样式的下拉框选项
        template_style_list = sorted(self.template_styles)
        self.hint_style_combo['values'] = template_style_list
        
        # 如果当前提示语样式不在模板样式中，设置为第一个样式或保持原值
        current_hint_style = self.hint_style.get()
        if current_hint_style not in template_style_list and template_style_list:
            self.hint_style.set(template_style_list[0])
        
        self.log(f"✓ 模板文档样式分析完成，共 {len(self.template_styles)} 种样式")
        self.root.after(500, lambda: self._stop_loading_progress())
    
    def _update_template_styles_ui(self, styles):
        """在主线程中更新模板样式列表UI"""
        self.template_styles = styles
        
        # 更新列表框
        self.template_listbox.delete(0, END)
        for style in sorted(styles):
            self.template_listbox.insert(END, style)
        
        # 更新提示语样式的下拉框选项
        template_style_list = sorted(styles)
        self.hint_style_combo['values'] = template_style_list
        self.log(f"已更新提示语样式下拉框，共 {len(template_style_list)} 个选项")
        
        # 如果当前提示语样式不在模板样式中，设置为第一个样式或保持原值
        current_hint_style = self.hint_style.get()
        if current_hint_style not in template_style_list and template_style_list:
            self.hint_style.set(template_style_list[0])
    
    def analyze_source_styles(self, files):
        """分析源文档样式（含虚拟样式检测）"""
        self.source_styles = set()
        try:
            for file in files:
                doc = Document(file)
                # 收集实际样式名
                for para in doc.paragraphs:
                    if para.style and para.style.name:
                        self.source_styles.add(para.style.name)
                # 收集列表段落虚拟样式
                list_styles = self.converter.get_list_virtual_styles(doc)
                self.source_styles.update(list_styles)
            
            # 更新列表框
            self.source_listbox.delete(0, END)
            for style in sorted(self.source_styles):
                self.source_listbox.insert(END, style)
            
            self.log(f"源文档样式分析完成，共 {len(self.source_styles)} 种样式")
        except Exception as e:
            messagebox.showerror("错误", f"分析源文档样式失败: {e}")
    
    def analyze_template_styles(self, template_file):
        """分析模板文档样式"""
        self.template_styles = set()
        try:
            doc = Document(template_file)
            for style in doc.styles:
                if style.type == WD_STYLE_TYPE.PARAGRAPH:
                    self.template_styles.add(style.name)
            
            # 更新列表框
            self.template_listbox.delete(0, END)
            for style in sorted(self.template_styles):
                self.template_listbox.insert(END, style)
            
            # 更新应答句样式的下拉框选项
            template_style_list = sorted(self.template_styles)
            self.answer_style_combo['values'] = template_style_list
            self.answer_source_combo['values'] = template_style_list
            self.answer_copy_combo['values'] = template_style_list
            self.hint_style_combo['values'] = template_style_list
            self.log(f"已更新应答样式下拉框，共 {len(template_style_list)} 个选项")
            
            # 如果当前应答样式不在模板样式中，设置为第一个样式或保持原值
            current_answer_style = self.answer_style.get()
            if current_answer_style not in template_style_list and template_style_list:
                self.answer_style.set(template_style_list[0])
                self.log(f"应答样式已自动设置为: {template_style_list[0]}")
            # 如果当前原文/应答原文样式不在模板样式中，设置为第一个样式或保持原值
            current_answer_source = self.answer_source_style.get()
            if current_answer_source not in template_style_list and template_style_list:
                self.answer_source_style.set(template_style_list[0])
            current_answer_copy = self.answer_copy_style.get()
            if current_answer_copy not in template_style_list and template_style_list:
                self.answer_copy_style.set(template_style_list[0])
            
            # 如果当前提示语样式不在模板样式中，设置为第一个样式或保持原值
            current_hint_style = self.hint_style.get()
            if current_hint_style not in template_style_list and template_style_list:
                self.hint_style.set(template_style_list[0])
            
            self.log(f"模板文档样式分析完成，共 {len(self.template_styles)} 种样式")
        except Exception as e:
            messagebox.showerror("错误", f"分析模板文档样式失败: {e}")
    
    def open_style_mapping(self):
        """打开样式映射配置对话框"""
        if not self.template_styles:
            messagebox.showwarning("警告", "请先选择模板文档以加载样式信息")
            return
        
        # 确定要配置的文件
        target_file = None
        target_styles = None
        
        if self.selected_source_file and self.selected_source_file in self.file_style_maps:
            # 有选中的文件，针对该文件配置
            target_file = self.selected_source_file
            target_styles = self.file_style_maps[target_file]['styles']
            current_mapping = self.file_style_maps[target_file].get('mapping', {})
        elif len(self.source_files) == 1:
            # 单文件模式
            target_file = self.source_files[0]
            if target_file in self.file_style_maps:
                target_styles = self.file_style_maps[target_file]['styles']
                current_mapping = self.file_style_maps[target_file].get('mapping', {})
            else:
                target_styles = self.source_styles
                current_mapping = self.custom_style_map
        else:
            messagebox.showwarning("警告", "请先在源文件列表中选择一个文件")
            return
        
        if not target_styles:
            messagebox.showwarning("警告", "所选文件没有样式信息。\n\n请先点击「浏览...」加载源文档，等待样式分析完成后，再点击「配置样式映射」按钮。")
            return
        
        # 检查模板样式是否为空
        if not self.template_styles:
            messagebox.showwarning("警告", "模板文档的样式列表为空。\n\n请先选择模板文档，等待样式分析完成后，再点击「配置样式映射」按钮。")
            return
        
        # 获取该文件的当前表格/图片样式配置
        # 优先使用文件级配置，其次使用默认配置中的表格/图片样式定义
        current_tbl_img_config = {}
        if target_file in self.file_style_maps and self.file_style_maps[target_file].get('tbl_img_config'):
            current_tbl_img_config = self.file_style_maps[target_file]['tbl_img_config']
        else:
            # 从 default_config.json 中读取默认的表格/图片样式定义
            current_tbl_img_config = {
                'enable_table_style': self.default_config.get('enable_table_style', 0),
                'table_style': self.default_config.get('table_style', 'Body Text'),
                'table_answer_style': self.default_config.get('table_answer_style', 'Body Text'),
                'enable_image_style': self.default_config.get('enable_image_style', 0),
                'image_style': self.default_config.get('image_style', 'Body Text'),
                'image_answer_style': self.default_config.get('image_answer_style', 'Body Text')
            }
        
        # 获取当前应答句配置
        current_answer_config = self.default_config.get('answer_config', {})
        if not current_answer_config:
            current_answer_config = {
                'do_answer': self.default_config.get('do_answer', 0),
                'answer_text': self.default_config.get('answer_text', '应答：本投标人理解并满足要求。'),
                'answer_style': self.default_config.get('answer_style', '应答句'),
                'answer_mode': self.default_config.get('answer_mode', '章节标题后插入'),
                'answer_source_style': self.default_config.get('answer_source_style', ''),
                'answer_copy_style': self.default_config.get('answer_copy_style', ''),
            }
        
        # 获取当前列表段落兜底配置
        current_list_config = {
            'method': self.default_config.get('list_method', 'bullet'),
            'bullet': self.default_config.get('list_bullet', '● '),
            'style': self.default_config.get('list_style', 'Body Text'),
            'answer_method': self.default_config.get('list_answer_method', 'bullet'),
            'answer_bullet': self.default_config.get('list_answer_bullet', '● '),
            'answer_style': self.default_config.get('list_answer_style', 'Body Text'),
        }
        
        # 打开对话框（传入保存的默认映射、保存回调、应答句配置、表格/图片/列表配置、清除章节编号配置）
        current_remove_chapter_label = self.default_config.get('remove_chapter_label', 0)
        
        # 优先使用文件级配置
        if target_file in self.file_style_maps:
            if self.file_style_maps[target_file].get('answer_config'):
                current_answer_config = self.file_style_maps[target_file]['answer_config']
            if self.file_style_maps[target_file].get('list_config'):
                current_list_config = self.file_style_maps[target_file]['list_config']
        
        dialog = StyleMappingDialog(self.root, target_styles, self.template_styles, current_mapping,
                                     saved_default_mapping=self.default_style_map,
                                     save_default_callback=self.save_default_style_map,
                                     current_tbl_img_config=current_tbl_img_config,
                                     current_remove_chapter_label=current_remove_chapter_label,
                                     current_answer_config=current_answer_config,
                                     current_list_config=current_list_config)
        mapping_result, answer_config, tbl_img_config, list_config, remove_chapter_label = dialog.show()
        
        if mapping_result is not None:
            # 保存清除章节标题编号配置
            self.remove_chapter_label = remove_chapter_label
            
            # 保存映射配置到对应文件
            if target_file in self.file_style_maps:
                self.file_style_maps[target_file]['mapping'] = mapping_result
                self.file_style_maps[target_file]['tbl_img_config'] = tbl_img_config
                self.file_style_maps[target_file]['remove_chapter_label'] = remove_chapter_label
                self.file_style_maps[target_file]['answer_config'] = answer_config
                self.file_style_maps[target_file]['list_config'] = list_config
            else:
                self.custom_style_map = mapping_result
                self.custom_tbl_img_config = tbl_img_config or {}
                self.custom_remove_chapter_label = remove_chapter_label
                self.custom_answer_config = answer_config or {}
                self.custom_list_config = list_config or {}
            
            configured_count = sum(1 for v in mapping_result.values() if v)
            filename = os.path.basename(target_file) if target_file else "当前文件"
            self.log(f"已为 {filename} 配置样式映射，共 {configured_count} 个映射关系")
            # 打印清除章节编号的配置状态
            ch_label_info = "已勾选" if remove_chapter_label else "未勾选"
            self.log(f"  清除章节标题编号: {ch_label_info}")
            # 记录应答句和列表配置状态
            do_answer = answer_config.get('do_answer', 0) if answer_config else 0
            answer_mode = answer_config.get('answer_mode', '') if answer_config else ''
            self.log(f"  应答句: {'已启用(' + answer_mode + ')' if do_answer else '未启用'}")
            self.log(f"  列表段落兜底: {list_config.get('method', 'bullet') if list_config else 'bullet'}")
            messagebox.showinfo("成功", f"样式映射配置完成！\n\n文件: {filename}\n共 {configured_count} 个样式映射关系\n清除章节标题编号: {ch_label_info}")
    
    def log(self, message, detailed=False):
        """添加日志"""
        # 如果是详细日志，检查是否启用详细日志模式
        if detailed and not self.detailed_log_var.get():
            return
        
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.log_text.insert(END, f"[{timestamp}] {message}\n")
        self.log_text.see(END)
        self.root.update_idletasks()
    
    def clear_log(self):
        """清空日志"""
        self.log_text.delete(1.0, END)
    
    def validate_inputs(self):
        """验证输入"""
        if not self.source_files:
            messagebox.showwarning("警告", "请选择至少一个源文档")
            return False
        
        if not self.template_file.get():
            messagebox.showwarning("警告", "请选择模板文档")
            return False
        
        if not os.path.exists(self.template_file.get()):
            messagebox.showerror("错误", f"模板文件不存在: {self.template_file.get()}")
            return False
        
        for sf in self.source_files:
            if not os.path.exists(sf):
                messagebox.showerror("错误", f"源文件不存在: {sf}")
                return False
        
        return True
    
    def do_com_conversion(self):
        """使用 Word COM 进行转换"""
        try:
            # 检查 pywin32 是否安装
            try:
                import win32com.client
            except ImportError:
                self.log("\n✗ 错误：未安装 pywin32")
                self.log("  请运行: pip install pywin32")
                self.root.after(0, lambda: messagebox.showerror(
                    "错误", 
                    "未安装 pywin32\n\n请运行命令安装:\npip install pywin32"
                ))
                return
            
            self.log("\n正在启动 Word...")
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            
            success_count = 0
            fail_count = 0
            
            for idx, source in enumerate(self.source_files, 1):
                base = os.path.splitext(source)[0]
                output = f"{base}_converted.docx"
                
                self.root.after(0, lambda i=idx, t=len(self.source_files), n=os.path.basename(source): 
                    self.log(f"\n[{i}/{t}] 处理: {n}"))
                
                doc = None
                try:
                    self.root.after(0, lambda: self.log("  正在打开文档..."))
                    doc = word.Documents.Open(os.path.abspath(source))
                    
                    self.root.after(0, lambda: self.log("  正在保存..."))
                    doc.SaveAs(os.path.abspath(output))
                    
                    self.output_files.append(output)
                    self.output_listbox.insert(END, output)
                    success_count += 1
                    
                    self.root.after(0, lambda: self.log(f"  ✓ 转换成功"))
                    
                except Exception as e:
                    fail_count += 1
                    self.root.after(0, lambda err=str(e): self.log(f"  ✗ 转换失败: {err}"))
                    
                finally:
                    try:
                        if doc:
                            doc.Close(SaveChanges=0)
                    except:
                        pass
            
            # 关闭 Word
            try:
                word.Quit()
            except:
                pass
            
            self.log(f"\n{'=' * 60}")
            self.log(f"转换完成！")
            self.log(f"  成功: {success_count} 个")
            if fail_count > 0:
                self.log(f"  失败: {fail_count} 个")
            self.log(f"{'=' * 60}")
            
            self.root.after(0, lambda: messagebox.showinfo(
                "成功", 
                f"Word COM 转换完成！\n\n成功: {success_count} 个\n失败: {fail_count} 个\n\n所有 Visio 图已完整保留。"
            ))
            
        except Exception as e:
            self.log(f"\n✗ COM 转换失败: {e}")
            import traceback
            traceback.print_exc()
            self.root.after(0, lambda: messagebox.showerror("错误", f"COM 转换失败:\n{e}"))
    
    def start_conversion(self):
        """开始转换"""
        if not self.validate_inputs():
            return
        
        # 清空之前的输出文件列表
        self.output_files = []
        self.output_listbox.delete(0, END)
        
        # 重置停止标志
        self.stop_conversion_flag = False
        
        # 在新线程中执行转换
        import threading
        self.conversion_thread = threading.Thread(target=self.do_conversion)
        self.conversion_thread.daemon = True
        self.conversion_thread.start()
        
        # 启用停止按钮，禁用开始按钮
        self.stop_btn.config(state='normal')
        self.convert_btn.config(state='disabled')
    
    def do_conversion(self):
        """执行转换（在后台线程中运行）"""
        try:
            self.log("=" * 60)
            self.log("开始文档转换...")
            self.log(f"源文件: {len(self.source_files)} 个")
            self.log(f"模板文件: {os.path.basename(self.template_file.get())}")
            
            # 检查是否使用 Word COM 转换
            if self.use_word_com.get():
                self.log("\n⚠ 使用 Word COM 转换模式")
                self.log("  - 优点：完整保留 Visio 图和 OLE 对象")
                self.log("  - 缺点：需要安装 Word，速度较慢")
                self.do_com_conversion()
                return
            
            # 单文件模式
            if len(self.source_files) == 1:
                # 检查是否被停止
                if self.stop_conversion_flag:
                    self.log("\n⚠ 转换已停止")
                    return
                
                source = self.source_files[0]
                base = os.path.splitext(source)[0]
                output = f"{base}_converted.docx"
                
                self.log(f"\n处理文件: {os.path.basename(source)}")
                
                # 初始化进度条（固定 maximum = 100）
                self.loading_progress['maximum'] = 100
                self.loading_progress['value'] = 0
                self.root.update_idletasks()
                
                # 获取该文件的映射配置
                file_mapping = None
                file_tbl_img_config = {}
                file_answer_config = {}
                file_list_config = {}
                file_remove_chapter_label = None
                if source in self.file_style_maps and self.file_style_maps[source].get('mapping'):
                    file_mapping = self.file_style_maps[source]['mapping']
                    file_tbl_img_config = self.file_style_maps[source].get('tbl_img_config', {})
                    file_answer_config = self.file_style_maps[source].get('answer_config', {})
                    file_list_config = self.file_style_maps[source].get('list_config', {})
                    file_remove_chapter_label = self.file_style_maps[source].get('remove_chapter_label', None)
                    self.log(f"使用该文件的自定义映射: {len(file_mapping)} 个样式")
                elif self.custom_style_map:
                    file_mapping = self.custom_style_map
                    file_tbl_img_config = self.custom_tbl_img_config or {}
                    file_answer_config = getattr(self, 'custom_answer_config', {})
                    file_list_config = getattr(self, 'custom_list_config', {})
                    file_remove_chapter_label = getattr(self, 'custom_remove_chapter_label', None)
                    self.log(f"使用全局自定义映射: {len(file_mapping)} 个样式")
                elif self.default_style_map:
                    # 兜底：使用 default_config.json 中的默认映射
                    file_mapping = self.default_style_map
                    file_answer_config = {
                        'do_answer': self.default_config.get('do_answer', 0),
                        'answer_text': self.default_config.get('answer_text', '应答：本投标人理解并满足要求。'),
                        'answer_style': self.default_config.get('answer_style', '应答句'),
                        'answer_mode': self.default_config.get('answer_mode', '章节标题后插入'),
                        'answer_source_style': self.default_config.get('answer_source_style', ''),
                        'answer_copy_style': self.default_config.get('answer_copy_style', ''),
                    }
                    file_list_config = {
                        'method': self.default_config.get('list_method', 'bullet'),
                        'bullet': self.default_config.get('list_bullet', '● '),
                        'style': self.default_config.get('list_style', 'Body Text'),
                        'answer_method': self.default_config.get('list_answer_method', 'bullet'),
                        'answer_bullet': self.default_config.get('list_answer_bullet', '● '),
                        'answer_style': self.default_config.get('list_answer_style', 'Body Text'),
                    }
                    file_remove_chapter_label = self.default_config.get('remove_chapter_label', 0)
                    self.log(f"使用默认配置映射: {len(file_mapping)} 个样式")
                else:
                    self.log("使用系统默认映射规则")
                    file_remove_chapter_label = 0
                
                # 如果以上都没有设置清除章节编号配置，使用默认值
                if file_remove_chapter_label is None:
                    file_remove_chapter_label = self.default_config.get('remove_chapter_label', 0)
                
                # 进度回调 - 仅输出日志，不更新进度条
                def progress_callback(step, message):
                    self.log(f"  [{step}/7] {message}")
                
                # 警告回调 - 输出 OLE/VML 等警告信息
                def warning_callback(message):
                    def update_ui():
                        self.log(f"  ⚠ {message}")
                    self.root.after(0, update_ui)
                
                # 显示开始处理的提示
                self.root.after(0, lambda: self.log("  正在处理中，请稍候..."))
                if self.detailed_log_var.get():
                    self.log("  [详细] 已启用详细日志模式")
                
                # 输出表格/图片样式参数的日志
                tbl_style_val = file_tbl_img_config.get('table_style', 'Body Text')
                tbl_enable_val = file_tbl_img_config.get('enable_table_style', 0)
                img_style_val = file_tbl_img_config.get('image_style', 'Body Text')
                img_enable_val = file_tbl_img_config.get('enable_image_style', 0)
                self.log(f"  表格样式覆盖: {'启用(' + str(tbl_style_val) + ')' if tbl_enable_val else '未启用'} (enable={tbl_enable_val}, style={tbl_style_val})")
                self.log(f"  图片样式覆盖: {'启用(' + str(img_style_val) + ')' if img_enable_val else '未启用'} (enable={img_enable_val}, style={img_style_val})")
                self.log(f"  清除章节标题编号: {'是' if file_remove_chapter_label else '否'}")
                
                # 从配置中提取参数
                _do_answer = file_answer_config.get('do_answer', 0) if file_answer_config else 0
                _answer_text = file_answer_config.get('answer_text', '应答：本投标人理解并满足要求。') if file_answer_config else ''
                _answer_style = file_answer_config.get('answer_style', '应答句') if file_answer_config else ''
                _answer_source = file_answer_config.get('answer_source_style', '') if file_answer_config else ''
                _answer_copy = file_answer_config.get('answer_copy_style', '') if file_answer_config else ''
                _answer_mode_label = file_answer_config.get('answer_mode', '章节标题后插入') if file_answer_config else '章节标题后插入'
                _answer_mode = self.answer_mode_options.get(_answer_mode_label, "before_heading")
                _list_method = file_list_config.get('method', 'bullet') if file_list_config else 'bullet'
                _list_bullet = file_list_config.get('bullet', '● ') if file_list_config else '● '
                _list_style = file_list_config.get('style', 'Body Text') if file_list_config else 'Body Text'
                # 当 do_answer=0 时，应答原文的列表配置不应生效
                if _do_answer:
                    _list_answer_method = file_list_config.get('answer_method', 'bullet') if file_list_config else 'bullet'
                    _list_answer_bullet = file_list_config.get('answer_bullet', '● ') if file_list_config else '● '
                    _list_answer_style = file_list_config.get('answer_style', 'Body Text') if file_list_config else 'Body Text'
                else:
                    _list_answer_method = 'bullet'
                    _list_answer_bullet = '● '
                    _list_answer_style = 'Body Text'
                
                # 执行转换
                success, actual_file, msg = self.converter.full_convert(
                    source_file=source,
                    template_file=self.template_file.get(),
                    output_file=output,
                    custom_style_map=file_mapping,
                    do_mood=self.do_mood_conversion.get(),
                    answer_text=_answer_text,
                    answer_style=_answer_style,
                    answer_source_style=_answer_source,
                    answer_copy_style=_answer_copy,
                    table_answer_style=file_tbl_img_config.get('table_answer_style', ''),
                    list_bullet=_list_bullet,
                    list_method=_list_method,
                    list_style=_list_style,
                    list_answer_method=_list_answer_method,
                    list_answer_style=_list_answer_style,
                    list_answer_bullet=_list_answer_bullet,
                    do_answer_insertion=_do_answer,
                    answer_mode=_answer_mode,
                    do_hint_insertion=self.do_hint_insertion.get(),
                    hint_type=self.hint_type.get(),
                    hint_text=self.hint_text.get(),
                    hint_image_path=self.hint_image_path.get(),
                    hint_style=self.hint_style.get(),
                    progress_callback=progress_callback,
                    warning_callback=warning_callback,
                    table_style_override=file_tbl_img_config.get('table_style', 'Body Text'),
                    enable_table_style=file_tbl_img_config.get('enable_table_style', 0),
                    image_style_override=file_tbl_img_config.get('image_style', 'Body Text'),
                    enable_image_style=file_tbl_img_config.get('enable_image_style', 0),
                    remove_chapter_label=file_remove_chapter_label
                )
                
                if success:
                    # 使用实际文件名（可能包含时间戳）
                    self.output_files.append(actual_file)
                    self.output_listbox.insert(END, os.path.basename(actual_file))
                    self.log(f"\n✓ 转换成功！")
                    self.log(f"  输出文件: {actual_file}")
                    messagebox.showinfo("成功", f"转换完成！\n\n输出文件: {os.path.basename(actual_file)}")
                else:
                    self.log(f"\n✗ 转换失败: {msg}")
                    messagebox.showerror("错误", f"转换失败: {msg}")
            
            # 多文件模式
            else:
                self.log("\n多文件模式：每个文件使用各自的映射配置")
                
                # 初始化进度条（固定 maximum = 100）
                self.loading_progress['maximum'] = 100
                self.loading_progress['value'] = 0
                self.root.update_idletasks()
                
                success_count = 0
                fail_count = 0
                
                for idx, source in enumerate(self.source_files, 1):
                    # 检查是否被停止
                    if self.stop_conversion_flag:
                        self.log("\n⚠ 转换已停止")
                        break
                    
                    base = os.path.splitext(source)[0]
                    output = f"{base}_converted.docx"
                    
                    self.root.after(0, lambda i=idx, t=len(self.source_files), n=os.path.basename(source): 
                        self.log(f"\n[{i}/{t}] 处理: {n}"))
                    
                    # 获取该文件的映射配置
                    file_mapping = None
                    file_tbl_img_config = {}
                    file_answer_config = {}
                    file_list_config = {}
                    file_remove_chapter_label = None
                    if source in self.file_style_maps and self.file_style_maps[source].get('mapping'):
                        file_mapping = self.file_style_maps[source]['mapping']
                        file_tbl_img_config = self.file_style_maps[source].get('tbl_img_config', {})
                        file_answer_config = self.file_style_maps[source].get('answer_config', {})
                        file_list_config = self.file_style_maps[source].get('list_config', {})
                        file_remove_chapter_label = self.file_style_maps[source].get('remove_chapter_label', None)
                        self.root.after(0, lambda: self.log("    使用该文件的自定义映射"))
                    elif self.default_style_map:
                        # 兜底：使用 default_config.json 中的默认映射
                        file_mapping = self.default_style_map
                        file_answer_config = {
                            'do_answer': self.default_config.get('do_answer', 0),
                            'answer_text': self.default_config.get('answer_text', '应答：本投标人理解并满足要求。'),
                            'answer_style': self.default_config.get('answer_style', '应答句'),
                            'answer_mode': self.default_config.get('answer_mode', '章节标题后插入'),
                            'answer_source_style': self.default_config.get('answer_source_style', ''),
                            'answer_copy_style': self.default_config.get('answer_copy_style', ''),
                        }
                        file_list_config = {
                            'method': self.default_config.get('list_method', 'bullet'),
                            'bullet': self.default_config.get('list_bullet', '● '),
                            'style': self.default_config.get('list_style', 'Body Text'),
                            'answer_method': self.default_config.get('list_answer_method', 'bullet'),
                            'answer_bullet': self.default_config.get('list_answer_bullet', '● '),
                            'answer_style': self.default_config.get('list_answer_style', 'Body Text'),
                        }
                        file_remove_chapter_label = self.default_config.get('remove_chapter_label', 0)
                        self.root.after(0, lambda: self.log("    使用默认配置映射"))
                    else:
                        file_remove_chapter_label = 0
                        self.root.after(0, lambda: self.log("    使用系统默认映射"))
                    
                    if file_remove_chapter_label is None:
                        file_remove_chapter_label = self.default_config.get('remove_chapter_label', 0)
                    
                    # 进度回调 - 仅输出日志，不更新进度条
                    def make_progress_callback(file_idx):
                        def progress_callback(step, message):
                            self.log(f"    [{step}/7] {message}")
                        return progress_callback
                    
                    progress_callback = make_progress_callback(idx)
                    
                    # 警告回调 - 输出 OLE/VML 等警告信息
                    def warning_callback(message):
                        def update_ui():
                            self.log(f"    ⚠ {message}")
                        self.root.after(0, update_ui)
                    
                    # 显示开始处理的提示
                    self.root.after(0, lambda: self.log("    正在处理中，请稍候..."))
                    if self.detailed_log_var.get():
                        self.log("    [详细] 已启用详细日志模式")
                    
                    # 从配置中提取参数
                    _do_answer = file_answer_config.get('do_answer', 0) if file_answer_config else 0
                    _answer_text = file_answer_config.get('answer_text', '应答：本投标人理解并满足要求。') if file_answer_config else ''
                    _answer_style = file_answer_config.get('answer_style', '应答句') if file_answer_config else ''
                    _answer_source = file_answer_config.get('answer_source_style', '') if file_answer_config else ''
                    _answer_copy = file_answer_config.get('answer_copy_style', '') if file_answer_config else ''
                    _answer_mode_label = file_answer_config.get('answer_mode', '章节标题后插入') if file_answer_config else '章节标题后插入'
                    _answer_mode = self.answer_mode_options.get(_answer_mode_label, "before_heading")
                    _list_method = file_list_config.get('method', 'bullet') if file_list_config else 'bullet'
                    _list_bullet = file_list_config.get('bullet', '● ') if file_list_config else '● '
                    _list_style = file_list_config.get('style', 'Body Text') if file_list_config else 'Body Text'
                    _list_answer_method = file_list_config.get('answer_method', 'bullet') if file_list_config else 'bullet'
                    _list_answer_bullet = file_list_config.get('answer_bullet', '● ') if file_list_config else '● '
                    _list_answer_style = file_list_config.get('answer_style', 'Body Text') if file_list_config else 'Body Text'
                    
                    success, actual_file, msg = self.converter.full_convert(
                        source_file=source,
                        template_file=self.template_file.get(),
                        output_file=output,
                        custom_style_map=file_mapping,
                        do_mood=self.do_mood_conversion.get(),
                        answer_text=_answer_text,
                        answer_style=_answer_style,
                        answer_source_style=_answer_source,
                        answer_copy_style=_answer_copy,
                        table_answer_style=file_tbl_img_config.get('table_answer_style', ''),
                        list_bullet=_list_bullet,
                        list_method=_list_method,
                        list_style=_list_style,
                        list_answer_method=_list_answer_method,
                        list_answer_style=_list_answer_style,
                        list_answer_bullet=_list_answer_bullet,
                        do_answer_insertion=_do_answer,
                        answer_mode=_answer_mode,
                        do_hint_insertion=self.do_hint_insertion.get(),
                        hint_type=self.hint_type.get(),
                        hint_text=self.hint_text.get(),
                        hint_image_path=self.hint_image_path.get(),
                        hint_style=self.hint_style.get(),
                        progress_callback=progress_callback,
                        warning_callback=warning_callback,
                        table_style_override=file_tbl_img_config.get('table_style', 'Body Text'),
                        enable_table_style=file_tbl_img_config.get('enable_table_style', 0),
                        image_style_override=file_tbl_img_config.get('image_style', 'Body Text'),
                        enable_image_style=file_tbl_img_config.get('enable_image_style', 0),
                        remove_chapter_label=file_remove_chapter_label
                    )
                    
                    if success:
                        # 使用实际文件名（可能包含时间戳）
                        self.output_files.append(actual_file)
                        self.output_listbox.insert(END, os.path.basename(actual_file))
                        self.log(f"    ✓ 完成: {os.path.basename(actual_file)}")
                        success_count += 1
                    else:
                        self.log(f"    ✗ 失败: {msg}")
                        fail_count += 1
                
                self.log(f"\n{'='*60}")
                self.log(f"批量转换完成！成功: {success_count}, 失败: {fail_count}")
                messagebox.showinfo("完成", 
                    f"批量转换完成！\n\n成功: {success_count}\n失败: {fail_count}\n\n输出文件已显示在列表中，双击可打开")
        
        except Exception as e:
            self.log(f"\n✗ 发生错误: {str(e)}")
            import traceback
            traceback.print_exc()
            messagebox.showerror("错误", f"转换过程中发生错误: {str(e)}")
        
        finally:
            # 恢复按钮状态
            def restore_buttons():
                self.stop_btn.config(state='disabled')
                self.convert_btn.config(state='normal')
            self.root.after(0, restore_buttons)
    
    def open_selected_file(self, event=None):
        """打开选中的输出文件"""
        selection = self.output_listbox.curselection()
        if selection:
            index = selection[0]
            if 0 <= index < len(self.output_files):
                filepath = self.output_files[index]
                self.open_file(filepath)
    
    def open_file(self, filepath):
        """使用默认程序打开文件"""
        if os.path.exists(filepath):
            try:
                os.startfile(filepath)  # Windows
            except AttributeError:
                # macOS/Linux
                subprocess.call(['open', filepath])
        else:
            messagebox.showerror("错误", f"文件不存在: {filepath}")
    
    def open_output_folder(self):
        """打开输出文件所在文件夹"""
        if self.output_files:
            # 获取第一个输出文件的目录
            folder = os.path.dirname(os.path.abspath(self.output_files[0]))
            try:
                os.startfile(folder)
            except Exception as e:
                messagebox.showerror("错误", f"无法打开文件夹: {e}")
        else:
            messagebox.showwarning("警告", "还没有输出文件")


def main():
    """主函数"""
    root = Tk()
    
    # 创建应用
    app = DocumentConverterGUI(root)
    
    # 运行
    root.mainloop()


if __name__ == "__main__":
    main()
