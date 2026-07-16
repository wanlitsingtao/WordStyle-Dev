# -*- coding: utf-8 -*-
"""检查大纲级别标题"""
from docx import Document
from docx.oxml.ns import qn
import os

base = r'E:\LingMa\WordStyle\WordStyle-Dev\test'
files = ['1517_converted.docx', '1517_converted_154902.docx']

for fname in files:
    fpath = os.path.join(base, fname)
    if not os.path.exists(fpath):
        print(f'文件不存在: {fpath}')
        continue
    doc = Document(fpath)
    tag = '不勾选' if '154902' not in fname else '勾选'
    
    print(f'\n=== {tag}: {fname} ===')
    print(f'段落总数: {len(doc.paragraphs)}')
    
    # 统计大纲标题
    outline_count = 0
    heading_count = 0
    for i, p in enumerate(doc.paragraphs):
        pPr = p._element.find(qn('w:pPr'))
        if pPr is not None:
            olvl = pPr.find(qn('w:outlineLvl'))
            if olvl is not None:
                olvl_val = olvl.get(qn('w:val'))
                if outline_count < 10:
                    print(f'  大纲 §{i} Lv{olvl_val}: {p.text[:60]!r} style={p.style.name}')
                outline_count += 1
        
        if 'Heading' in p.style.name or p.style.name.startswith('BN_标题'):
            heading_count += 1
    
    print(f'大纲级别标题数: {outline_count}')
    print(f'Heading样式标题数: {heading_count}')
