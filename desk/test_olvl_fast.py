# -*- coding: utf-8 -*-
"""快速检查大纲标题 - 仅扫描前1000段"""
from docx import Document
from docx.oxml.ns import qn
import os, sys

base = r'E:\LingMa\WordStyle\WordStyle-Dev\test'
files = ['1517_converted.docx', '1517_converted_154902.docx']

for fname in files:
    fpath = os.path.join(base, fname)
    if not os.path.exists(fpath):
        continue
    doc = Document(fpath)
    tag = '不勾选' if '154902' not in fname else '勾选'
    
    print(f'\n=== {tag}: {fname} ===')
    para_count = len(doc.paragraphs)
    print(f'段落总数: {para_count}')
    
    # 只扫前2000段
    limit = min(2000, para_count)
    outline_count = 0
    for i in range(limit):
        p = doc.paragraphs[i]
        pPr = p._element.find(qn('w:pPr'))
        if pPr is not None:
            olvl = pPr.find(qn('w:outlineLvl'))
            if olvl is not None:
                olvl_val = olvl.get(qn('w:val'))
                if outline_count < 5:
                    text = p.text[:60].encode('utf-8', errors='replace').decode('utf-8', errors='replace')
                    print(f'  大纲 §{i} Lv{olvl_val}: {text!r} style={p.style.name}')
                outline_count += 1
    
    print(f'前{limit}段中大纲标题数: {outline_count}')
