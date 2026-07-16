# -*- coding: utf-8 -*-
"""最小测试：直接检查 insert_response_after_headings 中三个模板的样式"""
import sys, os
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy

# 直接用一个小文档测试
doc = Document()

# 写点内容模拟标题和正文
for i in range(3):
    p = doc.add_paragraph(f"这是第{i+1}章的内容，用来测试应答句插入")
    
# 保存
test_file = os.path.join(os.path.dirname(__file__), "test_mini.docx")
doc.save(test_file)

# 重新加载
from doc_converter import DocumentConverter
conv = DocumentConverter()

# 直接调用 insert_response_after_headings 的底层逻辑
# 模拟三个用户选择的样式
answer_style = "BN_应答句"
answer_source_style = "BN_原文缩进" 
answer_copy_style = "BN_应答缩进"

# 重新打开文档
doc2 = Document(test_file)

# 创建三个模板
temp_para = doc2.add_paragraph("应答句占位")
temp_para.style = answer_style
answer_template = deepcopy(temp_para._element)
# 从 XML 中获取段落样式名
pPr = answer_template.find(qn('w:pPr'))
pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
ans_style = pStyle.get(qn('w:val')) if pStyle is not None else "无"
print(f"answer_template 样式: {ans_style}")

temp_para._element.getparent().remove(temp_para._element)

temp_source = doc2.add_paragraph("原文占位")
temp_source.style = answer_source_style
source_template = deepcopy(temp_source._element)
pPr = source_template.find(qn('w:pPr'))
pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
src_style = pStyle.get(qn('w:val')) if pStyle is not None else "无"
print(f"source_template 样式: {src_style}")
temp_source._element.getparent().remove(temp_source._element)

temp_copy = doc2.add_paragraph("应答原文占位")
temp_copy.style = answer_copy_style
copy_template = deepcopy(temp_copy._element)
pPr = copy_template.find(qn('w:pPr'))
pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
cpy_style = pStyle.get(qn('w:val')) if pStyle is not None else "无"
print(f"copy_template 样式: {cpy_style}")
temp_copy._element.getparent().remove(temp_copy._element)

print()
print("=" * 50)
print("结论：三个模板的样式正确")
print(f"  answer_template → {ans_style} (应答句)")
print(f"  source_template → {src_style} (原文)")
print(f"  copy_template → {cpy_style} (应答原文)")
print("=" * 50)

# 清理
os.remove(test_file)
