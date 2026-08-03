# -*- coding: utf-8 -*-
"""
文档格式转换程序 - 核心处理模块
整合样式转换、祈使语气转换、标题后插入应答句功能
"""
import os
import sys
import re
import io
import logging
from datetime import datetime
from copy import deepcopy
from collections import defaultdict

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement, parse_xml
    from lxml import etree
    from docx.shared import Emu
    from docx.enum.style import WD_STYLE_TYPE
    from docx.image.exceptions import UnrecognizedImageError
except ImportError:
    print("错误：未安装 python-docx 库。请运行: pip install python-docx")
    sys.exit(1)

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False


# ==================== 配置常量 ====================
DEFAULT_TARGET = "Normal"
DEFAULT_TABLE_STYLE = "Body Text"
DEFAULT_IMAGE_STYLE = "Body Text"
IMAGE_SCALE_RATIO = 2 / 3
LIST_BULLET_SYMBOL = "● "
TABLE_BORDER_SIZE = '4'
TABLE_BORDER_COLOR = '000000'
ANSWER_TEXT = "应答：本投标人理解并满足要求。"
ANSWER_STYLE = "应答句"
HEADING_STYLE_IDS = {'1', '2', '3', '4', '5', '6', '7', '8', '9'}
HEADING_STYLES = {f"Heading {i}" for i in range(1, 10)}

# 样式映射表
STYLE_MAP = {
    "Heading 1": "Heading 1",
    "Heading 2": "Heading 2",
    "Heading 3": "Heading 3",
    "Heading 4": "Heading 4",
    "Heading 5": "Heading 5",
    "Heading 6": "Heading 6",
    "List Paragraph": "List Paragraph",
}

OUTLINE_STYLE_MAP = {
    1: "Heading 1", 2: "Heading 2", 3: "Heading 3",
    4: "Heading 4", 5: "Heading 5", 6: "Heading 6",
    7: "Heading 7", 8: "Heading 8", 9: "Heading 9",
}

# 祈使语气替换规则
MULTI_IMPERATIVE_TO_STATEMENT = {
    "必须": "将", "不得": "不会", "不应": "不会", "不可": "不会",
    "不能": "无法", "切勿": "不要", "严禁": "禁止", "请勿": "请避免",
    "不许": "不允许",
}

MULTI_EXCEPTIONS = [
    "不可抗力", "不得已", "不由得", "不可通行", "不可开交", "不可理喻", "不可或缺",
    "不得少于", "不得大于", "不得超过", "不得低于", "不得高于", "不得小于", "不得用于",
    "不可否认", "不可避免", "不可逆", "不可分割",
]

SINGLE_REPLACE = {"应": "将", "须": "将"}

EXCEPTION_WORDS_YING = [
    "响应", "应用", "适应", "相应", "供应", "反应", "效应", "对应", "有应", "报应",
    "呼应", "感应", "应邀", "应酬", "应允", "应声", "应景", "应试", "应变", "应付",
    "应急", "应验", "应战", "应征", "应运", "应答", "应对", "应接", "应诺", "应求",
    "应时", "应需",
]
EXCEPTION_WORDS_XU = ["必须", "无须", "无需","须知"]
REPLACE_MAP = {
    "投标人需要": "本投标人",
    "投标人需": "本投标人",
    "投标人": "本投标人",
}


def build_word_pattern(word):
    return r'(?<![a-zA-Z0-9])' + re.escape(word) + r'(?![a-zA-Z0-9])'


def clean_list_numbering(text):
    """清理开头数字编号：1、 1） (1) 等"""
    pattern = r'^\s*(?:\d+[、\)）]|[（(]\d+[）\)])\s*'
    cleaned = re.sub(pattern, '', text, count=1)
    return cleaned


MULTI_IMPERATIVE_PATTERNS = [build_word_pattern(w) for w in MULTI_IMPERATIVE_TO_STATEMENT.keys()]
MULTI_IMPERATIVE_REGEX = re.compile('|'.join(MULTI_IMPERATIVE_PATTERNS))

SINGLE_IMPERATIVE_PATTERNS = [build_word_pattern(w) for w in SINGLE_REPLACE.keys()]
SINGLE_IMPERATIVE_REGEX = re.compile('|'.join(SINGLE_IMPERATIVE_PATTERNS))

REPLACE_REGEX = None
if REPLACE_MAP:
    patterns = []
    for word, repl in REPLACE_MAP.items():
        if word.startswith("投标人"):
            pat = r'(?<![本])' + re.escape(word) + r'(?![a-zA-Z0-9])'
        else:
            pat = build_word_pattern(word)
        patterns.append(pat)
    REPLACE_REGEX = re.compile('|'.join(patterns))


class DocumentConverter:
    """文档转换器主类"""
    
    def __init__(self):
        self.logger = None
        self.stats = {"para": 0, "table": 0, "heading": 0}
        self.source_styles = set()  # 源文档中使用的样式
        self.template_styles = set()  # 模板文档中的样式
        self.list_bullet = LIST_BULLET_SYMBOL  # 列表段落符号，默认为配置常量
        
    def setup_logger(self, source_file):
        log_filename = os.path.splitext(source_file)[0] + "_err.log"
        logger = logging.getLogger(f"converter_{os.path.basename(source_file)}")
        logger.setLevel(logging.WARNING)
        if not logger.handlers:
            handler = logging.FileHandler(log_filename, encoding='utf-8')
            formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
            handler.setFormatter(formatter)
            logger.addHandler(handler)
        self.logger = logger
        return logger
    
    def get_list_virtual_styles(self, doc):
        """检测文档中具有编号/符号的列表段落，返回分类后的虚拟样式名称集合。

        对于有自动编号（numPr）的段落：
          - 如果能获取编号格式信息，根据格式类型生成：
            '1 列表段落'（数字编号）、'● 列表段落'（符号编号）
          - 如果无法获取具体格式信息，统一归为 '● 列表段落'

        同一种格式只出现一个虚拟样式名（不按序号区分）。
        """
        virtual_styles = set()
        # 先收集所有有 numPr 的段落
        for para in doc.paragraphs:
            if self.has_numbering(para):
                # 尝试获取编号格式信息
                fmt = self._detect_numbering_format(para)
                virtual_styles.add(fmt)
        return virtual_styles

    def _detect_numbering_format(self, paragraph):
        """检测段落的编号格式类型，返回分类标识字符串。

        返回格式：
          - '1 列表段落'：数字编号（1. / 1) / (1) / ①）
          - '● 列表段落'：符号编号（bullet）
        """
        # 尝试通过 numPr/numId 获取编号定义来判断是 bullet 还是 decimal
        try:
            pPr = paragraph._element.find(qn('w:pPr'))
            if pPr is not None:
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    numId_elem = numPr.find(qn('w:numId'))
                    if numId_elem is not None:
                        numId = numId_elem.get(qn('w:val'))
                        if numId:
                            # 通过文档的 numbering 部分查找编号格式
                            from docx.oxml.ns import qn as _qn
                            doc_element = paragraph._element.getroottree().getroot()
                            numbering_part = doc_element.find('.//' + _qn('w:numbering'))
                            if numbering_part is not None:
                                # 查找 num 元素
                                num_elem = numbering_part.find(f'.//' + _qn('w:num') + f'[@' + _qn('w:numId') + f'="{numId}"]')
                                if num_elem is not None:
                                    abstractNumId_elem = num_elem.find(_qn('w:abstractNumId'))
                                    if abstractNumId_elem is not None:
                                        abstractNumId = abstractNumId_elem.get(_qn('w:val'))
                                        if abstractNumId:
                                            # 查找 abstractNum 定义
                                            abs_num = numbering_part.find(f'.//' + _qn('w:abstractNum') + f'[@' + _qn('w:abstractNumId') + f'="{abstractNumId}"]')
                                            if abs_num is not None:
                                                # 查找级别定义中的 numFmt
                                                lvl = abs_num.find('.//' + _qn('w:lvl'))
                                                if lvl is not None:
                                                    numFmt = lvl.find(_qn('w:numFmt'))
                                                    if numFmt is not None:
                                                        fmt_val = numFmt.get(_qn('w:val'))
                                                        if fmt_val == 'bullet':
                                                            return '● 列表段落'
                                                        elif fmt_val == 'decimal':
                                                            return '1 列表段落'
                                            # 查找替代格式：通过 numStyleLink
                                            styleLink = abs_num.find(_qn('w:numStyleLink'))
                                            if styleLink is not None:
                                                val = styleLink.get(_qn('w:val'))
                                                if val:
                                                    return '1 列表段落'
        except Exception:
            pass

        # 如果上述方式不可行，回退到通过文本前缀检测
        text = paragraph.text.strip() if paragraph.text else ''
        if text:
            # 常见数字编号前缀
            if text[0].isdigit():
                return '1 列表段落'
            # 常见符号前缀
            if text[0] in ('●', '◆', '▪', '▸', '➢', '○', '·', '', '-', '–', '*', '+', '·'):
                return '● 列表段落'
            if text.startswith('(') or text.startswith('（'):
                return '1 列表段落'
            if text.startswith('①') or text.startswith('②') or text.startswith('③'):
                return '1 列表段落'
            # 中文数字编号
            if text[0] in '一二三四五六七八九十':
                if len(text) > 1 and text[1] in ('、', '，', '　', '.'):
                    return '1 列表段落'

        # 无法判断，默认为符号
        return '● 列表段落'

    def get_all_styles_from_doc(self, doc_or_path):
        """获取文档中使用的所有样式（包括虚拟大纲级别样式和列表段落虚拟样式）
        
        Args:
            doc_or_path: Document对象或文件路径
        Returns:
            set: 样式名集合
        """
        if isinstance(doc_or_path, str):
            from docx import Document
            doc = Document(doc_or_path)
        else:
            doc = doc_or_path
        styles = set()
        for para in doc.paragraphs:
            if para.style and para.style.name:
                styles.add(para.style.name)
        # 额外收集具有 outlineLvl 但样式为 Normal 的段落，生成虚拟大纲样式名
        outline_styles = self.get_outline_virtual_styles(doc)
        styles.update(outline_styles)
        # 额外收集具有编号/符号的列表段落，生成列表虚拟样式名
        list_styles = self.get_list_virtual_styles(doc)
        styles.update(list_styles)
        return styles
    
    def get_outline_virtual_styles(self, doc):
        """检测文档中通过大纲级别（outlineLvl）标记但无独立样式的段落，
        返回虚拟样式名称集合（如 '[大纲级别 1]'、'[大纲级别 2]'）。
        仅统计那些段落应用的样式名称为 'Normal' 或其他无大纲级别的普通样式，
        且段落自身有 outlineLvl 属性（直接设置）的段落。"""
        virtual_styles = set()
        # 收集所有已确认为标题的样式名（如 Heading 1, 2 等，这些不需要虚拟化）
        actual_heading_style_names = set()
        for style_name in HEADING_STYLES:
            actual_heading_style_names.add(style_name)
        # 加上常见的内置标题样式
        for i in range(1, 10):
            actual_heading_style_names.add(f'heading {i}')
            actual_heading_style_names.add(f'Heading{i}')
        
        for para in doc.paragraphs:
            para_style_name = para.style.name if para.style and para.style.name else 'Normal'
            # 如果段落已经有已知的标题样式名，跳过
            if para_style_name in actual_heading_style_names:
                continue
            
            elem = para._element
            pPr = elem.find(qn('w:pPr'))
            if pPr is not None:
                outline = pPr.find(qn('w:outlineLvl'))
                if outline is not None:
                    val = outline.get(qn('w:val'))
                    if val is not None:
                        try:
                            level = int(val) + 1  # 转为 1-9 级别
                            if 1 <= level <= 9:
                                virtual_styles.add(f'[大纲级别 {level}]')
                        except ValueError:
                            pass
        return virtual_styles
    
    def get_template_styles(self, template_doc):
        """获取模板文档中的所有可用样式"""
        styles = set()
        for style in template_doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                styles.add(style.name)
        return styles
    
    def clear_document_content(self, doc):
        """清空文档内容但保留样式"""
        for table in doc.tables:
            table._element.getparent().remove(table._element)
        for para in doc.paragraphs:
            p = para._element
            p.getparent().remove(p)
        doc.add_paragraph()
    
    def get_outline_level(self, paragraph_or_elem, doc=None):
        """获取段落的大纲级别
        
        参数可以是：
        - python-docx 的 Paragraph 对象
        - lxml 的 XML 元素
        """
        # 判断输入类型
        if hasattr(paragraph_or_elem, '_element'):
            # 这是 Paragraph 对象
            elem = paragraph_or_elem._element
        else:
            # 这是 XML 元素
            elem = paragraph_or_elem
        
        # 安全检查：确保elem是XML元素
        if not hasattr(elem, 'tag'):
            return 0
        if elem.tag != qn('w:p'):
            return 0
        
        # 1. 从段落自身 w:pPr/w:outlineLvl 获取
        pPr = elem.find(qn('w:pPr'))
        if pPr is not None:
            outline = pPr.find(qn('w:outlineLvl'))
            if outline is not None:
                val = outline.get(qn('w:val'))
                if val is not None:
                    try:
                        level = int(val) + 1
                        return level
                    except ValueError:
                        pass
        
        # 2. 从段落应用的样式中获取（仅当提供doc时）
        if doc is not None and pPr is not None:
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                style_id = pStyle.get(qn('w:val'))
                if style_id:
                    try:
                        style = doc.styles[style_id]
                        style_elem = style._element
                        style_pPr = style_elem.find(qn('w:pPr'))
                        if style_pPr is not None:
                            outline = style_pPr.find(qn('w:outlineLvl'))
                            if outline is not None:
                                val = outline.get(qn('w:val'))
                                if val is not None:
                                    try:
                                        return int(val) + 1
                                    except ValueError:
                                        pass
                    except KeyError:
                        pass
        
        return 0
    
    def is_toc_paragraph(self, paragraph):
        """检查段落是否为目录（TOC）"""
        # 只通过域代码来判断，不依赖样式名称，避免误判
        elem = paragraph._element
        
        # 检查是否包含 TOC 域指令
        # ★ 修复：TOC 域指令必须以 "TOC" 开头（如 TOC \o "1-3"），
        # 不能仅因为域代码中含 "TOC" 子串就误判（如 HYPERLINK \l "_Toc27841" 中的书签名 _Toc27841）
        has_toc_instr = False
        for instr_text in elem.findall('.//' + qn('w:instrText')):
            if instr_text.text and instr_text.text.strip().upper().startswith('TOC'):
                has_toc_instr = True
                break
        
        if has_toc_instr:
            return True
        
        # 检查是否有 PAGEREF 域（TOC 中的页码引用）且有超链接
        has_pageref = False
        for instr_text in elem.findall('.//' + qn('w:instrText')):
            if instr_text.text and 'PAGEREF' in instr_text.text.upper():
                has_pageref = True
                break
        
        has_hyperlink = len(elem.findall('.//' + qn('w:hyperlink'))) > 0
        
        # 只有同时有 PAGEREF 和超链接时，才判定为目录
        if has_pageref and has_hyperlink:
            return True
        
        return False
    
    def has_numbering(self, paragraph):
        """检查段落是否有编号
        numId=0 表示无编号（Word 中的特殊值），应视为没有编号。
        """
        pPr = paragraph._element.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                # 检查 numId 是否为 0（0 表示无编号）
                numId = numPr.find(qn('w:numId'))
                if numId is not None:
                    val = numId.get(qn('w:val'))
                    if val == '0':
                        return False
                return True
        return False
    
    def remove_auto_numbering(self, paragraph):
        """移除自动编号"""
        pPr = paragraph._element.get_or_add_pPr()
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            pPr.remove(numPr)
    
    def remove_manual_numbering(self, text):
        """移除手动编号（智能判断中文数字是否为编号）
        
        仅匹配独立的编号格式，不会误伤文本中的数字（如"17号线"中的"17"）。
        改进：支持"1 .总则"（数字与分隔符之间有空格）等变体格式。
        """
        fragment_patterns = [
            # 字母+数字多级编号：如"C5.1.1.1"、"A1.2"、"B3.4.5"等附录/章节编号
            # 允许点之间有空格：如"C5. 1.1.1"
            r'[A-Za-z]\d+(?:\s*\.\s*\d+)*(?:\s*[\.、，,）\)\s]|(?=\D|$))?',
            # 多级阿拉伯数字编号："1.1"、"1.1.1"、"14.2.1"等，点分隔，后面可选分隔符或直接跟文字
            # 允许数字与点之间有空格，如"1 .总则"
            r'\d+\s*\.\s*\d+(?:\s*\.\s*\d+)*(?:\s*[\.、，,）\)\s]|(?=\D|$))',
            # 单级阿拉伯数字编号："1."、"1、"、"1）"等，必须有显式分隔符，避免误伤"17号线"
            # 允许数字与分隔符之间有空格：如"1 .总则"、"1 、总则"
            r'\d+\s*[\.、，,）\)\s]',
            # 中文数字后带分隔符才视为编号
            r'[一二三四五六七八九十]+\s*[、.．)）]',
            # 括号内的中/阿拉伯数字编号（如"（二）"、"（1）"）
            r'（[一二三四五六七八九十0-9]+）',
            r'\([0-9]+\)',
            r'[①-⑩]',
            r'[A-Za-z]\.[\s、，]',
        ]
        pattern = r'^\s*(' + '|'.join(fragment_patterns) + r')[\s、，]*'
        compiled = re.compile(pattern)
        cleaned = text
        while True:
            m = compiled.match(cleaned)
            if m:
                cleaned = cleaned[m.end():]
            else:
                break
        return cleaned
    
    def _clean_residual_numbering_artifacts(self, text):
        """清理自动编号段落中残留的手动编号痕迹
        
        当段落具有自动编号（numPr）时，原文中可能残留：
        - 前导点号：如".总则"（原始文档中自动编号提供"1"，文本残留"."）
        - 空格+点号：如" .总则"
        - 前导空格：如"  概述"
        - 其他分隔符残留
        
        此方法仅在段落有自动编号时调用，用于清理这些残留字符。
        """
        if not text:
            return text
        
        # 清理前导的点号、空格、分隔符等（这些是手动编号被自动编号取代后残留的）
        cleaned = re.sub(r'^[\s\.．、，,）\)]+', '', text)
        
        return cleaned
    
    def remove_chapter_section_marking(self, text):
        """移除"第X章/第X节/第X篇/第X部分"等章节标记
        
        匹配如：第一章、第一节、第一篇、第二部分、第二章、第二节等。
        仅对文本中实际存在的章节标记进行清理，不影响自动编号。
        """
        if not text:
            return text
        # 匹配 "第[一二三四五六七八九十]+[章节篇部分][\s、，]*" 开头
        # 注意："部分"为两字词，使用 (?:部分|[章节篇]) 来同时匹配单字和双字
        chapter_pattern = r'^\s*第[一二三四五六七八九十]+(?:部分|[章节篇])[\s、，]*'
        cleaned = re.sub(chapter_pattern, '', text).strip()
        return cleaned

    def _copy_numPr_to_paragraph(self, target_para, src_numPr):
        """复制自动编号定义(numPr)到目标段落"""
        if src_numPr is None:
            return
        new_pPr = target_para._element.find(qn('w:pPr'))
        if new_pPr is None:
            new_pPr = etree.SubElement(target_para._element, qn('w:pPr'))
            target_para._element.insert(0, new_pPr)
        new_numPr = etree.SubElement(new_pPr, qn('w:numPr'))
        for child_tag in ['w:numId', 'w:ilvl']:
            src_child = src_numPr.find(qn(child_tag))
            if src_child is not None:
                val = src_child.get(qn('w:val'))
                if val is not None:
                    new_child = etree.SubElement(new_numPr, qn(child_tag))
                    new_child.set(qn('w:val'), val)

    def _is_chapter_style_numbering(self, para):
        """判断段落的自动编号是否是章节样式（如"第%1节"、"第%1章"等）
        
        返回 True 如果编号模板中包含"第"字，表示是章节标记类编号。
        用于在不勾选"清除章/节/篇"时，仅对章节类编号进行解析和保留。
        """
        result = self._get_numbering_lvl_text(para)
        if result and '第' in result:
            return True
        return False

    def _get_numbering_lvl_text(self, para):
        """获取段落自动编号的 lvlText 模板
        
        返回 lvlText 的 val 属性值（如"第%1节"、"%1"、"%1.%2"等），
        如果无法获取则返回空字符串。
        """
        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            return ''
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            return ''
        numId_elem = numPr.find(qn('w:numId'))
        ilvl_elem = numPr.find(qn('w:ilvl'))
        if numId_elem is None or ilvl_elem is None:
            return ''
        numId = numId_elem.get(qn('w:val'))
        ilvl = ilvl_elem.get(qn('w:val'))
        if numId is None or ilvl is None:
            return ''
        
        try:
            doc = para.part.document
            numbering_part = doc.part.numbering_part
            root = numbering_part._element
            nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            abstractNumId = None
            for num in root.findall('.//w:num', nsmap):
                nid = num.get(qn('w:numId'))
                if nid == numId:
                    abs_ref = num.find('w:abstractNumId', nsmap)
                    if abs_ref is not None:
                        abstractNumId = abs_ref.get(qn('w:val'))
                    break
            if abstractNumId is None:
                return ''
            
            abstractNum = None
            for an in root.findall('.//w:abstractNum', nsmap):
                aid = an.get(qn('w:abstractNumId'))
                if aid == abstractNumId:
                    abstractNum = an
                    break
            if abstractNum is None:
                return ''
            
            current_lvl = None
            for l in abstractNum.findall('w:lvl', nsmap):
                li = l.get(qn('w:ilvl'))
                if li == ilvl:
                    current_lvl = l
                    break
            if current_lvl is None:
                return ''
            
            lvlText_elem = current_lvl.find('w:lvlText', nsmap)
            if lvlText_elem is None:
                return ''
            return lvlText_elem.get(qn('w:val'), '')
        except Exception:
            return ''

    def _resolve_auto_numbering_text(self, para):
        """解析段落自动编号的文本表示（如numId=4, ilvl=0 → "第二节"）
        
        通过访问源文档的 numbering part（内存中的 XML），查找 numId 对应的
        abstractNumId，再找到对应级别的 lvlText 和 numFmt，结合编号实例的
        当前值，生成完整的编号文本。如果无法解析，返回空字符串。
        
        改进：支持多级编号占位符（%1、%2、%3...）的完整解析。
        例如 lvlText='%1.%2.%3' 时，会分别解析级别0、1、2的编号值，
        正确替换所有占位符，避免出现未替换的"%2"等残留字符。
        """
        # 获取段落的 numPr
        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            return ''
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            return ''
        
        numId_elem = numPr.find(qn('w:numId'))
        ilvl_elem = numPr.find(qn('w:ilvl'))
        if numId_elem is None or ilvl_elem is None:
            return ''
        
        numId = numId_elem.get(qn('w:val'))
        ilvl = ilvl_elem.get(qn('w:val'))
        if numId is None or ilvl is None:
            return ''
        
        try:
            # 获取文档的 numbering part（内存中的 XML）
            doc = para.part.document
            numbering_part = doc.part.numbering_part
            root = numbering_part._element
            
            nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            # 查找 numId 对应的抽象编号
            abstractNumId = None
            for num in root.findall('.//w:num', nsmap):
                nid = num.get(qn('w:numId'))
                if nid == numId:
                    abs_ref = num.find('w:abstractNumId', nsmap)
                    if abs_ref is not None:
                        abstractNumId = abs_ref.get(qn('w:val'))
                    break
            
            if abstractNumId is None:
                return ''
            
            # 查找抽象编号定义
            abstractNum = None
            for an in root.findall('.//w:abstractNum', nsmap):
                aid = an.get(qn('w:abstractNumId'))
                if aid == abstractNumId:
                    abstractNum = an
                    break
            
            if abstractNum is None:
                return ''
            
            # 查找当前级别的定义，获取 lvlText 模板
            current_lvl = None
            for l in abstractNum.findall('w:lvl', nsmap):
                li = l.get(qn('w:ilvl'))
                if li == ilvl:
                    current_lvl = l
                    break
            
            if current_lvl is None:
                return ''
            
            # 获取文本模板
            lvlText_elem = current_lvl.find('w:lvlText', nsmap)
            if lvlText_elem is None:
                return ''
            lvlText = lvlText_elem.get(qn('w:val'), '')
            
            if not lvlText:
                return ''
            
            # 查找所有占位符 %1, %2, %3, ...
            import re
            placeholders = set(re.findall(r'%(\d+)', lvlText))
            
            # 对每个占位符级别，解析对应的编号值
            result = lvlText
            for ph in placeholders:
                level = int(ph) - 1  # %1 → level 0, %2 → level 1, ...
                level_str = str(level)
                
                # 查找该级别的定义
                lvl_def = None
                for l in abstractNum.findall('w:lvl', nsmap):
                    li = l.get(qn('w:ilvl'))
                    if li == level_str:
                        lvl_def = l
                        break
                
                if lvl_def is None:
                    # 级别定义不存在，用当前级别的 fmt 替代
                    lvl_def = current_lvl
                
                # 获取起始值（默认1）
                start_elem = lvl_def.find('w:start', nsmap)
                start_val = 1
                if start_elem is not None:
                    start_val = int(start_elem.get(qn('w:val'), '1'))
                
                # 检查 num 元素是否有该级别的 lvlOverride
                for num in root.findall('.//w:num', nsmap):
                    nid = num.get(qn('w:numId'))
                    if nid == numId:
                        for lo in num.findall('w:lvlOverride', nsmap):
                            loi = lo.get(qn('w:ilvl'))
                            if loi == level_str:
                                so = lo.find('w:startOverride', nsmap)
                                if so is not None:
                                    start_val = int(so.get(qn('w:val'), str(start_val)))
                        break
                
                # 统计该 numId+ilvl 在此段落之前出现的次数
                count = 0
                for p in doc.paragraphs:
                    # 先检查是否已到达当前段落（基于 element identity）
                    if p._element is para._element:
                        break
                    ppPr = p._element.find(qn('w:pPr'))
                    if ppPr is not None:
                        pnumPr = ppPr.find(qn('w:numPr'))
                        if pnumPr is not None:
                            pnid = pnumPr.find(qn('w:numId'))
                            if pnid is not None and pnid.get(qn('w:val')) == numId:
                                pilvl = pnumPr.find(qn('w:ilvl'))
                                if pilvl is not None and pilvl.get(qn('w:val')) == level_str:
                                    count += 1
                
                current_num = start_val + count
                
                # 获取编号格式
                numFmt_elem = lvl_def.find('w:numFmt', nsmap)
                if numFmt_elem is None:
                    continue
                numFmt = numFmt_elem.get(qn('w:val'), '')
                
                # 将编号值转换为对应格式的文本
                num_text = self._format_numbering_value(current_num, numFmt)
                if num_text is None:
                    continue
                
                # 替换占位符
                result = result.replace('%' + ph, num_text)
            
            # 编号文本后加一个空格，与段落文本分隔
            if result:
                result = result + ' '
            
            return result
        except Exception:
            return ''
    
    def _format_numbering_value(self, num, fmt):
        """将数字编号值转换为指定格式的文本"""
        if fmt == 'decimal':
            return str(num)
        elif fmt in ('upperRoman', 'upperLetter'):
            # 大写罗马数字/字母 - 简化处理
            return str(num)
        elif fmt in ('lowerRoman', 'lowerLetter'):
            return str(num)
        elif fmt in ('chineseCounting', 'chineseCountingThousand', 'japaneseCounting'):
            # 中文数字：一、二、三...
            chinese_nums = '一二三四五六七八九十'
            if num <= 10:
                return chinese_nums[num - 1]
            elif num < 100:
                tens = num // 10
                ones = num % 10
                if tens == 1:
                    result = '十'
                else:
                    result = chinese_nums[tens - 1] + '十'
                if ones > 0:
                    result += chinese_nums[ones - 1]
                return result
            return str(num)
        elif fmt == 'bullet':
            # 项目符号 - 返回空或符号本身
            return ''
        elif fmt == 'none':
            return ''
        else:
            return str(num)

    def get_target_style(self, original_style_name, template_doc, source_file=""):
        """获取目标样式名称"""
        # 使用实例变量中的样式映射，避免使用全局变量
        style_map = getattr(self, 'current_style_map', STYLE_MAP)
        target = style_map.get(original_style_name)
        if target is not None:
            try:
                template_doc.styles[target]
                return target
            except KeyError:
                try:
                    template_doc.styles[original_style_name]
                    return original_style_name
                except KeyError:
                    return DEFAULT_TARGET
        else:
            try:
                template_doc.styles[original_style_name]
                return original_style_name
            except KeyError:
                return DEFAULT_TARGET
    
    def get_image_size(self, image_bytes):
        """获取图片尺寸"""
        if not PIL_AVAILABLE:
            return None, None
        try:
            img = Image.open(io.BytesIO(image_bytes))
            return img.width, img.height
        except:
            return None, None
    
    def get_image_extent(self, blip_element):
        """从 blip 元素向上查找 wp:inline，并返回图片的原始显示尺寸 (cx, cy) 单位 EMU，若失败则返回 (None, None)"""
        parent = blip_element.getparent()
        while parent is not None:
            if parent.tag == qn('wp:inline'):
                extent = parent.find(qn('wp:extent'))
                if extent is not None:
                    cx = int(extent.get('cx', '0'))
                    cy = int(extent.get('cy', '0'))
                    return (cx, cy)
                break
            parent = parent.getparent()
        return (None, None)
    
    def resize_image_to_fixed_width(self, image_bytes, target_width_emu, dpi=96):
        """调整图片大小"""
        w_px, h_px = self.get_image_size(image_bytes)
        if w_px is None or h_px is None:
            return None, None
        w_emu = int(w_px / dpi * 914400)
        if w_emu <= target_width_emu:
            return w_emu, int(h_px / dpi * 914400)
        scale = target_width_emu / w_emu
        return int(w_emu * scale), int(h_px / dpi * 914400 * scale)
    
    def add_picture(self, run, img_bytes, page_width_emu, available_width_emu, emu_width=None, emu_height=None):
        """
        添加图片。
        若提供了源文档的显示尺寸 (emu_width, emu_height)，则优先使用；
        当图片宽度超出可用宽度时，缩放到页面宽度的 IMAGE_SCALE_RATIO，高度等比缩放。
        如果未提供尺寸，则使用图片的像素尺寸（96 DPI 计算）并做相同处理。
        """
        if not PIL_AVAILABLE:
            run.add_picture(io.BytesIO(img_bytes))
            return

        # 如果有源尺寸，直接使用；否则从像素计算
        if emu_width is not None and emu_height is not None:
            w_emu = emu_width
            h_emu = emu_height
        else:
            try:
                img = Image.open(io.BytesIO(img_bytes))
                w_px, h_px = img.size
                w_emu = int(w_px / 96 * 914400)
                h_emu = int(h_px / 96 * 914400)
            except:
                run.add_picture(io.BytesIO(img_bytes))
                return

        # 如果宽度超出可用宽度，则按页面宽度的 IMAGE_SCALE_RATIO 缩放
        if w_emu > available_width_emu:
            target_w = int(page_width_emu * IMAGE_SCALE_RATIO)
            scale = target_w / w_emu
            new_w = int(w_emu * scale)
            new_h = int(h_emu * scale)
            run.add_picture(io.BytesIO(img_bytes), width=Emu(new_w), height=Emu(new_h))
        else:
            run.add_picture(io.BytesIO(img_bytes), width=Emu(w_emu), height=Emu(h_emu))
    
    def set_table_width(self, table, width_emu):
        """设置表格宽度"""
        width_dxa = int(width_emu / 635)
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = parse_xml('<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            tbl.insert(0, tblPr)
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = parse_xml('<w:tblW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            tblPr.append(tblW)
        tblW.set(qn('w:w'), str(width_dxa))
        tblW.set(qn('w:type'), 'dxa')
    
    def set_table_borders(self, table):
        """为表格添加边框"""
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                for border_name in ['top', 'left', 'bottom', 'right']:
                    existing = tcPr.find(qn(f'w:{border_name}'))
                    if existing is not None:
                        tcPr.remove(existing)
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), TABLE_BORDER_SIZE)
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), TABLE_BORDER_COLOR)
                    tcPr.append(border)
    
    def copy_element_with_objects(self, source_elem, target_doc, target_style_name,
                                  page_width_emu, available_width_emu, para_idx=None, source_file="",
                                  warning_callback=None):
        """复制元素（包含图片、Visio图、OLE对象等）
        :param warning_callback: 警告回调函数 callback(message)
        """
        # 检查是否为段落
        if hasattr(source_elem, 'tag') and source_elem.tag == qn('w:p'):
            return self.copy_paragraph_with_images(
                source_elem, target_doc, target_style_name,
                page_width_emu, available_width_emu, 
                para_idx if para_idx is not None else 0, source_file,
                warning_callback
            )
        
        # 检查是否为表格
        elif hasattr(source_elem, 'tag') and source_elem.tag == qn('w:tbl'):
            # 这里需要找到对应的表格索引
            return None
        
        # 处理其他类型的元素（如OLE对象、Visio图等）
        else:
            return self.copy_special_element(source_elem, target_doc, target_style_name)
    
    def copy_special_element(self, source_elem, target_doc, target_style_name):
        """复制特殊元素（OLE对象、Visio图等）
        注意：OLE/VML对象的关系ID(rId)在新文档中无效，直接复制XML会导致文档损坏。
        因此OLE/VML对象只添加占位提示，不复制其XML结构。
        """
        try:
            # 创建一个新的段落来容纳特殊对象
            new_para = target_doc.add_paragraph()
            try:
                new_para.style = target_style_name
            except KeyError:
                new_para.style = target_doc.styles['Normal']
            
            # 检查是否包含OLE对象或形状
            objects = source_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object')
            shapes = source_elem.findall('.//{urn:schemas-microsoft-com:vml}shape')
            
            if objects or shapes:
                # 不复制OLE/VML的XML（会导致文档损坏），只输出提示
                new_para.add_run("[OLE对象已跳过，请手动复制]")
                return new_para
            else:
                # 如果没有特殊对象，返回空段落
                return new_para
        except Exception as e:
            print(f"警告：复制特殊元素时出错: {e}")
            # 出错时返回一个空段落
            new_para = target_doc.add_paragraph()
            try:
                new_para.style = target_style_name
            except KeyError:
                new_para.style = target_doc.styles['Normal']
            return new_para
    
    def copy_paragraph_with_images(self, source_para, target_doc, target_style_name,
                                   page_width_emu, available_width_emu, para_idx, source_file="",
                                   warning_callback=None, image_style_override=None, enable_image_style=False,
                                   remove_chapter_label=False,
                                   list_method='bullet', list_style='Body Text',
                                   enable_list_style=True):
        """复制段落（包含图片、Visio图、OLE对象等）
        :param warning_callback: 警告回调函数 callback(message)
        :param image_style_override: 图片样式覆盖（当enable_image_style=True时使用）
        :param enable_image_style: 是否启用图片样式覆盖
        """
        # 调试：检查大纲级别
        outline_level = self.get_outline_level(source_para)
        
        # 检查是否为目录段落，如果是则保持原样式
        if self.is_toc_paragraph(source_para):
            new_para = target_doc.add_paragraph()
            # 保持原始样式或应用目标样式
            try:
                new_para.style = target_style_name
            except KeyError:
                new_para.style = target_doc.styles['Normal']
            
            # 复制内容但不修改
            for run in source_para.runs:
                new_run = new_para.add_run(run.text)
                # 复制格式
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                
                # 复制图片
                blips = run._element.findall('.//' + qn('a:blip'))
                for blip in blips:
                    rId = blip.get(qn('r:embed'))
                    if rId:
                        try:
                            img_part = source_para.part.related_parts[rId]
                            img_bytes = img_part.blob
                            emu_w, emu_h = self.get_image_extent(blip)
                            pic_run = new_para.add_run()
                            self.add_picture(pic_run, img_bytes, page_width_emu, available_width_emu, emu_w, emu_h)
                        except Exception:
                            pass
            return new_para
        
        # 检查是否包含 OLE 对象或 VML 形状
        has_ole_objects = source_para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object')
        has_vml_shapes = source_para._element.findall('.//{urn:schemas-microsoft-com:vml}shape')
        
        # ★ 修复：包含 OLE/VML 对象的段落，按源文档顺序重建内容：
        # OLE对象所在run的位置插入占位提示，文本保持原位置。
        # 不能直接深度复制OLE的XML到新文档，因为OLE引用的关系ID(rId)在新文档中无效，
        # 会导致文档打开报错。文本内容必须保留，避免用户信息丢失。
        if has_ole_objects or has_vml_shapes:
            warning_msg = f"[WARNING] 段落 {para_idx} 包含 OLE/VML 对象\n  - OLE 对象数: {len(has_ole_objects)}\n  - VML 形状数: {len(has_vml_shapes)}\n  文本内容已保留，OLE对象请在原文中手动复制。"
            print(warning_msg)
            if warning_callback:
                try:
                    warning_callback(warning_msg)
                except:
                    pass
            
            # 创建新段落，设置目标样式
            # ★ 修复：OLE提示语段落使用图片兜底样式（如果启用了图片样式覆盖）
            ole_para_style = target_style_name
            if enable_image_style and image_style_override:
                try:
                    target_doc.styles[image_style_override]
                    ole_para_style = image_style_override
                except KeyError:
                    pass
            new_para = target_doc.add_paragraph()
            try:
                new_para.style = ole_para_style
            except KeyError:
                new_para.style = target_doc.styles['Normal']
            
            # 按源段落的XML子元素顺序重建内容
            # run的XML顺序与 source_para.runs 的顺序一致
            for run in source_para.runs:
                # 检查这个run是否包含OLE对象
                run_has_ole = bool(run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object'))
                run_has_vml = bool(run._element.findall('.//{urn:schemas-microsoft-com:vml}shape'))
                
                if run_has_ole or run_has_vml:
                    # OLE对象所在位置：插入占位提示（替代原OLE对象）
                    # ★ 修复：OLE占位提示应用图片兜底样式格式
                    ole_run = new_para.add_run("[OLE对象，请手动复制]")
                    if enable_image_style and image_style_override:
                        try:
                            img_style = target_doc.styles[image_style_override]
                            if img_style.font:
                                ole_run.font.bold = img_style.font.bold
                                ole_run.font.italic = img_style.font.italic
                                ole_run.font.underline = img_style.font.underline
                                if img_style.font.size:
                                    ole_run.font.size = img_style.font.size
                                if img_style.font.color and img_style.font.color.rgb:
                                    try:
                                        ole_run.font.color.rgb = img_style.font.color.rgb
                                    except:
                                        pass
                        except KeyError:
                            pass
                elif run.text:
                    # 普通文本：复制文本及格式
                    new_run = new_para.add_run(run.text)
                    # 复制基本格式
                    if run.font:
                        new_run.font.bold = run.font.bold
                        new_run.font.italic = run.font.italic
                        new_run.font.underline = run.font.underline
                        if run.font.size:
                            new_run.font.size = run.font.size
                        if run.font.color and run.font.color.rgb:
                            try:
                                new_run.font.color.rgb = run.font.color.rgb
                            except:
                                pass
            
            return new_para
        
        # 普通段落处理（原有逻辑）
        has_image = any(run._element.findall('.//' + qn('a:blip')) for run in source_para.runs)
        new_para = target_doc.add_paragraph()
        src_style_name = source_para.style.name
        
        if outline_level > 0:
            # 优先使用用户自定义样式映射
            style_map = getattr(self, 'current_style_map', STYLE_MAP)
            # 生成虚拟大纲样式名，用于查找用户映射
            virtual_style_name = f'[大纲级别 {outline_level}]'
            mapped_style = style_map.get(virtual_style_name)
            if mapped_style is None:
                # 回退：用原始样式名查找
                mapped_style = style_map.get(src_style_name)
            if mapped_style is not None:
                final_style = mapped_style
            else:
                final_style = OUTLINE_STYLE_MAP.get(outline_level)
                if final_style is None:
                    final_style = f"Heading {outline_level}"
            
            # 检查目标文档中是否存在该样式
            try:
                target_doc.styles[final_style]
            except KeyError:
                # 样式不存在，尝试从源文档复制
                print(f"[WARNING] 模板中缺少样式 '{final_style}'，尝试从源文档复制...")
                try:
                    source_style = source_para.part.document.styles[final_style]
                    # 复制样式到目标文档
                    new_style = target_doc.styles.add_style(final_style, WD_STYLE_TYPE.PARAGRAPH)
                    new_style.base_style = target_doc.styles['Normal']
                    # 复制基本格式
                    if hasattr(source_style, 'font'):
                        if source_style.font.bold:
                            new_style.font.bold = True
                        if source_style.font.size:
                            new_style.font.size = source_style.font.size
                        if source_style.font.color and source_style.font.color.rgb:
                            new_style.font.color.rgb = source_style.font.color.rgb
                    print(f"[INFO] 已从源文档复制样式 '{final_style}'")
                except Exception as e:
                    print(f"[WARNING] 无法复制样式 '{final_style}': {e}，将尝试创建基本样式")
                    try:
                        # 创建基本的标题样式
                        new_style = target_doc.styles.add_style(final_style, WD_STYLE_TYPE.PARAGRAPH)
                        new_style.base_style = target_doc.styles['Normal']
                        new_style.font.bold = True
                        # 根据大纲级别设置字体大小
                        size_map = {1: 24, 2: 20, 3: 16, 4: 14, 5: 13, 6: 12, 7: 12, 8: 12, 9: 12}
                        from docx.shared import Pt
                        new_style.font.size = Pt(size_map.get(outline_level, 12))
                        print(f"[INFO] 已创建基本样式 '{final_style}'")
                    except Exception as e2:
                        print(f"[ERROR] 无法创建样式 '{final_style}': {e2}，使用默认样式")
                        final_style = DEFAULT_TARGET
        else:
            if has_image:
                # 图片段落样式：不受样式映射影响，只按单独的图片样式定义处理
                # 1. enable_image_style=True → 使用image_style_override指定的样式
                # 2. 未启用 → 保留源样式名（模板中存在则用，否则DEFAULT_TARGET）
                if enable_image_style and image_style_override:
                    # 级别1：复选框选中，使用覆盖样式
                    final_style = image_style_override
                    try:
                        target_doc.styles[final_style]
                    except KeyError:
                        final_style = DEFAULT_TARGET
                else:
                    # 级别2：保留源样式名
                    try:
                        target_doc.styles[src_style_name]
                        final_style = src_style_name
                    except KeyError:
                        final_style = DEFAULT_TARGET
            else:
                final_style = target_style_name
        
        try:
            new_para.style = final_style
        except Exception:
            new_para.style = target_doc.styles['Normal']
        
        is_heading_by_outline = outline_level > 0
        is_heading_by_style = src_style_name in HEADING_STYLES
        is_heading_by_target = final_style in HEADING_STYLES
        is_custom_src_heading = '标题' in src_style_name or src_style_name.startswith('Heading')
        is_custom_tgt_heading = '标题' in final_style or final_style.startswith('Heading')
        
        if is_heading_by_outline or is_heading_by_style or is_heading_by_target or is_custom_src_heading or is_custom_tgt_heading:
            # 使用 para.text 而非手动拼接 runs，因为 runs 不包含超链接(w:hyperlink)内部的 run 文本
            # 例如"第六章  图纸（如有）"中，"第六章"和"图纸"在 hyperlink 内部，runs 无法获取
            full_text = source_para.text
            has_auto_numbering = self.has_numbering(source_para)
            
            if remove_chapter_label:
                # 勾选"清除第X章/第X节"：
                # 1. 移除自动编号（章节标记可能来自自动编号，如numId=17→"第二节"）
                self.remove_auto_numbering(new_para)
                # 2. 清理文本中的"第X章/第X节/第X篇"
                cleaned_text = self.remove_chapter_section_marking(full_text)
                # 3. 清理常规手动编号（"一、"、"3.1"等）
                cleaned_text = self.remove_manual_numbering(cleaned_text)
                # 4. 如果有自动编号，清理残留的手动编号痕迹（如".总则"中的前导点）
                if has_auto_numbering:
                    cleaned_text = self._clean_residual_numbering_artifacts(cleaned_text)
            else:
                # 未勾选"清除第X章/第X节"：
                # 保留章节标记，只清理常规手动编号（"一、"、"3.1"、"（1）"等）
                cleaned_text = self.remove_manual_numbering(full_text)
                # 如果有自动编号，清理残留的手动编号痕迹
                if has_auto_numbering:
                    cleaned_text = self._clean_residual_numbering_artifacts(cleaned_text)
                # ★ 修复：不直接复制 numPr（不同文档的 numId 映射不同，会导致错误编号），
                # 改为选择性解析自动编号：
                # 1. 如果编号是章节样式（如"第%1节"→"第二节"），解析并拼接到文本前
                # 2. 普通数字编号（如"%1"→"1"、"%1.%2"→"1.1"）不解析，避免文本中出现冗余编号
                if has_auto_numbering and self._is_chapter_style_numbering(source_para):
                    numbering_text = self._resolve_auto_numbering_text(source_para)
                    if numbering_text:
                        cleaned_text = numbering_text + cleaned_text
            new_para.clear()
            new_para.add_run(cleaned_text)
            for run_idx, run in enumerate(source_para.runs):
                blips = run._element.findall('.//' + qn('a:blip'))
                for blip in blips:
                    rId = blip.get(qn('r:embed'))
                    if rId:
                        try:
                            img_part = source_para.part.related_parts[rId]
                            img_bytes = img_part.blob
                            emu_w, emu_h = self.get_image_extent(blip)
                            pic_run = new_para.add_run()
                            self.add_picture(pic_run, img_bytes, page_width_emu, available_width_emu, emu_w, emu_h)
                        except Exception:
                            pass
            return new_para
        
        if self.has_numbering(source_para) and enable_list_style:
            if list_method == 'style':
                # "样式"模式：使用 Step 4 的 list_style 作为列表段落的样式。
                # 列表段落的样式由 Step 4 的"列表段落"配置区独立控制，不使用 Step 3 的样式映射结果。
                # 不加符号，不清除自动编号，只复制文本内容（保留原始编号）
                try:
                    target_doc.styles[list_style]
                    new_para.style = list_style
                except Exception:
                    try:
                        new_para.style = target_doc.styles['Normal']
                    except Exception:
                        pass
                # 检查目标样式本身是否已经包含 numPr 定义（如 BN_原文引用列表项目符号
                # 样式自带 numId=4 的项目符号编号）。如果样式自带编号，则不再从源段落
                # 拷贝 numPr，以免覆盖样式中定义的项目符号/编号格式。
                style_has_numPr = False
                try:
                    style_xml = target_doc.styles[list_style]._element.xml
                    if '<w:numPr>' in style_xml or '<w:numPr ' in style_xml:
                        style_has_numPr = True
                except Exception:
                    pass
                
                if not style_has_numPr:
                    # ★ 修复：如果目标样式没有自带编号定义（如 BN_正文），
                    # 说明用户希望将列表段落转为普通正文段落。
                    # 此时不应该从源段落复制自动编号，而应该移除原有编号。
                    # 清除源段落原有的自动编号，使其变为普通正文段落
                    self.remove_auto_numbering(new_para)
                # 使用 para.text 获取完整文本（包括超链接内部的run）
                list_runs_text = ''.join(run.text for run in source_para.runs)
                list_full_text = source_para.text
                if list_full_text and len(list_full_text) > len(list_runs_text):
                    # 有超链接内部文本，使用完整文本
                    # 但仍需处理图片
                    for run_idx, run in enumerate(source_para.runs):
                        blips = run._element.findall('.//' + qn('a:blip'))
                        if blips:
                            for blip in blips:
                                rId = blip.get(qn('r:embed'))
                                if rId:
                                    try:
                                        img_part = source_para.part.related_parts[rId]
                                        img_bytes = img_part.blob
                                        emu_w, emu_h = self.get_image_extent(blip)
                                        pic_run = new_para.add_run()
                                        self.add_picture(pic_run, img_bytes, page_width_emu, available_width_emu, emu_w, emu_h)
                                    except Exception:
                                        pass
                    if list_full_text.strip():
                        new_para.add_run(list_full_text.strip())
                else:
                    for run_idx, run in enumerate(source_para.runs):
                        blips = run._element.findall('.//' + qn('a:blip'))
                        if blips:
                            for blip in blips:
                                rId = blip.get(qn('r:embed'))
                                if rId:
                                    try:
                                        img_part = source_para.part.related_parts[rId]
                                        img_bytes = img_part.blob
                                        emu_w, emu_h = self.get_image_extent(blip)
                                        pic_run = new_para.add_run()
                                        self.add_picture(pic_run, img_bytes, page_width_emu, available_width_emu, emu_w, emu_h)
                                    except Exception:
                                        pass
                        else:
                            if run.text:
                                new_para.add_run(run.text)
                return new_para
            else:
                # "符号"模式：保留原有逻辑（添加 bullet 符号，清除编号）
                new_para.add_run(self.list_bullet)
                self.remove_auto_numbering(new_para)
                # 对于列表段落，使用专门的编号清理函数
                # 使用 para.text 获取完整文本（包括超链接内部的run）
                symbol_runs_text = ''.join(run.text for run in source_para.runs)
                symbol_full_text = source_para.text
                if symbol_full_text and len(symbol_full_text) > len(symbol_runs_text):
                    full_text = symbol_full_text
                else:
                    full_text = symbol_runs_text
                cleaned_text = clean_list_numbering(full_text)
                if cleaned_text:
                    new_para.add_run(cleaned_text)
                for run_idx, run in enumerate(source_para.runs):
                    blips = run._element.findall('.//' + qn('a:blip'))
                    for blip in blips:
                        rId = blip.get(qn('r:embed'))
                        if rId:
                            try:
                                img_part = source_para.part.related_parts[rId]
                                img_bytes = img_part.blob
                                emu_w, emu_h = self.get_image_extent(blip)
                                pic_run = new_para.add_run()
                                self.add_picture(pic_run, img_bytes, page_width_emu, available_width_emu, emu_w, emu_h)
                            except Exception:
                                pass
                return new_para
        
        # 获取段落完整文本（包括超链接内部的run文本）
        # python-docx 的 para.runs 不返回超链接(w:hyperlink)内部的run，需要用 para.text
        runs_text = ''.join(run.text for run in source_para.runs)
        full_para_text = source_para.text
        
        if full_para_text and len(full_para_text) > len(runs_text):
            # 段落中有超链接内部的隐藏文本，使用完整文本
            # 但仍需处理图片（图片在直接子级run中）
            for run_idx, run in enumerate(source_para.runs):
                blips = run._element.findall('.//' + qn('a:blip'))
                if blips:
                    for blip in blips:
                        rId = blip.get(qn('r:embed'))
                        if rId:
                            try:
                                img_part = source_para.part.related_parts[rId]
                                img_bytes = img_part.blob
                                emu_w, emu_h = self.get_image_extent(blip)
                                pic_run = new_para.add_run()
                                self.add_picture(pic_run, img_bytes, page_width_emu, available_width_emu, emu_w, emu_h)
                            except Exception:
                                pass
            # 使用完整段落文本
            cleaned_text = full_para_text.strip()
            if cleaned_text:
                new_para.add_run(cleaned_text)
        else:
            for run_idx, run in enumerate(source_para.runs):
                blips = run._element.findall('.//' + qn('a:blip'))
                if blips:
                    for blip in blips:
                        rId = blip.get(qn('r:embed'))
                        if rId:
                            try:
                                img_part = source_para.part.related_parts[rId]
                                img_bytes = img_part.blob
                                emu_w, emu_h = self.get_image_extent(blip)
                                pic_run = new_para.add_run()
                                self.add_picture(pic_run, img_bytes, page_width_emu, available_width_emu, emu_w, emu_h)
                            except Exception:
                                pass
                else:
                    if run.text:
                        new_para.add_run(run.text)
        
        return new_para
    
    def detect_merged_cells(self, table):
        """
        检测表格中的合并单元格
        :param table: python-docx 表格对象
        :return: 包含合并信息的字典 {'has_merge': bool, 'grid_span_count': int, 'v_merge_count': int}
        """
        grid_span_count = 0
        v_merge_count = 0
        
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._element.find(qn('w:tcPr'))
                if tc_pr is not None:
                    # 检测横向合并
                    grid_span_elem = tc_pr.find(qn('w:gridSpan'))
                    if grid_span_elem is not None:
                        span_val = grid_span_elem.get(qn('w:val'))
                        if span_val:
                            try:
                                span = int(span_val)
                                if span > 1:
                                    grid_span_count += 1
                            except ValueError:
                                pass
                    
                    # 检测纵向合并
                    v_merge_elem = tc_pr.find(qn('w:vMerge'))
                    if v_merge_elem is not None:
                        v_merge_count += 1
        
        has_merge = (grid_span_count > 0 or v_merge_count > 0)
        return {
            'has_merge': has_merge,
            'grid_span_count': grid_span_count,
            'v_merge_count': v_merge_count
        }
    
    def copy_table_with_images(self, source_table, target_doc, table_idx, available_width_emu, source_file="",
                               warning_callback=None, table_style_override=None, enable_table_style=False):
        """
        复制表格（包含图片、边框）
        注意：不支持合并单元格，会输出警告信息
        :param source_table: 源表格
        :param target_doc: 目标文档
        :param table_idx: 表格索引
        :param available_width_emu: 可用宽度
        :param source_file: 源文件名
        :param warning_callback: 警告回调函数
        :param table_style_override: 表格样式覆盖（当enable_table_style=True时使用）
        :param enable_table_style: 是否启用表格样式覆盖
        """
        # 检测合并单元格
        merge_info = self.detect_merged_cells(source_table)
        if merge_info['has_merge'] and warning_callback:
            warnings = []
            if merge_info['grid_span_count'] > 0:
                warnings.append(f"{merge_info['grid_span_count']}个横向合并")
            if merge_info['v_merge_count'] > 0:
                warnings.append(f"{merge_info['v_merge_count']}个纵向合并")
            warning_msg = f"表格 {table_idx} 包含合并单元格（{'、'.join(warnings)}），已跳过合并属性，请手动调整"
            warning_callback(warning_msg)
        
        # 获取源表格的行数和列数
        rows = len(source_table.rows)
        cols = len(source_table.columns)
        
        # 创建新表格
        new_table = target_doc.add_table(rows=rows, cols=cols)
        new_table.style = source_table.style
        
        # 表格单元格样式：两级决策辅助函数
        def _get_table_para_style(src_style_name):
            """决定表格内段落的目标样式
            表格不受样式映射影响，只按单独的表格样式定义处理：
            1. enable_table_style=True → 使用table_style_override指定的样式
            2. 未启用 → 保留源样式名（模板中存在则用，否则DEFAULT_TARGET）
            """
            if enable_table_style and table_style_override:
                # 级别1：复选框选中，使用覆盖样式
                try:
                    target_doc.styles[table_style_override]
                    return table_style_override
                except KeyError:
                    return DEFAULT_TARGET
            else:
                # 级别2：保留源样式名
                try:
                    target_doc.styles[src_style_name]
                    return src_style_name
                except KeyError:
                    return DEFAULT_TARGET
        
        self.set_table_width(new_table, available_width_emu)
        self.set_table_borders(new_table)
        
        # 复制单元格内容（简单的双层循环）
        for i, row in enumerate(source_table.rows):
            for j, cell in enumerate(row.cells):
                try:
                    new_cell = new_table.cell(i, j)
                except IndexError:
                    continue
                
                # 清空单元格内容
                new_cell._element.clear_content()
                
                # 复制段落内容
                for para_idx, para in enumerate(cell.paragraphs):
                    new_para = new_cell.add_paragraph()
                    src_para_style = para.style.name
                    new_para.style = _get_table_para_style(src_para_style)
                    
                    if self.has_numbering(para):
                        new_para.add_run(self.list_bullet)
                        self.remove_auto_numbering(new_para)
                        full_text = ''.join(run.text for run in para.runs)
                        cleaned_text = clean_list_numbering(full_text)
                        if cleaned_text:
                            new_para.add_run(cleaned_text)
                        for run_idx, run in enumerate(para.runs):
                            blips = run._element.findall('.//' + qn('a:blip'))
                            for blip in blips:
                                rId = blip.get(qn('r:embed'))
                                if rId:
                                    try:
                                        img_part = para.part.related_parts[rId]
                                        img_bytes = img_part.blob
                                        emu_w, emu_h = self.get_image_extent(blip)
                                        pic_run = new_para.add_run()
                                        self.add_picture(pic_run, img_bytes, available_width_emu, available_width_emu, emu_w, emu_h)
                                    except Exception:
                                        pass
                        continue
                    
                    # 检查是否包含特殊对象（Visio图、OLE对象等）
                    objects = para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object')
                    shapes = para._element.findall('.//{urn:schemas-microsoft-com:vml}shape')
                    
                    if objects or shapes:
                        # 复制特殊对象
                        for obj in objects + shapes:
                            new_obj = deepcopy(obj)
                            new_para._element.append(new_obj)
                        
                        # 输出警告
                        if warning_callback:
                            try:
                                warning_msg = f"表格 {table_idx} 单元格 [{i},{j}] 包含 OLE/VML 对象"
                                warning_callback(warning_msg)
                            except:
                                pass
                    else:
                        # 处理普通文本和图片
                        for run_idx, run in enumerate(para.runs):
                            blips = run._element.findall('.//' + qn('a:blip'))
                            if blips:
                                for blip in blips:
                                    rId = blip.get(qn('r:embed'))
                                    if rId:
                                        try:
                                            img_part = para.part.related_parts[rId]
                                            img_bytes = img_part.blob
                                            emu_w, emu_h = self.get_image_extent(blip)
                                            pic_run = new_para.add_run()
                                            self.add_picture(pic_run, img_bytes, available_width_emu, available_width_emu, emu_w, emu_h)
                                        except Exception:
                                            pass
                            else:
                                if run.text:
                                    new_para.add_run(run.text)
        
        return new_table
    
    def convert_styles(self, source_file, template_file, output_file, custom_style_map=None, list_bullet=None,
                       warning_callback=None,
                       table_style_override=None, enable_table_style=False,
                       image_style_override=None, enable_image_style=False,
                       remove_chapter_label=False,
                       list_method='bullet', list_style='Body Text',
                       enable_list_style=True):
        """
        样式转换主函数
        :param source_file: 源文件路径
        :param template_file: 模板文件路径
        :param output_file: 输出文件路径
        :param custom_style_map: 自定义样式映射表（可选）
        :param list_bullet: 列表段落符号（可选，默认为配置常量）
        :param warning_callback: 警告回调函数 callback(message)
        :param table_style_override: 表格样式覆盖（当enable_table_style=True时使用）
        :param enable_table_style: 是否启用表格样式覆盖
        :param image_style_override: 图片样式覆盖（当enable_image_style=True时使用）
        :param enable_image_style: 是否启用图片样式覆盖
        :param list_method: 列表段落处理方式 'bullet'（符号）或 'style'（样式）
        :param list_style: 列表段落兜底样式名（当list_method='style'时使用）
        :return: (success, actual_file, message)
        """
        # 使用局部样式映射副本，避免修改全局变量
        style_map = STYLE_MAP.copy()
        if custom_style_map:
            style_map.update(custom_style_map)
        
        # 将样式映射存储为实例变量，供子方法使用
        self.current_style_map = style_map
        
        # 设置列表符号
        if list_bullet is not None:
            self.list_bullet = list_bullet
        
        logger = self.setup_logger(source_file)
        
        try:
            template_doc = Document(template_file)
        except Exception as e:
            return False, f"加载模板文档失败: {e}"
        
        # 获取模板和源文档的样式
        self.template_styles = self.get_template_styles(template_doc)
        
        try:
            source_doc = Document(source_file)
        except Exception as e:
            return False, f"加载源文档失败: {e}"
        
        self.source_styles = self.get_all_styles_from_doc(source_doc)
        
        new_doc = Document(template_file)
        self.clear_document_content(new_doc)
        
        section = new_doc.sections[0]
        page_width = section.page_width
        left_margin = section.left_margin
        right_margin = section.right_margin
        available_width = page_width - left_margin - right_margin
        
        body = source_doc.element.body
        para_idx = 0
        table_idx = 0
        self.stats = {"para": 0, "table": 0, "heading": 0}
        
        for child in body:
            if child.tag == qn('w:p'):
                if para_idx < len(source_doc.paragraphs):
                    para = source_doc.paragraphs[para_idx]
                    src_style = para.style.name
                    
                    # ★ 修复：对于有编号的列表段落，使用虚拟样式名（如 "1 列表段落"）进行样式映射，
                    # 而不是用真实的样式名（通常是 "Normal"）。这样 Step 3 中用户配置的
                    # "1 列表段落" → "BN_原文引用列表项目符号" 映射才能生效。
                    # ★ 修复：标题段落（有outlineLevel或样式为Heading）即使有编号也不视为列表段落，
                    # 应走正常的标题样式映射路径。
                    is_heading_by_outline = self.get_outline_level(para) > 0
                    is_heading_by_style = src_style in HEADING_STYLES
                    style_map = getattr(self, 'current_style_map', STYLE_MAP)
                    is_heading_by_mapped = style_map.get(src_style) in HEADING_STYLES
                    is_custom_heading = '标题' in src_style or src_style.startswith('Heading')
                    if is_heading_by_outline or is_heading_by_style or is_heading_by_mapped or is_custom_heading:
                        # 标题段落：走标题样式映射
                        target_style = self.get_target_style(src_style, new_doc, source_file)
                    elif self.has_numbering(para):
                        # 非标题的列表段落：使用虚拟样式名进行映射
                        virtual_style = self._detect_numbering_format(para)
                        target_style = self.get_target_style(virtual_style, new_doc, source_file)
                    else:
                        target_style = self.get_target_style(src_style, new_doc, source_file)
                    
                    new_para = self.copy_paragraph_with_images(
                        para, new_doc, target_style,
                        page_width, available_width,
                        para_idx, source_file,
                        warning_callback,
                        image_style_override=image_style_override,
                        enable_image_style=enable_image_style,
                        remove_chapter_label=remove_chapter_label,
                        list_method=list_method,
                        list_style=list_style,
                        enable_list_style=enable_list_style
                    )
                    
                    if self.get_outline_level(para) > 0 or src_style in HEADING_STYLES or style_map.get(src_style) in HEADING_STYLES or '标题' in src_style or src_style.startswith('Heading'):
                        self.stats["heading"] += 1
                    self.stats["para"] += 1
                    para_idx += 1
            
            elif child.tag == qn('w:tbl'):
                if table_idx < len(source_doc.tables):
                    table = source_doc.tables[table_idx]
                    self.copy_table_with_images(table, new_doc, table_idx, available_width, source_file,
                                               warning_callback,
                                               table_style_override=table_style_override,
                                               enable_table_style=enable_table_style)
                    self.stats["table"] += 1
                    table_idx += 1
            
            else:
                # 处理其他类型的元素（如OLE对象、Visio图等）
                # 尝试获取样式名称，如果无法获取则使用默认样式
                try:
                    # 对于非段落/表格元素，尝试查找其所属段落的样式
                    parent_para = child.getparent()
                    while parent_para is not None and parent_para.tag != qn('w:p'):
                        parent_para = parent_para.getparent()
                    
                    if parent_para is not None:
                        # 找到父段落，尝试获取其样式
                        pPr = parent_para.find(qn('w:pPr'))
                        if pPr is not None:
                            pStyle = pPr.find(qn('w:pStyle'))
                            if pStyle is not None:
                                style_id = pStyle.get(qn('w:val'))
                                if style_id:
                                    target_style = self.get_target_style(style_id, new_doc, source_file)
                                else:
                                    target_style = DEFAULT_TARGET
                            else:
                                target_style = DEFAULT_TARGET
                        else:
                            target_style = DEFAULT_TARGET
                    else:
                        target_style = DEFAULT_TARGET
                except:
                    target_style = DEFAULT_TARGET
                
                # 复制特殊元素
                special_para = self.copy_special_element(child, new_doc, target_style)
                if special_para is not None:
                    self.stats["para"] += 1  # 计入统计
        
        # 使用重试机制保存文档
        success, actual_file, msg = self.save_with_retry(new_doc, output_file)
        if success:
            return True, actual_file, f"转换完成！段落: {self.stats['para']}, 表格: {self.stats['table']}, 标题: {self.stats['heading']}。{msg}"
        else:
            return False, output_file, msg
    
    def is_part_of_exception(self, full_text, match_start, match_end, word):
        """判断单字词是否属于例外词"""
        if word == "应":
            exceptions = EXCEPTION_WORDS_YING
        elif word == "须":
            exceptions = EXCEPTION_WORDS_XU
        else:
            return False
        
        start = max(0, match_start - 20)
        end = min(len(full_text), match_end + 20)
        substr = full_text[start:end]
        for exc in exceptions:
            if exc in substr:
                pos = substr.find(exc)
                while pos != -1:
                    exc_start = start + pos
                    exc_end = exc_start + len(exc)
                    if exc_start <= match_start < exc_end:
                        return True
                    pos = substr.find(exc, pos+1)
        return False
    
    def is_multi_exception(self, full_text, match_start, match_end, word):
        """判断多字祈使词是否属于例外词"""
        start = max(0, match_start - 20)
        end = min(len(full_text), match_end + 20)
        substr = full_text[start:end]
        for exc in MULTI_EXCEPTIONS:
            if exc in substr:
                pos = substr.find(exc)
                while pos != -1:
                    exc_start = start + pos
                    exc_end = exc_start + len(exc)
                    if exc_start <= match_start < exc_end:
                        return True
                    pos = substr.find(exc, pos+1)
        return False
    
    def replace_multiple_imperative(self, run_text, full_text, run_start_offset):
        """替换多字祈使词"""
        if not run_text:
            return run_text
        def repl(match):
            word = match.group(0)
            abs_start = run_start_offset + match.start()
            abs_end = run_start_offset + match.end()
            if self.is_multi_exception(full_text, abs_start, abs_end, word):
                return word
            return MULTI_IMPERATIVE_TO_STATEMENT.get(word, word)
        return MULTI_IMPERATIVE_REGEX.sub(repl, run_text)
    
    def replace_single_imperative(self, run_text, full_text, run_start_offset):
        """替换单字祈使词"""
        if not run_text:
            return run_text
        def repl(match):
            word = match.group(0)
            abs_start = run_start_offset + match.start()
            abs_end = run_start_offset + match.end()
            if self.is_part_of_exception(full_text, abs_start, abs_end, word):
                return word
            return SINGLE_REPLACE.get(word, word)
        return SINGLE_IMPERATIVE_REGEX.sub(repl, run_text)
    
    def process_paragraph_mood(self, para):
        """处理段落语气转换"""
        full_text = ''.join(run.text for run in para.runs)
        modified = False
        current_offset = 0
        
        for run in para.runs:
            text = run.text
            if not text:
                current_offset += len(text)
                continue
            
            new_text = text
            if REPLACE_REGEX:
                new_text = REPLACE_REGEX.sub(lambda m: REPLACE_MAP.get(m.group(0), m.group(0)), new_text)
            new_text = self.replace_multiple_imperative(new_text, full_text, current_offset)
            new_text = self.replace_single_imperative(new_text, full_text, current_offset)
            new_text = new_text.replace('将将', '将把')
            
            if new_text != text:
                run.text = new_text
                modified = True
            
            current_offset += len(text)
        
        return modified
    
    def convert_mood(self, input_file, output_file=None):
        """
        祈使语气转换
        :param input_file: 输入文件
        :param output_file: 输出文件（如果为None则覆盖原文件）
        :return: (success, actual_output_file, message)
        """
        if output_file is None:
            output_file = input_file
        
        try:
            doc = Document(input_file)
        except Exception as e:
            return False, output_file, f"加载文档失败: {e}"
        
        modified_count = 0
        para_count = 0
        
        for para in doc.paragraphs:
            para_count += 1
            # 跳过标记为 keepOriginal 的段落（copy_chapter 模式的第一份副本）
            if self._is_keep_original_paragraph(para._element):
                continue
            if self.process_paragraph_mood(para):
                modified_count += 1
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para_count += 1
                        if self._is_keep_original_paragraph(para._element):
                            continue
                        if self.process_paragraph_mood(para):
                            modified_count += 1
        
        # 清除所有 _keepOriginal_ 书签标记
        self._remove_keep_original_markers(doc)
        
        # 使用重试机制保存文档
        success, actual_file, msg = self.save_with_retry(doc, output_file)
        if success:
            return True, actual_file, f"语气转换完成！处理段落: {para_count}, 修改: {modified_count}。{msg}"
        else:
            return False, output_file, msg
    
    def _is_keep_original_paragraph(self, elem):
        """检查段落是否标记为 keepOriginal（不做语气转换）"""
        if not hasattr(elem, 'tag') or elem.tag != qn('w:p'):
            return False
        # 查找 bookmarkStart 元素，检查是否有 _keepOriginal_ 书签
        for child in elem:
            if child.tag == qn('w:bookmarkStart'):
                if child.get(qn('w:name')) == '_keepOriginal_':
                    return True
        return False
    
    def _remove_keep_original_markers(self, doc):
        """清除文档中所有 _keepOriginal_ 书签标记"""
        body = doc.element.body
        # 遍历所有段落元素
        for elem in body.iter(qn('w:p')):
            # 移除 bookmarkStart 和对应的 bookmarkEnd
            bookmark_ids_to_remove = set()
            starts_to_remove = []
            ends_to_remove = []
            
            for child in elem:
                if child.tag == qn('w:bookmarkStart'):
                    if child.get(qn('w:name')) == '_keepOriginal_':
                        bookmark_ids_to_remove.add(child.get(qn('w:id')))
                        starts_to_remove.append(child)
                elif child.tag == qn('w:bookmarkEnd'):
                    if child.get(qn('w:id')) in bookmark_ids_to_remove:
                        ends_to_remove.append(child)
            
            for start in starts_to_remove:
                elem.remove(start)
            for end in ends_to_remove:
                elem.remove(end)

    def _is_hint_paragraph(self, elem):
        """检查段落是否标记为提示语（hint）"""
        if not hasattr(elem, 'tag') or elem.tag != qn('w:p'):
            return False
        for child in elem:
            if child.tag == qn('w:bookmarkStart'):
                if child.get(qn('w:name')) == '_hint_':
                    return True
        return False

    def _is_ole_placeholder_paragraph(self, elem):
        """检查段落是否包含OLE占位提示文本 [OLE对象，请手动复制]"""
        if not hasattr(elem, 'tag') or elem.tag != qn('w:p'):
            return False
        for run in elem.findall('.//' + qn('w:r')):
            text_elem = run.find(qn('w:t'))
            if text_elem is not None and text_elem.text and '[OLE对象，请手动复制]' in text_elem.text:
                return True
        return False

    def _remove_hint_markers(self, doc):
        """清除文档中所有 _hint_ 书签标记"""
        body = doc.element.body
        for elem in body.iter(qn('w:p')):
            bookmark_ids_to_remove = set()
            starts_to_remove = []
            ends_to_remove = []
            for child in elem:
                if child.tag == qn('w:bookmarkStart'):
                    if child.get(qn('w:name')) == '_hint_':
                        bookmark_ids_to_remove.add(child.get(qn('w:id')))
                        starts_to_remove.append(child)
                elif child.tag == qn('w:bookmarkEnd'):
                    if child.get(qn('w:id')) in bookmark_ids_to_remove:
                        ends_to_remove.append(child)
            for start in starts_to_remove:
                elem.remove(start)
            for end in ends_to_remove:
                elem.remove(end)
    
    def ensure_style_exists(self, doc, style_name):
        """确保文档中存在指定样式"""
        try:
            doc.styles[style_name]
        except KeyError:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = doc.styles['Normal']
    
    def get_style_id(self, elem):
        """获取段落元素的样式ID"""
        pPr = elem.find(qn('w:pPr'))
        if pPr is None:
            return None
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is None:
            return None
        return pStyle.get(qn('w:val'))
    
    def get_style_id_by_name(self, doc, style_name):
        """通过样式名称获取样式ID（处理 name 与 style_id 不一致的情况）"""
        if not style_name:
            return None
        try:
            style = doc.styles[style_name]
            return style.style_id
        except KeyError:
            return None
    
    def is_plain_paragraph(self, elem):
        """判断是否为无样式的普通正文段落"""
        if not hasattr(elem, 'tag'):
            return False
        if elem.tag != qn('w:p'):
            return False
        pPr = elem.find(qn('w:pPr'))
        if pPr is None:
            return True
        pStyle = pPr.find(qn('w:pStyle'))
        return pStyle is None
    
    def contains_image(self, elem):
        """检查段落是否包含图片"""
        if not hasattr(elem, 'tag'):
            return False
        if elem.tag != qn('w:p'):
            return False
        blips = elem.findall('.//' + qn('a:blip'))
        return len(blips) > 0
    
    def is_table_elem(self, elem):
        """判断是否为表格"""
        if not hasattr(elem, 'tag'):
            return False
        return elem.tag == qn('w:tbl')
    
    def create_answer_paragraph_element(self, doc, answer_text, answer_style):
        """创建应答句段落XML元素"""
        temp_para = doc.add_paragraph(answer_text)
        temp_para.style = answer_style
        para_elem = deepcopy(temp_para._element)
        temp_para._element.getparent().remove(temp_para._element)
        return para_elem
    
    def is_heading_paragraph(self, elem, doc=None):
        """判断段落是否为标题（通过大纲级别 + 样式名/ID判断）
        
        支持标准 Heading 1-9 样式、自定义"标题"样式（如 BN_标题0），
        以及具有大纲级别的段落。
        
        兼容两种调用场景：
        - python-docx Paragraph 对象：.style 返回 Style 对象，有 .name 属性
        - lxml/CT_P XML 元素：.style 返回样式 ID 字符串（如 '1', 'a3', 'BN0'）
        """
        if self.get_outline_level(elem, doc) > 0:
            return True
        
        if hasattr(elem, 'style') and elem.style:
            s = elem.style
            if isinstance(s, str):
                # XML元素 (CT_P等)：.style 返回样式ID字符串
                style_id = s
                # 检查是否为内置标题样式ID (1-9)
                if style_id in HEADING_STYLE_IDS:
                    return True
                # 尝试通过doc解析样式名称，识别自定义标题样式
                if doc is not None:
                    try:
                        style_obj = doc.styles[style_id]
                        style_name = style_obj.name
                        if style_name in HEADING_STYLES:
                            return True
                        if '标题' in style_name or style_name.startswith('Heading'):
                            return True
                    except (KeyError, AttributeError):
                        pass
            elif hasattr(s, 'name'):
                # Paragraph对象：.style 返回Style对象
                style_name = s.name
                if style_name in HEADING_STYLES:
                    return True
                if '标题' in style_name or style_name.startswith('Heading'):
                    return True
        
        return False
    
    def insert_hint_paragraph(self, input_file, output_file=None,
                              hint_type='text', hint_text='招标文件原文',
                              hint_image_path=None, hint_style='Normal'):
        """
        在每个章节标题后（正文开始前）插入提示语
        :param input_file: 输入文件
        :param output_file: 输出文件（如果为None则覆盖原文件）
        :param hint_type: 提示语类型 'text' 或 'image'
        :param hint_text: 提示语文本内容（hint_type='text' 时使用）
        :param hint_image_path: 提示语图片文件路径（hint_type='image' 时使用）
        :param hint_style: 提示语段落样式名称
        :return: (success, actual_output_file, message)
        """
        if output_file is None:
            output_file = input_file

        try:
            doc = Document(input_file)
        except Exception as e:
            return False, output_file, f"加载文档失败: {e}"

        # 确保提示语样式存在
        self.ensure_style_exists(doc, hint_style)

        # 计算可用页面宽度（用于图片提示语）
        section = doc.sections[0]
        available_width = section.page_width - section.left_margin - section.right_margin

        body = doc.element.body
        children = list(body)
        new_children = []
        insert_count = 0
        total_heading_count = 0
        hint_bookmark_id = 0

        i = 0
        while i < len(children):
            child = children[i]

            if not hasattr(child, 'tag'):
                i += 1
                continue

            new_children.append(child)

            # 检查是否为标题
            if child.tag == qn('w:p') and self.is_heading_paragraph(child, doc):
                total_heading_count += 1

                # 检查下一个元素：如果不是标题，则在标题后插入提示语
                if i + 1 < len(children):
                    next_elem = children[i + 1]
                    if hasattr(next_elem, 'tag') and not self.is_heading_paragraph(next_elem, doc):
                        if hint_type == 'text':
                            # 文本提示语：创建段落并应用样式
                            hint_para = doc.add_paragraph(hint_text)
                            hint_para.style = hint_style
                            hint_elem = deepcopy(hint_para._element)
                            hint_para._element.getparent().remove(hint_para._element)
                            # 标记为提示语，避免 copy_chapter 模式重复复制
                            self._add_hint_marker(hint_elem, hint_bookmark_id)
                            hint_bookmark_id += 1
                            new_children.append(hint_elem)
                            insert_count += 1
                        elif hint_type == 'image' and hint_image_path:
                            # 图片提示语：创建空段落，应用样式，插入嵌入式图片
                            hint_para = doc.add_paragraph()
                            hint_para.style = hint_style
                            # 图片宽度设为版心宽度（页面宽度减去左右边距），高度按比例缩放
                            try:
                                picture = hint_para.add_run().add_picture(hint_image_path)
                                picture.width = available_width
                                # 使用 PIL 按原始宽高比显式计算高度，避免比例异常
                                try:
                                    from PIL import Image
                                    with Image.open(hint_image_path) as img:
                                        img_w, img_h = img.size
                                    if img_w and img_w > 0:
                                        picture.height = int(int(picture.width) * img_h / img_w)
                                except Exception:
                                    pass
                            except Exception as e:
                                print(f"插入提示语图片失败: {e}")
                                # 回退为文本提示语
                                hint_para.text = hint_text
                            hint_elem = deepcopy(hint_para._element)
                            hint_para._element.getparent().remove(hint_para._element)
                            # 标记为提示语，避免 copy_chapter 模式重复复制
                            self._add_hint_marker(hint_elem, hint_bookmark_id)
                            hint_bookmark_id += 1
                            new_children.append(hint_elem)
                            insert_count += 1

            i += 1

        # 清空并重组body
        for child in list(body):
            body.remove(child)
        for elem in new_children:
            body.append(elem)

        # 使用重试机制保存文档
        success, actual_file, msg = self.save_with_retry(doc, output_file)
        if success:
            return True, actual_file, f"已插入 {insert_count} 个章节提示语，共发现标题 {total_heading_count} 个。{msg}"
        else:
            return False, output_file, msg

    def _add_hint_marker(self, elem, bookmark_id):
        """为提示语段落添加 _hint_ 书签标记"""
        if not hasattr(elem, 'tag') or elem.tag != qn('w:p'):
            return
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), str(bookmark_id))
        bookmark_start.set(qn('w:name'), '_hint_')
        elem.insert(0, bookmark_start)
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), str(bookmark_id))
        elem.append(bookmark_end)
    
    def insert_response_after_headings(self, input_file, output_file=None, 
                                       answer_text=None, answer_style=None,
                                       answer_mode='before_heading',
                                       answer_source_style=None,
                                       answer_copy_style=None,
                                       table_answer_style=None,
                                       list_method='bullet',
                                       list_style='Body Text',
                                       list_answer_method='bullet',
                                       list_answer_style='Body Text',
                                       list_answer_bullet='● ',
                                       enable_image_style=False,
                                       image_style_override=None):
        """
        插入应答句（支持5种模式）
        :param input_file: 输入文件
        :param output_file: 输出文件（如果为None则覆盖原文件）
        :param answer_text: 应答文本
        :param answer_style: 应答样式
        :param answer_mode: 插入模式
            - 'before_heading': 章节标题后插入（默认）
            - 'after_heading': 章节末尾插入
            - 'copy_chapter': 原文+应答句+应答原文
            - 'before_paragraph': 逐段前插入
            - 'after_paragraph': 逐段后插入
        :param answer_source_style: 应答原文样式（仅 copy_chapter 模式使用）
        :param enable_image_style: 是否启用图片样式覆盖
        :param image_style_override: 图片样式覆盖名称
        :return: (success, actual_output_file, message)
        """
        if output_file is None:
            output_file = input_file
        if answer_text is None:
            answer_text = ANSWER_TEXT
        if answer_style is None:
            answer_style = ANSWER_STYLE
        if answer_source_style is None:
            answer_source_style = answer_style  # 默认与应答句样式相同
        if answer_copy_style is None:
            answer_copy_style = answer_source_style  # 默认与应答原文样式相同
        
        try:
            doc = Document(input_file)
        except Exception as e:
            return False, output_file, f"加载文档失败: {e}"
        
        self.ensure_style_exists(doc, answer_style)
        self.ensure_style_exists(doc, answer_source_style)
        self.ensure_style_exists(doc, answer_copy_style)
        if table_answer_style:
            self.ensure_style_exists(doc, table_answer_style)
        if list_style:
            self.ensure_style_exists(doc, list_style)
        if list_answer_style:
            self.ensure_style_exists(doc, list_answer_style)
        if enable_image_style and image_style_override:
            self.ensure_style_exists(doc, image_style_override)
        
        # 预创建应答段落模板
        temp_para = doc.add_paragraph(answer_text)
        temp_para.style = answer_style
        answer_template = deepcopy(temp_para._element)
        temp_para._element.getparent().remove(temp_para._element)
        
        # 预创建应答原文段落模板（copy_chapter 模式使用，用应答文本占位，后面会改内容）
        temp_source = doc.add_paragraph(answer_text)
        temp_source.style = answer_source_style
        source_template = deepcopy(temp_source._element)
        temp_source._element.getparent().remove(temp_source._element)
        
        # 预创建应答原文副本段落模板（copy_chapter 模式使用，用于应答原文副本，使用 answer_copy_style）
        temp_copy = doc.add_paragraph(answer_text)
        temp_copy.style = answer_copy_style
        copy_template = deepcopy(temp_copy._element)
        temp_copy._element.getparent().remove(temp_copy._element)
        
        body = doc.element.body
        children = list(body)
        new_children = []
        
        # 根据模式选择不同的处理逻辑
        if answer_mode == 'before_heading':
            insert_count, total_heading_count = self._insert_before_headings(
                children, new_children, answer_template, doc
            )
        elif answer_mode == 'after_heading':
            insert_count, total_heading_count = self._insert_after_headings(
                children, new_children, answer_template, doc
            )
        elif answer_mode == 'copy_chapter':
            insert_count, total_heading_count = self._insert_with_copy_chapter(
                children, new_children, answer_template, source_template, copy_template, doc,
                table_answer_style=table_answer_style,
                list_method=list_method,
                list_style=list_style,
                list_answer_method=list_answer_method,
                list_answer_style=list_answer_style,
                list_answer_bullet=list_answer_bullet,
                enable_image_style=enable_image_style,
                image_style_override=image_style_override
            )
        elif answer_mode == 'before_paragraph':
            insert_count, total_heading_count = self._insert_before_paragraphs(
                children, new_children, answer_template, doc
            )
        elif answer_mode == 'after_paragraph':
            insert_count, total_heading_count = self._insert_after_paragraphs(
                children, new_children, answer_template, doc
            )
        else:
            # 默认使用章节标题后插入
            insert_count, total_heading_count = self._insert_before_headings(
                children, new_children, answer_template, doc
            )
        
        # 清空并重组body
        for child in list(body):
            body.remove(child)
        for elem in new_children:
            body.append(elem)
        
        # 清除提示语标记（避免残留到最终文档）
        self._remove_hint_markers(doc)
        
        # 使用重试机制保存文档
        success, actual_file, msg = self.save_with_retry(doc, output_file)
        if success:
            return True, actual_file, f"已插入 {insert_count} 个应答句，共发现标题 {total_heading_count} 个。{msg}"
        else:
            return False, output_file, msg

    def _insert_before_headings(self, children, new_children, answer_template, doc):
        """
        章节标题后插入应答句（模式1：before_heading）
        判断条件：如果标题后下一个元素不是标题，则在该标题后插入应答句
        :return: (insert_count, total_heading_count)
        """
        insert_count = 0
        total_heading_count = 0
        i = 0
        
        while i < len(children):
            child = children[i]
            
            # 安全检查
            if not hasattr(child, 'tag'):
                i += 1
                continue
            
            new_children.append(child)
            
            # 检查是否为标题
            if child.tag == qn('w:p') and self.is_heading_paragraph(child, doc):
                total_heading_count += 1
                
                # 检查下一个元素
                if i + 1 < len(children):
                    next_elem = children[i + 1]
                    
                    if hasattr(next_elem, 'tag'):
                        # 如果下一个不是标题，则插入应答句
                        if not self.is_heading_paragraph(next_elem, doc):
                            answer_elem = deepcopy(answer_template)
                            new_children.append(answer_elem)
                            insert_count += 1
            
            i += 1
        
        return insert_count, total_heading_count
    
    def _insert_after_headings(self, children, new_children, answer_template, doc):
        """
        章节末尾插入应答句（模式2：after_heading）
        判断条件：如果当前不是标题 + 下一个元素是标题，则在当前元素后插入应答句
        特殊情况：
        - 全文档第一个标题前不插入（因为前面没有章节）
        - 两个标题之间不插入
        - 文章最后一段如果是正文，在其后插入
        :return: (insert_count, total_heading_count)
        """
        insert_count = 0
        total_heading_count = 0
        
        # 第一步：遍历所有元素，添加元素并统计标题
        for i, child in enumerate(children):
            # 安全检查
            if not hasattr(child, 'tag'):
                new_children.append(child)
                continue
            
            # 统计标题数量
            if child.tag == qn('w:p') and self.is_heading_paragraph(child, doc):
                total_heading_count += 1
            
            # 添加当前元素
            new_children.append(child)
            
            # 第二步：判断是否需要在当前元素后插入应答句
            # 条件：当前不是标题 + 下一个元素是标题
            if i + 1 < len(children):
                next_elem = children[i + 1]
                
                # 检查下一个元素是否为标题
                if hasattr(next_elem, 'tag') and next_elem.tag == qn('w:p'):
                    if self.is_heading_paragraph(next_elem, doc):
                        # 下一个是标题，检查当前元素是否不是标题
                        is_not_heading = True
                        
                        # 检查当前是否为标题
                        if hasattr(child, 'tag') and child.tag == qn('w:p'):
                            if self.is_heading_paragraph(child, doc):
                                is_not_heading = False
                        
                        # 如果当前不是标题（可以是正文、表格、图片等），则在其后插入应答句
                        if is_not_heading:
                            answer_elem = deepcopy(answer_template)
                            new_children.append(answer_elem)
                            insert_count += 1
            else:
                # 当前是最后一个元素
                # 如果当前不是标题，在其后插入应答句
                is_not_heading = True
                if hasattr(child, 'tag') and child.tag == qn('w:p'):
                    if self.is_heading_paragraph(child, doc):
                        is_not_heading = False
                
                if is_not_heading:
                    answer_elem = deepcopy(answer_template)
                    new_children.append(answer_elem)
                    insert_count += 1
        
        return insert_count, total_heading_count
    
    def _insert_with_copy_chapter(self, children, new_children, answer_template, source_template, copy_template, doc,
                                  table_answer_style=None,
                                  list_method='bullet', list_style='Body Text',
                                  list_answer_method='bullet', list_answer_style='Body Text',
                                  list_answer_bullet='● ',
                                  enable_image_style=False, image_style_override=None):
        """
        原文+应答句+应答原文（模式3：copy_chapter）
        最终效果：标题 → 提示语 → 原文（未转换，标记为 keepOriginal）→ 应答句 → 应答原文（语气转换后）
        注意：此模式在 full_convert 中会调换流水线顺序（先插入应答句，后语气转换），
              因此原始正文在插入应答句时仍是未转换状态，加上 keepOriginal 标记后会被跳过。
        :param answer_template: 应答句段落模板（使用 answer_style）
        :param source_template: 应答原文段落模板（使用 answer_source_style）
        :param copy_template: 应答原文副本段落模板（使用 answer_copy_style）
        :param table_answer_style: 表格应答样式（用于应答原文副本中的表格段落，若为None则使用 answer_copy_style）
        :param list_answer_method: 应答原文列表段落处理方式 'bullet'（符号）或 'style'（样式）
        :param list_answer_style: 应答原文列表段落兜底样式名
        :param list_answer_bullet: 应答原文列表段落符号
        :return: (insert_count, total_heading_count)
        """
        insert_count = 0
        total_heading_count = 0
        bookmark_id = 0  # 书签 ID 计数器
        
        def is_heading(elem):
            """判断元素是否为标题段落"""
            if not hasattr(elem, 'tag'):
                return False
            if elem.tag != qn('w:p'):
                return False
            return self.is_heading_paragraph(elem, doc)
        
        def remove_keep_original_from_element(elem):
            """移除元素中的 keepOriginal 书签标记"""
            if not hasattr(elem, 'tag') or elem.tag != qn('w:p'):
                return
            bookmark_ids = set()
            starts_to_remove = []
            ends_to_remove = []
            for child in elem:
                if child.tag == qn('w:bookmarkStart'):
                    if child.get(qn('w:name')) == '_keepOriginal_':
                        bookmark_ids.add(child.get(qn('w:id')))
                        starts_to_remove.append(child)
                elif child.tag == qn('w:bookmarkEnd'):
                    if child.get(qn('w:id')) in bookmark_ids:
                        ends_to_remove.append(child)
            for start in starts_to_remove:
                elem.remove(start)
            for end in ends_to_remove:
                elem.remove(end)
        
        def _elem_has_numbering(elem):
            """检查XML元素（w:p）是否有编号
            检查两种方式：
            1. 段落元素自身是否有 w:numPr（直接定义的编号）
            2. 段落样式（pStyle）是否在样式定义中包含 w:numPr（样式自带的编号）
            注意：numId=0 表示无编号，应视为没有编号。
            """
            pPr = elem.find(qn('w:pPr'))
            if pPr is not None:
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    # 检查 numId 是否为 0（0 表示无编号）
                    numId = numPr.find(qn('w:numId'))
                    if numId is not None:
                        val = numId.get(qn('w:val'))
                        if val == '0':
                            return False
                    return True
                # 检查段落样式是否自带编号
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    pStyle_val = pStyle.get(qn('w:val'))
                    if pStyle_val:
                        try:
                            style_xml = doc.styles[pStyle_val]._element.xml
                            if '<w:numPr>' in style_xml:
                                return True
                        except Exception:
                            pass
            return False
        
        def _get_list_answer_style_id():
            """获取应答原文列表段落的样式ID"""
            if list_answer_method == 'style' and list_answer_style:
                sid = self.get_style_id_by_name(doc, list_answer_style)
                if sid:
                    return sid
            return None

        def _apply_paragraph_style(elem, style_id):
            """将段落样式ID写入段落的 pStyle，保留已有 numPr/numId 等编号信息。"""
            if not hasattr(elem, 'tag') or elem.tag != qn('w:p') or not style_id:
                return
            pPr = elem.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                elem.insert(0, pPr)
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is None:
                pStyle = OxmlElement('w:pStyle')
                pPr.append(pStyle)
            pStyle.set(qn('w:val'), style_id)
        
        # 当前章节标题索引，None 表示尚未进入任何章节
        current_chapter_heading = None
        # 章节内容缓冲区：保存当前章节内的非提示语元素，用于生成第二份副本
        chapter_buffer = []
        
        i = 0
        while i < len(children):
            child = children[i]
            
            # 安全检查：非 XML 元素直接输出
            if not hasattr(child, 'tag'):
                new_children.append(child)
                i += 1
                continue
            
            # 遇到标题：先刷新前一个章节，再输出当前标题
            if is_heading(child):
                # 刷新前一个章节的内容
                if current_chapter_heading is not None and chapter_buffer:
                    # 检查是否有非提示语内容
                    has_real_content = any(not self._is_hint_paragraph(elem) for elem in chapter_buffer)
                    if has_real_content:
                        # 插入应答句
                        answer_elem = deepcopy(answer_template)
                        new_children.append(answer_elem)
                        insert_count += 1
                        
                        # 复制应答原文副本（用 copy_template 样式，对应应答原文格式），跳过提示语
                        for elem in chapter_buffer:
                            if self._is_hint_paragraph(elem):
                                continue
                            # 表格元素：深拷贝原始表格，并将所有段落的样式改为 table_answer_style（优先）或 answer_copy_style
                            if elem.tag == qn('w:tbl'):
                                source_elem = deepcopy(elem)
                                # 获取表格应答样式ID（优先使用 table_answer_style）
                                tbl_style_id = None
                                if table_answer_style:
                                    tbl_style_id = self.get_style_id_by_name(doc, table_answer_style)
                                if not tbl_style_id:
                                    tbl_style_id = self.get_style_id(copy_template)
                                if tbl_style_id:
                                    # 遍历表格内所有段落，修改样式
                                    for p_elem in source_elem.iter(qn('w:p')):
                                        pPr = p_elem.find(qn('w:pPr'))
                                        if pPr is None:
                                            pPr = OxmlElement('w:pPr')
                                            p_elem.insert(0, pPr)
                                        pStyle = pPr.find(qn('w:pStyle'))
                                        if pStyle is None:
                                            pStyle = OxmlElement('w:pStyle')
                                            pPr.append(pStyle)
                                        pStyle.set(qn('w:val'), tbl_style_id)
                                new_children.append(source_elem)
                                continue
                            # ★ 修复：列表段落用 deepcopy(elem) 保留原始列表结构
                            if _elem_has_numbering(elem):
                                source_elem = deepcopy(elem)
                                remove_keep_original_from_element(source_elem)
                                # 如果指定了列表样式，应用它；否则保留原始列表样式
                                if list_answer_method == 'style' and list_answer_style:
                                    list_sid = self.get_style_id_by_name(doc, list_answer_style)
                                    if list_sid:
                                        _apply_paragraph_style(source_elem, list_sid)
                                new_children.append(source_elem)
                            else:
                                # ★ 修复：OLE占位段落应用图片兜底样式
                                is_ole_placeholder = self._is_ole_placeholder_paragraph(elem)
                                if is_ole_placeholder and enable_image_style and image_style_override:
                                    # OLE占位段落使用图片兜底样式
                                    img_sid = self.get_style_id_by_name(doc, image_style_override)
                                    if img_sid:
                                        source_elem = deepcopy(elem)
                                        _apply_paragraph_style(source_elem, img_sid)
                                        remove_keep_original_from_element(source_elem)
                                    else:
                                        source_elem = deepcopy(copy_template)
                                        source_runs = source_elem.findall('.//' + qn('w:r'))
                                        orig_runs = elem.findall('.//' + qn('w:r'))
                                        for r in source_runs:
                                            source_elem.remove(r)
                                        for r in orig_runs:
                                            source_elem.append(deepcopy(r))
                                        remove_keep_original_from_element(source_elem)
                                else:
                                    # 用 copy_template 替换内容，保持应答原文副本样式（answer_copy_style）
                                    source_elem = deepcopy(copy_template)
                                    # 复制原段落的文本内容到 copy_template
                                    source_runs = source_elem.findall('.//' + qn('w:r'))
                                    orig_runs = elem.findall('.//' + qn('w:r'))
                                    # 清除模板中的 runs
                                    for r in source_runs:
                                        source_elem.remove(r)
                                    # 复制原段落的 runs
                                    for r in orig_runs:
                                        source_elem.append(deepcopy(r))
                                    remove_keep_original_from_element(source_elem)
                                    # ★ 修复：应答原文副本图片段落应用图片兜底样式
                                    if enable_image_style and image_style_override:
                                        has_img = (elem.find('.//' + qn('w:drawing')) is not None or
                                                   elem.find('.//' + qn('w:pict')) is not None or
                                                   elem.find('.//' + qn('pic:pic')) is not None or
                                                   elem.find('.//' + qn('wp:inline')) is not None or
                                                   elem.find('.//' + qn('wp:anchor')) is not None)
                                        if has_img:
                                            img_sid = self.get_style_id_by_name(doc, image_style_override)
                                            if img_sid:
                                                _apply_paragraph_style(source_elem, img_sid)
                                new_children.append(source_elem)
                
                chapter_buffer.clear()
                new_children.append(child)
                current_chapter_heading = i
                total_heading_count += 1
                i += 1
                continue
            
            # 非标题元素
            if current_chapter_heading is not None:
                # 在章节内：提示语直接输出，其他内容标记为 keepOriginal 后输出，并加入缓冲区
                if self._is_hint_paragraph(child):
                    new_children.append(child)
                else:
                    # 用 source_template 替换内容，应用原文格式（answer_source_style）
                    if child.tag == qn('w:p'):
                        child_pPr = child.find(qn('w:pPr'))
                        child_numPr = child_pPr.find(qn('w:numPr')) if child_pPr is not None else None
                        
                        # 获取段落样式ID
                        _pStyle_val = None
                        if child_pPr is not None:
                            _pStyle_elem = child_pPr.find(qn('w:pStyle'))
                            if _pStyle_elem is not None:
                                _pStyle_val = _pStyle_elem.get(qn('w:val'))
                        
                        # ★ 修复：标题段落（包括带编号的标题）不视为列表段落
                        if _pStyle_val and self.is_heading_paragraph(child, doc):
                            _is_list_para = False
                        else:
                            # 判断是否为列表段落（使用 _elem_has_numbering 函数，自动排除 numId=0 的情况）
                            _is_list_para = _elem_has_numbering(child)
                            if not _is_list_para:
                                # 补充检查1：段落样式名是否为虚拟列表样式（样式名包含"列表段落"）
                                if _pStyle_val and ('列表段落' in _pStyle_val):
                                    _is_list_para = True
                                if not _is_list_para:
                                    # 补充检查2：list_method='style'时，检查段落的 style_id 是否等于 list_style 的 style_id
                                    if _pStyle_val and list_method == 'style' and list_style:
                                        _list_sid = self.get_style_id_by_name(doc, list_style)
                                        if _list_sid and _pStyle_val == _list_sid:
                                            _is_list_para = True
                        
                        # ★ 修复：列表段落用 deepcopy(child) 保留原始列表结构，再用指定样式覆盖
                        if _is_list_para:
                            source_elem = deepcopy(child)
                            # 添加 keepOriginal 标记，使其在语气转换时保留未转换状态
                            bookmark_start = OxmlElement('w:bookmarkStart')
                            bookmark_start.set(qn('w:id'), str(bookmark_id))
                            bookmark_start.set(qn('w:name'), '_keepOriginal_')
                            source_elem.insert(0, bookmark_start)
                            bookmark_end = OxmlElement('w:bookmarkEnd')
                            bookmark_end.set(qn('w:id'), str(bookmark_id))
                            source_elem.append(bookmark_end)
                            bookmark_id += 1
                            # 如果指定了列表样式，应用它；否则保留原始列表样式
                            if list_method == 'style' and list_style:
                                list_sid = self.get_style_id_by_name(doc, list_style)
                                if list_sid:
                                    _apply_paragraph_style(source_elem, list_sid)
                                # ★ 修复：如果目标样式没有自带编号，移除此段落的自动编号
                                try:
                                    _style_xml2 = doc.styles[list_style]._element.xml
                                    if not ('<w:numPr>' in _style_xml2 or '<w:numPr ' in _style_xml2):
                                        self.remove_auto_numbering(source_elem)
                                except Exception:
                                    pass
                            new_children.append(source_elem)
                            chapter_buffer.append(source_elem)
                            # ★ 修复：图片段落应用兜底样式
                            if enable_image_style and image_style_override:
                                has_img = (child.find('.//' + qn('w:drawing')) is not None or
                                           child.find('.//' + qn('w:pict')) is not None or
                                           child.find('.//' + qn('pic:pic')) is not None or
                                           child.find('.//' + qn('wp:inline')) is not None or
                                           child.find('.//' + qn('wp:anchor')) is not None)
                                if has_img:
                                    img_sid = self.get_style_id_by_name(doc, image_style_override)
                                    if img_sid:
                                        _apply_paragraph_style(source_elem, img_sid)
                            i += 1
                            continue
                        
                        # ★ 修复：OLE占位段落使用图片兜底样式，而非source_template样式
                        is_ole_ph = self._is_ole_placeholder_paragraph(child)
                        if is_ole_ph and enable_image_style and image_style_override:
                            img_sid = self.get_style_id_by_name(doc, image_style_override)
                            if img_sid:
                                source_elem = deepcopy(child)
                                _apply_paragraph_style(source_elem, img_sid)
                            else:
                                source_elem = deepcopy(source_template)
                                source_runs = source_elem.findall('.//' + qn('w:r'))
                                orig_runs = child.findall('.//' + qn('w:r'))
                                for r in source_runs:
                                    source_elem.remove(r)
                                for r in orig_runs:
                                    source_elem.append(deepcopy(r))
                        else:
                            source_elem = deepcopy(source_template)
                            # 复制原段落的文本内容到 source_template
                            source_runs = source_elem.findall('.//' + qn('w:r'))
                            orig_runs = child.findall('.//' + qn('w:r'))
                            for r in source_runs:
                                source_elem.remove(r)
                            for r in orig_runs:
                                source_elem.append(deepcopy(r))
                        # 给原始正文段落添加 keepOriginal 标记，使其在语气转换时保留未转换状态
                        bookmark_start = OxmlElement('w:bookmarkStart')
                        bookmark_start.set(qn('w:id'), str(bookmark_id))
                        bookmark_start.set(qn('w:name'), '_keepOriginal_')
                        source_elem.insert(0, bookmark_start)
                        bookmark_end = OxmlElement('w:bookmarkEnd')
                        bookmark_end.set(qn('w:id'), str(bookmark_id))
                        source_elem.append(bookmark_end)
                        bookmark_id += 1
                        # ★ 关键修复：检查子元素是否有 numPr（列表编号），如果有则在 source_elem 中保留标记
                        if child_numPr is not None:
                            # 判断目标样式是否自带编号（list_method='style' 时）
                            _style_has_numPr = False
                            if list_method == 'style' and list_style:
                                try:
                                    _style_xml = doc.styles[list_style]._element.xml
                                    if '<w:numPr>' in _style_xml or '<w:numPr ' in _style_xml:
                                        _style_has_numPr = True
                                except Exception:
                                    pass
                            # ★ 修复：如果目标样式没有自带编号（如 BN_正文），
                            # 说明用户希望将列表段落转为普通正文段落，不应复制 numPr
                            if not _style_has_numPr:
                                # 不复制 numPr，并移除已有编号（直接操作lxml元素）
                                src_pPr2 = source_elem.find(qn('w:pPr'))
                                if src_pPr2 is not None:
                                    old_numPr = src_pPr2.find(qn('w:numPr'))
                                    if old_numPr is not None:
                                        src_pPr2.remove(old_numPr)
                            else:
                                # 在 source_elem 的 pPr 中加入 numPr 标记，供后续副本拷贝时检测列表段落
                                src_pPr = source_elem.find(qn('w:pPr'))
                                if src_pPr is None:
                                    src_pPr = OxmlElement('w:pPr')
                                    source_elem.insert(0, src_pPr)
                                src_numPr = src_pPr.find(qn('w:numPr'))
                                if src_numPr is None:
                                    # 复制完整的 numPr 元素（包括 numId 等），而不仅仅是添加空标记
                                    src_numPr = deepcopy(child_numPr)
                                    src_pPr.append(src_numPr)
                        if child_numPr is not None and list_method == 'style' and list_style and _is_list_para:
                            list_sid = self.get_style_id_by_name(doc, list_style)
                            if list_sid:
                                _apply_paragraph_style(source_elem, list_sid)
                        # ★ 修复：图片段落应用兜底样式
                        if enable_image_style and image_style_override:
                            has_img = (child.find('.//' + qn('w:drawing')) is not None or
                                       child.find('.//' + qn('w:pict')) is not None or
                                       child.find('.//' + qn('pic:pic')) is not None or
                                       child.find('.//' + qn('wp:inline')) is not None or
                                       child.find('.//' + qn('wp:anchor')) is not None)
                            if has_img:
                                img_sid = self.get_style_id_by_name(doc, image_style_override)
                                if img_sid:
                                    _apply_paragraph_style(source_elem, img_sid)
                        new_children.append(source_elem)
                        chapter_buffer.append(source_elem)
                    else:
                        new_children.append(child)
                        chapter_buffer.append(child)
            else:
                # 不在任何章节内（文档开头无标题），直接输出
                new_children.append(child)
            
            i += 1
        
        # 处理最后一个章节
        if current_chapter_heading is not None and chapter_buffer:
            has_real_content = any(not self._is_hint_paragraph(elem) for elem in chapter_buffer)
            if has_real_content:
                # 插入应答句
                answer_elem = deepcopy(answer_template)
                new_children.append(answer_elem)
                insert_count += 1
                
                # 复制第二份副本（应答原文，用 copy_template 样式），跳过提示语
                for elem in chapter_buffer:
                    if self._is_hint_paragraph(elem):
                        continue
                    # 表格元素：深拷贝原始表格，并将所有段落的样式改为 table_answer_style（优先）或 answer_copy_style
                    if elem.tag == qn('w:tbl'):
                        source_elem = deepcopy(elem)
                        # 获取表格应答样式ID（优先使用 table_answer_style）
                        tbl_style_id = None
                        if table_answer_style:
                            tbl_style_id = self.get_style_id_by_name(doc, table_answer_style)
                        if not tbl_style_id:
                            tbl_style_id = self.get_style_id(copy_template)
                        if tbl_style_id:
                            # 遍历表格内所有段落，修改样式
                            for p_elem in source_elem.iter(qn('w:p')):
                                pPr = p_elem.find(qn('w:pPr'))
                                if pPr is None:
                                    pPr = OxmlElement('w:pPr')
                                    p_elem.insert(0, pPr)
                                pStyle = pPr.find(qn('w:pStyle'))
                                if pStyle is None:
                                    pStyle = OxmlElement('w:pStyle')
                                    pPr.append(pStyle)
                                pStyle.set(qn('w:val'), tbl_style_id)
                        new_children.append(source_elem)
                        continue
                    # 判断是否为列表段落：原段落有 numPr 且启用了 style 模式
                    # 补充检查：如果段落样式为虚拟列表样式也视为列表段落
                    _has_numpr = _elem_has_numbering(elem)
                    _has_list_style = False
                    _is_style_match = False
                    if not _has_numpr:
                        _pPr_el = elem.find(qn('w:pPr'))
                        if _pPr_el is not None:
                            _pS_el = _pPr_el.find(qn('w:pStyle'))
                            if _pS_el is not None:
                                _ps_val = _pS_el.get(qn('w:val'))
                                # 检查1：虚拟列表样式
                                if _ps_val and '列表段落' in str(_ps_val):
                                    _has_list_style = True
                                # 检查2：list_answer_method='style'时匹配 style_id
                                if not _has_list_style and _ps_val and list_answer_method == 'style' and list_answer_style:
                                    _ans_sid = self.get_style_id_by_name(doc, list_answer_style)
                                    if _ans_sid and _ps_val == _ans_sid:
                                        _is_style_match = True
                    is_list_para = ((_has_numpr or _has_list_style or _is_style_match) and 
                                    list_answer_method == 'style' and 
                                    list_answer_style)
                    if is_list_para:
                        # 列表段落：使用指定的列表样式
                        list_sid = self.get_style_id_by_name(doc, list_answer_style)
                        if list_sid:
                            # ★ 修复：基于 elem（有 numPr）创建，而不是 copy_template
                            source_elem = deepcopy(elem)
                            # 修改样式为指定的列表样式
                            _apply_paragraph_style(source_elem, list_sid)
                            remove_keep_original_from_element(source_elem)
                            new_children.append(source_elem)
                        else:
                            # 样式名称无效，回退到普通模板
                            source_elem = deepcopy(copy_template)
                            source_runs = source_elem.findall('.//' + qn('w:r'))
                            orig_runs = elem.findall('.//' + qn('w:r'))
                            for r in source_runs:
                                source_elem.remove(r)
                            for r in orig_runs:
                                source_elem.append(deepcopy(r))
                            remove_keep_original_from_element(source_elem)
                            # ★ 修复：应答原文副本图片段落应用图片兜底样式
                            if enable_image_style and image_style_override:
                                has_img = (elem.find('.//' + qn('w:drawing')) is not None or
                                           elem.find('.//' + qn('w:pict')) is not None or
                                           elem.find('.//' + qn('pic:pic')) is not None or
                                           elem.find('.//' + qn('wp:inline')) is not None or
                                           elem.find('.//' + qn('wp:anchor')) is not None)
                                if has_img:
                                    img_sid = self.get_style_id_by_name(doc, image_style_override)
                                    if img_sid:
                                        _apply_paragraph_style(source_elem, img_sid)
                            new_children.append(source_elem)
                    else:
                        # 非列表段落：用 copy_template 替换内容
                        source_elem = deepcopy(copy_template)
                        source_runs = source_elem.findall('.//' + qn('w:r'))
                        orig_runs = elem.findall('.//' + qn('w:r'))
                        for r in source_runs:
                            source_elem.remove(r)
                        for r in orig_runs:
                            source_elem.append(deepcopy(r))
                        remove_keep_original_from_element(source_elem)
                        # ★ 修复：应答原文副本图片段落应用图片兜底样式
                        if enable_image_style and image_style_override:
                            has_img = (elem.find('.//' + qn('w:drawing')) is not None or
                                       elem.find('.//' + qn('w:pict')) is not None or
                                       elem.find('.//' + qn('pic:pic')) is not None or
                                       elem.find('.//' + qn('wp:inline')) is not None or
                                       elem.find('.//' + qn('wp:anchor')) is not None)
                            if has_img:
                                img_sid = self.get_style_id_by_name(doc, image_style_override)
                                if img_sid:
                                    _apply_paragraph_style(source_elem, img_sid)
                        new_children.append(source_elem)
        
        return insert_count, total_heading_count
    
    def _insert_before_paragraphs(self, children, new_children, answer_template, doc):
        """
        逐段前插入应答句（模式4：before_paragraph）- 支持语义段落分组
        逻辑：
        1. 将连续的语义相关段落分组（短句、引号上下文、列表）
        2. 在每个语义单元前插入一个应答句
        :return: (insert_count, total_heading_count)
        """
        insert_count = 0
        total_heading_count = 0
        
        # 第一步：将元素分组为语义单元
        semantic_groups = self._group_semantic_units(children, doc)
        
        # 第二步：遍历每个语义单元，在单元前插入应答句
        for group in semantic_groups:
            if not group:
                continue
            
            first_elem = group[0]
            
            # 统计标题数量
            if hasattr(first_elem, 'tag') and first_elem.tag == qn('w:p'):
                if self.is_heading_paragraph(first_elem, doc):
                    total_heading_count += len([e for e in group if hasattr(e, 'tag') and e.tag == qn('w:p') and self.is_heading_paragraph(e, doc)])
            
            # 判断是否为需要插入应答句的语义单元
            should_insert = self._should_insert_answer_for_group(group, doc)
            
            if should_insert:
                # 在语义单元前插入应答句
                answer_elem = deepcopy(answer_template)
                new_children.append(answer_elem)
                insert_count += 1
            
            # 添加语义单元中的所有元素
            for elem in group:
                new_children.append(elem)
        
        return insert_count, total_heading_count

    def _insert_after_paragraphs(self, children, new_children, answer_template, doc):
        """
        逐段后插入应答句（模式5：after_paragraph）- 支持语义段落分组
        逻辑：
        1. 将连续的语义相关段落分组（短句、引号上下文、列表）
        2. 在每个语义单元后插入一个应答句
        :return: (insert_count, total_heading_count)
        """
        insert_count = 0
        total_heading_count = 0
        
        # 第一步：将元素分组为语义单元
        semantic_groups = self._group_semantic_units(children, doc)
        
        # 第二步：遍历每个语义单元，在单元后插入应答句
        for group in semantic_groups:
            if not group:
                continue
            
            first_elem = group[0]
            
            # 统计标题数量
            if hasattr(first_elem, 'tag') and first_elem.tag == qn('w:p'):
                if self.is_heading_paragraph(first_elem, doc):
                    total_heading_count += len([e for e in group if hasattr(e, 'tag') and e.tag == qn('w:p') and self.is_heading_paragraph(e, doc)])
            
            # 先添加语义单元中的所有元素
            for elem in group:
                new_children.append(elem)
            
            # 判断是否为需要插入应答句的语义单元
            should_insert = self._should_insert_answer_for_group(group, doc)
            
            if should_insert:
                # 在语义单元后插入应答句
                answer_elem = deepcopy(answer_template)
                new_children.append(answer_elem)
                insert_count += 1
        
        return insert_count, total_heading_count
    
    # ==================== 语义分组辅助方法 ====================
    
    def _is_list_paragraph(self, elem):
        """判断段落是否是列表（有编号或项目符号）
        numId=0 表示无编号，应视为没有列表。
        """
        if not hasattr(elem, 'tag') or elem.tag != qn('w:p'):
            return False
        
        pPr = elem.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                numId = numPr.find(qn('w:numId'))
                if numId is not None:
                    val = numId.get(qn('w:val'))
                    if val == '0':
                        return False
                return True
        
        return False
    
    def _get_paragraph_text(self, elem):
        """获取段落的文本内容"""
        if not hasattr(elem, 'tag') or elem.tag != qn('w:p'):
            return ""
        
        text_elems = elem.findall('.//' + qn('w:t'))
        return ''.join([t.text for t in text_elems if t.text])
    
    def _ends_with_colon_or_quote(self, text):
        """判断文本是否以冒号或引号结尾（需要与下一段合并）"""
        if not text:
            return False
        
        text = text.rstrip()
        
        if text.endswith('\uff1a') or text.endswith(':'):
            return True
        if text.endswith('\u201d') or text.endswith('"'):
            if len(text) > 1 and (text[-2] == '\uff1a' or text[-2] == ':'):
                return True
        
        return False
    
    def _is_short_paragraph(self, text, threshold=20):
        """判断是否为短段落"""
        if not text:
            return True
        return len(text.strip()) < threshold
    
    def _is_manual_numbered_paragraph(self, text):
        """判断段落是否是手动编号（如1.、2）、a.等）"""
        if not text:
            return False
        
        text = text.strip()
        
        patterns = [
            r'^\d+[、\.．]',
            r'^\d+）',
            r'^\d+\)',
            r'^[（(]\d+[）)]',
            r'^[一二三四五六七八九十]+[、\.．]',
            r'^[a-zA-Z][、\.．]',
            r'^[a-zA-Z]）',
            r'^[a-zA-Z]\)',
        ]
        
        for pattern in patterns:
            if re.match(pattern, text):
                return True
        
        return False
    
    def _is_bullet_point_paragraph(self, text):
        """判断段落是否是项目符号列表（如●、■、◆等）"""
        if not text:
            return False
        
        text = text.strip()
        
        bullet_symbols = ['\u25cf', '\u25cb', '\u25a0', '\u25a1', '\u25c6', '\u25c7',
                          '\u25b2', '\u25b3', '\u25ba', '\u25b6', '\u2022', '-', '*']
        
        for symbol in bullet_symbols:
            if text.startswith(symbol):
                return True
        
        return False
    
    def _is_empty_paragraph(self, text):
        """判断段落是否为空行"""
        if not text:
            return True
        return len(text.strip()) == 0
    
    def _group_semantic_units(self, children, doc):
        """
        将元素分组为语义单元
        规则：
        1. 标题单独成组
        2. 连续的列表项合并为一个组
        3. 以冒号/引号结尾的段落与下一段合并
        4. 连续的短段落合并为一个组
        :return: 分组后的列表 [[elem1, elem2], [elem3], ...]
        """
        groups = []
        current_group = []
        
        for i, child in enumerate(children):
            if not hasattr(child, 'tag'):
                if current_group:
                    groups.append(current_group)
                    current_group = []
                groups.append([child])
                continue
            
            is_heading = False
            if child.tag == qn('w:p') and self.is_heading_paragraph(child, doc):
                is_heading = True
            
            text = self._get_paragraph_text(child) if child.tag == qn('w:p') else ""
            
            is_empty = False
            if child.tag == qn('w:p'):
                is_empty = self._is_empty_paragraph(text)
            
            if is_empty:
                continue
            
            is_list = False
            if child.tag == qn('w:p'):
                is_list = self._is_list_paragraph(child)
            
            is_manual_numbered = False
            if child.tag == qn('w:p'):
                is_manual_numbered = self._is_manual_numbered_paragraph(text)
            
            is_bullet_point = False
            if child.tag == qn('w:p'):
                is_bullet_point = self._is_bullet_point_paragraph(text)
            
            if is_heading:
                if current_group:
                    groups.append(current_group)
                    current_group = []
                groups.append([child])
            elif is_list or is_manual_numbered or is_bullet_point:
                should_merge = False
                
                if current_group:
                    if (self._is_last_group_list(current_group, doc) or 
                        self._is_last_group_manual_numbered(current_group) or
                        self._is_last_group_bullet_point(current_group)):
                        should_merge = True
                    else:
                        prev_text = self._get_last_paragraph_text(current_group)
                        if self._ends_with_colon_or_quote(prev_text):
                            should_merge = True
                
                if should_merge:
                    current_group.append(child)
                else:
                    if current_group:
                        groups.append(current_group)
                    current_group = [child]
            elif self._ends_with_colon_or_quote(text):
                current_group.append(child)
            elif self._is_short_paragraph(text) and current_group:
                prev_text = self._get_last_paragraph_text(current_group)
                if self._is_short_paragraph(prev_text) or self._ends_with_colon_or_quote(prev_text):
                    current_group.append(child)
                else:
                    groups.append(current_group)
                    current_group = [child]
            else:
                if current_group:
                    prev_text = self._get_last_paragraph_text(current_group)
                    prev_is_numbered_or_bullet = self._is_last_group_manual_numbered(current_group) or self._is_last_group_bullet_point(current_group)
                    
                    if self._ends_with_colon_or_quote(prev_text):
                        current_group.append(child)
                    elif prev_is_numbered_or_bullet:
                        current_group.append(child)
                    else:
                        groups.append(current_group)
                        current_group = [child]
                else:
                    current_group = [child]
        
        if current_group:
            groups.append(current_group)
        
        return groups
    
    def _is_last_group_list(self, group, doc):
        """检查组中最后一个元素是否是列表"""
        if not group:
            return False
        last_elem = group[-1]
        if hasattr(last_elem, 'tag') and last_elem.tag == qn('w:p'):
            return self._is_list_paragraph(last_elem)
        return False
    
    def _is_last_group_manual_numbered(self, group):
        """检查组中最后一个元素是否是手动编号段落"""
        if not group:
            return False
        last_elem = group[-1]
        if hasattr(last_elem, 'tag') and last_elem.tag == qn('w:p'):
            text = self._get_paragraph_text(last_elem)
            return self._is_manual_numbered_paragraph(text)
        return False
    
    def _is_last_group_bullet_point(self, group):
        """检查组中最后一个元素是否是项目符号段落"""
        if not group:
            return False
        last_elem = group[-1]
        if hasattr(last_elem, 'tag') and last_elem.tag == qn('w:p'):
            text = self._get_paragraph_text(last_elem)
            return self._is_bullet_point_paragraph(text)
        return False
    
    def _get_last_paragraph_text(self, group):
        """获取组中最后一个段落的文本"""
        for elem in reversed(group):
            if hasattr(elem, 'tag') and elem.tag == qn('w:p'):
                return self._get_paragraph_text(elem)
        return ""
    
    def _should_insert_answer_for_group(self, group, doc):
        """判断是否应该为该语义单元插入应答句"""
        if not group:
            return False
        
        first_elem = group[0]
        
        # 排除标题
        if hasattr(first_elem, 'tag') and first_elem.tag == qn('w:p'):
            if self.is_heading_paragraph(first_elem, doc):
                return False
        
        # 排除图片
        if hasattr(first_elem, 'tag') and first_elem.tag == qn('w:p'):
            if len(first_elem.findall('.//' + qn('a:blip'))) > 0:
                return False
        
        return True
    
    def full_convert(self, source_file, template_file, output_file, 
                     custom_style_map=None, do_mood=True, 
                     answer_text=None, answer_style=None,
                     answer_source_style=None,
                     answer_copy_style=None,
                     table_answer_style=None,
                     list_bullet=None, do_answer_insertion=True,
                     answer_mode='before_heading',
                     do_hint_insertion=False, hint_type='text',
                     hint_text='招标文件原文', hint_image_path=None,
                     hint_style='Normal',
                     progress_callback=None, warning_callback=None,
                     table_style_override=None, enable_table_style=False,
                     image_style_override=None, enable_image_style=False,
                     remove_chapter_label=False,
                     list_method='bullet', list_style='Body Text',
                     list_answer_method='bullet', list_answer_style='Body Text',
                     list_answer_bullet='● ',
                     enable_list_style=True):
        """
        完整转换流程：样式转换 -> 提示语插入 -> 语气转换 -> 插入应答句
        固定为7个步骤，跳过的步骤也会计入进度
        
        特殊处理：copy_chapter 模式时，调换语气转换和应答句插入的顺序，
        使第一份副本保留原文（祈使语气），第二份副本完成语气转换。
        
        :param source_file: 源文件
        :param template_file: 模板文件
        :param output_file: 最终输出文件
        :param custom_style_map: 自定义样式映射
        :param do_mood: 是否进行语气转换
        :param answer_text: 应答文本
        :param answer_style: 应答样式
        :param answer_source_style: 应答原文样式（copy_chapter 模式使用）
        :param list_bullet: 列表段落符号
        :param do_answer_insertion: 是否插入应答句
        :param answer_mode: 应答句插入模式
            - 'before_heading': 章节标题后插入（默认）
            - 'after_heading': 章节末尾插入
            - 'copy_chapter': 原文+应答句+应答原文
            - 'before_paragraph': 逐段前插入
            - 'after_paragraph': 逐段后插入
        :param do_hint_insertion: 是否插入章节提示语
        :param hint_type: 提示语类型 'text' 或 'image'
        :param hint_text: 提示语文本内容
        :param hint_image_path: 提示语图片文件路径
        :param hint_style: 提示语段落样式
        :param progress_callback: 进度回调函数 callback(step, message)
        :param warning_callback: 警告回调函数 callback(message)
        :param table_style_override: 表格样式覆盖（当enable_table_style=True时使用）
        :param enable_table_style: 是否启用表格样式覆盖
        :param image_style_override: 图片样式覆盖（当enable_image_style=True时使用）
        :param enable_image_style: 是否启用图片样式覆盖
        :param remove_chapter_label: 是否清除章/节标题编号（如"第一章"、"第一节"）
        :return: (success, actual_output_file, message)
        """
        # 固定7个步骤，确保进度条能正确填满
        if progress_callback:
            progress_callback(1, "开始样式转换...")
        
        # 步骤1：样式转换
        temp_file_1 = output_file.rsplit('.', 1)[0] + "_temp1.docx"
        success, actual_file, msg = self.convert_styles(source_file, template_file, temp_file_1, custom_style_map, list_bullet,
                                           warning_callback,
                                           table_style_override, enable_table_style,
                                           image_style_override, enable_image_style,
                                           remove_chapter_label,
                                           list_method=list_method, list_style=list_style,
                                           enable_list_style=enable_list_style)
        if not success:
            return False, output_file, f"样式转换失败: {msg}"
        
        # 若因重名保存到了备用文件，后续步骤需使用实际文件路径
        temp_file_1 = actual_file
        
        if progress_callback:
            progress_callback(2, f"样式转换完成: {msg}")
        
        # ========== 章节提示语插入（在语气转换之前） ==========
        # 提示语插入在章节标题后、正文开始前，不受语气转换影响
        if do_hint_insertion:
            if progress_callback:
                progress_callback(2.5, "开始插入章节提示语...")
            temp_hint = output_file.rsplit('.', 1)[0] + "_temp_hint.docx"
            success, actual_file, msg = self.insert_hint_paragraph(
                temp_file_1, temp_hint, hint_type, hint_text, hint_image_path, hint_style
            )
            if not success:
                return False, output_file, f"插入提示语失败: {msg}"
            os.remove(temp_file_1)
            temp_file_1 = actual_file  # 若因重名保存到备用文件，需使用实际路径
            print(f"章节提示语插入完成: {msg}")
        
        # ========== 根据 answer_mode 决定流水线顺序 ==========
        # copy_chapter 模式：先插入应答句 → 后语气转换（第一份副本不做语气转换）
        # 其他模式：先语气转换 → 后插入应答句（标准流水线）
        
        actual_output_file = output_file  # 默认使用原始输出文件名
        
        if answer_mode == 'copy_chapter' and do_answer_insertion and do_mood:
            # ===== copy_chapter 模式专用流水线 =====
            # 步骤2-3：插入应答句（在语气转换之前，此时原文未转换）
            if progress_callback:
                progress_callback(3, "开始插入应答句（保留原文模式）...")
            temp_file_2 = output_file.rsplit('.', 1)[0] + "_temp2.docx"
            success, actual_file, msg = self.insert_response_after_headings(
                temp_file_1, temp_file_2, answer_text, answer_style, answer_mode,
                answer_source_style=answer_source_style,
                answer_copy_style=answer_copy_style,
                table_answer_style=table_answer_style,
                list_method=list_method,
                list_style=list_style,
                list_answer_method=list_answer_method,
                list_answer_style=list_answer_style,
                list_answer_bullet=list_answer_bullet,
                enable_image_style=enable_image_style,
                image_style_override=image_style_override
            )
            if not success:
                return False, output_file, f"插入应答句失败: {msg}"
            actual_output_file = actual_file
            
            # 更新 temp_file_1 为应答句插入后的文件
            os.remove(temp_file_1)
            temp_file_1 = actual_file  # 若因重名保存到备用文件，需使用实际路径
            
            if progress_callback:
                progress_callback(4, f"插入应答句完成: {msg}")
            
            # 步骤4-5：语气转换（跳过标记为 keepOriginal 的第一份副本段落）
            # 直接输出到最终文件，由 save_with_retry 处理重名，避免后续清理误删中间文件
            if progress_callback:
                progress_callback(5, "开始语气转换（跳过原文副本）...")
            success, actual_file, msg = self.convert_mood(temp_file_1, output_file)
            if not success:
                return False, output_file, f"语气转换失败: {msg}"
            
            # 语气转换后的实际输出文件才是最终结果
            actual_output_file = actual_file
            
            # 删除插入应答句后的中间临时文件
            try:
                if os.path.exists(temp_file_1):
                    os.remove(temp_file_1)
            except:
                pass
            temp_file_1 = actual_file
            
            if progress_callback:
                progress_callback(6, f"语气转换完成: {msg}")
        
        else:
            # ===== 标准流水线（其他模式或无语气转换） =====
            # 步骤2-3：语气转换（占用2个步骤槽位）
            if do_mood:
                if progress_callback:
                    progress_callback(3, "开始语气转换...")
                temp_file_2 = output_file.rsplit('.', 1)[0] + "_temp2.docx"
                success, actual_file, msg = self.convert_mood(temp_file_1, temp_file_2)
                if not success:
                    return False, output_file, f"语气转换失败: {msg}"
                if progress_callback:
                    progress_callback(4, f"语气转换完成: {msg}")
                os.remove(temp_file_1)  # 清理临时文件
                temp_file_1 = actual_file  # 若因重名保存到备用文件，需使用实际路径
            else:
                # 跳过语气转换，但仍然占用步骤3和4
                if progress_callback:
                    progress_callback(3, "跳过语气转换")
                    progress_callback(4, "已跳过语气转换")
            
            # 步骤5-6：插入应答句（占用2个步骤槽位）
            if do_answer_insertion:
                if progress_callback:
                    progress_callback(5, "开始插入应答句...")
                success, actual_file, msg = self.insert_response_after_headings(
                    temp_file_1, output_file, answer_text, answer_style, answer_mode,
                    list_method=list_method,
                    list_style=list_style,
                    list_answer_method=list_answer_method,
                    list_answer_style=list_answer_style,
                    list_answer_bullet=list_answer_bullet
                )
                if not success:
                    return False, output_file, f"插入应答句失败: {msg}"
                
                actual_output_file = actual_file  # 更新为实际文件名
                
                if progress_callback:
                    progress_callback(6, f"插入应答句完成: {msg}")
            else:
                # 不插入应答句，直接复制文件，但仍然占用步骤5和6
                if progress_callback:
                    progress_callback(5, "跳过应答句插入")
                import shutil
                actual_output_file = self._generate_unique_filename(output_file)
                if actual_output_file != output_file:
                    print(f"  检测到重名文件，使用备用文件名: {actual_output_file}")
                shutil.copy2(temp_file_1, actual_output_file)
                if progress_callback:
                    progress_callback(6, "已跳过应答句插入")
        
        # 清理临时文件（只删除文件名包含 _temp 的中间文件，避免误删最终输出文件）
        try:
            if os.path.exists(temp_file_1) and "_temp" in os.path.basename(temp_file_1):
                os.remove(temp_file_1)
        except:
            pass
        
        # 步骤7：完成
        if progress_callback:
            progress_callback(7, "转换全部完成！")
        
        return True, actual_output_file, "转换成功完成！"
    
    def _generate_unique_filename(self, output_file):
        """
        生成不重复的文件名：若原始文件名已存在，则追加 _HHMMSS 时间戳；
        若同一秒内仍有冲突，再追加三位序号（_001）。
        :param output_file: 原始输出文件路径
        :return: 可用的文件路径
        """
        import os

        if not os.path.exists(output_file):
            return output_file

        base, ext = os.path.splitext(output_file)
        time_suffix = datetime.now().strftime("_%H%M%S")
        candidate = f"{base}{time_suffix}{ext}"

        if not os.path.exists(candidate):
            return candidate

        # 极端情况：同一秒内仍有冲突，追加序号
        for i in range(1, 1000):
            candidate = f"{base}{time_suffix}_{i:03d}{ext}"
            if not os.path.exists(candidate):
                return candidate

        # 兜底，理论上不会走到这里
        return output_file

    def save_with_retry(self, doc, output_file, max_retries=10):
        """
        智能保存文档：若目标文件已存在（重名），自动追加 _HHMMSS 时间戳；
        保存过程中如遇占用，也会自动生成新的备用文件名并重试。
        :param doc: Document对象
        :param output_file: 原始输出文件路径
        :param max_retries: 最大重试次数
        :return: (success, actual_output_file, message)
        """
        import os
        import time

        # 首次保存：若存在重名文件则生成带时间戳的备用文件名
        current_file = self._generate_unique_filename(output_file)
        if current_file != output_file:
            print(f"  检测到重名文件，使用备用文件名: {current_file}")

        for attempt in range(max_retries):
            try:
                doc.save(current_file)
                if current_file == output_file:
                    return True, current_file, f"文档已保存到 {current_file}"
                else:
                    return True, current_file, f"检测到重名，文档已保存到: {current_file}"
            except (PermissionError, OSError, IOError) as e:
                # 文件在保存过程中被占用（罕见情况）
                if attempt == 0:
                    print(f"  警告：保存文档失败（文件可能被占用）: {e}")

                # 生成新的文件名
                base, ext = os.path.splitext(output_file)
                time_suffix = datetime.now().strftime("_%H%M%S")
                current_file = f"{base}{time_suffix}{ext}"

                # 如果新文件名仍冲突，追加序号避免覆盖
                idx = 1
                while os.path.exists(current_file) and idx < 1000:
                    current_file = f"{base}{time_suffix}_{idx:03d}{ext}"
                    idx += 1

                print(f"  尝试备用文件名: {current_file}")

                # 稍等片刻再重试
                time.sleep(0.3)
            except Exception as e:
                # 其他异常直接返回失败
                return False, output_file, f"保存文档失败: {e}"

        # 重试次数用尽
        return False, output_file, f"无法保存文档，已尝试 {max_retries} 次"


if __name__ == "__main__":
    # 测试代码
    converter = DocumentConverter()
    print("文档转换器模块加载成功")
    print(f"Pillow可用: {PIL_AVAILABLE}")
