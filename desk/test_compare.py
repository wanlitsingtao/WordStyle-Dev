# -*- coding: utf-8 -*-
"""对比分析两个转换后的文档：不勾选 vs 勾选"清除第X章/第X节"
1517_converted.docx = 不勾选
1517_converted_154902.docx = 勾选
"""
from docx import Document
from docx.oxml.ns import qn
import os

base = r'E:\LingMa\WordStyle\WordStyle-Dev\test'
f1 = os.path.join(base, '1517_converted.docx')
f2 = os.path.join(base, '1517_converted_154902.docx')

if not os.path.exists(f1):
    print(f'文件不存在: {f1}')
    # 查找最新文件
    for f in os.listdir(base):
        if '1517_converted' in f and f.endswith('.docx'):
            print(f'  找到: {f}')
            f1 = os.path.join(base, f)
            break
if not os.path.exists(f2):
    print(f'文件不存在: {f2}')
    for f in os.listdir(base):
        if '1517_converted_154902' in f:
            print(f'  找到: {f}')
            f2 = os.path.join(base, f)
            break

doc1 = Document(f1)
doc2 = Document(f2)

print(f'文件1(不勾选): {os.path.basename(f1)}')
print(f'  段落数: {len(doc1.paragraphs)}')
print(f'  表格数: {len(doc1.tables)}')
print(f'文件2(勾选): {os.path.basename(f2)}')
print(f'  段落数: {len(doc2.paragraphs)}')
print(f'  表格数: {len(doc2.tables)}')

# 1. 对比三个关键标题段落
print('\n' + '='*60)
print('1. 关键标题段落对比')
print('='*60)

# 从源文档获取这三个段落的文本
src = Document(os.path.join(base, '1517.docx'))
target_indices = []
for i, p in enumerate(src.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    # 找15号线、17号线、第六章相关标题
    if '15号线乘客信息系统' in p.text and len(p.text) < 40:
        target_indices.append((i, p.text[:60], '15号线'))
    elif '17号线一期' in p.text and len(p.text) < 40:
        target_indices.append((i, p.text[:60], '17号线'))
    elif '第六章' in p.text or '图纸（如有）' in p.text:
        target_indices.append((i, p.text[:60], '第六章'))

# 现在在转换后的文档中找这些段落
for src_idx, src_text, label in target_indices:
    print(f'\n--- {label} (源文档§{src_idx}): {src_text!r} ---')
    
    # 在doc1（不勾选）中找
    found1 = False
    for i, p in enumerate(doc1.paragraphs):
        if label == '15号线' and '15号线' in p.text:
            # 检查大纲级别
            pPr = p._element.find(qn('w:pPr'))
            olvl = pPr.find(qn('w:outlineLvl')) if pPr is not None else None
            numPr = pPr.find(qn('w:numPr')) if pPr is not None else None
            olvl_val = olvl.get(qn('w:val')) if olvl is not None else None
            print(f'  不勾选 §{i}: {p.text[:80]!r}  style={p.style.name} outlineLvl={olvl_val} numPr={numPr is not None}')
            found1 = True
            break
    if not found1:
        # 用style名称搜Heading
        for i, p in enumerate(doc1.paragraphs):
            if 'Heading' in p.style.name and ('15号线' in p.text or '乘客信息系统' in p.text):
                pPr = p._element.find(qn('w:pPr'))
                olvl = pPr.find(qn('w:outlineLvl')) if pPr is not None else None
                olvl_val = olvl.get(qn('w:val')) if olvl is not None else None
                print(f'  不勾选 §{i}: {p.text[:80]!r}  style={p.style.name} outlineLvl={olvl_val}')
                found1 = True
                break
    
    # 在doc2（勾选）中找
    found2 = False
    for i, p in enumerate(doc2.paragraphs):
        if label == '15号线' and '15号线' in p.text:
            pPr = p._element.find(qn('w:pPr'))
            olvl = pPr.find(qn('w:outlineLvl')) if pPr is not None else None
            numPr = pPr.find(qn('w:numPr')) if pPr is not None else None
            olvl_val = olvl.get(qn('w:val')) if olvl is not None else None
            print(f'  勾选   §{i}: {p.text[:80]!r}  style={p.style.name} outlineLvl={olvl_val} numPr={numPr is not None}')
            found2 = True
            break
    if not found2:
        for i, p in enumerate(doc2.paragraphs):
            if 'Heading' in p.style.name and ('15号线' in p.text or '乘客信息系统' in p.text):
                pPr = p._element.find(qn('w:pPr'))
                olvl = pPr.find(qn('w:outlineLvl')) if pPr is not None else None
                olvl_val = olvl.get(qn('w:val')) if olvl is not None else None
                print(f'  勾选   §{i}: {p.text[:80]!r}  style={p.style.name} outlineLvl={olvl_val}')
                found2 = True
                break

# 2. 检查大纲级别的标题列表
print('\n' + '='*60)
print('2. 所有大纲标题(不勾选)')
print('='*60)
count1 = 0
for i, p in enumerate(doc1.paragraphs):
    pPr = p._element.find(qn('w:pPr'))
    olvl = pPr.find(qn('w:outlineLvl')) if pPr is not None else None
    if olvl is not None and p.text.strip():
        olvl_val = olvl.get(qn('w:val'))
        print(f'  §{i} [Lv{olvl_val}] style={p.style.name}: {p.text[:80]!r}')
        count1 += 1
        if count1 >= 20:
            print('  ... (只显示前20个)')
            break

print(f'\n所有大纲标题(勾选)')
count2 = 0
for i, p in enumerate(doc2.paragraphs):
    pPr = p._element.find(qn('w:pPr'))
    olvl = pPr.find(qn('w:outlineLvl')) if pPr is not None else None
    if olvl is not None and p.text.strip():
        olvl_val = olvl.get(qn('w:val'))
        print(f'  §{i} [Lv{olvl_val}] style={p.style.name}: {p.text[:80]!r}')
        count2 += 1
        if count2 >= 20:
            print('  ... (只显示前20个)')
            break

print(f'\n大纲标题数: 不勾选={count1}, 勾选={count2}')
