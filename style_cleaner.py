# -*- coding: utf-8 -*-
"""
模板样式精简引擎（T04 / 工具箱 Tab B）
无 Streamlit 依赖，可独立测试。

职责：
1. 分析与文档转换模块一致的模板段落样式及使用次数（含表格内段落）
2. 对 Word 内置样式做最小保留（Normal / Heading / Body Text / List Paragraph / Header / Footer），其余可删
3. 删除样式：对 basedOn/next 引用自动重指向到首个保留祖先，清理 link 悬空引用，保证文档可正常打开
"""
from typing import List, Dict, Set, Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

# 样式类型中文名
STYLE_TYPE_LABELS = {
    WD_STYLE_TYPE.PARAGRAPH: "段落",
    WD_STYLE_TYPE.CHARACTER: "字符",
    WD_STYLE_TYPE.TABLE: "表格",
    WD_STYLE_TYPE.LIST: "列表",
}

# Word 内置样式的最小保留集：这些样式即使未使用也必须保留，
# 否则会影响文档默认格式或文档转换模块的硬编码映射目标。
ESSENTIAL_STYLE_NAMES = {
    'Normal',
    'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5',
    'Heading 6', 'Heading 7', 'Heading 8', 'Heading 9',
    'Body Text',
    'List Paragraph',
    'Header', 'Footer',
}


class StyleCleaner:
    """模板样式精简引擎（无状态工具类）"""

    @staticmethod
    def _iter_all_paragraphs(doc):
        """遍历主体 + 表格 + 文本框段落。"""
        paragraphs = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
        try:
            from docx.text.paragraph import Paragraph
            body_parent = doc.paragraphs[0]._parent if doc.paragraphs else None
            if body_parent is not None:
                for txbx in doc.element.body.iter(qn('w:txbxContent')):
                    for p in txbx.findall(qn('w:p')):
                        paragraphs.append(Paragraph(p, body_parent))
        except Exception:
            pass
        return paragraphs

    @staticmethod
    def analyze_styles(docx_file) -> Dict:
        """分析模板文档样式及使用次数。

        Returns:
            {
                "total": int,
                "builtin_count": int,   # Word 内置样式数
                "custom_count": int,    # 自定义样式数
                "used": int,
                "unused": int,
                "cleanable": int,       # 可清理（未使用、非必需保留、且无 basedOn/next 依赖）
                "styles": [
                    {
                        "style_id": str,
                        "name": str,
                        "type": str,          # "段落"/"字符"/"表格"/"列表"
                        "usage_count": int,
                        "builtin": bool,      # 是否内置样式
                        "protected": bool,    # 是否属于最小保留集（默认不可删）
                    }, ...
                ],
            }
        """
        doc = Document(docx_file)

        # 统计使用次数（按 style_id）
        usage = {}
        for para in StyleCleaner._iter_all_paragraphs(doc):
            if para.style is not None:
                sid = para.style.style_id
                if sid:
                    usage[sid] = usage.get(sid, 0) + 1

        # 与文档转换模块一致：仅统计模板中的"段落样式且有名"
        # （见 doc_converter.get_template_styles / components.upload.get_template_styles_list）。
        styles = []
        for style in doc.styles:
            try:
                stype = style.type
            except Exception:
                stype = None
            if stype != WD_STYLE_TYPE.PARAGRAPH:
                continue
            name = style.name or ""
            if not name:
                continue
            sid = style.style_id
            type_label = STYLE_TYPE_LABELS.get(stype, "未知")
            is_builtin = bool(getattr(style, 'builtin', False))
            count = usage.get(sid, 0)
            # 最小保留集：仅内置且属于必需样式的才默认保护，
            # 其余内置样式若未使用也可清理。
            protected = is_builtin and (name in ESSENTIAL_STYLE_NAMES)
            styles.append({
                "style_id": sid,
                "name": name,
                "type": type_label,
                "usage_count": count,
                "builtin": is_builtin,
                "protected": protected,
            })

        used_count = sum(1 for s in styles if s["usage_count"] > 0)
        builtin_count = sum(1 for s in styles if s["builtin"])
        # 可清理数：未使用且非最小保留集的样式数（依赖引用会被自动重指向）
        cleanable = sum(
            1 for s in styles
            if s["usage_count"] == 0 and not s["protected"]
        )
        return {
            "total": len(styles),
            "builtin_count": builtin_count,
            "custom_count": len(styles) - builtin_count,
            "used": used_count,
            "unused": len(styles) - used_count,
            "cleanable": cleanable,
            "styles": styles,
        }

    @staticmethod
    def _resolve_reference(sid, delete_set: Set[str], ref_map: Dict, default: Optional[str] = None):
        """沿引用链找到第一个不在删除集合中的样式 id。

        Args:
            sid: 起始样式 id（通常在删除集合中）。
            delete_set: 待删除样式 id 集合。
            ref_map: styleId -> 引用目标（如 basedOn / next）。
            default: 找不到可保留祖先时返回的默认值。

        Returns:
            第一个非删除样式 id；若引用链断裂或成环，返回 default。
        """
        seen = set()
        cur = sid
        while cur is not None:
            if cur not in delete_set:
                return cur
            if cur in seen:
                return default
            seen.add(cur)
            cur = ref_map.get(cur)
        return default

    @staticmethod
    def cleanup_styles(docx_file, output_file, delete_style_ids: List[str]) -> Dict:
        """删除指定样式并保存新文档。

        Args:
            docx_file: 输入 docx 路径。
            output_file: 输出 docx 路径。
            delete_style_ids: 要删除的 styleId 列表（内部会过滤最小保留集，
                并对 basedOn/next 引用自动重指向，对 link 引用清理）。

        Returns:
            {"deleted": int, "skipped_protected": int, "repointed": int, "message": str}
        """
        doc = Document(docx_file)
        delete_set = set(delete_style_ids or [])

        # 1. 过滤最小保留集样式（内置且必需：Normal / Heading 1-9 / Body Text / List Paragraph / Header / Footer）
        skipped_protected = 0
        for style in doc.styles:
            if style.style_id in delete_set:
                if bool(getattr(style, 'builtin', False)) and (style.name or "") in ESSENTIAL_STYLE_NAMES:
                    delete_set.discard(style.style_id)
                    skipped_protected += 1

        # 构建删除前的引用链（用于重指向）
        styles_elm = doc.styles.element
        basedon_map = {}
        next_map = {}
        for style_elm in styles_elm.findall(qn('w:style')):
            sid = style_elm.get(qn('w:styleId'))
            b = style_elm.find(qn('w:basedOn'))
            n = style_elm.find(qn('w:next'))
            basedon_map[sid] = b.get(qn('w:val')) if b is not None else None
            next_map[sid] = n.get(qn('w:val')) if n is not None else None

        # 2. 物理删除样式元素
        deleted = 0
        for style_elm in list(styles_elm.findall(qn('w:style'))):
            sid = style_elm.get(qn('w:styleId'))
            if sid in delete_set:
                styles_elm.remove(style_elm)
                deleted += 1

        # 3. 重指向 / 清理保留样式中指向已删除样式的引用
        repointed = 0
        for style_elm in styles_elm.findall(qn('w:style')):
            # basedOn：重指向到被删样式的第一个非删除祖先；无祖先则移除（等价于基于 Normal）
            b = style_elm.find(qn('w:basedOn'))
            if b is not None and b.get(qn('w:val')) in delete_set:
                new_base = StyleCleaner._resolve_reference(
                    b.get(qn('w:val')), delete_set, basedon_map
                )
                if new_base:
                    b.set(qn('w:val'), new_base)
                else:
                    style_elm.remove(b)
                repointed += 1
            # next：重指向到第一个非删除样式，找不到则移除（移除后 Word 默认使用 Normal）
            n = style_elm.find(qn('w:next'))
            if n is not None and n.get(qn('w:val')) in delete_set:
                new_next = StyleCleaner._resolve_reference(
                    n.get(qn('w:val')), delete_set, next_map
                )
                if new_next:
                    n.set(qn('w:val'), new_next)
                else:
                    style_elm.remove(n)
                repointed += 1
            # link：直接移除指向已删除样式的悬空链接
            link = style_elm.find(qn('w:link'))
            if link is not None and link.get(qn('w:val')) in delete_set:
                style_elm.remove(link)
                repointed += 1

        doc.save(output_file)

        message = f"已删除 {deleted} 个样式"
        if skipped_protected:
            message += f"，跳过 {skipped_protected} 个最小保留样式"
        if repointed:
            message += f"，重指向 {repointed} 处引用"
        return {
            "deleted": deleted,
            "skipped_protected": skipped_protected,
            "repointed": repointed,
            "message": message,
        }
